#!/usr/bin/env python3
"""
Система створення Word альбомів з точними розмірами
Генерує професійні альбоми з титульною сторінкою, описом та сторінками зображень
"""

import os
import json
import tempfile
from datetime import datetime
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass
from PIL import Image

# Word документи
try:
    from docx import Document
    from docx.shared import Cm, Mm, Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_STYLE_TYPE
    from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.enum.section import WD_SECTION_START
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn
    from docx.oxml.shared import qn as shared_qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from PyQt5.QtCore import QObject, pyqtSignal

from core.constants import ALBUM, WORD_STYLES, TEMPLATE_DEFAULTS
from utils.file_utils import get_templates_directory, ensure_directory_exists


@dataclass
class ImageData:
    """Дані про оброблене зображення для альбому"""
    filename: str                    # Назва файлу
    image_path: str                  # Шлях до оригінального зображення  
    processed_image_path: str        # Шлях до обробленого зображення
    target_number: str               # Номер цілі
    azimuth: float                   # Азимут в градусах
    range_km: float                  # Дальність в кілометрах
    height: str                      # Висота цілі
    obstacles: str                   # Перешкоди
    detection: str                   # Тип виявлення
    timestamp: datetime              # Час аналізу
    
    def get_range_formatted(self) -> str:
        """Отримання відформатованої дальності"""
        return f"{self.range_km:.1f}"
    
    def get_azimuth_formatted(self) -> str:
        """Отримання відформатованого азимуту"""
        return f"{self.azimuth:.0f}°"


@dataclass
class TitlePageData:
    """Дані для титульної сторінки альбому"""
    document_date: str               # Дата документу
    unit_info: str                   # Інформація про підрозділ
    commander_rank: str              # Звання командира
    commander_name: str              # Ім'я командира
    chief_of_staff_rank: str        # Звання начальника штабу
    chief_of_staff_name: str        # Ім'я начальника штабу
    template_name: str = "default"   # Назва використаного шаблону


class AlbumCreator(QObject):
    """
    Створювач Word альбомів з професійним форматуванням
    
    Функціональність:
    - Титульна сторінка з підтримкою шаблонів
    - Сторінка опису документів з підписами
    - Сторінки з зображеннями (точні розміри таблиць)
    - Критично важливі константи розмірів:
      * Таблиці: 205мм ширина × 130мм висота  
      * Поля сторінок: 2.5мм ліворуч для таблиць
      * Пропорції опису РЛС: 28.60% × 19.54%
    """
    
    # Сигнали для відстеження прогресу
    progress_updated = pyqtSignal(int, str)    # прогрес%, повідомлення
    album_created = pyqtSignal(str)            # шлях до створеного альбому
    error_occurred = pyqtSignal(str)           # текст помилки
    
    def __init__(self):
        super().__init__()
        
        # Перевірка доступності Word бібліотек
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx library not available")
        
        # Поточний стан
        self.document: Optional[Document] = None
        self.current_images: List[ImageData] = []
        self.title_data: Optional[TitlePageData] = None
        self.output_path: Optional[str] = None
        
        # Шаблони титульних сторінок
        self.templates: Dict[str, dict] = {}
        self._load_templates()
        
        print("AlbumCreator ініціалізовано")
    
    # ===============================
    # ОСНОВНІ МЕТОДИ СТВОРЕННЯ АЛЬБОМУ
    # ===============================
    
    def create_album(self, images_data: List[ImageData], title_data: TitlePageData, 
                    output_path: str) -> bool:
        """
        Створення повного альбому з усіма сторінками
        
        Args:
            images_data: Список даних про оброблені зображення
            title_data: Дані для титульної сторінки
            output_path: Шлях для збереження альбому
            
        Returns:
            True якщо альбом створено успішно
        """
        try:
            self.current_images = images_data
            self.title_data = title_data
            self.output_path = output_path
            
            print(f"=== Створення Word альбому ===")
            print(f"Зображень: {len(images_data)}")
            print(f"Вихідний файл: {os.path.basename(output_path)}")
            
            # Ініціалізація документу
            self.progress_updated.emit(5, "Ініціалізація документу...")
            self._initialize_document()
            
            # 1. Титульна сторінка
            self.progress_updated.emit(15, "Створення титульної сторінки...")
            self._create_title_page()
            
            # 2. Сторінка опису документів
            self.progress_updated.emit(25, "Створення сторінки опису...")
            self._create_description_page()
            
            # 3. Сторінки з зображеннями
            self.progress_updated.emit(35, "Створення сторінок зображень...")
            self._create_image_pages()
            
            # 4. Збереження документу
            self.progress_updated.emit(90, "Збереження документу...")
            self._save_document()
            
            self.progress_updated.emit(100, "Альбом створено успішно!")
            self.album_created.emit(output_path)
            
            print(f"✅ Альбом створено: {os.path.basename(output_path)}")
            return True
            
        except Exception as e:
            error_msg = f"Помилка створення альбому: {str(e)}"
            print(f"❌ {error_msg}")
            self.error_occurred.emit(error_msg)
            return False
    
    def _initialize_document(self):
        """Ініціалізація нового документу"""
        self.document = Document()
        
        # Налаштування базових стилів документу
        self._setup_document_styles()
        
        print("Документ ініціалізовано")
    
    def _save_document(self):
        """Збереження документу у файл"""
        if not self.document or not self.output_path:
            raise ValueError("Документ або шлях не встановлені")
        
        # Переконуємося що директорія існує
        output_dir = os.path.dirname(self.output_path)
        if output_dir:
            ensure_directory_exists(output_dir)
        
        # Збереження
        self.document.save(self.output_path)
        print(f"Документ збережено: {self.output_path}")
    
    # ===============================
    # СТИЛІ ТА ФОРМАТУВАННЯ ДОКУМЕНТУ
    # ===============================
    
    def _setup_document_styles(self):
        """Налаштування стилів документу"""
        styles = self.document.styles
        
        # Базовий стиль документу (Arial, як в Legacy)
        try:
            normal_style = styles['Normal']
            normal_font = normal_style.font
            normal_font.name = WORD_STYLES.FONT_NAME
            normal_font.size = Pt(WORD_STYLES.FONT_SIZE_NORMAL)
        except Exception as e:
            print(f"Помилка налаштування базового стилю: {e}")
        
        print("Стилі документу налаштовано")
    
    def _set_page_margins(self, section, margin_type: str):
        """
        Встановлення полів сторінки
        
        Args:
            section: Секція документу
            margin_type: Тип полів ('standard' або 'table')
        """
        if margin_type == 'table':
            # КРИТИЧНІ поля для сторінок з таблицями
            section.left_margin = Mm(ALBUM.TABLE_PAGES_LEFT_MARGIN)    # 2.5мм ліворуч!
            section.right_margin = Mm(ALBUM.TABLE_PAGES_RIGHT_MARGIN)  # 5мм праворуч
            section.top_margin = Mm(ALBUM.TABLE_PAGES_TOP_MARGIN)      # 20мм зверху
            section.bottom_margin = Mm(ALBUM.TABLE_PAGES_BOTTOM_MARGIN) # 5мм знизу
        else:
            # Стандартні поля для титульної та описової сторінок
            section.left_margin = Mm(ALBUM.STANDARD_LEFT_MARGIN)       # 25мм
            section.right_margin = Mm(ALBUM.STANDARD_RIGHT_MARGIN)     # 25мм
            section.top_margin = Mm(ALBUM.STANDARD_TOP_MARGIN)         # 10мм
            section.bottom_margin = Mm(ALBUM.STANDARD_BOTTOM_MARGIN)   # 10мм
    
    # ===============================
    # ТИТУЛЬНА СТОРІНКА
    # ===============================
    
    def _create_title_page(self):
        """Створення титульної сторінки з використанням шаблону"""
        if not self.title_data:
            raise ValueError("Дані титульної сторінки не встановлені")
        
        # Встановлення стандартних полів для титульної сторінки
        section = self.document.sections[0]
        self._set_page_margins(section, 'standard')
        
        # Завантаження шаблону
        template = self._get_template(self.title_data.template_name)
        
        if template:
            self._create_title_from_template(template)
        else:
            self._create_default_title_page()
        
        print(f"Титульна сторінка створена (шаблон: {self.title_data.template_name})")
    
    def _create_title_from_template(self, template: dict):
        """
        Створення титульної сторінки на основі шаблону
        
        Args:
            template: Словник з конфігурацією шаблону
        """
        elements = template.get('elements', [])
        
        for element in elements:
            element_type = element.get('type')
            
            if element_type == 'heading':
                self._add_title_heading(element)
            elif element_type == 'text':
                self._add_title_text(element)
            elif element_type == 'table':
                self._add_title_table(element)
            elif element_type == 'spacing':
                self._add_title_spacing(element)
    
    def _add_title_heading(self, element: dict):
        """Додавання заголовку титульної сторінки"""
        text = self._substitute_template_variables(element.get('text', ''))
        
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = paragraph.add_run(text)
        run.font.size = Pt(element.get('font_size', 16))
        run.font.bold = element.get('bold', True)
        run.font.name = WORD_STYLES.FONT_NAME
    
    def _add_title_text(self, element: dict):
        """Додавання тексту титульної сторінки"""
        text = self._substitute_template_variables(element.get('text', ''))
        
        paragraph = self.document.add_paragraph()
        alignment_map = {
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'right': WD_ALIGN_PARAGRAPH.RIGHT
        }
        paragraph.alignment = alignment_map.get(element.get('alignment', 'center'))
        
        run = paragraph.add_run(text)
        run.font.size = Pt(element.get('font_size', 12))
        run.font.bold = element.get('bold', False)
        run.font.name = WORD_STYLES.FONT_NAME
    
    def _add_title_table(self, element: dict):
        """Додавання таблиці на титульну сторінку"""
        rows_data = element.get('rows', [])
        if not rows_data:
            return
        
        # Створення таблиці
        table = self.document.add_table(rows=len(rows_data), cols=2)
        table.autofit = False
        
        # Налаштування ширини колонок
        col_widths = element.get('column_widths', [0.3, 0.7])
        table.columns[0].width = Cm(21 * col_widths[0])  # Загальна ширина ~21см
        table.columns[1].width = Cm(21 * col_widths[1])
        
        # Заповнення рядків
        for i, row_data in enumerate(rows_data):
            row = table.rows[i]
            
            # Перша колонка (назва поля)
            cell1 = row.cells[0]
            cell1.text = row_data.get('label', '')
            cell1.paragraphs[0].runs[0].font.bold = True
            cell1.paragraphs[0].runs[0].font.size = Pt(11)
            
            # Друга колонка (значення)
            cell2 = row.cells[1]
            value = self._substitute_template_variables(row_data.get('value', ''))
            cell2.text = value
            cell2.paragraphs[0].runs[0].font.size = Pt(11)
    
    def _add_title_spacing(self, element: dict):
        """Додавання відступу на титульну сторінку"""
        lines = element.get('lines', 1)
        for _ in range(lines):
            self.document.add_paragraph()
    
    def _substitute_template_variables(self, text: str) -> str:
        """
        Підстановка змінних у текст шаблону
        
        Args:
            text: Текст з змінними у форматі {variable_name}
            
        Returns:
            Текст з підставленими значеннями
        """
        if not self.title_data:
            return text
        
        variables = {
            'document_date': self.title_data.document_date,
            'unit_info': self.title_data.unit_info,
            'commander_rank': self.title_data.commander_rank,
            'commander_name': self.title_data.commander_name,
            'chief_of_staff_rank': self.title_data.chief_of_staff_rank,
            'chief_of_staff_name': self.title_data.chief_of_staff_name,
            'total_images': str(len(self.current_images)),
            'current_date': datetime.now().strftime('%d.%m.%Y')
        }
        
        try:
            return text.format(**variables)
        except KeyError as e:
            print(f"Невідома змінна в шаблоні: {e}")
            return text
    
    def _create_default_title_page(self):
        """Створення титульної сторінки за замовчуванням"""
        # Заголовок
        title_paragraph = self.document.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run("АЛЬБОМ ЗОБРАЖЕНЬ")
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run.font.name = WORD_STYLES.FONT_NAME
        
        # Відступ
        self.document.add_paragraph()
        
        # Дата
        if self.title_data.document_date:
            date_paragraph = self.document.add_paragraph()
            date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_run = date_paragraph.add_run(f"Дата: {self.title_data.document_date}")
            date_run.font.size = Pt(12)
            date_run.font.name = WORD_STYLES.FONT_NAME
    
    # ===============================
    # СТОРІНКА ОПИСУ ДОКУМЕНТІВ
    # ===============================
    
    def _create_description_page(self):
        """Створення сторінки опису документів з підписами"""
        # Додавання нової сторінки
        self.document.add_section(WD_SECTION_START.NEW_PAGE)
        current_section = self.document.sections[-1]
        self._set_page_margins(current_section, 'standard')
        
        # Заголовок сторінки
        heading = self.document.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading_run = heading.add_run("ОПИС ДОКУМЕНТІВ")
        heading_run.font.size = Pt(14)
        heading_run.font.bold = True
        heading_run.font.name = WORD_STYLES.FONT_NAME
        
        # Відступ
        self.document.add_paragraph()
        
        # Опис альбому
        description = self.document.add_paragraph()
        desc_text = (f"Альбом містить {len(self.current_images)} оброблених зображень "
                    f"з азимутальною сіткою для аналізу координат та дальності цілей.")
        description.add_run(desc_text).font.name = WORD_STYLES.FONT_NAME
        
        # Додаткові відступи для підписів
        for _ in range(10):
            self.document.add_paragraph()
        
        # Підписи (як в Legacy)
        self._add_signature_lines()
        
        print("Сторінка опису створена")
    
    def _add_signature_lines(self):
        """Додавання ліній для підписів"""
        if not self.title_data:
            return
        
        # Відступ перед підписами
        self.document.add_paragraph()
        
        # Підпис командира
        if self.title_data.commander_rank and self.title_data.commander_name:
            commander_para = self.document.add_paragraph()
            commander_text = f"{self.title_data.commander_rank}    __________________    {self.title_data.commander_name}"
            commander_para.add_run(commander_text).font.name = WORD_STYLES.FONT_NAME
        
        # Відступ між підписами
        self.document.add_paragraph()
        self.document.add_paragraph()
        
        # Підпис начальника штабу
        if self.title_data.chief_of_staff_rank and self.title_data.chief_of_staff_name:
            chief_para = self.document.add_paragraph()
            chief_text = f"{self.title_data.chief_of_staff_rank}    __________________    {self.title_data.chief_of_staff_name}"
            chief_para.add_run(chief_text).font.name = WORD_STYLES.FONT_NAME
    
    # ===============================
    # СТОРІНКИ З ЗОБРАЖЕННЯМИ
    # ===============================
    
    def _create_image_pages(self):
        """Створення сторінок з зображеннями (по 2 на сторінку)"""
        if not self.current_images:
            return
        
        # Групування зображень по 2 на сторінку
        images_per_page = 2
        page_count = 0
        
        for i in range(0, len(self.current_images), images_per_page):
            page_images = self.current_images[i:i + images_per_page]
            page_count += 1
            
            # Оновлення прогресу
            progress = 35 + int((i / len(self.current_images)) * 50)
            self.progress_updated.emit(progress, f"Створення сторінки {page_count}...")
            
            # Створення сторінки
            self._create_single_image_page(page_images, page_count)
        
        print(f"Створено {page_count} сторінок зображень")
    
    def _create_single_image_page(self, images: List[ImageData], page_number: int):
        """
        Створення однієї сторінки з зображеннями
        
        Args:
            images: Список зображень для сторінки (1-2 зображення)
            page_number: Номер сторінки
        """
        # Додавання нової сторінки з полями для таблиць
        self.document.add_section(WD_SECTION_START.NEW_PAGE)
        current_section = self.document.sections[-1]
        self._set_page_margins(current_section, 'table')  # КРИТИЧНІ поля!
        
        # Створення таблиці для кожного зображення
        for i, image_data in enumerate(images):
            if i > 0:
                # Відступ між таблицями на одній сторінці
                self.document.add_paragraph()
            
            self._create_image_table(image_data)
    
    def _create_image_table(self, image_data: ImageData):
        """
        Створення таблиці для одного зображення з ТОЧНИМИ розмірами
        
        Args:
            image_data: Дані зображення
        """
        # КРИТИЧНІ РОЗМІРИ: 205мм ширина × 130мм висота
        table_width_mm = ALBUM.TABLE_WIDTH_MM      # 205мм
        table_height_mm = ALBUM.TABLE_HEIGHT_MM    # 130мм
        
        # Структура таблиці: 9 рядків × 9 колонок (як в Legacy)
        table = self.document.add_table(rows=9, cols=9)
        table.autofit = False
        
        # Встановлення точної ширини таблиці
        table.width = Mm(table_width_mm)
        
        # Розподіл ширини колонок
        col_width = Mm(table_width_mm / 9)  # ~22.78мм на колонку
        for col in table.columns:
            col.width = col_width
        
        # Встановлення висоти рядків
        row_height = Mm(table_height_mm / 9)  # ~14.44мм на рядок
        for row in table.rows:
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            row.height = row_height
        
        # Заповнення таблиці
        self._fill_image_table(table, image_data)
        
        # Налаштування стилів таблиці
        self._style_image_table(table)
    
    def _fill_image_table(self, table, image_data: ImageData):
        """
        Заповнення таблиці даними зображення
        
        Args:
            table: Таблиця Word
            image_data: Дані зображення
        """
        # Головна область зображення (рядки 0-6, колонки 0-6) - 7×7 = 77.78%
        self._add_image_to_table(table, image_data, start_row=0, end_row=6, 
                                start_col=0, end_col=6)
        
        # Область опису РЛС (рядки 0-1, колонки 7-8) - як в Legacy: 28.60% × 19.54%
        self._add_radar_description(table, image_data, start_row=0, end_row=1,
                                  start_col=7, end_col=8)
        
        # Дані аналізу (рядки 7-8, колонки 0-8)
        self._add_analysis_data(table, image_data, start_row=7, end_row=8)
    
    def _add_image_to_table(self, table, image_data: ImageData, 
                           start_row: int, end_row: int, start_col: int, end_col: int):
        """Додавання зображення до таблиці"""
        # Об'єднання клітинок для зображення
        top_left_cell = table.cell(start_row, start_col)
        bottom_right_cell = table.cell(end_row, end_col)
        merged_cell = top_left_cell.merge(bottom_right_cell)
        
        # Додавання зображення
        if os.path.exists(image_data.processed_image_path):
            try:
                # Розрахунок розміру зображення для вміщення в клітинку
                available_width_mm = ALBUM.TABLE_WIDTH_MM * 0.7778  # 7/9 від ширини
                available_height_mm = ALBUM.TABLE_HEIGHT_MM * 0.7778  # 7/9 від висоти
                
                paragraph = merged_cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Додавання зображення з автоматичним масштабуванням
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.add_picture(image_data.processed_image_path, 
                              width=Mm(available_width_mm - 2),  # -2мм для відступів
                              height=Mm(available_height_mm - 2))
                
            except Exception as e:
                # Якщо зображення не можна завантажити, додаємо текст
                print(f"Помилка додавання зображення: {e}")
                merged_cell.text = f"Зображення: {image_data.filename}"
        else:
            merged_cell.text = f"Зображення не знайдено: {image_data.filename}"
        
        # Вирівнювання в центрі
        merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    def _add_radar_description(self, table, image_data: ImageData,
                             start_row: int, end_row: int, start_col: int, end_col: int):
        """Додавання опису РЛС (пропорції 28.60% × 19.54%)"""
        # Об'єднання клітинок для опису РЛС
        top_left_cell = table.cell(start_row, start_col)
        bottom_right_cell = table.cell(end_row, end_col)
        radar_cell = top_left_cell.merge(bottom_right_cell)
        
        # Текст опису РЛС
        radar_text = (
            f"Ціль №{image_data.target_number}\n"
            f"Азимут: {image_data.get_azimuth_formatted()}\n"
            f"Дальність: {image_data.get_range_formatted()} км\n"
            f"Висота: {image_data.height}\n"
            f"{image_data.obstacles}\n"
            f"{image_data.detection}"
        )
        
        radar_cell.text = radar_text
        radar_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Форматування тексту
        paragraph = radar_cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.size = Pt(9)
            run.font.name = WORD_STYLES.FONT_NAME
    
    def _add_analysis_data(self, table, image_data: ImageData, 
                          start_row: int, end_row: int):
        """Додавання рядків з даними аналізу"""
        # Рядок 7: Назва файлу та час
        filename_cell = table.cell(start_row, 0)
        for col in range(1, 9):
            filename_cell = filename_cell.merge(table.cell(start_row, col))
        
        filename_text = f"Файл: {image_data.filename} | Час аналізу: {image_data.timestamp.strftime('%d.%m.%Y %H:%M')}"
        filename_cell.text = filename_text
        filename_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Рядок 8: Додаткова інформація
        info_cell = table.cell(end_row, 0)
        for col in range(1, 9):
            info_cell = info_cell.merge(table.cell(end_row, col))
        
        info_text = f"Координати: Азимут {image_data.get_azimuth_formatted()}, Дальність {image_data.get_range_formatted()} км | Статус: {image_data.detection}"
        info_cell.text = info_text
        info_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Форматування текстів
        for row_idx in [start_row, end_row]:
            paragraph = table.cell(row_idx, 0).paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.size = Pt(8)
                run.font.name = WORD_STYLES.FONT_NAME
    
    def _style_image_table(self, table):
        """Стилізація таблиці зображення"""
        # Встановлення бордерів для всіх клітинок
        for row in table.rows:
            for cell in row.cells:
                # Встановлення бордерів через XML (найнадійніший спосіб)
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                
                # Додавання бордерів
                borders = OxmlElement('w:tcBorders')
                
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(shared_qn('w:val'), 'single')
                    border.set(shared_qn('w:sz'), '4')  # Товщина 0.5pt
                    border.set(shared_qn('w:space'), '0')
                    border.set(shared_qn('w:color'), '000000')  # Чорний колір
                    borders.append(border)
                
                tcPr.append(borders)
        
        # Вирівнювання таблиці по центру
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # ===============================
    # УПРАВЛІННЯ ШАБЛОНАМИ
    # ===============================
    
    def _load_templates(self):
        """Завантаження шаблонів титульних сторінок"""
        templates_dir = get_templates_directory()
        
        if not os.path.exists(templates_dir):
            print(f"Папка шаблонів не існує: {templates_dir}")
            self._create_default_templates()
            return
        
        # Пошук файлів шаблонів
        template_files = [f for f in os.listdir(templates_dir) 
                         if f.endswith('.json')]
        
        if not template_files:
            print("Файли шаблонів не знайдено, створюємо за замовчуванням")
            self._create_default_templates()
            return
        
        # Завантаження шаблонів
        for template_file in template_files:
            template_path = os.path.join(templates_dir, template_file)
            template_name = os.path.splitext(template_file)[0]
            
            try:
                with open(template_path, 'r', encoding='utf-8') as f:
                    template_data = json.load(f)
                
                self.templates[template_name] = template_data
                print(f"Завантажено шаблон: {template_name}")
                
            except Exception as e:
                print(f"Помилка завантаження шаблону {template_file}: {e}")
        
        print(f"Завантажено {len(self.templates)} шаблонів")
    
    def _create_default_templates(self):
        """Створення шаблонів за замовчуванням"""
        templates_dir = get_templates_directory()
        ensure_directory_exists(templates_dir)
        
        # Базовий шаблон
        default_template = {
            "name": "Базовий шаблон",
            "description": "Стандартний шаблон для створення альбомів",
            "elements": [
                {
                    "type": "heading",
                    "text": "АЛЬБОМ ФОТОГРАФІЙ АЗИМУТАЛЬНОГО АНАЛІЗУ",
                    "font_size": 16,
                    "bold": True,
                    "alignment": "center"
                },
                {
                    "type": "spacing",
                    "lines": 2
                },
                {
                    "type": "text",
                    "text": "Дата створення: {document_date}",
                    "font_size": 12,
                    "alignment": "center"
                },
                {
                    "type": "spacing",
                    "lines": 3
                },
                {
                    "type": "table",
                    "column_widths": [0.4, 0.6],
                    "rows": [
                        {"label": "Підрозділ:", "value": "{unit_info}"},
                        {"label": "Командир:", "value": "{commander_rank} {commander_name}"},
                        {"label": "Начальник штабу:", "value": "{chief_of_staff_rank} {chief_of_staff_name}"},
                        {"label": "Кількість зображень:", "value": "{total_images}"},
                        {"label": "Дата створення:", "value": "{current_date}"}
                    ]
                }
            ]
        }
        
        # Військовий шаблон
        military_template = {
            "name": "Військовий шаблон",
            "description": "Шаблон для військових підрозділів",
            "elements": [
                {
                    "type": "text",
                    "text": "ЗБРОЙНІ СИЛИ УКРАЇНИ",
                    "font_size": 14,
                    "bold": True,
                    "alignment": "center"
                },
                {
                    "type": "spacing",
                    "lines": 1
                },
                {
                    "type": "heading",
                    "text": "АЛЬБОМ РОЗВІДУВАЛЬНИХ ФОТОГРАФІЙ",
                    "font_size": 18,
                    "bold": True,
                    "alignment": "center"
                },
                {
                    "type": "text",
                    "text": "Азимутальний аналіз цілей",
                    "font_size": 12,
                    "alignment": "center"
                },
                {
                    "type": "spacing",
                    "lines": 4
                },
                {
                    "type": "table",
                    "column_widths": [0.35, 0.65],
                    "rows": [
                        {"label": "Підрозділ:", "value": "{unit_info}"},
                        {"label": "Командир підрозділу:", "value": "{commander_rank} {commander_name}"},
                        {"label": "Начальник штабу:", "value": "{chief_of_staff_rank} {chief_of_staff_name}"},
                        {"label": "Дата документу:", "value": "{document_date}"},
                        {"label": "Проаналізовано цілей:", "value": "{total_images}"}
                    ]
                },
                {
                    "type": "spacing",
                    "lines": 8
                },
                {
                    "type": "text",
                    "text": "м. ____________________",
                    "font_size": 12,
                    "alignment": "center"
                },
                {
                    "type": "text",
                    "text": "{current_date}",
                    "font_size": 12,
                    "alignment": "center"
                }
            ]
        }
        
        # Збереження шаблонів
        templates_to_create = [
            ("default", default_template),
            ("military", military_template)
        ]
        
        for template_name, template_data in templates_to_create:
            template_path = os.path.join(templates_dir, f"{template_name}.json")
            
            try:
                with open(template_path, 'w', encoding='utf-8') as f:
                    json.dump(template_data, f, indent=2, ensure_ascii=False)
                
                self.templates[template_name] = template_data
                print(f"Створено шаблон: {template_name}")
                
            except Exception as e:
                print(f"Помилка створення шаблону {template_name}: {e}")
    
    def _get_template(self, template_name: str) -> Optional[dict]:
        """
        Отримання шаблону за назвою
        
        Args:
            template_name: Назва шаблону
            
        Returns:
            Словник шаблону або None
        """
        return self.templates.get(template_name)
    
    def get_available_templates(self) -> Dict[str, str]:
        """
        Отримання списку доступних шаблонів
        
        Returns:
            Словник {template_name: template_description}
        """
        templates_info = {}
        
        for name, template in self.templates.items():
            description = template.get('description', name)
            templates_info[name] = description
        
        return templates_info
    
    # ===============================
    # УТИЛІТАРНІ МЕТОДИ
    # ===============================
    
    def validate_images_data(self, images_data: List[ImageData]) -> List[str]:
        """
        Валідація даних зображень перед створенням альбому
        
        Args:
            images_data: Список даних зображень
            
        Returns:
            Список помилок (порожній якщо все ОК)
        """
        errors = []
        
        if not images_data:
            errors.append("Список зображень порожній")
            return errors
        
        for i, image_data in enumerate(images_data, 1):
            # Перевірка існування файлів
            if not os.path.exists(image_data.processed_image_path):
                errors.append(f"Зображення {i}: файл не існує - {image_data.processed_image_path}")
            
            # Перевірка обов'язкових полів
            if not image_data.target_number:
                errors.append(f"Зображення {i}: не вказано номер цілі")
            
            if not image_data.filename:
                errors.append(f"Зображення {i}: не вказано назву файлу")
            
            # Перевірка числових значень
            if not isinstance(image_data.azimuth, (int, float)) or image_data.azimuth < 0 or image_data.azimuth >= 360:
                errors.append(f"Зображення {i}: некоректний азимут - {image_data.azimuth}")
            
            if not isinstance(image_data.range_km, (int, float)) or image_data.range_km < 0:
                errors.append(f"Зображення {i}: некоректна дальність - {image_data.range_km}")
        
        return errors
    
    def validate_title_data(self, title_data: TitlePageData) -> List[str]:
        """
        Валідація даних титульної сторінки
        
        Args:
            title_data: Дані титульної сторінки
            
        Returns:
            Список помилок (порожний якщо все ОК)
        """
        errors = []
        
        if not title_data.document_date:
            errors.append("Не вказано дату документу")
        
        if not title_data.unit_info:
            errors.append("Не вказано інформацію про підрозділ")
        
        # Перевірка шаблону
        if title_data.template_name and title_data.template_name not in self.templates:
            errors.append(f"Шаблон не знайдено: {title_data.template_name}")
        
        return errors
    
    def estimate_processing_time(self, images_count: int) -> int:
        """
        Оцінка часу обробки в секундах
        
        Args:
            images_count: Кількість зображень
            
        Returns:
            Приблизний час обробки в секундах
        """
        # Базовий час: титульна + опис + збереження
        base_time = 5
        
        # Час на одне зображення (включаючи стиснення та форматування)
        time_per_image = 3
        
        return base_time + (images_count * time_per_image)
    
    def get_album_statistics(self) -> Dict[str, Any]:
        """
        Отримання статистики альбому
        
        Returns:
            Словник зі статистикою
        """
        if not self.current_images:
            return {}
        
        # Базова статистика
        stats = {
            'total_images': len(self.current_images),
            'title_template': self.title_data.template_name if self.title_data else 'unknown',
            'date_created': datetime.now().isoformat(),
            'estimated_pages': 2 + ((len(self.current_images) + 1) // 2)  # титульна + опис + сторінки зображень
        }
        
        # Статистика по цілях
        targets = [img.target_number for img in self.current_images if img.target_number]
        stats['unique_targets'] = len(set(targets))
        
        # Статистика по типах виявлення
        detections = [img.detection for img in self.current_images if img.detection]
        detection_counts = {}
        for detection in detections:
            detection_counts[detection] = detection_counts.get(detection, 0) + 1
        stats['detection_types'] = detection_counts
        
        # Діапазони азимуту та дальності
        azimuths = [img.azimuth for img in self.current_images]
        ranges = [img.range_km for img in self.current_images]
        
        if azimuths:
            stats['azimuth_range'] = {'min': min(azimuths), 'max': max(azimuths)}
        if ranges:
            stats['distance_range'] = {'min': min(ranges), 'max': max(ranges)}
        
        return stats
    
    def cleanup_temp_files(self):
        """Очищення тимчасових файлів"""
        # Тут можна додати логіку очищення тимчасових файлів
        # якщо вони створювались в процесі роботи
        pass


# ===============================
# ДОПОМІЖНІ ФУНКЦІЇ
# ===============================

def create_sample_images_data(count: int = 3) -> List[ImageData]:
    """
    Створення тестових даних зображень
    
    Args:
        count: Кількість тестових зображень
        
    Returns:
        Список ImageData з тестовими даними
    """
    sample_images = []
    
    for i in range(count):
        # Створення тестового зображення
        test_image = Image.new('RGB', (400, 300), color=(100 + i * 50, 150, 200))
        
        # Збереження у тимчасовий файл
        temp_file = tempfile.NamedTemporaryFile(suffix='.jpg', delete=False)
        test_image.save(temp_file.name, 'JPEG')
        
        image_data = ImageData(
            filename=f"test_image_{i+1:02d}.jpg",
            image_path=temp_file.name,
            processed_image_path=temp_file.name,
            target_number=f"Ціль-{i+1:02d}",
            azimuth=45.0 + i * 30,  # 45°, 75°, 105°...
            range_km=2.5 + i * 0.8,  # 2.5км, 3.3км, 4.1км...
            height=f"{150 + i * 50}м",
            obstacles="без перешкод" if i % 2 == 0 else "з перешкодами",
            detection="Виявлення" if i % 3 == 0 else "Супроводження",
            timestamp=datetime.now()
        )
        sample_images.append(image_data)
    
    return sample_images


def create_sample_title_data() -> TitlePageData:
    """
    Створення тестових даних титульної сторінки
    
    Returns:
        TitlePageData з тестовими даними
    """
    return TitlePageData(
        document_date=datetime.now().strftime('%d.%m.%Y'),
        unit_info="1-й батальйон, 2-га рота",
        commander_rank="капітан",
        commander_name="Іванов І.І.",
        chief_of_staff_rank="старший лейтенант", 
        chief_of_staff_name="Петров П.П.",
        template_name="default"
    )


if __name__ == "__main__":
    # Тестування створювача альбомів
    import sys
    from PyQt5.QtWidgets import QApplication, QMainWindow, QVBoxLayout, QHBoxLayout, QWidget, QPushButton, QLabel, QProgressBar, QTextEdit, QComboBox
    from PyQt5.QtCore import QThread, pyqtSignal
    
    class AlbumCreationThread(QThread):
        """Потік для створення альбому без блокування UI"""
        finished = pyqtSignal(bool, str)
        
        def __init__(self, creator, images_data, title_data, output_path):
            super().__init__()
            self.creator = creator
            self.images_data = images_data
            self.title_data = title_data
            self.output_path = output_path
        
        def run(self):
            try:
                success = self.creator.create_album(
                    self.images_data, self.title_data, self.output_path
                )
                self.finished.emit(success, self.output_path if success else "Помилка створення")
            except Exception as e:
                self.finished.emit(False, str(e))
    
    class AlbumCreatorTestWindow(QMainWindow):
        def __init__(self):
            super().__init__()
            self.setWindowTitle("Тестування AlbumCreator")
            self.setGeometry(100, 100, 800, 600)
            
            # Ініціалізація створювача альбомів
            try:
                self.album_creator = AlbumCreator()
            except ImportError as e:
                print(f"Критична помилка: {e}")
                sys.exit(1)
            
            # Центральний віджет
            central_widget = QWidget()
            self.setCentralWidget(central_widget)
            layout = QVBoxLayout(central_widget)
            
            # Заголовок
            title_label = QLabel("Тестування системи створення Word альбомів")
            title_label.setStyleSheet("font-size: 16px; font-weight: bold; margin: 10px;")
            layout.addWidget(title_label)
            
            # Вибір шаблону
            template_layout = QHBoxLayout()
            template_layout.addWidget(QLabel("Шаблон титульної сторінки:"))
            
            self.template_combo = QComboBox()
            available_templates = self.album_creator.get_available_templates()
            for template_name, description in available_templates.items():
                self.template_combo.addItem(f"{template_name} - {description}", template_name)
            template_layout.addWidget(self.template_combo)
            template_layout.addStretch()
            layout.addLayout(template_layout)
            
            # Кнопки управління
            buttons_layout = QHBoxLayout()
            
            self.create_sample_btn = QPushButton("Створити тестовий альбом (3 зображення)")
            self.create_sample_btn.clicked.connect(self.create_sample_album)
            buttons_layout.addWidget(self.create_sample_btn)
            
            self.create_large_btn = QPushButton("Створити великий альбом (10 зображень)")
            self.create_large_btn.clicked.connect(self.create_large_album)
            buttons_layout.addWidget(self.create_large_btn)
            
            buttons_layout.addStretch()
            layout.addLayout(buttons_layout)
            
            # Прогрес-бар
            self.progress_bar = QProgressBar()
            self.progress_bar.setVisible(False)
            layout.addWidget(self.progress_bar)
            
            # Статус
            self.status_label = QLabel("Готовий до створення альбому")
            layout.addWidget(self.status_label)
            
            # Лог подій
            layout.addWidget(QLabel("Лог подій:"))
            self.log_text = QTextEdit()
            self.log_text.setMaximumHeight(200)
            self.log_text.setStyleSheet("""
                QTextEdit {
                    font-family: 'Courier New', monospace;
                    font-size: 10px;
                    background-color: #f9f9f9;
                    border: 1px solid #ccc;
                }
            """)
            layout.addWidget(self.log_text)
            
            # Підключення сигналів
            self.album_creator.progress_updated.connect(self.update_progress)
            self.album_creator.album_created.connect(self.album_created)
            self.album_creator.error_occurred.connect(self.album_error)
            
            # Потік для створення альбому
            self.creation_thread = None
            
            self.log("AlbumCreator тестове вікно готове")
            self.log(f"Доступно шаблонів: {len(available_templates)}")
        
        def create_sample_album(self):
            """Створення тестового альбому з 3 зображеннями"""
            self.create_album_with_count(3)
        
        def create_large_album(self):
            """Створення великого альбому з 10 зображеннями"""
            self.create_album_with_count(10)
        
        def create_album_with_count(self, count: int):
            """Створення альбому з вказаною кількістю зображень"""
            if self.creation_thread and self.creation_thread.isRunning():
                self.log("Альбом вже створюється, зачекайте...")
                return
            
            try:
                # Підготовка даних
                self.log(f"Підготовка тестових даних для {count} зображень...")
                
                images_data = create_sample_images_data(count)
                title_data = create_sample_title_data()
                
                # Встановлення обраного шаблону
                selected_template = self.template_combo.currentData()
                title_data.template_name = selected_template
                
                # Валідація даних
                image_errors = self.album_creator.validate_images_data(images_data)
                title_errors = self.album_creator.validate_title_data(title_data)
                
                if image_errors or title_errors:
                    error_msg = "Помилки валідації:\n" + "\n".join(image_errors + title_errors)
                    self.log(error_msg)
                    return
                
                # Вихідний файл
                output_path = os.path.join(tempfile.gettempdir(), f"test_album_{count}_images.docx")
                
                # Показ прогресу
                self.progress_bar.setVisible(True)
                self.progress_bar.setValue(0)
                self.create_sample_btn.setEnabled(False)
                self.create_large_btn.setEnabled(False)
                
                # Оцінка часу
                estimated_time = self.album_creator.estimate_processing_time(count)
                self.log(f"Очікуваний час створення: {estimated_time} секунд")
                
                # Запуск в окремому потоці
                self.creation_thread = AlbumCreationThread(
                    self.album_creator, images_data, title_data, output_path
                )
                self.creation_thread.finished.connect(self.creation_finished)
                self.creation_thread.start()
                
                self.log(f"Розпочато створення альбому: {os.path.basename(output_path)}")
                
            except Exception as e:
                self.log(f"Помилка підготовки: {e}")
                self.reset_ui()
        
        def update_progress(self, progress: int, message: str):
            """Оновлення прогресу створення"""
            self.progress_bar.setValue(progress)
            self.status_label.setText(message)
            self.log(f"[{progress}%] {message}")
        
        def album_created(self, output_path: str):
            """Обробка успішного створення альбому"""
            self.log(f"✅ Альбом створено: {output_path}")
            
            # Статистика
            stats = self.album_creator.get_album_statistics()
            self.log(f"📊 Статистика: {stats}")
        
        def album_error(self, error_message: str):
            """Обробка помилки створення альбому"""
            self.log(f"❌ Помилка: {error_message}")
        
        def creation_finished(self, success: bool, result: str):
            """Завершення створення альбому"""
            self.reset_ui()
            
            if success:
                self.log(f"🎉 Альбом успішно створено: {result}")
                
                # Пропозиція відкрити файл
                try:
                    import os
                    if os.name == 'nt':  # Windows
                        os.startfile(result)
                    elif os.name == 'posix':  # macOS/Linux
                        os.system(f'open "{result}"')
                except Exception as e:
                    self.log(f"Не вдалося відкрити файл: {e}")
            else:
                self.log(f"💥 Помилка створення альбому: {result}")
        
        def reset_ui(self):
            """Скидання UI після завершення"""
            self.progress_bar.setVisible(False)
            self.create_sample_btn.setEnabled(True)
            self.create_large_btn.setEnabled(True)
            self.status_label.setText("Готовий до створення альбому")
        
        def log(self, message: str):
            """Додавання повідомлення в лог"""
            timestamp = datetime.now().strftime('%H:%M:%S')
            self.log_text.append(f"[{timestamp}] {message}")
            print(f"[{timestamp}] {message}")
        
        def closeEvent(self, event):
            """Очищення при закритті"""
            if self.creation_thread and self.creation_thread.isRunning():
                self.creation_thread.quit()
                self.creation_thread.wait()
            
            # Очищення тимчасових файлів
            self.album_creator.cleanup_temp_files()
            
            super().closeEvent(event)
    
    # Запуск тесту
    app = QApplication(sys.argv)
    
    try:
        window = AlbumCreatorTestWindow()
        window.show()
        
        print("=== Тестування AlbumCreator ===")
        print("Функції для тестування:")
        print("1. 'Створити тестовий альбом' - 3 зображення з базовими даними")
        print("2. 'Створити великий альбом' - 10 зображень для тестування продуктивності")
        print("3. Вибір шаблону титульної сторінки")
        print("4. Валідація даних перед створенням")
        print("5. Прогрес-бар та статистика в реальному часі")
        print("6. Автоматичне відкриття створеного файлу")
        print("\nКритичні розміри (як в Legacy):")
        print(f"• Таблиці: {ALBUM.TABLE_WIDTH_MM}мм × {ALBUM.TABLE_HEIGHT_MM}мм")
        print(f"• Поля сторінок: {ALBUM.TABLE_PAGES_LEFT_MARGIN}мм ліворуч")
        
        sys.exit(app.exec_())
        
    except Exception as e:
        print(f"Критична помилка запуску: {e}")
        sys.exit(1)