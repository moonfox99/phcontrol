#!/usr/bin/env python3
"""
Azimuth Image Processor - PyQt5 Version with Ukrainian Translation
Professional interface for processing azimuth grid images
COMPACT VERSION - uses separate modules for translations and widgets
"""

import sys
import os
import math
import tempfile
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QHBoxLayout, 
                             QVBoxLayout, QGridLayout, QPushButton, QLabel, 
                             QComboBox, QTextEdit, QScrollArea, QFrame,
                             QFileDialog, QMessageBox, QSizePolicy, QSplitter, QToolTip, QLineEdit)
from PyQt5.QtCore import Qt, QSize, pyqtSignal, QTimer, QPoint
from PyQt5.QtGui import QPixmap, QFont
from PIL import Image, ImageDraw

# Import custom modules
from translations import Translations
from widgets import ClickableLabel, VerticalThumbnailWidget
from image_processor import AzimuthImageProcessor

from docx.oxml.ns import qn
from docx.oxml import parse_xml

# Try to import python-docx for Word document generation
try:
    from docx import Document
    from docx.shared import Inches, Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. Install with: pip install python-docx")

class AzimuthGUI(QMainWindow):
    def __init__(self):
        super().__init__()
        self.processor = None
        self.current_image_path = None
        self.current_click = None
        self.scale_factor_x = 1.0
        self.scale_factor_y = 1.0
        self.offset_x = 0
        self.offset_y = 0
        self.current_folder = None
        
        # Language settings
        self.current_language = 'ENGLISH'
        self.translations = Translations()
        
        # Scale edge settings
        self.scale_edge_mode = False
        self.scale_edge_point = None
        self.custom_scale_distance = None
        
        # NEW: Center setting mode
        self.center_setting_mode = False
        
        # Report data settings
        self.current_target_number = "0001"
        self.current_height = "0.0"
        self.current_obstacles = "без перешкод"
        self.current_detection = "Виявлення"
        
        # Word document settings
        self.current_word_document = None
        self.word_document_path = None

        # NEW: Додати змінні для пакетної обробки
        self.processed_images = []  # Список оброблених зображень з даними
        self.current_folder_images = []  # Список всіх зображень в папці
        self.current_image_index = -1  # Поточний індекс зображення
        
        self.init_ui()
    
    def tr(self, key):
        """Get translation for the current language"""
        return self.translations.get(self.current_language, key)
    
    def set_language(self, language):
        """Change the interface language"""
        self.current_language = language
        self.update_interface_text()
        
        # Update menu checkboxes
        for lang, action in self.language_actions.items():
            action.setChecked(lang == language)
    
    def update_interface_text(self):
        """Update all interface text with current language"""
        self.setWindowTitle(self.tr("window_title"))
        self.create_menu_bar()
        
        # Update UI elements
        if hasattr(self, 'control_title'):
            self.control_title.setText(self.tr("controls"))
        if hasattr(self, 'report_title'):
            self.report_title.setText(self.tr("report_data"))
        if hasattr(self, 'browser_label'):
            self.browser_label.setText(self.tr("photo_browser"))
        
        # Update buttons
        if hasattr(self, 'open_image_btn'):
            self.open_image_btn.setText(self.tr("open_image"))
        if hasattr(self, 'open_folder_btn'):
            self.open_folder_btn.setText(self.tr("open_folder"))
        if hasattr(self, 'save_image_btn'):
            self.save_image_btn.setText(self.tr("save_current_image"))
        if hasattr(self, 'export_new_btn'):
            self.export_new_btn.setText(self.tr("create_new_album"))
        if hasattr(self, 'export_add_btn'):
            self.export_add_btn.setText(self.tr("add_to_existing_album"))
        if hasattr(self, 'clear_btn'):
            self.clear_btn.setText(self.tr("clear_analysis_point"))
        if hasattr(self, 'scale_edge_btn'):
            self.scale_edge_btn.setText(self.tr("set_scale_edge"))
        
        # Update labels
        if hasattr(self, 'file_ops_label'):
            self.file_ops_label.setText(self.tr("file_operations"))
        if hasattr(self, 'azimuth_grid_label'):
            self.azimuth_grid_label.setText(self.tr("azimuth_grid"))
        if hasattr(self, 'move_center_label'):
            self.move_center_label.setText(self.tr("move_center"))
        if hasattr(self, 'scale_edge_label'):
            self.scale_edge_label.setText(f"{self.tr('scale_edge_mode')}:")
        if hasattr(self, 'results_label'):
            self.results_label.setText(self.tr("results"))
        
        # Update image label if no image loaded
        if not self.current_image_path:
            self.image_label.setText(self.tr("open_instruction"))
        
        # Update results if there's current data
        if self.processor:
            self.update_results_display()
    
    def update_translations_file():
        """Додаткові українські переклади"""
        additional_translations = {
            'UKRAINIAN': {
                # Нові переклади для інтерфейсу
                "set_center_mode": "Режим вибору центру сітки",
                "set_center": "Встановити центр сітки", 
                "batch_processing": "Пакетна обробка",
                "save_current_image_data": "Зберегти дані зображення",
                "previous": "Попереднє",
                "next": "Наступне",
                
                # Повідомлення про збереження налаштувань
                "grid_settings_saved": "Налаштування сітки збережено для наступних зображень",
                "grid_settings_applied": "Застосовано збережені налаштування сітки",
                "no_grid_settings": "Немає збережених налаштувань сітки"
            }
        }
        return additional_translations

    def init_ui(self):
        """Initialize the user interface"""
        self.setWindowTitle(self.tr("window_title"))
        
        # Set reasonable default window size (but allow resizing)
        default_width = 1400
        default_height = 900
        
        # Set minimum size to ensure usability
        min_width = 1000  
        min_height = 700
        self.setMinimumSize(min_width, min_height)
        
        # Allow normal window behavior - user can move, resize, maximize
        self.setWindowFlags(Qt.Window)
        
        # Set default size and let OS decide position
        self.resize(default_width, default_height)
        self.showMaximized()
        
        # Create menu bar
        self.create_menu_bar()
        
        # Central widget
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        # Main layout
        main_layout = QHBoxLayout()
        central_widget.setLayout(main_layout)
        
        # Create splitter for resizable panels
        main_splitter = QSplitter(Qt.Horizontal)
        main_layout.addWidget(main_splitter)
        
        # Left panel - Controls (fixed width)
        self.create_control_panel(main_splitter)
        
        # Vertical image browser panel (initially hidden)
        self.create_vertical_browser_panel(main_splitter)
        
        # Center panel - Image only (expandable but constrained)
        self.create_image_panel(main_splitter)
        
        # Right panel - Report data (fixed width)
        self.create_report_panel(main_splitter)
        
        # Set splitter proportions and constraints
        main_splitter.setStretchFactor(0, 0)  # Controls panel - never stretch
        main_splitter.setStretchFactor(1, 0)  # Browser panel - fixed size
        main_splitter.setStretchFactor(2, 1)  # Image panel - can stretch
        main_splitter.setStretchFactor(3, 0)  # Report panel - never stretch
        
        # Set initial sizes: Controls(300) | Browser(0 - hidden) | Image(800) | Report(300)
        main_splitter.setSizes([300, 0, 800, 300])
        
        # Store reference to splitter for proportion management
        self.main_splitter = main_splitter
    
    def resizeEvent(self, event):
        """Handle window resize - simplified version that doesn't force proportions"""
        super().resizeEvent(event)
        
        # Only update image display if we have an image loaded
        # Remove all the complex proportion calculations - let the user resize freely
        if hasattr(self, 'processor') and self.processor:
            # Just refresh the image display with a small delay
            QTimer.singleShot(50, self.update_image_display_after_resize)
        
        # Optional: Adjust splitter if browser is visible to maintain reasonable proportions
        if hasattr(self, 'main_splitter') and hasattr(self, 'browser_widget'):
            current_sizes = self.main_splitter.sizes()
            total_width = sum(current_sizes)
            
            # Only adjust if window is reasonably sized
            if total_width > 800:
                if self.browser_widget.isVisible():
                    # Controls(300) | Browser(180) | Image(rest-300) | Report(300)
                    new_image_width = max(400, total_width - 300 - 180 - 300)
                    self.main_splitter.setSizes([300, 180, new_image_width, 300])
                else:
                    # Controls(300) | Browser(0) | Image(rest-300) | Report(300)  
                    new_image_width = max(400, total_width - 300 - 300)
                    self.main_splitter.setSizes([300, 0, new_image_width, 300])
    
    def update_image_display_after_resize(self):
        """Update image display after window resize to maintain proportions"""
        if hasattr(self, 'processor') and self.processor:
            self.display_image()
    
    def create_menu_bar(self):
        """Create menu bar with Settings"""
        menubar = self.menuBar()
        menubar.clear()
        
        # Settings menu
        settings_menu = menubar.addMenu(self.tr("settings"))
        
        # Language submenu
        language_menu = settings_menu.addMenu(self.tr("language"))
        
        # Language actions
        english_action = language_menu.addAction("English")
        english_action.triggered.connect(lambda: self.set_language('ENGLISH'))
        
        ukrainian_action = language_menu.addAction("Українська")
        ukrainian_action.triggered.connect(lambda: self.set_language('UKRAINIAN'))
        
        # Make actions checkable
        english_action.setCheckable(True)
        ukrainian_action.setCheckable(True)
        
        # Set current language as checked
        english_action.setChecked(self.current_language == 'ENGLISH')
        ukrainian_action.setChecked(self.current_language == 'UKRAINIAN')
        
        # Store actions for later updates
        self.language_actions = {
            'ENGLISH': english_action,
            'UKRAINIAN': ukrainian_action
        }
    
    def create_control_panel(self, parent):
        """Create the left control panel"""
        control_widget = QWidget()
        control_widget.setFixedWidth(300)
        control_widget.setStyleSheet("""
            QWidget {
                background-color: #f5f5f5; 
                border-right: 1px solid #ccc;
            }
            QLabel {
                background: none;
                border: none;
                color: #333;
            }
            QPushButton {
                background-color: #e1e1e1;
                border: 1px solid #c1c1c1;
                border-radius: 4px;
                padding: 8px 12px;
                font-size: 11px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #d1d1d1;
                border: 1px solid #b1b1b1;
            }
            QPushButton:pressed {
                background-color: #c1c1c1;
            }
            QComboBox {
                border: 1px solid #c1c1c1;
                border-radius: 4px;
                padding: 4px 8px;
                background-color: white;
                font-size: 11px;
            }
        """)
        
        layout = QVBoxLayout()
        layout.setContentsMargins(15, 15, 15, 15)
        layout.setSpacing(10)
        control_widget.setLayout(layout)
        
        # Title
        self.control_title = QLabel(self.tr("controls"))
        self.control_title.setFont(QFont("Arial", 14, QFont.Bold))
        self.control_title.setAlignment(Qt.AlignCenter)
        self.control_title.setStyleSheet("font-size: 14px; font-weight: bold; margin-bottom: 10px;")
        layout.addWidget(self.control_title)
        
        # Image Selection section
        self.file_ops_label = QLabel(self.tr("file_operations"))
        self.file_ops_label.setFont(QFont("Arial", 10, QFont.Bold))
        self.file_ops_label.setStyleSheet("color: #666; margin-top: 10px; margin-bottom: 5px;")
        layout.addWidget(self.file_ops_label)
        
        # Buttons
        self.open_image_btn = QPushButton(self.tr("open_image"))
        self.open_image_btn.clicked.connect(self.open_image)
        layout.addWidget(self.open_image_btn)
        
        self.open_folder_btn = QPushButton(self.tr("open_folder"))
        self.open_folder_btn.clicked.connect(self.open_folder)
        layout.addWidget(self.open_folder_btn)
        
        self.save_image_btn = QPushButton(self.tr("save_current_image"))
        self.save_image_btn.clicked.connect(self.save_current_image)
        layout.addWidget(self.save_image_btn)
        
        # Export to Word buttons
        self.export_new_btn = QPushButton(self.tr("create_new_album"))
        self.export_new_btn.clicked.connect(self.create_new_word_album)
        layout.addWidget(self.export_new_btn)
        
        self.export_add_btn = QPushButton(self.tr("add_to_existing_album"))
        self.export_add_btn.clicked.connect(self.add_to_existing_album)
        layout.addWidget(self.export_add_btn)
        
        # Separator line
        separator1 = QFrame()
        separator1.setFrameShape(QFrame.HLine)
        separator1.setFrameShadow(QFrame.Sunken)
        separator1.setStyleSheet("color: #ccc; margin: 10px 0px;")
        layout.addWidget(separator1)
        
        # Azimuth Grid section
        self.azimuth_grid_label = QLabel(self.tr("azimuth_grid"))
        self.azimuth_grid_label.setFont(QFont("Arial", 10, QFont.Bold))
        self.azimuth_grid_label.setStyleSheet("color: #666; margin-top: 5px; margin-bottom: 5px;")
        layout.addWidget(self.azimuth_grid_label)
        
        # Scale setting
        scale_layout = QHBoxLayout()
        scale_label = QLabel(self.tr("scale_setting"))
        scale_label.setStyleSheet("color: #333; margin-right: 5px;")
        scale_layout.addWidget(scale_label)
        self.scale_combo = QComboBox()
        self.scale_combo.addItems(["50", "100", "150", "200", "300"])
        self.scale_combo.setCurrentText("300")
        self.scale_combo.currentTextChanged.connect(self.update_scale)
        scale_layout.addWidget(self.scale_combo)
        scale_layout.addStretch()
        
        scale_widget = QWidget()
        scale_widget.setLayout(scale_layout)
        layout.addWidget(scale_widget)
        
        # Scale edge tool
        self.scale_edge_label = QLabel(f"{self.tr('scale_edge_mode')}:")
        self.scale_edge_label.setStyleSheet("color: #333; margin-top: 5px;")
        layout.addWidget(self.scale_edge_label)
        self.scale_edge_btn = QPushButton(self.tr("set_scale_edge"))
        self.scale_edge_btn.setCheckable(True)
        self.scale_edge_btn.clicked.connect(self.toggle_scale_edge_mode)
        layout.addWidget(self.scale_edge_btn)
        
        # NEW: Center setting tool
        self.center_setting_label = QLabel("Set Center Mode:")
        self.center_setting_label.setStyleSheet("color: #333; margin-top: 5px;")
        layout.addWidget(self.center_setting_label)
        self.set_center_btn = QPushButton("Set Center")
        self.set_center_btn.setCheckable(True)
        self.set_center_btn.clicked.connect(self.toggle_center_setting_mode)
        layout.addWidget(self.set_center_btn)
        
        # Center movement
        self.move_center_label = QLabel(self.tr("move_center"))
        self.move_center_label.setStyleSheet("color: #333; margin-top: 10px;")
        layout.addWidget(self.move_center_label)
        
        move_layout = QGridLayout()
        move_layout.setSpacing(2)

        btn_up = self.create_center_move_button_with_hover("↑", lambda: self.move_center(0, -2))
        move_layout.addWidget(btn_up, 0, 1)

        btn_left = self.create_center_move_button_with_hover("←", lambda: self.move_center(-2, 0))
        move_layout.addWidget(btn_left, 1, 0)

        btn_center = QPushButton("●")
        btn_center.setMaximumWidth(40)
        btn_center.clicked.connect(self.show_center_preview)
        # Add hover for center button too
        def center_enter(event):
            self.show_center_zoom()
            QPushButton.enterEvent(btn_center, event)
        def center_leave(event):
            self.hide_center_zoom()
            QPushButton.leaveEvent(btn_center, event)
        btn_center.enterEvent = center_enter
        btn_center.leaveEvent = center_leave
        move_layout.addWidget(btn_center, 1, 1)

        btn_right = self.create_center_move_button_with_hover("→", lambda: self.move_center(2, 0))
        move_layout.addWidget(btn_right, 1, 2)

        btn_down = self.create_center_move_button_with_hover("↓", lambda: self.move_center(0, 2))
        move_layout.addWidget(btn_down, 2, 1)
        
        move_widget = QWidget()
        move_widget.setLayout(move_layout)
        layout.addWidget(move_widget)
        
        # Separator line
        separator2 = QFrame()
        separator2.setFrameShape(QFrame.HLine)
        separator2.setFrameShadow(QFrame.Sunken)
        separator2.setStyleSheet("color: #ccc; margin: 10px 0px;")
        layout.addWidget(separator2)
        
        # Clear button
        self.clear_btn = QPushButton(self.tr("clear_analysis_point"))
        self.clear_btn.clicked.connect(self.clear_results)
        layout.addWidget(self.clear_btn)
        
        # Results area
        self.results_label = QLabel(self.tr("results"))
        self.results_label.setFont(QFont("Arial", 10, QFont.Bold))
        self.results_label.setStyleSheet("color: #666; margin-top: 10px; margin-bottom: 5px;")
        layout.addWidget(self.results_label)
        
        self.results_text = QTextEdit()
        self.results_text.setMaximumHeight(200)
        self.results_text.setStyleSheet("""
            QTextEdit {
                background-color: white;
                border: 1px solid #ccc;
                font-family: 'Courier New', monospace;
                font-size: 10px;
                border-radius: 4px;
            }
        """)
        layout.addWidget(self.results_text)

        # Batch processing buttons
        batch_separator = QFrame()
        batch_separator.setFrameShape(QFrame.HLine)
        batch_separator.setFrameShadow(QFrame.Sunken)
        batch_separator.setStyleSheet("color: #ccc; margin: 10px 0px;")
        layout.addWidget(batch_separator)

        batch_label = QLabel("Batch Processing:")
        batch_label.setFont(QFont("Arial", 10, QFont.Bold))
        batch_label.setStyleSheet("color: #666; margin-top: 5px; margin-bottom: 5px;")
        layout.addWidget(batch_label)

        self.save_current_btn = QPushButton("Save Current Image Data")
        self.save_current_btn.clicked.connect(self.save_current_image_data)
        layout.addWidget(self.save_current_btn)

        nav_layout = QHBoxLayout()
        self.prev_btn = QPushButton("← Previous")
        self.prev_btn.clicked.connect(self.previous_image_in_batch)
        nav_layout.addWidget(self.prev_btn)

        self.next_btn = QPushButton("Next →")
        self.next_btn.clicked.connect(self.next_image_in_batch)
        nav_layout.addWidget(self.next_btn)

        nav_widget = QWidget()
        nav_widget.setLayout(nav_layout)
        layout.addWidget(nav_widget)

        self.create_batch_btn = QPushButton("Create Batch Album")
        self.create_batch_btn.clicked.connect(self.create_batch_album)
        self.create_batch_btn.setStyleSheet("font-weight: bold; background-color: #4CAF50; color: white;")
        layout.addWidget(self.create_batch_btn)

        # Оновлені переклади для batch processing
        batch_label = QLabel(self.tr("batch_processing"))  # Тепер з перекладом 
        self.save_current_btn = QPushButton(self.tr("save_current_image_data"))  # З перекладом
        self.prev_btn = QPushButton(f"← {self.tr('previous')}")  # З перекладом
        self.next_btn = QPushButton(f"{self.tr('next')} →")  # З перекладом
        
        # Stretch at bottom
        layout.addStretch()
        
        parent.addWidget(control_widget)
    
    def toggle_center_setting_mode(self):
        """Toggle center setting mode"""
        self.center_setting_mode = self.set_center_btn.isChecked()
        
        # Disable scale edge mode if center mode is active
        if self.center_setting_mode and self.scale_edge_mode:
            self.scale_edge_mode = False
            self.scale_edge_btn.setChecked(False)
            self.scale_edge_btn.setStyleSheet("")
        
        # Update cursor and visual feedback
        self.image_label.set_center_setting_mode(self.center_setting_mode)
        
        if self.center_setting_mode:
            self.set_center_btn.setStyleSheet("background-color: #FF9800; color: white;")  # Orange color
            self.add_result("Center setting mode active - click to set new center")
        else:
            self.set_center_btn.setStyleSheet("")
    
    def set_center_point(self, x, y):
        """Set new center point at specified coordinates"""
        if not self.processor:
            return
        
        # Ensure coordinates are within bounds
        x = max(0, min(x, self.processor.image.width - 1))
        y = max(0, min(y, self.processor.image.height - 1))
        
        # Calculate the offset needed to move center to this point
        current_center_x = self.processor.center_x
        current_center_y = self.processor.center_y
        
        dx = x - current_center_x
        dy = y - current_center_y
        
        # Move center to the clicked point
        self.processor.move_center(dx, dy)
        
        # Auto-disable center setting mode after setting
        self.center_setting_mode = False
        self.set_center_btn.setChecked(False)
        self.set_center_btn.setStyleSheet("")
        self.image_label.set_center_setting_mode(False)
        
        # Update display
        self.display_image()
        
        # Recalculate current analysis point if exists
        if self.current_click:
            azimuth, range_val = self.calculate_azimuth_range(
                self.current_click['x'], self.current_click['y']
            )
            self.current_click['azimuth'] = azimuth
            self.current_click['range'] = range_val
        
        # Update results
        self.update_results_display()
        self.update_report_data()
        self.add_result(f"Center moved to: ({self.processor.center_x}, {self.processor.center_y})")

    def create_vertical_browser_panel(self, parent):
        """Create the vertical image browser panel"""
        browser_widget = QWidget()
        browser_widget.setFixedWidth(180)
        browser_widget.setStyleSheet("""
            QWidget {
                background-color: #f8f8f8; 
                border-right: 1px solid #ccc;
                border-left: 1px solid #ccc;
            }
            QLabel {
                background: none;
                border: none;
                color: #333;
            }
        """)
        
        layout = QVBoxLayout()
        layout.setContentsMargins(0, 15, 0, 15)
        layout.setSpacing(10)
        browser_widget.setLayout(layout)
        
        # Browser title
        self.browser_label = QLabel(self.tr("photo_browser"))
        self.browser_label.setFont(QFont("Arial", 10, QFont.Bold))
        self.browser_label.setStyleSheet("color: #666; margin-bottom: 5px; padding: 0 10px;")
        self.browser_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(self.browser_label)
        
        # Create vertical scroll area for thumbnails
        self.thumbnail_scroll = QScrollArea()
        self.thumbnail_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.thumbnail_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self.thumbnail_scroll.setWidgetResizable(False)
        self.thumbnail_scroll.setStyleSheet("border: none; background: transparent;")
        
        # Create vertical thumbnail widget
        self.thumbnail_widget = VerticalThumbnailWidget()
        self.thumbnail_widget.image_selected.connect(self.load_image_from_browser)
        self.thumbnail_scroll.setWidget(self.thumbnail_widget)
        
        layout.addWidget(self.thumbnail_scroll)
        
        # Initially hide the browser panel
        browser_widget.hide()
        self.browser_widget = browser_widget
        
        parent.addWidget(browser_widget)
    
    def create_image_panel(self, parent):
        """Create the center panel for image display only"""
        image_widget = QWidget()
        layout = QVBoxLayout()
        image_widget.setLayout(layout)
        
        # Image display area
        self.image_label = ClickableLabel()
        self.image_label.clicked.connect(self.on_image_click)
        self.image_label.dragged.connect(self.on_image_drag)
        self.image_label.mouse_moved.connect(self.on_mouse_hover)
        self.image_label.setText(self.tr("open_instruction"))
        
        # Scroll area for image
        scroll_area = QScrollArea()
        scroll_area.setWidget(self.image_label)
        scroll_area.setWidgetResizable(True)
        scroll_area.setMinimumHeight(400)
        layout.addWidget(scroll_area)
        
        parent.addWidget(image_widget)
    
    def create_report_panel(self, parent):
        """Create the right panel for report data entry"""
        report_widget = QWidget()
        report_widget.setFixedWidth(300)
        report_widget.setStyleSheet("background-color: #f9f9f9; border-left: 1px solid #ccc;")
        
        layout = QVBoxLayout()
        report_widget.setLayout(layout)
        
        # Title
        self.report_title = QLabel(self.tr("report_data"))
        self.report_title.setFont(QFont("Arial", 14, QFont.Bold))
        self.report_title.setAlignment(Qt.AlignCenter)
        layout.addWidget(self.report_title)
 
        # Manual input section
        manual_group = QFrame()
        manual_group.setStyleSheet("background-color: #fff; border: 1px solid #ddd; border-radius: 5px; padding: 10px;")
        manual_layout = QVBoxLayout()
        manual_group.setLayout(manual_layout)
        
        # Target number input
        self.target_number_input = QLineEdit()
        self.target_number_input.setPlaceholderText("Номер цілі")
        self.target_number_input.setText(self.current_target_number)
        self.target_number_input.textChanged.connect(self.update_target_number)
        manual_layout.addWidget(self.target_number_input)
        
        # Auto azimuth label
        self.auto_azimuth_label = QLabel("β - --°")
        manual_layout.addWidget(self.auto_azimuth_label)
        
        # Auto distance label  
        self.auto_distance_label = QLabel("D - -- км")
        manual_layout.addWidget(self.auto_distance_label)
        
        # Height input with H prefix
        height_layout = QHBoxLayout()
        height_layout.addWidget(QLabel("H –"))
        self.height_input = QLineEdit(self.current_height)
        self.height_input.setMaximumWidth(80)
        self.height_input.textChanged.connect(self.update_height)
        height_layout.addWidget(self.height_input)
        height_layout.addWidget(QLabel(self.tr("km_unit")))
        height_layout.addStretch()
        manual_layout.addLayout(height_layout)
        
        # Obstacles dropdown
        self.obstacles_combo = QComboBox()
        self.obstacles_combo.addItems([self.tr("no_obstacles"), self.tr("with_obstacles")])
        self.obstacles_combo.currentTextChanged.connect(self.update_obstacles)
        manual_layout.addWidget(self.obstacles_combo)
        
        # Detection type dropdown
        self.detection_combo = QComboBox()
        self.detection_combo.addItems([self.tr("detection"), self.tr("tracking"), self.tr("loss")])
        self.detection_combo.currentTextChanged.connect(self.update_detection)
        manual_layout.addWidget(self.detection_combo)
        
        # Auto scale label
        self.auto_scale_label = QLabel("M = --")
        manual_layout.addWidget(self.auto_scale_label)
        
        layout.addWidget(manual_group)

        self.center_setting_label = QLabel(self.tr("set_center_mode"))  # З перекладом
        self.set_center_btn = QPushButton(self.tr("set_center"))  # З перекладом
        
        # Stretch to push everything to top
        layout.addStretch()
        
        # Initially hide the panel
        report_widget.hide()
        self.report_widget = report_widget
        
        parent.addWidget(report_widget)
    
    def add_result(self, text):
        """Add text to results area"""
        self.results_text.append(text)
    
    def open_image(self):
        """Open a single image file"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, self.tr("select_image"), "",
            f"{self.tr('image_files')} (*.jpg *.jpeg *.png *.bmp *.gif *.tiff);;{self.tr('all_files')} (*.*)"
        )
        
        if file_path:
            self.load_image(file_path)
            # Show report panel when image is loaded
            self.report_widget.show()
    
    def open_folder(self):
        """Open a folder and display thumbnails"""
        folder_path = QFileDialog.getExistingDirectory(self, self.tr("select_folder"))
        
        if folder_path:
            self.current_folder = folder_path
            self.load_folder_thumbnails()
            
            # Show vertical browser panel
            self.browser_widget.show()
            print("Made vertical browser visible")
            
            # Update splitter sizes: Controls(300) | Browser(180) | Image(720) | Report(300)
            self.main_splitter.setSizes([300, 180, 720, 300])
            
            # Count images
            image_extensions = ('.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff', '.tif')
            image_count = sum(1 for f in os.listdir(folder_path) 
                             if f.lower().endswith(image_extensions))
            
            self.add_result(f"{self.tr('loaded_folder')}: {os.path.basename(folder_path)}")
            self.add_result(self.tr("found_images").format(count=image_count))
            
            # Show report panel when folder is loaded
            self.report_widget.show()

    def load_folder_thumbnails(self):
        """Load thumbnails for all images in the current folder"""
        if not self.current_folder:
            return
        
        print(f"Loading thumbnails from folder: {self.current_folder}")
        
        # Clear existing thumbnails
        self.thumbnail_widget.clear_thumbnails()
        
        # Find image files
        image_extensions = ('.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff', '.tif')
        image_files = []
        
        try:
            for filename in os.listdir(self.current_folder):
                if filename.lower().endswith(image_extensions):
                    full_path = os.path.join(self.current_folder, filename)
                    image_files.append(full_path)
        except Exception as e:
            print(f"Error reading folder: {e}")
            return

        # Sort files
        image_files.sort()
        print(f"Total images found: {len(image_files)}")
        
        if len(image_files) == 0:
            print("No image files found in folder!")
            # Add a message widget
            no_images_label = QLabel(self.tr("no_images_found"))
            no_images_label.setAlignment(Qt.AlignCenter)
            no_images_label.setStyleSheet("color: gray; font-size: 14px; padding: 20px;")
            no_images_label.setWordWrap(True)
            self.thumbnail_widget.layout.addWidget(no_images_label)
            return
        
        # Create thumbnails
        success_count = 0
        for i, image_path in enumerate(image_files):
            print(f"Creating thumbnail {i+1}/{len(image_files)}: {os.path.basename(image_path)}")
            try:
                self.thumbnail_widget.add_thumbnail(image_path)
                success_count += 1
            except Exception as e:
                print(f"Failed to create thumbnail for {os.path.basename(image_path)}: {e}")
        
        print(f"Successfully created {success_count}/{len(image_files)} thumbnails")
        
        # Update widget size for vertical layout
        widget_height = len(image_files) * 130 + 20  # 120 + 10 spacing per thumbnail + margins
        self.thumbnail_widget.setMinimumHeight(widget_height)
        self.thumbnail_widget.resize(160, widget_height)

    def load_image_from_browser(self, file_path):
        """Load image selected from thumbnail browser"""
        self.load_image(file_path)
        self.add_result(self.tr("loaded_from_browser").format(name=os.path.basename(file_path)))
    
    def load_image(self, file_path):
        """Load and display an image with proper format handling"""
        try:
            self.current_image_path = file_path
            scale_value = int(self.scale_combo.currentText())
            
            # Initialize processor
            self.processor = AzimuthImageProcessor(file_path, scale=scale_value)
            
            # Force convert processor image to RGB immediately after loading
            if hasattr(self.processor, 'image') and self.processor.image:
                original_mode = self.processor.image.mode
                print(f"Original image mode from processor: {original_mode}")
                
                if original_mode != 'RGB':
                    print(f"Converting processor image from {original_mode} to RGB")
                    if original_mode == 'RGBA':
                        # Handle transparency
                        rgb_image = Image.new('RGB', self.processor.image.size, (255, 255, 255))
                        rgb_image.paste(self.processor.image, mask=self.processor.image.split()[-1])
                        self.processor.image = rgb_image
                    else:
                        # Convert any other mode to RGB
                        self.processor.image = self.processor.image.convert('RGB')
                    print(f"Processor image converted to: {self.processor.image.mode}")
            
            self.display_image()
            self.current_click = None
            self.update_results_display()
            self.update_report_data()
            
            self.add_result(f"{self.tr('loaded')}: {os.path.basename(file_path)}")
            
        except Exception as e:
            print(f"Error loading image: {e}")
            QMessageBox.critical(self, self.tr("error"), 
                               self.tr("could_not_load").format(error=str(e)))
    
    def display_image(self):
        """Display the current image maintaining strict 15:13 proportions"""
        if not self.processor:
            return
        
        # Get PIL image and convert to QPixmap
        pil_image = self.processor.image.copy()
        
        # Draw center crosshairs and current click
        draw = ImageDraw.Draw(pil_image)
        
        # Center crosshairs
        center_x, center_y = self.processor.center_x, self.processor.center_y
        cross_size = 15
        
        draw.line([center_x - cross_size, center_y, center_x + cross_size, center_y], 
                  fill='red', width=2)
        draw.line([center_x, center_y - cross_size, center_x, center_y + cross_size], 
                  fill='red', width=2)
        draw.ellipse([center_x - 3, center_y - 3, center_x + 3, center_y + 3], 
                     fill='red', outline='white')
        
        # Current analysis point
        if self.current_click:
            click_x, click_y = self.current_click['x'], self.current_click['y']
            
            # Draw analysis point
            draw.ellipse([click_x - 4, click_y - 4, click_x + 4, click_y + 4], 
                         fill='blue', outline='white', width=1)
            
            # Draw line to top-right corner
            draw.line([click_x, click_y, pil_image.width - 1, 0], fill='blue', width=3)
        
        # Draw scale edge point if set
        if self.scale_edge_point:
            edge_x, edge_y = self.scale_edge_point['x'], self.scale_edge_point['y']
            
            # Draw scale edge point in green
            draw.ellipse([edge_x - 5, edge_y - 5, edge_x + 5, edge_y + 5], 
                         fill='green', outline='white', width=2)
            
            # Draw line from center to scale edge
            draw.line([center_x, center_y, edge_x, edge_y], fill='green', width=2)
            
            # Draw perpendicular lines for better visibility
            dx = edge_x - center_x
            dy = edge_y - center_y
            length = math.sqrt(dx*dx + dy*dy)
            if length > 0:
                nx, ny = -dy/length, dx/length
                perp_size = 8
                draw.line([
                    edge_x + nx*perp_size, edge_y + ny*perp_size,
                    edge_x - nx*perp_size, edge_y - ny*perp_size
                ], fill='green', width=2)
        
        # Convert PIL to QPixmap using temporary file method
        with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as temp_file:
            temp_path = temp_file.name
            pil_image.save(temp_path, 'JPEG', quality=95)
        
        pixmap = QPixmap(temp_path)
        
        # Clean up temp file
        try:
            os.remove(temp_path)
        except:
            pass
        
        # Get widget dimensions
        widget_width = self.image_label.width()
        widget_height = self.image_label.height()
        
        # Scale pixmap to fit widget while maintaining aspect ratio
        scaled_pixmap = pixmap.scaled(widget_width, widget_height, 
                                     Qt.KeepAspectRatio, Qt.SmoothTransformation)
        
        # Calculate scale factors and offsets
        original_width = pixmap.width()
        original_height = pixmap.height()
        scaled_width = scaled_pixmap.width()
        scaled_height = scaled_pixmap.height()
        
        # Calculate scale factors
        self.scale_factor_x = scaled_width / original_width
        self.scale_factor_y = scaled_height / original_height
        
        # Calculate centering offsets within widget
        self.offset_x = (widget_width - scaled_width) // 2
        self.offset_y = (widget_height - scaled_height) // 2
        
        # CRITICAL: Update ClickableLabel with geometry info
        self.image_label.update_image_geometry(
            original_width, original_height,
            self.scale_factor_x, self.scale_factor_y,
            self.offset_x, self.offset_y
        )
        
        # NEW: Set source image for zoom widget
        self.image_label.set_zoom_source_image(pixmap)
        
        # Set the pixmap
        self.image_label.setPixmap(scaled_pixmap)
        
        print(f"Image displayed: original({original_width}x{original_height}) -> "
              f"scaled({scaled_width}x{scaled_height}) -> offset({self.offset_x}, {self.offset_y})")
    
    def on_image_click(self, x, y):
        """Handle click on image with coordinates already converted by ClickableLabel"""
        if not self.processor:
            return
    
        print(f"on_image_click received: image coordinates ({x}, {y})")
    
        # Handle scale edge mode
        if self.scale_edge_mode:
            self.set_scale_edge_point(x, y)
            return
    
        # NEW: Handle center setting mode
        if self.center_setting_mode:
            self.set_center_point(x, y)
            return
    
        # Check if clicking near existing point (for drag mode)
        if self.current_click:
            # Convert existing point to widget coordinates for distance calculation
            existing_widget_x = self.current_click['x'] * self.scale_factor_x + self.offset_x
            existing_widget_y = self.current_click['y'] * self.scale_factor_y + self.offset_y
        
            # Convert clicked image coords back to widget coords for distance check
            clicked_widget_x = x * self.scale_factor_x + self.offset_x
            clicked_widget_y = y * self.scale_factor_y + self.offset_y
        
            distance = ((clicked_widget_x - existing_widget_x)**2 + (clicked_widget_y - existing_widget_y)**2)**0.5
        
            if distance <= 15:  # Within 15 pixels - start dragging
                print("Starting drag mode")
                return
    
        # Place analysis point at these coordinates
        self.place_analysis_point(x, y)
    
    def on_image_drag(self, x, y):
        """Handle drag on image with coordinates already converted by ClickableLabel"""
        if not self.processor:
            return
        
        print(f"on_image_drag received: image coordinates ({x}, {y})")
        self.place_analysis_point(x, y)
    
    def on_mouse_hover(self, x, y):
        """Handle mouse hover over image - show tooltip near analysis point"""
        if not self.processor or not self.current_click:
            QToolTip.hideText()
            return
        
        # Check if hovering near analysis point (convert to widget coords for distance)
        hover_widget_x = x * self.scale_factor_x + self.offset_x
        hover_widget_y = y * self.scale_factor_y + self.offset_y
        
        existing_widget_x = self.current_click['x'] * self.scale_factor_x + self.offset_x
        existing_widget_y = self.current_click['y'] * self.scale_factor_y + self.offset_y
        
        distance = ((hover_widget_x - existing_widget_x)**2 + (hover_widget_y - existing_widget_y)**2)**0.5
        
        if distance <= 15:  # Within 15 pixels of analysis point
            # Show tooltip with current analysis data
            azimuth = self.current_click['azimuth']
            range_val = self.current_click['range']
            tooltip_text = f"{self.tr('azimuth')}: {azimuth:.0f}°\n{self.tr('range')}: {range_val:.0f} км"
            
            # Calculate global position for tooltip
            point = self.image_label.mapToGlobal(self.image_label.rect().topLeft())
            tooltip_x = point.x() + hover_widget_x + 15
            tooltip_y = point.y() + hover_widget_y - 10
            
            QToolTip.showText(QPoint(int(tooltip_x), int(tooltip_y)), tooltip_text)
        else:
            QToolTip.hideText()
    
    def place_analysis_point(self, x, y):
        """Place analysis point at specified coordinates (already in image coords)"""
        if not self.processor:
            return
        
        # Ensure coordinates are within bounds
        x = max(0, min(x, self.processor.image.width - 1))
        y = max(0, min(y, self.processor.image.height - 1))
        
        print(f"Placing analysis point at image coordinates: ({x}, {y})")
        
        try:
            # Calculate azimuth and range
            azimuth, range_val = self.calculate_azimuth_range(x, y)
            
            # Update current click
            self.current_click = {
                'x': x, 'y': y,
                'azimuth': azimuth, 'range': range_val
            }
            
            # Redraw the display
            self.display_image()
            
            # Update results and report
            self.update_results_display()
            self.update_report_data()
            
        except Exception as e:
            QMessageBox.critical(self, self.tr("error"), f"Could not process point: {str(e)}")
    
    def calculate_azimuth_range(self, x, y):
        """Calculate azimuth and range using custom scale if available"""
        # Calculate relative coordinates from center
        dx = x - self.processor.center_x
        dy = self.processor.center_y - y  # Invert Y axis
        
        # Calculate range (distance from center)
        range_pixels = math.sqrt(dx**2 + dy**2)
        
        # Use custom scale distance if available
        if self.custom_scale_distance:
            scale_value = int(self.scale_combo.currentText())
            range_actual = (range_pixels / self.custom_scale_distance) * scale_value
        else:
            # Fall back to original bottom edge calculation
            bottom_edge_distance = self.processor.image.height - self.processor.center_y
            scale_value = int(self.scale_combo.currentText())
            range_actual = (range_pixels / bottom_edge_distance) * scale_value
        
        # Calculate azimuth (angle from north, clockwise)
        azimuth_radians = math.atan2(dx, dy)
        azimuth_degrees = math.degrees(azimuth_radians)
        
        # Normalize azimuth to 0-360 degrees
        if azimuth_degrees < 0:
            azimuth_degrees += 360
            
        return azimuth_degrees, range_actual
    
    def update_report_data(self):
        """Update the report data panel"""
        if not self.processor:
            self.auto_azimuth_label.setText("β - --°")
            self.auto_distance_label.setText("D - -- км")
            self.auto_scale_label.setText("M = --")
            return
        if self.current_click:
            # Update auto-filled data
            azimuth = self.current_click['azimuth']
            distance = self.current_click['range']
            scale = int(self.scale_combo.currentText())
            
            # Remove the status message - just update the data fields
            self.auto_azimuth_label.setText(f"β - {azimuth:.0f}°")
            self.auto_distance_label.setText(f"D - {distance:.1f} км")
            self.auto_scale_label.setText(f"M = {scale}")
        else:
            self.auto_azimuth_label.setText("β - --°")
            self.auto_distance_label.setText("D - -- км")
            
            # Scale is still available even without analysis point
            if hasattr(self, 'scale_combo'):
                scale = int(self.scale_combo.currentText())
                self.auto_scale_label.setText(f"M = {scale}")
            else:
                self.auto_scale_label.setText("M = --")
                
    def show_center_zoom(self):
        """Show zoom focused on current center position"""
        if not self.processor:
            return
        
        # Enable zoom widget and focus on center
        self.image_label.zoom_widget.show_zoom()
        
        # Update zoom to show center
        self.image_label.zoom_widget.update_cursor_position(self.processor.center_x, self.processor.center_y)

    def hide_center_zoom(self):
        """Hide zoom widget"""
        # Only hide if no special modes are active
        if not self.scale_edge_mode and not self.center_setting_mode:
            self.image_label.zoom_widget.hide_zoom()

    def create_center_move_button_with_hover(self, text, move_func):
        """Create a movement button with hover zoom functionality"""
        btn = QPushButton(text)
        btn.setMaximumWidth(40)
        btn.clicked.connect(move_func)
        
        # Add hover events for zoom
        def on_enter(event):
            self.show_center_zoom()
            QPushButton.enterEvent(btn, event)
        
        def on_leave(event):
            self.hide_center_zoom()
            QPushButton.leaveEvent(btn, event)
        
        btn.enterEvent = on_enter
        btn.leaveEvent = on_leave
        
        return btn

    def update_target_number(self, text):
        """Update persistent target number"""
        self.current_target_number = text
    
    def update_height(self, text):
        """Update height value"""
        self.current_height = text
    
    def update_obstacles(self, text):
        """Update obstacles status"""
        self.current_obstacles = text
    
    def update_detection(self, text):
        """Update detection type"""
        self.current_detection = text
    
    def save_current_image_data(self):
        """Зберегти дані поточного зображення для пакетної обробки"""
        if not self.processor or not self.current_click:
            QMessageBox.warning(self, "Warning", "No image or analysis point to save")
            return False
        
        # Створюємо дані для збереження
        image_data = {
            'image_path': self.current_image_path,
            'image_name': os.path.basename(self.current_image_path),
            'analysis_point': {
                'x': self.current_click['x'],
                'y': self.current_click['y'],
                'azimuth': self.current_click['azimuth'],
                'range': self.current_click['range']
            },
            'target_data': {
                'number': self.current_target_number,
                'height': self.current_height,
                'obstacles': self.current_obstacles,
                'detection': self.current_detection,
                'scale': int(self.scale_combo.currentText())
            },
            'processor_settings': {
                'center_x': self.processor.center_x,
                'center_y': self.processor.center_y,
                'scale_edge_point': self.scale_edge_point,
                'custom_scale_distance': self.custom_scale_distance
            }
        }
        
        # Перевіряємо чи вже існує запис для цього зображення
        existing_index = -1
        for i, saved_data in enumerate(self.processed_images):
            if saved_data['image_path'] == self.current_image_path:
                existing_index = i
                break
        
        if existing_index >= 0:
            # Оновлюємо існуючий запис
            self.processed_images[existing_index] = image_data
            self.add_result(f"Updated data for: {image_data['image_name']}")
        else:
            # Додаємо новий запис
            self.processed_images.append(image_data)
            self.add_result(f"Saved data for: {image_data['image_name']}")
        
        self.add_result(f"Total processed images: {len(self.processed_images)}")
        return True
    
    def load_folder_for_batch_processing(self):
        """Завантажити папку для пакетної обробки"""
        folder_path = QFileDialog.getExistingDirectory(self, "Select Folder for Batch Processing")
        
        if folder_path:
            self.current_folder = folder_path
            
            # Знайти всі зображення в папці
            image_extensions = ('.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff', '.tif')
            self.current_folder_images = []
            
            try:
                for filename in os.listdir(folder_path):
                    if filename.lower().endswith(image_extensions):
                        full_path = os.path.join(folder_path, filename)
                        self.current_folder_images.append(full_path)
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Could not read folder: {e}")
                return
            
            self.current_folder_images.sort()
            self.current_image_index = -1
            
            # Завантажити browser
            self.load_folder_thumbnails()
            self.browser_widget.show()
            self.main_splitter.setSizes([300, 180, 720, 300])
            
            self.add_result(f"Loaded folder for batch processing: {os.path.basename(folder_path)}")
            self.add_result(f"Found {len(self.current_folder_images)} images")
            self.add_result("Process each image, then use 'Create Batch Album'")

    def next_image_in_batch(self):
        """Перейти до наступного зображення в папці"""
        if not self.current_folder_images:
            QMessageBox.warning(self, "Warning", "No folder loaded for batch processing")
            return
        
        # Зберегти дані поточного зображення якщо є
        if self.processor and self.current_click:
            self.save_current_image_data()
        
        # Перейти до наступного
        self.current_image_index = (self.current_image_index + 1) % len(self.current_folder_images)
        next_image = self.current_folder_images[self.current_image_index]
        
        self.load_image(next_image)
        self.add_result(f"Loaded image {self.current_image_index + 1}/{len(self.current_folder_images)}: {os.path.basename(next_image)}")

    def previous_image_in_batch(self):
        """Перейти до попереднього зображення в папці"""
        if not self.current_folder_images:
            QMessageBox.warning(self, "Warning", "No folder loaded for batch processing")
            return
        
        # Зберегти дані поточного зображення якщо є
        if self.processor and self.current_click:
            self.save_current_image_data()
        
        # Перейти до попереднього
        self.current_image_index = (self.current_image_index - 1) % len(self.current_folder_images)
        prev_image = self.current_folder_images[self.current_image_index]
        
        self.load_image(prev_image)
        self.add_result(f"Loaded image {self.current_image_index + 1}/{len(self.current_folder_images)}: {os.path.basename(prev_image)}")

    def create_batch_album(self):
        """Створити Word альбом з усіх оброблених зображень"""
        if not DOCX_AVAILABLE:
            QMessageBox.warning(self, "Warning", "python-docx library not installed")
            return
        
        if not self.processed_images:
            QMessageBox.warning(self, "Warning", "No processed images to export")
            return
        
        # Зберегти дані поточного зображення якщо є
        if self.processor and self.current_click:
            self.save_current_image_data()
        
        # Вибрати місце збереження
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Create Batch Album", "",
            "Word Documents (*.docx);;All files (*.*)"
        )
        
        if file_path:
            try:
                # Створити новий документ
                doc = Document()
                
                # Встановити Arial як основний шрифт
                self.set_document_arial_font(doc)
                
                # Встановити поля документу як у JS коді
                self.set_document_margins(doc)

                # Створити таблицю для кожного обробленого зображення
                for i, image_data in enumerate(self.processed_images):
                    print(f"Processing image {i+1}/{len(self.processed_images)}: {image_data['image_name']}")
                    
                    # Додати пустий параграф перед таблицею (як у JS)
                    if i > 0:  # Не для першої таблиці
                        doc.add_paragraph("")
                    # Створити таблицю для цього зображення
                    self.create_js_style_table(doc, image_data)

                # Зберегти документ
                doc.save(file_path)
                
                QMessageBox.information(self, "Success", 
                                    f"Batch album created with {len(self.processed_images)} images!\nSaved: {os.path.basename(file_path)}")
                self.add_result(f"Created batch album: {os.path.basename(file_path)}")
                
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Could not create batch album: {str(e)}")
                import traceback
                traceback.print_exc()

    def set_document_arial_font(self, doc):
        """Встановити Arial як основний шрифт документу"""
        try:
            # Встановити стиль Normal з Arial
            styles = doc.styles
            normal_style = styles['Normal']
            normal_style.font.name = 'Arial'
            normal_style.font.size = Pt(12)
            
            # Встановити Arial для всіх стилів таблиць
            for style in styles:
                if hasattr(style, 'font'):
                    style.font.name = 'Arial'
                    
        except Exception as e:
            print(f"Error setting Arial font: {e}")


    def set_table_borders(self, table):
        """Встановити межі таблиці як у шаблоні"""
        try:
            tbl = table._tbl
            tblPr = tbl.tblPr
            
            # Створити межі
            tblBorders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tblBorders.append(border)
            tblPr.append(tblBorders)
            
            # Встановити ширину таблиці
            tblW = OxmlElement('w:tblW')
            tblW.set(qn('w:type'), 'pct')
            tblW.set(qn('w:w'), '5000')
            tblPr.append(tblW)
            
        except Exception as e:
            print(f"Error setting table borders: {e}")

    def create_processed_image_from_data(self, image_data):
        """Створити оброблене зображення з збережених даних"""
        try:
            # Завантажити зображення
            with Image.open(image_data['image_path']) as original_image:
                # Конвертувати в RGB
                if original_image.mode != 'RGB':
                    if original_image.mode == 'RGBA':
                        rgb_image = Image.new('RGB', original_image.size, (255, 255, 255))
                        rgb_image.paste(original_image, mask=original_image.split()[-1])
                        final_image = rgb_image
                    else:
                        final_image = original_image.convert('RGB')
                else:
                    final_image = original_image.copy()
            
            # Намалювати ЧОРНУ лінію від точки аналізу до правого верхнього кута
            draw = ImageDraw.Draw(final_image)
            analysis_point = image_data['analysis_point']
            
            draw.line([
                (analysis_point['x'], analysis_point['y']), 
                (final_image.width - 1, 0)
            ], fill='black', width=3)
            
            return final_image
            
        except Exception as e:
            print(f"Error creating processed image from data: {e}")
            return None

    def set_document_margins(self, doc):
        """Встановити поля документу як у JS коді"""
        try:
            section = doc.sections[0]
            section.top_margin = Cm(1)      # top: "1cm"
            section.bottom_margin = Cm(1)   # bottom: "1cm" 
            section.left_margin = Cm(0.25)  # left: "0.25cm"
            section.right_margin = Cm(0.5)  # right: "0.5cm"
            
        except Exception as e:
            print(f"Error setting document margins: {e}")

    def create_js_style_table(self, doc, image_data):
        """Створити таблицю точно як у JS коді"""
        try:
            # Створити таблицю 1x3 з фіксованою шириною
            table = doc.add_table(rows=1, cols=3)
            table.style = None
            
            # Встановити загальні властивості таблиці
            tbl = table._tbl
            tblPr = tbl.tblPr
            
            # Ширина таблиці 100%
            tblW = OxmlElement('w:tblW')
            tblW.set(qn('w:type'), 'pct')
            tblW.set(qn('w:w'), '5000')  # 100% ширини
            tblPr.append(tblW)
            
            # Layout fixed (як у JS)
            tblLayout = OxmlElement('w:tblLayout')
            tblLayout.set(qn('w:type'), 'fixed')
            tblPr.append(tblLayout)
            
            # Отримати комірки
            row = table.rows[0]
            label_cell = row.cells[0]
            image_cell = row.cells[1] 
            data_cell = row.cells[2]
            
            # Встановити точні розміри комірок як у JS
            self.set_js_cell_widths(label_cell, image_cell, data_cell)
            
            # Встановити висоту рядка (із JS: TABLE_IMAGE_CELL.HEIGHT)
            # Припускаємо що це приблизно 13cm для пропорцій 15:13
            row.height = Cm(13)
            
            # Заповнити комірки
            self.fill_js_label_cell(label_cell)
            self.fill_js_image_cell(image_cell, image_data)
            self.fill_js_data_cell(data_cell, image_data)
            
        except Exception as e:
            print(f"Error creating JS style table: {e}")
            import traceback
            traceback.print_exc()

    def set_js_cell_widths(self, label_cell, image_cell, data_cell):
        """Встановити ширини комірок точно як у JS коді"""
        try:
            # Ліва комірка: 2.55cm (як у JS)
            label_cell.width = Cm(2.55)
            
            # Права комірка: 3cm (як у JS)  
            data_cell.width = Cm(3)
            
            # Середня комірка: решта простору (буде автоматично розрахована)
            # У JS це TABLE_IMAGE_CELL.WIDTH - припустимо що це близько 15cm
            image_cell.width = Cm(15)
            
        except Exception as e:
            print(f"Error setting cell widths: {e}")

    def fill_js_label_cell(self, cell):
        """Заповнити ліву комірку точно як у JS"""
        try:
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = para.add_run("Індикатор ЗРЛ")
            run.font.name = 'Arial'
            run.font.size = Pt(12)  # font: { size: 12} з JS
            
            # Вертикальне центрування
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Встановити межі комірки як у JS коді
            self.set_js_label_cell_borders(cell)
            
        except Exception as e:
            print(f"Error filling JS label cell: {e}")

    def set_js_label_cell_borders(self, cell):
        """Встановити межі лівої комірки як у JS коді"""
        try:
            tc = cell._tc
            tcPr = tc.tcPr
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.append(tcPr)
            
            # Створити межі комірки
            tcBorders = OxmlElement('w:tcBorders')
            
            # left: { style: "none", size: 0, color: "ffffff" }
            left_border = OxmlElement('w:left')
            left_border.set(qn('w:val'), 'none')
            tcBorders.append(left_border)
            
            # right: { style: "single", size: 1, color: "000000" }
            right_border = OxmlElement('w:right')
            right_border.set(qn('w:val'), 'single')
            right_border.set(qn('w:sz'), '4')
            right_border.set(qn('w:color'), '000000')
            tcBorders.append(right_border)
            
            # top: { style: "none", size: 0, color: "ffffff" }
            top_border = OxmlElement('w:top')
            top_border.set(qn('w:val'), 'none')
            tcBorders.append(top_border)
            
            # bottom: { style: "none", size: 0, color: "ffffff" }
            bottom_border = OxmlElement('w:bottom')
            bottom_border.set(qn('w:val'), 'none')
            tcBorders.append(bottom_border)
            
            tcPr.append(tcBorders)
            
        except Exception as e:
            print(f"Error setting JS label cell borders: {e}")

    def fill_js_image_cell(self, cell, image_data):
        """Заповнити середню комірку зображенням як у JS"""
        try:
            # Створити оброблене зображення з чорною лінією
            processed_image = self.create_processed_image_from_data(image_data)
            
            if not processed_image:
                print("Failed to create processed image")
                return
            
            # Зберегти у тимчасовий файл
            with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as temp_file:
                temp_path = temp_file.name
                processed_image.save(temp_path, 'JPEG', quality=95)
            
            # Додати зображення в комірку з центруванням
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Розрахувати розмір зображення для пропорцій 15:13
            # TABLE_IMAGE_CELL.WIDTH із JS коду - припускаємо 15cm
            img_width = Cm(14.5)  # Трохи менше ширини комірки
            img_height = Cm(14.5 * 13/15)  # Підтримка пропорцій 15:13
            
            run = para.add_run()
            run.add_picture(temp_path, width=img_width, height=img_height)
            
            # Вертикальне центрування
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Білий фон комірки (як у JS: shading: { fill: "FFFFFF" })
            self.set_cell_shading(cell, "FFFFFF")

            # NEW: Встановити межі для комірки зображення
            self.set_js_image_cell_borders(cell)
            
            # Видалити тимчасовий файл
            try:
                os.remove(temp_path)
            except:
                pass
            
        except Exception as e:
            print(f"Error filling JS image cell: {e}")
            import traceback
            traceback.print_exc()

    def set_cell_shading(self, cell, color):
        """Встановити фон комірки"""
        try:
            tc = cell._tc
            tcPr = tc.tcPr
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.append(tcPr)
            
            # Додати білий фон
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), color)
            tcPr.append(shd)
            
        except Exception as e:
            print(f"Error setting cell shading: {e}")

    def set_js_image_cell_borders(self, cell):
        """Встановити межі для комірки зображення"""
        try:
            tc = cell._tc
            tcPr = tc.tcPr
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.append(tcPr)
            
            # Створити межі комірки - всі видимі
            tcBorders = OxmlElement('w:tcBorders')
            
            for border_name in ['left', 'right', 'top', 'bottom']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:color'), '000000')
                tcBorders.append(border)
            
            tcPr.append(tcBorders)
            
        except Exception as e:
            print(f"Error setting image cell borders: {e}")

    def set_js_data_cell_borders(self, cell):
        """Встановити межі для комірки даних"""
        try:
            tc = cell._tc
            tcPr = tc.tcPr
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.append(tcPr)
            
            # Створити межі комірки - всі видимі крім лівої (яка вже є від середньої комірки)
            tcBorders = OxmlElement('w:tcBorders')
            
            # Ліва межа невидима (оскільки середня комірка вже має праву межу)
            left_border = OxmlElement('w:left')
            left_border.set(qn('w:val'), 'none')
            tcBorders.append(left_border)
            
            # Решта меж видимі
            for border_name in ['right', 'top', 'bottom']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:color'), '000000')
                tcBorders.append(border)
            
            tcPr.append(tcBorders)
            
        except Exception as e:
            print(f"Error setting data cell borders: {e}")

    def fill_js_data_cell(self, cell, image_data):
        """Заповнити праву комірку даними ТОЧНО як у JS коді"""
        try:
            # Очистити комірку
            cell.text = ""
            
            # Отримати дані
            target_data = image_data['target_data']
            analysis_point = image_data['analysis_point']
            
            # Створити параграф з усім текстом
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # 1. ВАЖЛИВО: 2 переноси рядків зверху (як у JS: break: 2)
            run1 = para.add_run(" ")
            run1.font.name = 'Arial'
            run1.add_break()
            run1.add_break()
            
            # 2. НОМЕР ЦІЛІ (підкреслений, курсив, 12pt)
            run2 = para.add_run(target_data['number'])
            run2.font.name = 'Arial'
            run2.font.size = Pt(12)
            run2.italic = True
            run2.underline = True
            run2.add_break()
            
            # 3. АЗИМУТ (курсив, 12pt)
            run3 = para.add_run(f"β – {analysis_point['azimuth']:.0f}º")
            run3.font.name = 'Arial'
            run3.font.size = Pt(12)
            run3.italic = True
            run3.add_break()
            
            # 4. ДАЛЬНІСТЬ (курсив, 12pt)
            run4 = para.add_run(f"D – {analysis_point['range']:.1f} км")
            run4.font.name = 'Arial'
            run4.font.size = Pt(12)
            run4.italic = True
            run4.add_break()
            
            # 5. ВИСОТА (курсив, 12pt) - пропускаємо, бо не було в JS
            
            # 6. ПЕРЕШКОДИ (курсив, 9pt) - фіксований текст як у JS
            run5 = para.add_run("без перешкод")
            run5.font.name = 'Arial'
            run5.font.size = Pt(9)
            run5.italic = True
            run5.add_break()
            
            # 7. ТИП ВИЯВЛЕННЯ (курсив, 9pt)
            # У JS це COORDINATE_TYPE_EQUIVALENT[targetProps.type] - використовуємо наш detection
            run6 = para.add_run(target_data['detection'])
            run6.font.name = 'Arial'
            run6.font.size = Pt(9)
            run6.italic = True
            run6.add_break()
            
            # 8. МАСШТАБ (курсив, 9pt)
            run7 = para.add_run(f"М – {target_data['scale']}")
            run7.font.name = 'Arial'
            run7.font.size = Pt(9)
            run7.italic = True
            
            # NEW: Встановити межі для комірки даних  
            self.set_js_data_cell_borders(cell)

            # Без вертикального вирівнювання - за замовчуванням top
            
        except Exception as e:
            print(f"Error filling JS data cell: {e}")
            import traceback
            traceback.print_exc()

    def toggle_scale_edge_mode(self):
        """Toggle scale edge setting mode"""
        self.scale_edge_mode = self.scale_edge_btn.isChecked()
        
        # Update cursor and visual feedback
        self.image_label.set_scale_edge_mode(self.scale_edge_mode)
        
        if self.scale_edge_mode:
            self.scale_edge_btn.setStyleSheet("background-color: #4CAF50; color: white;")
            self.add_result(self.tr("scale_edge_active"))
        else:
            self.scale_edge_btn.setStyleSheet("")
    
    def set_scale_edge_point(self, x, y):
        """Set the scale edge point"""
        # Ensure coordinates are within bounds
        x = max(0, min(x, self.processor.image.width - 1))
        y = max(0, min(y, self.processor.image.height - 1))
        
        # Calculate distance from center
        dx = x - self.processor.center_x
        dy = y - self.processor.center_y
        distance = math.sqrt(dx*dx + dy*dy)
        
        # Store the scale edge point and distance
        self.scale_edge_point = {'x': x, 'y': y}
        self.custom_scale_distance = distance
        
        # Auto-disable scale edge mode after setting
        self.scale_edge_mode = False
        self.scale_edge_btn.setChecked(False)
        self.scale_edge_btn.setStyleSheet("")
        self.image_label.set_scale_edge_mode(False)
        
        # Update display
        self.display_image()
        
        # Recalculate current analysis point if exists
        if self.current_click:
            azimuth, range_val = self.calculate_azimuth_range(
                self.current_click['x'], self.current_click['y']
            )
            self.current_click['azimuth'] = azimuth
            self.current_click['range'] = range_val
        
        # Update results
        self.update_results_display()
        self.update_report_data()
        self.add_result(self.tr("scale_edge_set").format(distance=distance))
    
    def update_results_display(self):
        """Update the results area"""
        self.results_text.clear()
        
        if self.processor:
            self.add_result(self.tr("image_info").format(name=os.path.basename(self.current_image_path)))
            self.add_result(self.tr("size").format(width=self.processor.image.width, 
                                                  height=self.processor.image.height))
            self.add_result(self.tr("scale_info").format(scale=int(self.scale_combo.currentText())))
            self.add_result(self.tr("center_info").format(x=self.processor.center_x, 
                                                         y=self.processor.center_y))
            
            # Show scale reference information
            if self.custom_scale_distance:
                self.add_result(f"Custom scale edge: {self.custom_scale_distance:.1f} px = {self.scale_combo.currentText()} units")
            else:
                bottom_distance = self.processor.image.height - self.processor.center_y
                self.add_result(self.tr("bottom_edge").format(scale=int(self.scale_combo.currentText())))
                self.add_result(self.tr("pixels_south").format(pixels=bottom_distance))
            self.add_result("")
        
        if self.current_click:
            self.add_result(self.tr("analysis_point"))
            self.add_result(f"{self.tr('position')}: ({self.current_click['x']}, {self.current_click['y']})")
            self.add_result(f"{self.tr('azimuth')}: {self.current_click['azimuth']:.0f}°")
            self.add_result(f"{self.tr('range')}: {self.current_click['range']:.0f} км")
            self.add_result("")
            self.add_result(self.tr("click_to_place"))
            self.add_result(self.tr("drag_to_move"))
            self.add_result(self.tr("line_connects"))
        else:
            self.add_result(self.tr("click_on_image"))
    
    def update_scale(self):
        """Update scale value"""
        if self.processor:
            new_scale = int(self.scale_combo.currentText())
            
            if self.current_click:
                azimuth, range_val = self.calculate_azimuth_range(
                    self.current_click['x'], self.current_click['y']
                )
                self.current_click['azimuth'] = azimuth
                self.current_click['range'] = range_val
            
            self.update_results_display()
            self.update_report_data()
            self.add_result(self.tr("scale_updated").format(scale=new_scale))
    
    def move_center(self, dx, dy):
        """Move center point"""
        if not self.processor:
            QMessageBox.warning(self, self.tr("warning"), self.tr("no_image_first"))
            return
        
        self.processor.move_center(dx, dy)
        
        if self.current_click:
            azimuth, range_val = self.calculate_azimuth_range(
                self.current_click['x'], self.current_click['y']
            )
            self.current_click['azimuth'] = azimuth
            self.current_click['range'] = range_val
        
        self.display_image()
        self.update_results_display()
        self.update_report_data()
        self.add_result(self.tr("center_moved").format(x=self.processor.center_x, 
                                                    y=self.processor.center_y))
    
        # NEW: Update zoom to show new center position if zoom is visible
        if self.image_label.zoom_widget.isVisible():
            self.image_label.zoom_widget.update_cursor_position(self.processor.center_x, self.processor.center_y)
    
    def show_center_preview(self):
        """Show center position"""
        if self.processor:
            self.display_image()
            self.update_results_display()
    
    def clear_results(self):
        """Clear analysis point"""
        self.current_click = None
        if self.processor:
            self.display_image()
        self.update_results_display()
        self.update_report_data()
    
    def save_current_image(self):
        """Save processed image"""
        if not self.processor:
            QMessageBox.warning(self, self.tr("warning"), self.tr("no_image_first"))
            return
        
        if not self.current_click:
            QMessageBox.warning(self, self.tr("warning"), self.tr("no_analysis_point"))
            return
        
        file_path, _ = QFileDialog.getSaveFileName(
            self, self.tr("save_processed_image"), "",
            f"{self.tr('jpeg_files')} (*.jpg);;{self.tr('png_files')} (*.png);;{self.tr('all_files')} (*.*)"
        )
        
        if file_path:
            try:
                # Create final image with ONLY the line
                final_image = self.processor.image.copy()
                
                # Force convert to RGB for JPEG
                if file_path.lower().endswith(('.jpg', '.jpeg')):
                    if final_image.mode != 'RGB':
                        if final_image.mode == 'RGBA':
                            rgb_image = Image.new('RGB', final_image.size, (255, 255, 255))
                            rgb_image.paste(final_image, mask=final_image.split()[-1])
                            final_image = rgb_image
                        else:
                            final_image = final_image.convert('RGB')
                
                draw = ImageDraw.Draw(final_image)
                
                # Draw ONLY the line to top-right corner
                draw.line([
                    (self.current_click['x'], self.current_click['y']), 
                    (final_image.width - 1, 0)
                ], fill='blue', width=3)
                
                # Save with appropriate format
                if file_path.lower().endswith('.png'):
                    final_image.save(file_path, 'PNG')
                else:
                    final_image.save(file_path, 'JPEG', quality=95)
                
                self.add_result(f"{self.tr('saved')}: {os.path.basename(file_path)}")
                self.add_result(self.tr("saved_clean"))
                QMessageBox.information(self, self.tr("success"), self.tr("image_saved_successfully"))
                
            except Exception as e:
                QMessageBox.critical(self, self.tr("error"), 
                                   self.tr("could_not_save").format(error=str(e)))
    
    def create_new_word_album(self):
        """Create a new Word album document"""
        if not DOCX_AVAILABLE:
            QMessageBox.warning(self, self.tr("warning"), 
                               self.tr("docx_not_available"))
            return
        
        if not self.processor or not self.current_click:
            QMessageBox.warning(self, self.tr("warning"), 
                               self.tr("load_image_and_point"))
            return
        
        # Ask user for save location
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Create New Word Album", "",
            "Word Documents (*.docx);;All files (*.*)"
        )
        
        if file_path:
            try:
                # Create new document
                self.current_word_document = Document()
                self.word_document_path = file_path
                
                # Set default font and landscape orientation
                self.set_document_defaults()
                
                # Add current processed image to document
                self.add_image_to_word_document()
                
                # Save document
                self.current_word_document.save(file_path)
                
                QMessageBox.information(self, self.tr("success"), 
                                       f"New album created: {os.path.basename(file_path)}")
                self.add_result(f"Created new album: {os.path.basename(file_path)}")
                
            except Exception as e:
                QMessageBox.critical(self, self.tr("error"), 
                                   f"Could not create Word album: {str(e)}")
    
    def add_to_existing_album(self):
        """Add current image to existing Word album"""
        if not DOCX_AVAILABLE:
            QMessageBox.warning(self, self.tr("warning"), 
                               self.tr("docx_not_available"))
            return
        
        if not self.processor or not self.current_click:
            QMessageBox.warning(self, self.tr("warning"), 
                               self.tr("load_image_and_point"))
            return
        
        # Ask user to select existing document
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Select Existing Word Album", "",
            "Word Documents (*.docx);;All files (*.*)"
        )
        
        if file_path:
            try:
                # Open existing document
                self.current_word_document = Document(file_path)
                self.word_document_path = file_path
                
                # Ensure Arial is set as default
                self.set_document_defaults()
                
                # Add current processed image to document
                self.add_image_to_word_document()
                
                # Save document
                self.current_word_document.save(file_path)
                
                QMessageBox.information(self, self.tr("success"), 
                                       f"Added to album: {os.path.basename(file_path)}")
                self.add_result(f"Added to album: {os.path.basename(file_path)}")
                
            except Exception as e:
                QMessageBox.critical(self, self.tr("error"), 
                                   f"Could not add to Word album: {str(e)}")
    
    def set_document_defaults(self):
        """Set Arial as default font and landscape orientation"""
        try:
            # Set default paragraph font to Arial
            styles = self.current_word_document.styles
            default_style = styles['Normal']
            default_style.font.name = 'Arial'
            default_style.font.size = Pt(12)
            
            # Set document to landscape orientation
            section = self.current_word_document.sections[0]
            section.orientation = 1  # Landscape
            
            # Swap width and height for landscape
            new_width, new_height = section.page_height, section.page_width
            section.page_width = new_width
            section.page_height = new_height
            
        except Exception as e:
            print(f"Error setting document defaults: {e}")
    
    def add_image_to_word_document(self):
        """Add current processed image and data to Word document"""
        if not self.current_word_document or not self.processor or not self.current_click:
            return
        
        # Create table with 1 row and 3 columns
        table = self.current_word_document.add_table(rows=1, cols=3)
        table.style = None
        
        # Set table properties
        tbl = table._tbl
        tblPr = tbl.tblPr
        
        # Set table borders
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        tblPr.append(tblBorders)

        # Set table width to 100%
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:type'), 'pct')
        tblW.set(qn('w:w'), '5000')
        tblPr.append(tblW)
        
        # Get cells
        label_cell = table.cell(0, 0)    # Left cell for "Індикатор ЗРЛ"
        image_cell = table.cell(0, 1)    # Middle cell for image
        data_cell = table.cell(0, 2)     # Right cell for data
        
        # Set cell widths for landscape orientation
        label_cell.width = Cm(3)   # Narrow for label
        image_cell.width = Cm(20)  # Much larger for image in landscape
        data_cell.width = Cm(4)    # Slightly wider for data
        
        # Configure label cell (make borders invisible except right border)
        self.configure_label_cell(label_cell)
        
        # Set standard cell properties for image and data cells
        for cell in [image_cell, data_cell]:
            cellPr = cell._tc.tcPr
            if cellPr is None:
                cellPr = OxmlElement('w:tcPr')
                cell._tc.append(cellPr)
            
            # Set small cell margins
            tcMar = OxmlElement('w:tcMar')
            for margin in ['top', 'left', 'bottom', 'right']:
                mar = OxmlElement(f'w:{margin}')
                mar.set(qn('w:w'), '108')
                mar.set(qn('w:type'), 'dxa')
                tcMar.append(mar)
            cellPr.append(tcMar)
            
            # Center vertical alignment
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            cellPr.append(vAlign)
        
        # Add content to cells
        self.add_label_to_cell(label_cell)
        self.add_processed_image_to_cell(image_cell)
        self.add_data_to_cell(data_cell)
        
        # Add spacing after table
        if len(self.current_word_document.tables) > 1:
            para = self.current_word_document.add_paragraph()
            para.text = ""
    
    def configure_label_cell(self, cell):
        """Configure label cell with invisible borders except right border"""
        cellPr = cell._tc.tcPr
        if cellPr is None:
            cellPr = OxmlElement('w:tcPr')
            cell._tc.append(cellPr)
        
        # Set cell borders - transparent except right border
        tcBorders = OxmlElement('w:tcBorders')
        
        # Invisible borders
        for border_name in ['top', 'left', 'bottom']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')  # Invisible
            tcBorders.append(border)
        
        # Visible right border
        right_border = OxmlElement('w:right')
        right_border.set(qn('w:val'), 'single')
        right_border.set(qn('w:sz'), '4')  # Thin border
        right_border.set(qn('w:space'), '0')
        right_border.set(qn('w:color'), '000000')
        tcBorders.append(right_border)
        
        cellPr.append(tcBorders)
        
        # Set small margins
        tcMar = OxmlElement('w:tcMar')
        for margin in ['top', 'left', 'bottom', 'right']:
            mar = OxmlElement(f'w:{margin}')
            mar.set(qn('w:w'), '108')
            mar.set(qn('w:type'), 'dxa')
            tcMar.append(mar)
        cellPr.append(tcMar)
        
        # Center vertical alignment
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        cellPr.append(vAlign)
    
    def add_label_to_cell(self, cell):
        """Add 'Індикатор ЗРЛ' label to left cell"""
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add the label text
        run = para.add_run("Індикатор ЗРЛ")
        run.font.size = Pt(12)  # Arial is inherited from document default
    
    def add_processed_image_to_cell(self, cell):
        """Add processed image to cell with bulletproof format handling"""
        try:
            # Get image from processor and ensure it's RGB
            if not self.processor or not hasattr(self.processor, 'image'):
                print("No processor or image available")
                return
            
            final_image = self.processor.image.copy()
            print(f"Image mode before processing: {final_image.mode}")
            
            # BULLETPROOF conversion to RGB
            if final_image.mode != 'RGB':
                print(f"Converting from {final_image.mode} to RGB")
                try:
                    if final_image.mode == 'RGBA':
                        # Create white background for transparency
                        rgb_image = Image.new('RGB', final_image.size, (255, 255, 255))
                        rgb_image.paste(final_image, mask=final_image.split()[-1])
                        final_image = rgb_image
                    elif final_image.mode in ('P', 'L', '1'):
                        # Palette, Grayscale, or Monochrome
                        final_image = final_image.convert('RGB')
                    else:
                        # Any other exotic mode
                        final_image = final_image.convert('RGB')
                    print(f"Successfully converted to: {final_image.mode}")
                except Exception as conv_error:
                    print(f"Conversion error: {conv_error}")
                    # Force conversion using PIL's most compatible method
                    final_image = final_image.convert('RGB')
            
            # Verify we have RGB before proceeding
            if final_image.mode != 'RGB':
                raise ValueError(f"Failed to convert image to RGB, still in {final_image.mode} mode")
            
            # Now safely draw on the RGB image
            draw = ImageDraw.Draw(final_image)
            
            # Draw ONLY the line to top-right corner
            draw.line([
                (self.current_click['x'], self.current_click['y']), 
                (final_image.width - 1, 0)
            ], fill='blue', width=3)
            
            # Save to temporary file as JPEG
            with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as temp_file:
                temp_path = temp_file.name
                print(f"Saving RGB image to: {temp_path}")
                final_image.save(temp_path, 'JPEG', quality=95)
                print(f"Successfully saved JPEG")
            
            # Add to Word document
            cell.text = ""
            top_spacing_para = cell.paragraphs[0]
            top_spacing_para.text = ""
            
            # Set spacing
            pPr = top_spacing_para._element.get_or_add_pPr()
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:before'), '0')
            spacing.set(qn('w:after'), '720')
            pPr.append(spacing)
            
            # Add image paragraph
            img_paragraph = cell.add_paragraph()
            img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add image to document
            run = img_paragraph.add_run()
            run.add_picture(temp_path, width=Cm(19))  # Larger image for landscape orientation
            print(f"Successfully added image to Word document")
            
            # Clean up temp file
            try:
                os.remove(temp_path)
                print(f"Cleaned up temp file")
            except:
                pass
                
        except Exception as e:
            print(f"Error in add_processed_image_to_cell: {e}")
            import traceback
            traceback.print_exc()  # Full error traceback for debugging
    
    def add_data_to_cell(self, cell):
        """Add analysis data to Word document cell with tight formatting"""
        try:
            # Clear existing content
            cell.text = ""
            
            # Get data
            target_num = self.current_target_number
            azimuth = self.current_click['azimuth']
            distance = self.current_click['range']
            height = self.current_height
            obstacles = self.current_obstacles
            detection = self.current_detection
            scale = int(self.scale_combo.currentText())
            
            # Add top spacing
            top_spacing_para = cell.paragraphs[0]
            top_spacing_para.text = ""
            pPr = top_spacing_para._element.get_or_add_pPr()
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:before'), '0')
            spacing.set(qn('w:after'), '720')  # Spacing equivalent to font size 30
            pPr.append(spacing)
            
            # Create data items WITHOUT empty lines between them
            # Format: (text, font_size, italic, underline)
            data_items = [
                (f"{target_num}", 12, True, True),          # Target number - 12pt, italic, underlined
                (f"β - {azimuth:.0f}°", 12, True, False),   # Azimuth - 12pt, italic
                (f"D - {distance:.1f} км", 12, True, False), # Distance - 12pt, italic
                (f"Н - {height} км", 12, True, False),     # Height - 12pt, italic
                (f"{obstacles}", 9, True, False),           # Obstacles - 9pt, italic
                (f"{detection}", 9, True, False),           # Detection type - 9pt, italic
                (f"М={scale}", 9, True, False)              # Scale - 9pt, italic
            ]
            
            # Add each line with tight formatting (no empty lines)
            for text, font_size, italic, underline in data_items:
                # Add new paragraph for each line
                p = cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Left alignment
                
                # Set tight paragraph spacing - no gaps between lines
                pPr = p._element.get_or_add_pPr()
                spacing = OxmlElement('w:spacing')
                spacing.set(qn('w:before'), '0')
                spacing.set(qn('w:after'), '0')  # No space after each line
                spacing.set(qn('w:line'), '200')    # Very tight line spacing
                spacing.set(qn('w:lineRule'), 'auto')
                pPr.append(spacing)
                
                # Add text with formatting
                run = p.add_run(text)
                run.font.size = Pt(font_size)
                run.italic = italic
                run.underline = underline
                
                # Set character spacing (slightly condensed)
                rPr = run._element.rPr
                if rPr is None:
                    rPr = OxmlElement('w:rPr')
                    run._element.insert(0, rPr)
                spacing = OxmlElement('w:spacing')
                spacing.set(qn('w:val'), '-5')  # Slightly condensed
                rPr.append(spacing)
            
        except Exception as e:
            print(f"Error adding data to cell: {e}")

def main():
    app = QApplication(sys.argv)
    window = AzimuthGUI()
    window.show()
    sys.exit(app.exec_())

if __name__ == "__main__":
    main()