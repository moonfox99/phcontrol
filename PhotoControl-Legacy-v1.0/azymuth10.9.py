#!/usr/bin/env python3
"""
Azimuth Image Processor - PyQt5 Version with Ukrainian Translation
Professional interface for processing azimuth grid images
"""

import sys
import os
import math
import tempfile
import json
from datetime import datetime
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QHBoxLayout, 
                             QVBoxLayout, QGridLayout, QPushButton, QLabel, 
                             QComboBox, QTextEdit, QScrollArea, QFrame,
                             QFileDialog, QMessageBox, QSplitter, QToolTip, QLineEdit,
                             QCheckBox, QDateEdit, QSizePolicy, QDialog, QFormLayout, QGroupBox, QDoubleSpinBox,
                             QDialogButtonBox, QTabWidget)
from PyQt5.QtCore import Qt, QSize, pyqtSignal, QTimer, QPoint, QDate
from PyQt5.QtGui import QPixmap, QFont, QIcon
from PIL import Image, ImageDraw
from PIL import ImageFont

from translations import Translations
from widgets import ClickableLabel, VerticalThumbnailWidget
from image_processor import AzimuthImageProcessor
from documentation import DocumentationManager
from help_dialogs import AboutDialog

try:
    from docx import Document
    from docx.shared import Inches, Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE  # ← ВИПРАВЛЕНО: перенесено сюди
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.section import WD_SECTION_START
    from docx.oxml.shared import OxmlElement, qn
    from docx.oxml.ns import qn
    from docx.oxml import parse_xml
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

ALBUM_LAYOUT = {
    # Розміри сторінки A4 в міліметрах
    'PAGE_WIDTH': 210,
    'PAGE_HEIGHT': 297,
    
    # Стандартні поля для титульної та описової сторінок
    'STANDARD_LEFT_MARGIN': 25,
    'STANDARD_RIGHT_MARGIN': 25, 
    'STANDARD_TOP_MARGIN': 10,
    'STANDARD_BOTTOM_MARGIN': 10,

    # ПОЛЯ ДЛЯ СТОРІНОК З ТАБЛИЦЯМИ (20мм зверху, 5мм справа/знизу, 2.5мм зліва)
    'TABLE_PAGES_LEFT_MARGIN': 2.5,    # 2.5мм зліва
    'TABLE_PAGES_RIGHT_MARGIN': 5,     # 5мм справа
    'TABLE_PAGES_TOP_MARGIN': 20,      # 20мм зверху
    'TABLE_PAGES_BOTTOM_MARGIN': 5,    # 5мм знизу
    
    # ПАРАМЕТРИ ДЛЯ ОКРЕМИХ ТАБЛИЦЬ
    'TABLE_ROWS': 1,  # Тепер тільки 1 рядок на таблицю
    'TABLE_COLS': 3,
    
    # Висоти елементів (в мм)
    'PARAGRAPH_HEIGHT': 5,    # Параграф-розділювач 10мм
    'TABLE_HEIGHT': 130,       # Висота таблиці 130мм
    
    # Ширини колонок (в мм) - залишаються без змін
    'COL_1_WIDTH': 25,      # Перша колонка
    'COL_2_WIDTH': 150,     # Друга колонка (для зображень)
    'COL_3_WIDTH': 30,      # Третя колонка
    
    # Параметри зображень - без змін
    'IMAGE_ASPECT_RATIO': 15/13,
    'IMAGE_WIDTH_CM': 14,  # Ширина зображення в см
}

class DefaultTemplateData:
    """Централізовані базові дані для шаблонів"""
    
    @staticmethod
    def get_all():
        return {
            'unit_info': 'А0000',
            'commander': {'rank': 'полковник', 'name': 'П.П. ПЕТРЕНКО'},
            'chief_of_staff': {'rank': 'підполковник', 'name': 'С.С. СИДОРЕНКО'},
            'signature_info': {'rank': 'головний сержант', 'name': 'Анна Петренко'},
            'margins': {'top': 2.25, 'left': 2.5, 'bottom': 0, 'right': 0.75}
        }
    
    @classmethod
    def get_unit_info(cls):
        return cls.get_all()['unit_info']
    
    @classmethod
    def get_commander_info(cls):
        return cls.get_all()['commander']
    
    @classmethod
    def get_chief_of_staff_info(cls):
        return cls.get_all()['chief_of_staff']
    
    @classmethod
    def get_signature_info(cls):
        return cls.get_all()['signature_info']
    
# 2. ВИНЕСЕНІ СПІЛЬНІ СТИЛІ
class UIStyles:
    """Централізовані стилі для UI компонентів"""
    
    INPUT_FIELD = """
        border: 1px solid #dee2e6;
        border-radius: 4px;
        padding: 6px 10px;
        background-color: white;
        font: 12pt "Segoe UI", Arial, sans-serif;
        color: #495057;
        min-height: 20px;
    """
    
    INPUT_HOVER = """
        border: 1px solid #adb5bd;
        background-color: #f8f9fa;
    """
    
    INPUT_FOCUS = """
        border: 1px solid #6c757d;
        background-color: white;
    """
    
    BUTTON_BASE = """
        background-color: #f8f9fa;
        border: 1px solid #dee2e6;
        border-radius: 6px;
        padding: 10px 14px;
        font: 500 12pt "Segoe UI", Arial, sans-serif;
        color: #495057;
    """
    
    BUTTON_HOVER = """
        background-color: #e9ecef;
        border: 1px solid #adb5bd;
        color: #343a40;
    """
    
    @classmethod
    def get_input_style(cls):
        return f"""
            QLineEdit, QComboBox, QDateEdit {{
                {cls.INPUT_FIELD}
            }}
            QLineEdit:hover, QComboBox:hover, QDateEdit:hover {{
                {cls.INPUT_HOVER}
            }}
            QLineEdit:focus, QComboBox:focus, QDateEdit:focus {{
                {cls.INPUT_FOCUS}
            }}
        """
    
    @classmethod
    def get_button_style(cls):
        return f"""
            QPushButton {{
                {cls.BUTTON_BASE}
            }}
            QPushButton:hover {{
                {cls.BUTTON_HOVER}
            }}
        """

def get_reliable_font(font_size):
    try:
        # Курсивні шляхи
        italic_paths = [
            "C:/Windows/Fonts/ariali.ttf",
            "C:/Windows/Fonts/Arial Italic.ttf",
            "/System/Library/Fonts/Arial Italic.ttf",
            "/usr/share/fonts/truetype/arial/ariali.ttf",
            "/usr/share/fonts/TTF/ariali.ttf",
        ]
        
        # Перевіряємо курсивні шрифти
        for font_path in italic_paths:
            if os.path.exists(font_path):
                try:
                    return ImageFont.truetype(font_path, font_size)
                except:
                    continue
        
        # Звичайний Arial як fallback
        regular_paths = [
            "C:/Windows/Fonts/arial.ttf",
            "C:/Windows/Fonts/Arial.ttf",
        ]
        
        for font_path in regular_paths:
            if os.path.exists(font_path):
                try:
                    return ImageFont.truetype(font_path, font_size)
                except:
                    continue
        
        # Default font
        return ImageFont.load_default()
        
    except:
        return ImageFont.load_default()

def mm_to_cm(mm):
    """Перетворення міліметрів в сантиметри для docx"""
    return mm / 10.0

def format_ukrainian_date(date_obj):
    """Форматування дати по-українськи для титульної сторінки"""
    try:
        if isinstance(date_obj, str):
            from datetime import datetime
            date_obj = datetime.strptime(date_obj, "%Y-%m-%d")
        
        months = {
            1: 'січня', 2: 'лютого', 3: 'березня', 4: 'квітня',
            5: 'травня', 6: 'червня', 7: 'липня', 8: 'серпня',
            9: 'вересня', 10: 'жовтня', 11: 'листопада', 12: 'грудня'
        }
        
        day = date_obj.day
        month = months[date_obj.month]
        year = date_obj.year
        
        return f'{day} {month} {year} року'
    except:
        return "дата не вказана"

def set_a4_page_format(section):
    """Встановлення точного формату A4 для секції документу"""
    try:
        from docx.shared import Cm
        from docx.enum.section import WD_ORIENTATION
        
        # Точні розміри A4 згідно з ISO 216
        section.page_width = Cm(21.0)   # 210мм = 21.0см
        section.page_height = Cm(29.7)  # 297мм = 29.7см
        section.orientation = WD_ORIENTATION.PORTRAIT
        
        print(f"✓ A4 format set: {section.page_width.cm:.1f} x {section.page_height.cm:.1f} cm")
        
    except Exception as e:
        print(f"✗ Error setting A4 format: {e}")

def create_complete_album(processed_images, title_data, file_path):
    """Створення повного альбому з правильними полями та підписом на сторінці опису"""
    try:
        doc = Document()
        
        print("=== Creating Complete Album with Description Page Signature ===")
        
        # 1. ТИТУЛЬНА СТОРІНКА (поля встановлюються в create_title_page)
        print("1. Creating title page...")
        create_title_page(doc, title_data)
        
        # 2. СТОРІНКА ОПИСУ З ПІДПИСОМ (НОВА СЕКЦІЯ з розривом сторінки)
        print("2. Creating description page with signature...")
        desc_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
        set_a4_page_format(desc_section)
        
        # Встановлюємо поля для сторінки опису
        desc_section.left_margin = Cm(2.0)    # 2см ліве поле
        desc_section.top_margin = Cm(2.0)     # 2см верхнє поле  
        desc_section.right_margin = Cm(1.0)   # 1см праве поле
        desc_section.bottom_margin = Cm(1.0)  # 1см нижнє поле
        
        # ВИКОРИСТОВУЄМО НОВУ ФУНКЦІЮ З ПІДПИСОМ
        create_description_page_with_signature(doc, title_data)
        
        # 3. СТОРІНКИ З ТАБЛИЦЯМИ
        print("3. Creating table pages...")
        create_new_structure_pages(doc, processed_images)
        
        doc.save(file_path)
        print(f"✓ Complete album with signature saved: {file_path}")
        
        return True
    except Exception as e:
        print(f"✗ Error creating complete album: {e}")
        return False

# 4. ДОПОМІЖНА ФУНКЦІЯ ДЛЯ СТВОРЕННЯ СТИЛІВ
def create_or_get_style(doc, name, style_type, **properties):
    """Створити або отримати існуючий стиль з властивостями"""
    try:
        return doc.styles[name]
    except KeyError:
        style = doc.styles.add_style(name, style_type)
        
        # Застосовуємо властивості
        for prop_path, value in properties.items():
            obj = style
            parts = prop_path.split('.')
            
            # Навігація до потрібного об'єкту
            for part in parts[:-1]:
                obj = getattr(obj, part)
            
            # Встановлення значення
            setattr(obj, parts[-1], value)
        
        return style

# 5. СПРОЩЕНА ФУНКЦІЯ СТВОРЕННЯ СТИЛІВ ДЛЯ ТИТУЛЬНОЇ СТОРІНКИ
def create_title_page_styles(doc):
    """Створення всіх стилів для титульної сторінки"""
    from docx.enum.style import WD_STYLE_TYPE  # ← Додайте цей імпорт
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    styles_config = {
        'TitleTopSpacer': {
            'paragraph_format.first_line_indent': Cm(1),
            'paragraph_format.line_spacing': 1.15,
            'paragraph_format.space_before': Pt(0),
            'paragraph_format.space_after': Pt(10),
            'font.size': Pt(31),
            'font.name': 'Arial'
        },
        'TitleMain': {
            'paragraph_format.first_line_indent': Cm(1),
            'paragraph_format.alignment': WD_ALIGN_PARAGRAPH.CENTER,
            'paragraph_format.space_before': Pt(0),
            'paragraph_format.space_after': Pt(0),
            'font.size': Pt(28),
            'font.name': 'Arial',
            'font.bold': True
        },
        'TitleSignature': {
            'paragraph_format.first_line_indent': Cm(1),
            'paragraph_format.left_indent': Pt(0),
            'paragraph_format.line_spacing': 1.15,
            'paragraph_format.space_before': Pt(0),
            'paragraph_format.space_after': Pt(0),
            'font.size': Pt(16),
            'font.name': 'Arial',
            'font.bold': True
        }
    }
    
    created_styles = {}
    for style_name, properties in styles_config.items():
        created_styles[style_name] = create_or_get_style(
            doc, style_name, WD_STYLE_TYPE.PARAGRAPH, **properties
        )
    
    return created_styles

def create_description_page_with_signature(doc, title_data):
    """Створення сторінки опису з підписом начальника групи"""
    try:
        # 1. Заголовок
        _add_description_heading(doc)
        
        # 2. Таблиця
        _create_description_table(doc)
        
        # 3. Розділювачі
        _add_description_spacers(doc)
        
        # 4. Підпис
        _add_description_signature(doc, title_data)
        
        print("✅ Description page with signature created successfully!")
        
    except Exception as e:
        print(f"✗ Error creating description page: {e}")
        import traceback
        traceback.print_exc()

def _add_description_heading(doc):
    """Додавання заголовку сторінки опису"""
    heading = doc.add_paragraph("Опис альбому фотознімків")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(0)
    
    if len(heading.runs) > 0:
        heading.runs[0].font.size = Pt(22)
        heading.runs[0].font.bold = True
        heading.runs[0].font.name = 'Arial'

def _create_description_table(doc):
    """Створення таблиці опису з точними розмірами"""
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn
    
    # Створення таблиці
    table = doc.add_table(rows=6, cols=6)
    table.style = None
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Відступ таблиці 1см зліва
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), str(int(1.0 * 567)))
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)
    
    # Заголовки
    headers = [
        "№ з/п", "Назва документа", "Обліковий номер",
        "Кількість аркушів (знімків)", "Гриф секретності", "Примітка"
    ]
    
    # Ширини колонок
    column_widths = [1.25, 3.0, 3.0, 3.0, 3.0, 3.0]
    
    # Налаштування заголовків (висота 2см)
    header_row = table.rows[0]
    header_row.height = Cm(2.0)
    
    # XML для точної висоти
    trPr = header_row._tr.trPr
    if trPr is None:
        trPr = OxmlElement('w:trPr')
        header_row._tr.append(trPr)
    
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(2.0 * 567)))
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)
    
    # Заповнення заголовків
    for i, (cell, header_text, col_width) in enumerate(
        zip(header_row.cells, headers, column_widths)
    ):
        cell.width = Cm(col_width)
        cell.text = header_text
        
        try:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        except AttributeError:
            pass
        
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if len(para.runs) > 0:
            para.runs[0].font.name = 'Arial'
            para.runs[0].font.size = Pt(14)
        
        _set_cell_borders(cell, True, True, True, True)
    
    # Рядки даних (висота 0.6см кожен)
    for row_index in range(1, 6):
        row = table.rows[row_index]
        row.height = Cm(0.6)
        
        # XML для точної висоти
        trPr = row._tr.trPr
        if trPr is None:
            trPr = OxmlElement('w:trPr')
            row._tr.append(trPr)
        
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(int(0.6 * 567)))
        trHeight.set(qn('w:hRule'), 'exact')
        trPr.append(trHeight)
        
        for col_index, cell in enumerate(row.cells):
            cell.width = Cm(column_widths[col_index])
            
            try:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            except AttributeError:
                pass
            
            if col_index == 0:
                cell.text = f"{row_index}."
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                if len(para.runs) > 0:
                    para.runs[0].font.name = 'Arial'
                    para.runs[0].font.size = Pt(14)
            
            _set_cell_borders(cell, True, True, True, True)

def _set_cell_borders(cell, top, bottom, left, right):
    """Налаштування рамок комірки"""
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn
    
    try:
        tc = cell._tc
        tcPr = tc.tcPr
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            tc.append(tcPr)
        
        tcBorders = OxmlElement('w:tcBorders')
        
        border_settings = [
            ('top', top), ('bottom', bottom), ('left', left), ('right', right)
        ]
        
        for position, show_border in border_settings:
            border_element = OxmlElement(f'w:{position}')
            if show_border:
                border_element.set(qn('w:val'), 'single')
                border_element.set(qn('w:sz'), '4')
                border_element.set(qn('w:color'), '000000')
            else:
                border_element.set(qn('w:val'), 'none')
            
            tcBorders.append(border_element)
        
        tcPr.append(tcBorders)
        
    except Exception as e:
        print(f"Warning: Could not set cell borders: {e}")

def _add_description_spacers(doc):
    """Додавання 12 абзаців-розділювачів"""
    try:
        from docx.enum.style import WD_STYLE_TYPE  # ← Додайте цей імпорт
        spacer_style = doc.styles.add_style('DescriptionSpacer', WD_STYLE_TYPE.PARAGRAPH)
        spacer_style.paragraph_format.first_line_indent = Cm(1)
        spacer_style.paragraph_format.line_spacing = 1.15
        spacer_style.paragraph_format.space_before = Pt(0)
        spacer_style.paragraph_format.space_after = Pt(0)
        spacer_style.font.size = Pt(28)
        spacer_style.font.name = 'Arial'
    except:
        pass  # Стиль вже існує
    
    for i in range(12):
        doc.add_paragraph(" ", style='DescriptionSpacer')

def _add_description_signature(doc, title_data):
    """Додавання блоку підпису"""
    unit_info = title_data.get('unit_info', 'А0000')
    signature_info = title_data.get('signature_info', {
        'rank': 'головний сержант',
        'name': 'Анна Петренко'
    })
    
    # Створення стилю для підпису
    try:
        signature_style = doc.styles.add_style('DescriptionSignature', WD_STYLE_TYPE.PARAGRAPH)
        signature_style.paragraph_format.first_line_indent = Cm(1)
        signature_style.paragraph_format.left_indent = Pt(0)
        signature_style.paragraph_format.line_spacing = 1.15
        signature_style.paragraph_format.space_before = Pt(0)
        signature_style.paragraph_format.space_after = Pt(0)
        signature_style.font.size = Pt(14)
        signature_style.font.name = 'Arial'
        signature_style.font.bold = False
    except:
        signature_style = doc.styles['DescriptionSignature']
    
    # Три абзаци підпису
    para1 = doc.add_paragraph("Начальник групи секретного документального", style=signature_style)
    para2 = doc.add_paragraph(f"забезпечення військової частини {unit_info}", style=signature_style)
    
    para3 = doc.add_paragraph(style=signature_style)
    rank_run = para3.add_run(signature_info['rank'])
    rank_run.font.name = 'Arial'
    rank_run.font.size = Pt(14)
    rank_run.font.bold = False
    
    try:
        tab_stops = para3.paragraph_format.tab_stops
        tab_stops.clear_all()
        tab_stops.add_tab_stop(Cm(13), WD_TAB_ALIGNMENT.RIGHT)
        
        name_run = para3.add_run(f"\t{signature_info['name']}")
        name_run.font.name = 'Arial'
        name_run.font.size = Pt(14)
        name_run.font.bold = False
        
    except Exception:
        arrow_run = para3.add_run(f" → → → → → → → {signature_info['name']}")
        arrow_run.font.name = 'Arial'
        arrow_run.font.size = Pt(14)
        arrow_run.font.bold = False

# ДЕКОРАТОР ДЛЯ ПЕРЕВІРОК
def requires_image_and_point(func):
    """Декоратор для функцій що потребують зображення та точку аналізу"""
    def wrapper(self, *args, **kwargs):
        if not self.processor:
            QMessageBox.warning(self, "Попередження", "Спочатку завантажте зображення")
            return False
        
        if not self.current_click:
            QMessageBox.warning(self, "Попередження", "Немає точки аналізу")
            return False
        
        return func(self, *args, **kwargs)
    return wrapper

def requires_image_only(func):
    """Декоратор для функцій що потребують тільки зображення"""
    def wrapper(self, *args, **kwargs):
        if not self.processor:
            QMessageBox.warning(self, "Попередження", "Спочатку завантажте зображення")
            return False
        
        return func(self, *args, **kwargs)
    return wrapper

def create_new_structure_pages(doc, processed_images):
    """Створення сторінок з новою структурою та таблицями БЕЗ лівого поля"""
    try:
        if not processed_images:
            print("No processed images to add")
            return
        
        # Додаємо нову секцію для таблиць
        table_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
        
        # Встановлюємо A4 формат
        set_a4_page_format(table_section)
        
        # ОНОВЛЕНІ ПОЛЯ: 0мм зліва (таблиці впритул до краю!), 20мм зверху, 5мм справа/знизу
        table_section.left_margin = Cm(mm_to_cm(ALBUM_LAYOUT['TABLE_PAGES_LEFT_MARGIN']))  # 0мм!
        table_section.right_margin = Cm(mm_to_cm(ALBUM_LAYOUT['TABLE_PAGES_RIGHT_MARGIN']))
        table_section.top_margin = Cm(mm_to_cm(ALBUM_LAYOUT['TABLE_PAGES_TOP_MARGIN']))
        table_section.bottom_margin = Cm(mm_to_cm(ALBUM_LAYOUT['TABLE_PAGES_BOTTOM_MARGIN']))
        
        print(f"✓ NEW margins set: left=0mm (no margin!), top=20mm, right=5mm, bottom=5mm")
        
        # Обробляємо зображення парами (по 2 на сторінку)
        for i in range(0, len(processed_images), 2):
            first_image = processed_images[i]
            second_image = processed_images[i + 1] if i + 1 < len(processed_images) else None
            
            print(f"\n=== Creating page for images {i+1}-{i+2 if second_image else i+1} ===")

            # 1. Параграф-розділювач 5мм
            create_spacer_paragraph(doc, ALBUM_LAYOUT['PARAGRAPH_HEIGHT'])
            print("✓ Added top spacer paragraph (5mm)")
            
            # 2. Перша таблиця 130мм (БЕЗ лівого відступу)
            create_single_image_table(doc, first_image)
            print("✓ Added first table (130mm, aligned to LEFT EDGE)")
            
            # 3. Параграф-розділювач 5мм
            create_spacer_paragraph(doc, ALBUM_LAYOUT['PARAGRAPH_HEIGHT'])
            print("✓ Added middle spacer paragraph (5mm)")
            
            # 4. Друга таблиця 130мм (якщо є друге зображення)
            if second_image:
                create_single_image_table(doc, second_image)
                print("✓ Added second table (130mm, aligned to LEFT EDGE)")
            else:
                # Якщо немає другого зображення, додаємо порожню таблицю
                create_empty_table_placeholder(doc)
                print("✓ Added empty table placeholder")
        
        total_pages = (len(processed_images) + 1) // 2
        print(f"\n✅ Final result: {len(processed_images)} images in {total_pages} pages with LEFT-ALIGNED tables")
        
    except Exception as e:
        print(f"✗ Error creating new structure pages: {e}")

def create_spacer_paragraph(doc, height_mm):
    """Створення параграфа-розділювача з точною висотою"""
    try:
        para = doc.add_paragraph()
        para.text = ""  # Порожній параграф
        
        # Встановлюємо точну висоту параграфа
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = Cm(mm_to_cm(height_mm))
        para.paragraph_format.line_spacing_rule = 0  # Точна висота
        
        print(f"✓ Created spacer paragraph: {height_mm}mm")
        
    except Exception as e:
        print(f"✗ Error creating spacer paragraph: {e}")

def create_single_image_table(doc, image_data):
    """Створення таблиці 1x3 для одного зображення"""
    try:
        print(f"🔨 Creating single image table...")
        
        # Створюємо таблицю 1x3 (1 рядок, 3 колонки)
        table = doc.add_table(rows=1, cols=ALBUM_LAYOUT['TABLE_COLS'])
        table.style = None
        
        # Встановлюємо точну ширину таблиці
        set_table_width(table)
        
        # Налаштовуємо колонки (ширини)
        setup_column_widths(table)
        
        # Налаштовуємо єдиний рядок з висотою 130мм
        setup_image_row(table, 0, image_data)
        print(f"✅ Single image table created successfully")
        
        return table
        
    except Exception as e:
        print(f"❌ Error creating single image table: {e}")
        return None

def create_empty_table_placeholder(doc):
    """Створення порожньої таблиці-заповнювача для випадків з непарною кількістю зображень"""
    try:
        table = doc.add_table(rows=1, cols=ALBUM_LAYOUT['TABLE_COLS'])
        table.style = None
        
        set_table_width(table)
        setup_column_widths(table)
        
        # Налаштовуємо порожній рядок
        row = table.rows[0]
        row.height = Cm(mm_to_cm(ALBUM_LAYOUT['TABLE_HEIGHT']))  # 130мм
        
        for col_idx in range(ALBUM_LAYOUT['TABLE_COLS']):
            cell = row.cells[col_idx]
            cell.text = ""
            set_cell_background(cell, "FFFFFF")
            # Прозорі границі для порожньої таблиці
            set_cell_borders(cell, top=False, bottom=False, left=False, right=False)
        
        print(f"✓ Created empty table placeholder")
        
    except Exception as e:
        print(f"✗ Error creating empty table placeholder: {e}")

def set_table_width(table):
    """Встановлення фіксованої ширини таблиці та ЛІВОГО вирівнювання"""
    try:
        tbl = table._tbl
        tblPr = tbl.tblPr
        
        # Фіксований layout
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)
        
        # ТОЧНА ширина таблиці: 30+145+30 = 205мм
        total_width_mm = ALBUM_LAYOUT['COL_1_WIDTH'] + ALBUM_LAYOUT['COL_2_WIDTH'] + ALBUM_LAYOUT['COL_3_WIDTH']
        width_dxa = int(total_width_mm * 56.7)  # 205мм в DXA
        
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:type'), 'dxa')
        tblW.set(qn('w:w'), str(width_dxa))
        tblPr.append(tblW)
        
        # ВАЖЛИВО: Вирівнювання по ЛІВОМУ краю (таблиця притиснута до лівого краю аркуша)
        tblJc = OxmlElement('w:jc')
        tblJc.set(qn('w:val'), 'left')  # Ліве вирівнювання
        tblPr.append(tblJc)
        
        # ДОДАТКОВО: Забезпечуємо що таблиця починається з самого лівого краю
        tblInd = OxmlElement('w:tblInd')
        tblInd.set(qn('w:w'), '0')  # Нульовий відступ зліва
        tblInd.set(qn('w:type'), 'dxa')
        tblPr.append(tblInd)
        
        print(f"✓ Table: {total_width_mm}mm width, LEFT-aligned to page edge (no left margin)")
        
    except Exception as e:
        print(f"✗ Error setting table width and alignment: {e}")
def setup_column_widths(table):
    """Налаштування ширин колонок для таблиці 1x3"""
    try:
        col_1_width = ALBUM_LAYOUT['COL_1_WIDTH']      # 30мм - Індикатор ЗРЛ
        col_2_width = ALBUM_LAYOUT['COL_2_WIDTH']      # 145мм - зображення  
        col_3_width = ALBUM_LAYOUT['COL_3_WIDTH']      # 30мм - дані
        
        # Встановлюємо ширини колонок
        for col_idx in range(ALBUM_LAYOUT['TABLE_COLS']):
            row = table.rows[0]  # Тільки один рядок
            cell = row.cells[col_idx]
            
            if col_idx == 0:
                set_cell_width_mm(cell, col_1_width)    # 30мм
            elif col_idx == 1:
                set_cell_width_mm(cell, col_2_width)    # 145мм
            elif col_idx == 2:
                set_cell_width_mm(cell, col_3_width)    # 30мм
        
        print(f"✓ Column widths set: 30mm, 145mm, 30mm")
        
    except Exception as e:
        print(f"✗ Error setting column widths: {e}")

def setup_image_row(table, row_idx, image_data):
    """Налаштування рядка з зображенням (висота 130мм)"""
    try:
        row = table.rows[row_idx]
        
        # Встановлюємо висоту рядка 130мм
        row.height = Cm(mm_to_cm(ALBUM_LAYOUT['TABLE_HEIGHT']))  # 130мм
        
        # Фіксована висота рядка
        trPr = row._tr.trPr
        if trPr is None:
            trPr = OxmlElement('w:trPr')
            row._tr.append(trPr)
        
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(int(130 * 56.7)))  # 130мм в DXA
        trHeight.set(qn('w:hRule'), 'exact')  # ТОЧНА висота
        trPr.append(trHeight)

        # Налаштовуємо комірки
        for col_idx in range(ALBUM_LAYOUT['TABLE_COLS']):
            cell = row.cells[col_idx]
            
            if col_idx == 0:
                setup_first_cell(cell, True)  # Показуємо границі
            elif col_idx == 1:
                setup_image_cell(cell, image_data, True)  # Показуємо границі
            elif col_idx == 2:
                setup_data_cell(cell, image_data, True)  # Показуємо границі
        
        print(f"✓ Image row {row_idx + 1} configured (130mm height)")
        
    except Exception as e:
        print(f"✗ Error setting up image row {row_idx}: {e}")

# ===== ДОПОМІЖНІ ФУНКЦІЇ =====

def setup_first_cell(cell, show_borders):
    """Налаштування першої комірки (ширина 30мм, текст 'Індикатор ЗРЛ')"""
    try:
        # Вертикальне центрування
        set_cell_vertical_center(cell)
        
        # Текст
        cell.text = ""

        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = para.add_run("Індикатор ЗРЛ")
        run.font.name = 'Arial'
        run.font.size = Pt(12)

        # ВНУТРІШНІ ПОЛЯ КОМІРКИ
        tc = cell._tc
        tcPr = tc.tcPr
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            tc.append(tcPr)
        
        tcMar = OxmlElement('w:tcMar')
        # Налаштування відступів для лівої комірки
        margins = {
            'top': '0',        # 0 зверху
            'left': '140',     # 140 DXA = ~2.5мм зліва
            'bottom': '0',     # 0 знизу  
            'right': '0'       # 0 справа
        }
        
        for margin_name, value in margins.items():
            margin = OxmlElement(f'w:{margin_name}')
            margin.set(qn('w:w'), value)
            margin.set(qn('w:type'), 'dxa')
            tcMar.append(margin)
        
        tcPr.append(tcMar)
        
        # Границі: лише права (якщо показуємо границі)
        if show_borders:
            set_cell_borders(cell, top=False, bottom=False, left=False, right=True)
        else:
            set_cell_borders(cell, top=False, bottom=False, left=False, right=False)
        
    except Exception as e:
        print(f"✗ Error configuring first cell: {e}")

def setup_image_cell(cell, image_data, show_borders):
    try:
        # Вертикальне центрування
        set_cell_vertical_center(cell)
        
        # Очищуємо комірку
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # НУЛЬОВІ ВІДСТУПИ У ПАРАГРАФА
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0) 
        para.paragraph_format.left_indent = Pt(0)
        para.paragraph_format.right_indent = Pt(0)
        para.paragraph_format.line_spacing = 1.0
        
        if image_data:
            # Створюємо оброблене зображення (ТЕПЕР З ОПИСОМ РЛС НА ЗОБРАЖЕННІ)
            processed_image = create_processed_image_from_data(image_data)
            
            if processed_image:
                # Зберігаємо у тимчасовий файл
                with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as temp_file:
                    temp_path = temp_file.name
                    processed_image.save(temp_path, 'JPEG', quality=95)
                
                # Розмір зображення (з відступом для границь)
                border_thickness_mm = 1.0
                effective_width = ALBUM_LAYOUT['COL_2_WIDTH'] - border_thickness_mm    # 149мм
                effective_height = ALBUM_LAYOUT['TABLE_HEIGHT'] - border_thickness_mm  # 129мм
                
                # Додаємо зображення
                run = para.add_run()
                inline_shape = run.add_picture(temp_path, 
                                             width=Cm(effective_width / 10.0),   # 14.9см
                                             height=Cm(effective_height / 10.0)) # 12.9см
                
                # Видаляємо тимчасовий файл
                try:
                    os.remove(temp_path)
                except:
                    pass
                
                print(f"✓ Image added: {effective_width}mm x {effective_height}mm with radar description")
        
        # Налаштування полів комірки (нульові для точного позиціонування)
        tc = cell._tc
        tcPr = tc.tcPr
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            tc.append(tcPr)
        
        tcMar = OxmlElement('w:tcMar')
        for margin_name in ['top', 'left', 'bottom', 'right']:
            margin = OxmlElement(f'w:{margin_name}')
            margin.set(qn('w:w'), '0')  # Нульові поля
            margin.set(qn('w:type'), 'dxa')
            tcMar.append(margin)
        tcPr.append(tcMar)
        
        # Границі
        if show_borders:
            set_cell_borders(cell, top=True, bottom=True, left=True, right=True)
        else:
            set_cell_borders(cell, top=False, bottom=False, left=False, right=False)
        
        # Білий фон
        set_cell_background(cell, "FFFFFF")
        
        print(f"✓ Image cell configured with radar description on image")
        
    except Exception as e:
        print(f"✗ Error configuring image cell: {e}")
        import traceback
        traceback.print_exc()

# ===== ДОДАТКОВІ УТИЛІТИ =====

def setup_data_cell(cell, image_data, show_borders):
    """Налаштування комірки з даними з лівим відступом"""
    try:
        # Очищуємо комірку
        cell.text = ""

        # ВСТАНОВЛЮЄМО ВНУТРІШНІ ПОЛЯ КОМІРКИ ЗАМІСТЬ ВІДСТУПІВ ПАРАГРАФІВ
        tc = cell._tc
        tcPr = tc.tcPr
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            tc.append(tcPr)
        
        tcMar = OxmlElement('w:tcMar')
        # Налаштування відступів для правої комірки
        margins = {
            'left': '140',     # 140 DXA = ~2.5мм (замість 8pt відступу параграфа)
            'bottom': '0',     # 0мм знизу
        }
        
        for margin_name, value in margins.items():
            margin = OxmlElement(f'w:{margin_name}')
            margin.set(qn('w:w'), value)
            margin.set(qn('w:type'), 'dxa')
            tcMar.append(margin)
        
        tcPr.append(tcMar)
        
        if image_data:
            # Перший параграф з відступами
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(30)  # 30pt зверху
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.15
            
            # НОВИЙ ЛІВИЙ ВІДСТУП
            para.paragraph_format.left_indent = Pt(0)  # 0pt = ~0мм лівий відступ
            para.paragraph_format.right_indent = Pt(0)  # Без правого відступу
            
            # Отримуємо дані
            target_data = image_data['target_data']
            analysis_point = image_data['analysis_point']
            target_no = target_data['number']
            
            # 🎯 ФОРМУЄМО СПИСОК З ВИСОТОЮ
            data_lines = [
                (f"{target_no}", 12, True, True),
                (f"β – {analysis_point['azimuth']:.0f}ᴼ", 12, True, False),
                (f"D – {analysis_point['range']:.0f} км", 12, True, False),
            ]
            
            # 🎯 ДОДАЄМО ВИСОТУ ТІЛЬКИ ЯКЩО НЕ ДОРІВНЮЄ "0.0"
            height_value = target_data.get('height', '0.0')
            try:
                height_float = float(height_value)
                if height_float != 0.0:
                    data_lines.append((f"Н – {height_value} м", 12, True, False))
            except (ValueError, TypeError):
                # Якщо висота не число, але не "0" або "0.0" - додаємо
                if height_value not in ['0', '0.0', '', None]:
                    data_lines.append((f"Н – {height_value} м", 12, True, False))
            
            # Продовжуємо з рештою даних
            data_lines.extend([
                ("без перешкод", 9, True, False),
                (f"{target_data['detection']}", 9, True, False),
                (f"М – {target_data['scale']}", 9, True, False)
            ])
            
            # Додаємо рядки
            for i, (text, font_size, italic, underline) in enumerate(data_lines):
                if i > 0:
                    para = cell.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    para.paragraph_format.line_spacing = 1.15
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    
                    # Лівий відступ для всіх параграфів
                    para.paragraph_format.left_indent = Pt(0)
                    para.paragraph_format.right_indent = Pt(0)
                
                run = para.add_run(text)
                run.font.name = 'Arial'
                run.font.size = Pt(font_size)
                run.italic = italic
                run.underline = underline
        
        # Границі: всі сторони крім лівої (якщо показуємо границі)
        if show_borders:
            set_cell_borders(cell, top=True, bottom=True, left=False, right=True)
        else:
            set_cell_borders(cell, top=False, bottom=False, left=False, right=False)
        
    except Exception as e:
        print(f"✗ Error configuring data cell: {e}")


def set_cell_width_mm(cell, width_mm):
    """Встановлення ФІКСОВАНОЇ ширини комірки БЕЗ внутрішніх відступів"""
    try:
        width_dxa = int(width_mm * 56.7)
        
        tc = cell._tc
        tcPr = tc.tcPr
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            tc.append(tcPr)
        
        # ФІКСОВАНА ширина
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:type'), 'dxa')
        tcW.set(qn('w:w'), str(width_dxa))
        tcPr.append(tcW)
        
        # ЗАБОРОНА автоматичного підгону
        tcFitText = OxmlElement('w:tcFitText')
        tcFitText.set(qn('w:val'), '0')
        tcPr.append(tcFitText)
        
        print(f"✓ Fixed cell width: {width_mm}mm (NO internal margins)")
        
    except Exception as e:
        print(f"✗ Error setting fixed cell width: {e}")

def set_cell_vertical_center(cell):
    """Вертикальне центрування комірки"""
    try:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    except:
        try:
            tc = cell._tc
            tcPr = tc.tcPr
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.append(tcPr)
            
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            tcPr.append(vAlign)
        except Exception as e:
            print(f"Warning: Could not set vertical alignment: {e}")

def set_cell_borders(cell, top=False, bottom=False, left=False, right=False):
    """Налаштування рамок комірки"""
    try:
        tc = cell._tc
        tcPr = tc.tcPr
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            tc.append(tcPr)
        
        tcBorders = OxmlElement('w:tcBorders')
        
        border_settings = [
            ('top', top), ('bottom', bottom), ('left', left), ('right', right)
        ]
        
        for position, show_border in border_settings:
            border_element = OxmlElement(f'w:{position}')
            if show_border:
                border_element.set(qn('w:val'), 'single')
                border_element.set(qn('w:sz'), '4')
                border_element.set(qn('w:color'), '000000')
            else:
                border_element.set(qn('w:val'), 'none')
            
            tcBorders.append(border_element)
        
        tcPr.append(tcBorders)
        
    except Exception as e:
        print(f"Warning: Could not set cell borders: {e}")

def set_cell_background(cell, color_hex):
    """Встановлення кольору фону комірки"""
    try:
        tc = cell._tc
        tcPr = tc.tcPr
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            tc.append(tcPr)
        
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color_hex)
        tcPr.append(shd)
        
    except Exception as e:
        print(f"Warning: Could not set cell background: {e}")

def create_processed_image_from_data(image_data):
    """Створення обробленого зображення з описом РЛС на зображенні"""
    try:
        with Image.open(image_data['image_path']) as original_image:
            if original_image.mode != 'RGB':
                if original_image.mode == 'RGBA':
                    rgb_image = Image.new('RGB', original_image.size, (255, 255, 255))
                    rgb_image.paste(original_image, mask=original_image.split()[-1])
                    final_image = rgb_image
                else:
                    final_image = original_image.convert('RGB')
            else:
                final_image = original_image.copy()
        
        draw = ImageDraw.Draw(final_image)
        analysis_point = image_data['analysis_point']
        
        # Розрахунок позиції кінця лінії
        image_width = final_image.width
        image_height = final_image.height
        
        # Позиція підкреслення номера цілі: 12% висоти зображення
        underline_y = int(image_height * 0.12)
        end_x = image_width - 1  # Самий край
        end_y = underline_y      # На висоті підкреслення
        
        # Малюємо лінію від точки аналізу до розрахованої позиції
        draw.line([
            (analysis_point['x'], analysis_point['y']), 
            (end_x, end_y)
        ], fill='black', width=3)
        
        # ===== ДОДАЄМО ОПИС РЛС НА ЗОБРАЖЕННЯ =====
        if 'radar_description' in image_data and image_data['radar_description']['enabled']:
            print("📝 Adding radar description to image")
            add_radar_description_to_image(draw, image_data['radar_description'], image_width, image_height)
        
        return final_image
        
    except Exception as e:
        print(f"Error creating processed image: {e}")
        return None

def add_radar_description_to_image(draw, radar_data, image_width, image_height):
    """
    Виправлена версія додавання опису РЛС з правильним розміщенням тексту
    """
    try:
        # ===== ТОЧНІ ПРОПОРЦІЙНІ РОЗРАХУНКИ =====
        
        # Табличка РЛС займає ці відсотки від зображення:
        RADAR_BOX_WIDTH_PERCENT = 28.60   # 4.29см / 15см * 100% = 28.60%
        RADAR_BOX_HEIGHT_PERCENT = 19.54  # 2.54см / 13см * 100% = 19.54%
        
        # Розраховуємо розміри таблички в пікселях
        rect_width = int((image_width * RADAR_BOX_WIDTH_PERCENT) / 100)
        rect_height = int((image_height * RADAR_BOX_HEIGHT_PERCENT) / 100)
        
        # ТОЧНІ внутрішні відступи згідно зі зразком Word
        PADDING_HORIZONTAL_PERCENT = 5.83  # 0,25см / 4,29см * 100% = 5.83%
        PADDING_VERTICAL_PERCENT = 5.12    # 0,13см / 2,54см * 100% = 5.12%
        
        padding_horizontal = int((rect_width * PADDING_HORIZONTAL_PERCENT) / 100)
        padding_vertical = int((rect_height * PADDING_VERTICAL_PERCENT) / 100)
        
        # Позиція таблички - лівий нижній кут з невеличким відступом від краю
        margin_from_edge = int(image_width * 0.01)  # 1% від ширини зображення
        rect_x = margin_from_edge
        rect_y = image_height - rect_height - margin_from_edge
        
        print(f"📏 Image: {image_width}x{image_height}px")
        print(f"📦 Radar box: {rect_width}x{rect_height}px ({RADAR_BOX_WIDTH_PERCENT:.1f}% x {RADAR_BOX_HEIGHT_PERCENT:.1f}%)")
        print(f"📍 Position: ({rect_x}, {rect_y})")
        print(f"📐 Padding: {padding_horizontal}x{padding_vertical}px")
        
        # ===== 1. МАЛЮЄМО ПРЯМОКУТНИК З ПРОЗОРИМ ФОНОМ =====
        border_width = max(2, int(rect_width * 0.008))
        
        draw.rectangle([rect_x, rect_y, rect_x + rect_width, rect_y + rect_height], 
                      fill=None, outline='black', width=border_width)
        
        # ===== 2. ГОТУЄМО ТЕКСТ =====
        text_lines = []
        
        # Дата з компактним форматуванням
        date_obj = radar_data.get('date')
        if date_obj:
            if hasattr(date_obj, 'toString'):
                # Використовуємо компактний формат без пробілів
                date_str = date_obj.toString('dd.MM.yyyy')
            elif hasattr(date_obj, 'strftime'):
                date_str = date_obj.strftime('%d.%m.%Y')
            else:
                date_str = str(date_obj)
            text_lines.append(date_str)
        
        # Позивний
        callsign = radar_data.get('callsign')
        if callsign and callsign.strip():
            text_lines.append(callsign.strip())
        
        # Назва РЛС (в дужках)
        name = radar_data.get('name')
        if name and name.strip():
            text_lines.append(f"({name.strip()})")
        
        # Номер РЛС (символ № + номер) - ВИПРАВЛЕНО
        number = radar_data.get('number')
        if number and number.strip():
            text_lines.append(f"№ {number.strip()}")  # Додано пробіл після №
        
        print(f"📝 Text lines: {text_lines}")
        
        if not text_lines:
            print("⚠️ No content for radar description")
            return
        
        # ===== 3. РОЗРАХУНОК РОЗМІРУ ШРИФТУ=====
        # Arial 12pt в зразку Word відповідає приблизно 17% від висоти таблички 2,54см
        FONT_SIZE_PERCENT = 17.0
        font_size = max(12, int((rect_height * FONT_SIZE_PERCENT) / 100))
        
        # ===== 4. ЗАВАНТАЖУЄМО ШРИФТ (КУРСИВ) =====
        font = get_reliable_font(font_size)
        
        # ===== 5. РОЗМІЩУЄМО ТЕКСТ З ТОЧНИМИ ПРОПОРЦІЯМИ =====
        text_x = rect_x + padding_horizontal
        text_y = rect_y + padding_vertical
        
        # Розраховуємо доступну висоту для тексту
        available_height = rect_height - (2 * padding_vertical)
        
        # ТОЧНИЙ міжрядковий інтервал
        line_height = int(font_size * 1.2)  # Одинарний відступ
        
        # Якщо всі рядки не поміщаються - автоматично зменшуємо інтервал
        total_text_height = len(text_lines) * line_height
        if total_text_height > available_height:
            line_height = available_height // len(text_lines)
        
        print(f"📝 Text area: {available_height}px height, {line_height}px line height")
        print(f"📏 Font size: {font_size}px (corresponds to Arial 12pt)")
        
        for i, line in enumerate(text_lines):
            current_y = text_y + (i * line_height)
            
            # Перевіряємо чи текст поміщається в прямокутник
            if current_y + font_size > rect_y + rect_height - padding_vertical:
                print(f"⚠️ Text line {i+1} exceeds rectangle bounds")
                break
            
            # СПРОЩЕНА обробка всіх рядків ОДНАКОВО
            try:
                draw.text((text_x, current_y), line, fill='black', font=font)
                print(f"✓ Line {i+1}: '{line}' at ({text_x}, {current_y})")
            except Exception as text_error:
                print(f"⚠️ Error drawing text line {i+1}: {text_error}")
        
        print("✅ Radar description added with exact Word proportions!")
        
        # ===== ПІДСУМКОВА ІНФОРМАЦІЯ ВІДПОВІДНОСТІ ЗРАЗКУ =====
        actual_width_percent = (rect_width / image_width) * 100
        actual_height_percent = (rect_height / image_height) * 100
        
        print(f"📊 Final proportions: {actual_width_percent:.1f}% x {actual_height_percent:.1f}%")
        print(f"📏 Margins: horizontal={padding_horizontal}px ({PADDING_HORIZONTAL_PERCENT:.1f}%), vertical={padding_vertical}px ({PADDING_VERTICAL_PERCENT:.1f}%)")
        print(f"🔤 Font: Arial {font_size}px italic (corresponds to 12pt in Word)")
        print(f"📐 Line height: {line_height}px (single spacing like Word)")
        
    except Exception as e:
        print(f"❌ Error adding radar description: {e}")
        import traceback
        traceback.print_exc()

def create_title_page(doc, title_data):
    """Створення титульної сторінки з окремими абзацами для кожного рядка заголовку"""
    try:
        print("=== Creating title page with separate paragraphs for each title line ===")
        
        # Отримуємо дані з шаблону
        date = title_data['date']
        unit_info = title_data.get('unit_info', 'А0000')
        commander_info = title_data.get('commander_info', DefaultTemplateData.get_commander_info())
        chief_of_staff_info = title_data.get('chief_of_staff_info', DefaultTemplateData.get_chief_of_staff_info())
        margins = title_data.get('margins', {
            'top': 2.25, 'left': 2.5, 'bottom': 0, 'right': 0.75
        })
        
        # Встановлюємо поля сторінки
        section = doc.sections[0]
        set_a4_page_format(section)
        
        section.left_margin = Cm(margins['left'])
        section.top_margin = Cm(margins['top'])
        section.right_margin = Cm(margins['right'])
        section.bottom_margin = Cm(margins['bottom'])
        
        print(f"✅ Margins set: left={margins['left']}cm, top={margins['top']}cm, right={margins['right']}cm, bottom={margins['bottom']}cm")
        
        # Створюємо стилі для титульної сторінки
        from docx.enum.style import WD_STYLE_TYPE
        
        # Стиль для верхніх абзаців (31pt + 1см відступ + 1,15 множна + 10pt після)
        try:
            top_spacer_style = doc.styles['TitleTopSpacer']
        except KeyError:
            top_spacer_style = doc.styles.add_style('TitleTopSpacer', WD_STYLE_TYPE.PARAGRAPH)
            top_spacer_style.paragraph_format.first_line_indent = Cm(1)
            top_spacer_style.paragraph_format.line_spacing = 1.15  # 1,15 як МНОЖНИК, не см!
            # НЕ встановлюємо line_spacing_rule - воно автоматично стає MULTIPLE
            top_spacer_style.paragraph_format.space_before = Pt(0)
            top_spacer_style.paragraph_format.space_after = Pt(10)     # 10pt після абзацу
            top_spacer_style.font.size = Pt(31)
            top_spacer_style.font.name = 'Arial'
        
        # Стиль для заголовку (28pt + центрування + 1см відступ першого рядка)
        try:
            title_main_style = doc.styles['TitleMain']
        except KeyError:
            title_main_style = doc.styles.add_style('TitleMain', WD_STYLE_TYPE.PARAGRAPH)
            title_main_style.paragraph_format.first_line_indent = Cm(1)
            title_main_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_main_style.paragraph_format.space_before = Pt(0)
            title_main_style.paragraph_format.space_after = Pt(0)
            title_main_style.font.size = Pt(28)
            title_main_style.font.name = 'Arial'
            title_main_style.font.bold = True
        
        # Стиль для середніх абзаців (28pt + 1см відступ + 1,15 множна + БЕЗ інтервалу після)
        try:
            middle_spacer_style = doc.styles['TitleMiddleSpacer']
        except KeyError:
            middle_spacer_style = doc.styles.add_style('TitleMiddleSpacer', WD_STYLE_TYPE.PARAGRAPH)
            middle_spacer_style.paragraph_format.first_line_indent = Cm(1)
            middle_spacer_style.paragraph_format.line_spacing = 1.15  # 1,15 як МНОЖНИК, не см!
            # НЕ встановлюємо line_spacing_rule - воно автоматично стає MULTIPLE
            middle_spacer_style.paragraph_format.space_before = Pt(0)
            middle_spacer_style.paragraph_format.space_after = Pt(0)     # БЕЗ інтервалу після
            middle_spacer_style.font.size = Pt(28)
            middle_spacer_style.font.name = 'Arial'
        
        # Стиль для підписів (16pt + БЕЗ лівого відступу + 1см відступ першого рядка)
        try:
            signature_style = doc.styles['TitleSignature']
        except KeyError:
            signature_style = doc.styles.add_style('TitleSignature', WD_STYLE_TYPE.PARAGRAPH)
            signature_style.paragraph_format.first_line_indent = Cm(1)      # 1см відступ першого рядка
            signature_style.paragraph_format.left_indent = Pt(0)            # БЕЗ лівого відступу абзацу
            signature_style.paragraph_format.line_spacing = 1.15            # 1,15 множна
            signature_style.paragraph_format.space_before = Pt(0)
            signature_style.paragraph_format.space_after = Pt(0)
            signature_style.font.size = Pt(16)
            signature_style.font.name = 'Arial'
            signature_style.font.bold = True
        
        # Стиль для розділювача між підписами (9pt + 1см відступ першого рядка + 1,15 множна)
        try:
            signature_spacer_style = doc.styles['TitleSignatureSpacer']
        except KeyError:
            signature_spacer_style = doc.styles.add_style('TitleSignatureSpacer', WD_STYLE_TYPE.PARAGRAPH)
            signature_spacer_style.paragraph_format.first_line_indent = Cm(1)   # 1см відступ першого рядка
            signature_spacer_style.paragraph_format.left_indent = Pt(0)         # БЕЗ лівого відступу абзацу
            signature_spacer_style.paragraph_format.line_spacing = 1.15         # 1,15 множна
            signature_spacer_style.paragraph_format.space_before = Pt(0)
            signature_spacer_style.paragraph_format.space_after = Pt(0)
            signature_spacer_style.font.size = Pt(9)                           # 9pt розмір
            signature_spacer_style.font.name = 'Arial'
        
        # Стиль для кінцевого абзацу (16pt + 1см відступ першого рядка)
        try:
            final_style = doc.styles['TitleFinal']
        except KeyError:
            final_style = doc.styles.add_style('TitleFinal', WD_STYLE_TYPE.PARAGRAPH)
            final_style.paragraph_format.first_line_indent = Cm(1)
            final_style.paragraph_format.line_spacing = Pt(16)
            final_style.paragraph_format.line_spacing_rule = 0
            final_style.paragraph_format.space_before = Pt(0)
            final_style.paragraph_format.space_after = Pt(0)
            final_style.font.size = Pt(16)
            final_style.font.name = 'Arial'
        
        print("✅ Title page styles created with correct spacing:")
        print("   • Top spacers: 1.15x line spacing (MULTIPLE) + 10pt after paragraph")
        print("   • Middle spacers: 1.15x line spacing (MULTIPLE) + 0pt after paragraph")
        print("   • Signatures: 1.15x line spacing, NO left indent, 1cm first line indent")
        print("   • Signature spacer: 9pt + 1.15x line spacing")
        
        # 4 абзаци зверху
        for i in range(4):
            doc.add_paragraph(" ", style=top_spacer_style)
        
        # ГОЛОВНИЙ ЗАГОЛОВОК: КОЖЕН РЯДОК = ОКРЕМИЙ АБЗАЦ
        doc.add_paragraph("АЛЬБОМ", style=title_main_style)
        doc.add_paragraph("фотознімків індикаторів РЛС", style=title_main_style)
        doc.add_paragraph(f"військової частини {unit_info}", style=title_main_style)
        doc.add_paragraph(format_ukrainian_date(date), style=title_main_style)
        
        print("✅ Signatures created as separate paragraphs:")
        print("   • Commander title paragraph")
        print("   • Commander name paragraph") 
        print("   • 9pt spacer paragraph")
        print("   • Chief title paragraph")
        print("   • Chief name paragraph")
        
        # 7 абзаців після заголовку
        for i in range(7):
            doc.add_paragraph(" ", style=middle_spacer_style)
        
        # ПІДПИСИ: КОЖЕН РЯДОК = ОКРЕМИЙ АБЗАЦ
        
        # 1. Перший абзац командира - посада
        commander_title_para = doc.add_paragraph(f"Командир військової частини {unit_info}", style=signature_style)
        
        # Налаштовуємо табуляцію для імені
        try:
            tab_stops = commander_title_para.paragraph_format.tab_stops
            tab_stops.clear_all()
            tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.RIGHT)
        except Exception as tab_error:
            print(f"Warning: Could not set tab stops: {tab_error}")
        
        # 2. Другий абзац командира - звання та ім'я
        commander_name_para = doc.add_paragraph(style=signature_style)
        commander_name_para.add_run(f"{commander_info['rank']}")
        commander_name_para.add_run(f"\t{commander_info['name']}")
        
        # Налаштовуємо табуляцію для імені
        try:
            tab_stops = commander_name_para.paragraph_format.tab_stops
            tab_stops.clear_all()
            tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.RIGHT)
        except:
            pass
        
        # 3. Розділювач між підписами (9pt абзац)
        spacer_para = doc.add_paragraph(" ", style=signature_spacer_style)
        
        # 4. Перший абзац начальника штабу - посада
        chief_title_para = doc.add_paragraph(f"Начальник штабу військової частини {unit_info}", style=signature_style)
        
        # Налаштовуємо табуляцію
        try:
            tab_stops = chief_title_para.paragraph_format.tab_stops
            tab_stops.clear_all()
            tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.RIGHT)
        except:
            pass
        
        # 5. Другий абзац начальника штабу - звання та ім'я
        chief_name_para = doc.add_paragraph(style=signature_style)
        chief_name_para.add_run(f"{chief_of_staff_info['rank']}")
        chief_name_para.add_run(f"\t{chief_of_staff_info['name']}")
        
        # Налаштовуємо табуляцію
        try:
            tab_stops = chief_name_para.paragraph_format.tab_stops
            tab_stops.clear_all()
            tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.RIGHT)
        except:
            pass
        
        # Кінцевий абзац
        doc.add_paragraph(" ", style=final_style)
        
        print("✅ Title page created successfully with separate paragraphs for title lines")
        
    except Exception as e:
        print(f"✗ Error creating title page: {e}")
        import traceback
        traceback.print_exc()

def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)


class TitlePageTemplateDialog(QDialog):
    """Діалогове вікно для створення/редагування шаблонів (ВИПРАВЛЕНА ВЕРСІЯ)"""
    
    def __init__(self, parent=None, mode='create', template_data=None):
        super().__init__(parent)
        self.mode = mode
        self.template_data = template_data or {}
        
        self.setWindowTitle("Створити новий шаблон" if mode == 'create' else "Редагувати шаблон")
        self.setModal(True)
        
        # ЗБІЛЬШЕНІ РОЗМІРИ ВІКНА
        self.setMinimumSize(520, 650)  # Ширша і вища
        self.setMaximumSize(520, 700)  # Максимум 700px висоти
        self.resize(520, 680)          # Початковий розмір
        
        # ОПТИМІЗОВАНІ СТИЛІ
        self.setStyleSheet("""
            QDialog {
                background-color: #f8f9fa;
                font-family: "Segoe UI", Arial, sans-serif;
            }
            QLabel {
                color: #343a40;
                font: 11pt "Segoe UI", Arial, sans-serif;
                background: none;
                border: none;
                padding: 2px 0px;
            }
            QLineEdit, QDoubleSpinBox {
                border: 1px solid #dee2e6;
                border-radius: 4px;
                padding: 8px 12px;
                background-color: white;
                font: 11pt "Segoe UI", Arial, sans-serif;
                color: #495057;
                min-height: 20px;
                margin: 2px 0px;
            }
            QLineEdit:hover, QDoubleSpinBox:hover {
                border: 1px solid #adb5bd;
                background-color: #f8f9fa;
            }
            QLineEdit:focus, QDoubleSpinBox:focus {
                border: 1px solid #495057;
                background-color: white;
            }
            QTabWidget {
                background-color: transparent;
            }
            QTabWidget::pane {
                border: 1px solid #dee2e6;
                border-radius: 6px;
                background-color: white;
                padding: 0px;
                margin-top: 30px;
            }
            QTabBar::tab {
                background-color: #e9ecef;
                color: #495057;
                border: 1px solid #dee2e6;
                border-bottom: none;
                border-top-left-radius: 6px;
                border-top-right-radius: 6px;
                padding: 8px 15px;
                margin-right: 2px;
                font: 500 10pt "Segoe UI", Arial, sans-serif;
                min-width: 130px;
                max-width: 150px;
                min-height: 35px;
            }
            QTabBar::tab:selected {
                background-color: white;
                color: #343a40;
                border: 1px solid #dee2e6;
                border-bottom: 1px solid white;
                font-weight: 600;
            }
            QTabBar::tab:hover:!selected {
                background-color: #f8f9fa;
                color: #495057;
            }
            QGroupBox {
                font-weight: 600;
                margin-top: 15px;
                padding: 15px;
                border: 1px solid #dee2e6;
                border-radius: 6px;
                background-color: white;
                color: #343a40;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 8px 0 8px;
                font: 600 11pt "Segoe UI", Arial, sans-serif;
                color: #495057;
                background-color: white;
            }
            QPushButton {
                background-color: #f8f9fa;
                border: 1px solid #dee2e6;
                border-radius: 6px;
                padding: 12px 18px;
                font: 500 11pt "Segoe UI", Arial, sans-serif;
                color: #495057;
                min-height: 16px;
            }
            QPushButton:hover {
                background-color: #e9ecef;
                border: 1px solid #adb5bd;
                color: #343a40;
            }
            QPushButton:pressed {
                background-color: #dee2e6;
                border: 1px solid #6c757d;
            }
            QPushButton[text="Зберегти"] {
                background-color: #495057;
                color: white;
                font-weight: 600;
                border: 1px solid #212529;
            }
            QPushButton[text="Зберегти"]:hover {
                background-color: #343a40;
            }
            QPushButton[text="Скасувати"] {
                background-color: #6c757d;
                color: white;
                border: 1px solid #495057;
            }
            QPushButton[text="Скасувати"]:hover {
                background-color: #495057;
                border: 1px solid #343a40;
            }
        """)
        
        # Основний layout
        layout = QVBoxLayout()
        layout.setContentsMargins(18, 15, 18, 15)
        layout.setSpacing(12)
        self.setLayout(layout)
        
        # Заголовок (КОМПАКТНІШИЙ)
        title_label = QLabel("Налаштування шаблону титульної сторінки")
        title_label.setFont(QFont("Segoe UI", 13, QFont.Bold))
        title_label.setAlignment(Qt.AlignCenter)
        title_label.setStyleSheet("""
            font-size: 13pt; 
            font-weight: bold; 
            margin-bottom: 8px; 
            padding: 12px; 
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #495057, stop:1 #343a40);
            color: white; 
            border-radius: 6px;
            margin: -3px -3px 8px -3px;
        """)
        layout.addWidget(title_label)
        
        # Табуляція для організації полів
        tab_widget = QTabWidget()
        tab_widget.tabBar().setExpanding(True)
        layout.addWidget(tab_widget)
        
        # Вкладки з оптимізованою висотою
        main_tab = self.create_main_tab()
        tab_widget.addTab(main_tab, "Основні\nдані")
        
        officials_tab = self.create_officials_tab()
        tab_widget.addTab(officials_tab, "Посадові\nособи")
        
        margins_tab = self.create_margins_tab()
        tab_widget.addTab(margins_tab, "Поля\nсторінки")
        
        # Кнопки
        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        
        button_box.button(QDialogButtonBox.Ok).setText("Зберегти")
        button_box.button(QDialogButtonBox.Cancel).setText("Скасувати")
        
        layout.addWidget(button_box)
        
        # Завантажуємо дані якщо редагуємо
        if self.mode == 'edit' and self.template_data:
            self.load_template_data()
    
    def create_main_tab(self):
        """Створення вкладки з основними даними (ОПТИМІЗОВАНА)"""
        widget = QWidget()
        layout = QFormLayout()
        layout.setSpacing(12)  # ЗБІЛЬШЕНИЙ простір між рядками
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setFieldGrowthPolicy(QFormLayout.ExpandingFieldsGrow)
        widget.setLayout(layout)
        
        # Назва шаблону
        self.name_edit = QLineEdit()
        self.name_edit.setPlaceholderText("Наприклад: Частина А1234")
        layout.addRow("Назва шаблону:", self.name_edit)
        
        # Назва військової частини
        self.unit_name_edit = QLineEdit()
        self.unit_name_edit.setPlaceholderText("А0000")
        layout.addRow("Військова частина:", self.unit_name_edit)
        
        return widget
    
    def create_officials_tab(self):
        """Створення вкладки з посадовими особами (ОПТИМІЗОВАНА)"""
        widget = QWidget()
        layout = QVBoxLayout()
        layout.setContentsMargins(15, 15, 15, 15)
        layout.setSpacing(15)  # ЗБІЛЬШЕНИЙ простір між групами
        widget.setLayout(layout)
        
        # Командир
        commander_group = QGroupBox("Командир")
        commander_layout = QFormLayout()
        commander_layout.setSpacing(10)
        commander_layout.setFieldGrowthPolicy(QFormLayout.ExpandingFieldsGrow)
        commander_group.setLayout(commander_layout)
        
        self.commander_rank_edit = QLineEdit()
        self.commander_rank_edit.setPlaceholderText("полковник")
        commander_layout.addRow("Звання:", self.commander_rank_edit)
        
        self.commander_name_edit = QLineEdit()
        self.commander_name_edit.setPlaceholderText("П.П. ПЕТРЕНКО")
        commander_layout.addRow("Ім'я:", self.commander_name_edit)
        
        layout.addWidget(commander_group)
        
        # Начальник штабу
        chief_group = QGroupBox("Начальник штабу")
        chief_layout = QFormLayout()
        chief_layout.setSpacing(10)
        chief_layout.setFieldGrowthPolicy(QFormLayout.ExpandingFieldsGrow)
        chief_group.setLayout(chief_layout)
        
        self.chief_rank_edit = QLineEdit()
        self.chief_rank_edit.setPlaceholderText("підполковник")
        chief_layout.addRow("Звання:", self.chief_rank_edit)
        
        self.chief_name_edit = QLineEdit()
        self.chief_name_edit.setPlaceholderText("С.С. СИДОРЕНКО")
        chief_layout.addRow("Ім'я:", self.chief_name_edit)
        
        layout.addWidget(chief_group)
        
        # Підпис
        signature_group = QGroupBox("Начальник групи секретного документального забезпечення")
        signature_group.setObjectName("signature_group")
        signature_layout = QFormLayout()
        signature_layout.setSpacing(10)
        signature_layout.setFieldGrowthPolicy(QFormLayout.ExpandingFieldsGrow)
        signature_group.setLayout(signature_layout)
        
        self.signature_rank_edit = QLineEdit()
        self.signature_rank_edit.setPlaceholderText("головний сержант")
        signature_layout.addRow("Звання:", self.signature_rank_edit)
        
        self.signature_name_edit = QLineEdit()
        self.signature_name_edit.setPlaceholderText("Анна Петренко")
        signature_layout.addRow("Ім'я:", self.signature_name_edit)
        
        layout.addWidget(signature_group)
        
        # Додаємо розтягування внизу
        layout.addStretch()
        
        return widget
    
    def create_margins_tab(self):
        """Створення вкладки з полями"""
        widget = QWidget()
        layout = QFormLayout()
        layout.setSpacing(15)
        layout.setContentsMargins(20, 25, 20, 25)
        layout.setFieldGrowthPolicy(QFormLayout.ExpandingFieldsGrow)
        widget.setLayout(layout)
        
        # Поля сторінки
        self.margin_top_spin = QDoubleSpinBox()
        self.margin_top_spin.setRange(0, 10)
        self.margin_top_spin.setSingleStep(0.25)
        self.margin_top_spin.setDecimals(2)
        self.margin_top_spin.setValue(2.25)
        self.margin_top_spin.setSuffix(" см")
        layout.addRow("Верхнє поле:", self.margin_top_spin)
        
        self.margin_left_spin = QDoubleSpinBox()
        self.margin_left_spin.setRange(0, 10)
        self.margin_left_spin.setSingleStep(0.25)
        self.margin_left_spin.setDecimals(2)
        self.margin_left_spin.setValue(2.5)
        self.margin_left_spin.setSuffix(" см")
        layout.addRow("Ліве поле:", self.margin_left_spin)
        
        self.margin_bottom_spin = QDoubleSpinBox()
        self.margin_bottom_spin.setRange(0, 10)
        self.margin_bottom_spin.setSingleStep(0.25)
        self.margin_bottom_spin.setDecimals(2)
        self.margin_bottom_spin.setValue(0)
        self.margin_bottom_spin.setSuffix(" см")
        layout.addRow("Нижнє поле:", self.margin_bottom_spin)
        
        self.margin_right_spin = QDoubleSpinBox()
        self.margin_right_spin.setRange(0, 10)
        self.margin_right_spin.setSingleStep(0.25)
        self.margin_right_spin.setDecimals(2)
        self.margin_right_spin.setValue(0.75)
        self.margin_right_spin.setSuffix(" см")
        layout.addRow("Праве поле:", self.margin_right_spin)
        
        return widget
    
    def load_template_data(self):
        """Завантаження даних шаблону у поля форми"""
        if not self.template_data:
            return
        
        # Основні дані
        self.name_edit.setText(self.template_data.get('name', ''))
        self.unit_name_edit.setText(self.template_data.get('unit_name', 'А0000'))
        
        # Командир
        commander = self.template_data.get('commander', {})
        self.commander_rank_edit.setText(commander.get('rank', 'полковник'))
        self.commander_name_edit.setText(commander.get('name', 'П.П. ПЕТРЕНКО'))
        
        # Начальник штабу
        chief = self.template_data.get('chief_of_staff', {})
        self.chief_rank_edit.setText(chief.get('rank', 'підполковник'))
        self.chief_name_edit.setText(chief.get('name', 'С.С. СИДОРЕНКО'))
        
        # Начальник групи секретного документального забезпечення
        signature = self.template_data.get('signature_info', {})
        self.signature_rank_edit.setText(signature.get('rank', 'головний сержант'))
        self.signature_name_edit.setText(signature.get('name', 'Анна Петренко'))
        
        # Поля
        margins = self.template_data.get('margins', {})
        self.margin_top_spin.setValue(margins.get('top', 2.25))
        self.margin_left_spin.setValue(margins.get('left', 2.5))
        self.margin_bottom_spin.setValue(margins.get('bottom', 0))
        self.margin_right_spin.setValue(margins.get('right', 0.75))

    def get_template_data(self):
        """Отримання даних з форми"""
        return {
            'name': self.name_edit.text().strip(),
            'unit_name': self.unit_name_edit.text().strip() or 'А0000',
            'commander': {
                'rank': self.commander_rank_edit.text().strip() or 'полковник',
                'name': self.commander_name_edit.text().strip() or 'П.П. ПЕТРЕНКО'
            },
            'chief_of_staff': {
                'rank': self.chief_rank_edit.text().strip() or 'підполковник',
                'name': self.chief_name_edit.text().strip() or 'С.С. СИДОРЕНКО'
            },
            'signature_info': {
                'rank': self.signature_rank_edit.text().strip() or 'головний сержант',
                'name': self.signature_name_edit.text().strip() or 'Анна Петренко'
            },
            'margins': {
                'top': self.margin_top_spin.value(),
                'left': self.margin_left_spin.value(),
                'bottom': self.margin_bottom_spin.value(),
                'right': self.margin_right_spin.value()
            },
            'created_date': datetime.now().isoformat(),
            'version': '1.1'
        }
        
    def accept(self):
        """Перевірка даних перед збереженням"""
        if not self.name_edit.text().strip():
            QMessageBox.warning(self, "Помилка", "Назва шаблону не може бути порожньою")
            return
        
        if not self.unit_name_edit.text().strip():
            QMessageBox.warning(self, "Помилка", "Назва військової частини не може бути порожньою")
            return
        
        super().accept()

# ===== ОСНОВНИЙ КЛАС GUI =====

class AzimuthGUI(QMainWindow):
    def __init__(self):
        super().__init__()
        self.processor = None
        self.current_image_path = None
        self.current_click = None
        self.scale_factor_x = 1.0
        self.scale_factor_y = 1.0
        self.offset_x = 0
        self.offset_y = 0
        self.current_folder = None
        
        self.current_language = 'UKRAINIAN'  # Default language
        self.translations = Translations()
        
        self.scale_edge_mode = False
        self.scale_edge_point = None
        self.custom_scale_distance = None
        self.center_setting_mode = False
        
        self.current_target_number = "0001"
        self.current_height = "0.0"
        self.current_obstacles = "без перешкод"
        self.current_detection = "Виявлення"

        self.radar_description_enabled = False
        self.radar_date = QDate.currentDate()
        self.radar_callsign = ""
        self.radar_name = ""
        self.radar_number = ""

        self.document_date = QDate.currentDate()
        
        self.processed_images = []

        # Ініціалізація документації
        self.doc_manager = DocumentationManager()
        self.doc_manager.create_documentation_files()  # Створити файли при запуску

        self.processed_image_paths = set()  # Множина шляхів оброблених зображень
        
        # Змінні для збереження налаштувань сітки між зображеннями
        self.saved_grid_settings = {
            'center_offset_x': 0,
            'center_offset_y': 0,
            'scale_edge_relative': None,
            'custom_scale_distance': None,
            'scale_value': "300"
        }
        self.setStyleSheet("""
            QDateEdit {
                border: 1px solid #dee2e6;
                border-radius: 4px;
                padding: 6px 10px;
                background-color: white;
                font: 12pt "Segoe UI", Arial, sans-serif;
                color: #495057;
                min-height: 22px;
            }
            QDateEdit:hover {
                border: 1px solid #adb5bd;
                background-color: #f8f9fa;
            }
            QDateEdit:focus {
                border: 1px solid #6c757d;
            }
            QDateEdit:disabled {
                background-color: #f5f5f5;
                color: #6c757d;
                border: 1px solid #e9ecef;
            }

            /* МІНІМАЛІСТИЧНИЙ ВИПАДНИЙ СПИСОК */
            QDateEdit::drop-down {
                border: none;
                background-color: transparent;
                width: 18px;
                margin-right: 4px;
            }
            QDateEdit::drop-down:hover {
                background-color: #f8f9fa;
                border-radius: 3px;
            }

            /* ПРОСТА СТРІЛКА */
            QDateEdit::down-arrow {
                image: none;
                border-left: 4px solid transparent;
                border-right: 4px solid transparent;
                border-top: 5px solid #6c757d;
                margin-top: 1px;
            }
            QDateEdit::down-arrow:hover {
                border-top-color: #495057;
            }
            QCalendarWidget {
                background-color: white;
                color: #495057;
                border: 1px solid #dee2e6;
                border-radius: 6px;
                font-family: "Segoe UI", Arial, sans-serif;
            }
            QCalendarWidget QWidget#qt_calendar_navigationbar {
                background-color: #f8f9fa;
                border-bottom: 1px solid #e9ecef;
                border-top-left-radius: 6px;
                border-top-right-radius: 6px;
                padding: 4px;
            }
            QCalendarWidget QToolButton {
                color: #6c757d;
                background-color: transparent;
                border: 1px solid transparent;
                border-radius: 4px;
                margin: 1px;
                padding: 4px 8px;
                font-weight: normal;
                min-width: 24px;
                min-height: 24px;
            }
            QCalendarWidget QToolButton:hover {
                background-color: #e9ecef;
                color: #495057;
                border: 1px solid #adb5bd;
            }
            QCalendarWidget QToolButton#qt_calendar_prevmonth {
                qproperty-text: "‹";
                font-size: 16pt;
                font-weight: bold;
            }
            QCalendarWidget QToolButton#qt_calendar_nextmonth {
                qproperty-text: "›";
                font-size: 16pt;
                font-weight: bold;
            }
            QCalendarWidget QSpinBox {
                color: #495057;
                background-color: white;
                border: 1px solid #dee2e6;
                border-radius: 3px;
                font-size: 12pt;
                font-weight: normal;
                padding: 2px 6px;
                min-width: 60px;
                margin: 1px;
            }
            QCalendarWidget QSpinBox:hover {
                border: 1px solid #adb5bd;
                background-color: #f8f9fa;
            }
            QCalendarWidget QHeaderView::section {
                color: #6c757d;
                background-color: #f8f9fa;
                border: none;
                border-bottom: 1px solid #e9ecef;
                font-weight: 500;
                padding: 6px 2px;
                font-size: 11pt;
            }
            QCalendarWidget QAbstractItemView {
                background-color: white;
                color: #495057;
                font-size: 12pt;
                border: none;
                selection-background-color: #495057;
                selection-color: white;
            }
            QCalendarWidget QAbstractItemView:item:selected {
                background-color: #495057;
                color: white;
                border-radius: 3px;
                font-weight: normal;
            }
            QCalendarWidget QAbstractItemView:item:focus {
                background-color: #e9ecef;
                color: #495057;
                border: 1px solid #adb5bd;
                border-radius: 3px;
            }""")
        # Встановлюємо іконку вікна
        icon_path = resource_path('netaz.ico')
        if os.path.exists(icon_path):
            self.setWindowIcon(QIcon(icon_path))
        
        self.init_ui()

        # ===== МЕТОДИ ДЛЯ ОБРОБКИ ОПИСУ РЛС =====
    def keyPressEvent(self, event):
        """Обробка натискань клавіш для переміщення центру/краю масштабу"""
        # ПЕРЕВІРЯЄМО ЧИ АКТИВНИЙ СПЕЦІАЛЬНИЙ РЕЖИМ
        if not (self.center_setting_mode or self.scale_edge_mode):
            super().keyPressEvent(event)
            return
        
        if not self.processor:
            super().keyPressEvent(event)
            return
        
        # ЗАБИРАЄМО ФОКУС З ПОЛІВ ВВЕДЕННЯ, АЛЕ ПЕРЕДАЄМО ЗОБРАЖЕННЮ
        focused_widget = self.focusWidget()
        if focused_widget and isinstance(focused_widget, (QLineEdit, QComboBox, QDateEdit)):
            # Передаємо фокус зображенню для збереження функціональності миші
            if hasattr(self, 'image_label'):
                self.image_label.setFocus()
            else:
                self.setFocus()
            print(f"🎯 Focus moved from {type(focused_widget).__name__} to image")
        
        # Визначаємо крок переміщення
        step = 1  # Базовий крок
        if event.modifiers() & Qt.ShiftModifier:
            step = 5  # Великий крок з Shift
        elif event.modifiers() & Qt.ControlModifier:
            step = 0.5  # Малий крок з Ctrl
        
        # Обробляємо стрілки
        dx, dy = 0, 0
        if event.key() == Qt.Key_Left:
            dx = -step
        elif event.key() == Qt.Key_Right:
            dx = step
        elif event.key() == Qt.Key_Up:
            dy = -step
        elif event.key() == Qt.Key_Down:
            dy = step
        elif event.key() == Qt.Key_Escape:
            # Escape для виходу з режиму
            if self.center_setting_mode:
                self.toggle_center_setting_mode()
            elif self.scale_edge_mode:
                self.toggle_scale_edge_mode()
            return
        else:
            super().keyPressEvent(event)
            return
        
        # Виконуємо переміщення
        if dx != 0 or dy != 0:
            if self.center_setting_mode:
                self.move_center_with_keyboard(dx, dy)
            elif self.scale_edge_mode:
                self.move_scale_edge_with_keyboard(dx, dy)
        
        event.accept()

    def move_center_with_keyboard(self, dx, dy):
        """Переміщення центру з клавіатури з зумом"""
        if not self.processor:
            return
        
        self.processor.move_center(dx, dy)
        self.save_current_grid_settings()
        
        # Оновлюємо точку аналізу якщо є
        if self.current_click:
            azimuth, range_val = self.calculate_azimuth_range(
                self.current_click['x'], self.current_click['y']
            )
            self.current_click['azimuth'] = azimuth
            self.current_click['range'] = range_val
        
        self.display_image()
        self.update_results_display()
        self.update_report_data()
        
        #  ПОКАЗУЄМО ЗУМ НА НОВІЙ ПОЗИЦІЇ ЦЕНТРУ
        if hasattr(self, 'image_label'):
            self.image_label.zoom_widget.show_zoom()
            self.image_label.zoom_widget.update_cursor_position(
                self.processor.center_x, self.processor.center_y
            )
    
        self.add_result(f"Центр: ({self.processor.center_x}, {self.processor.center_y}) | ←→↑↓")

    def move_scale_edge_with_keyboard(self, dx, dy):
        """Переміщення краю масштабу з клавіатури з зумом"""
        if not self.processor:
            return
        
        # Ініціалізуємо scale edge point якщо не існує
        if not self.scale_edge_point:
            self.scale_edge_point = {
                'x': self.processor.center_x + 50,
                'y': self.processor.center_y + 50
            }
        # Переміщуємо точку
        new_x = max(0, min(self.scale_edge_point['x'] + dx, self.processor.image.width - 1))
        new_y = max(0, min(self.scale_edge_point['y'] + dy, self.processor.image.height - 1))
        
        self.scale_edge_point['x'] = new_x
        self.scale_edge_point['y'] = new_y
        # Перераховуємо відстань
        dx_scale = new_x - self.processor.center_x
        dy_scale = new_y - self.processor.center_y
        self.custom_scale_distance = math.sqrt(dx_scale*dx_scale + dy_scale*dy_scale)
        self.save_current_grid_settings()
        
        # Оновлюємо точку аналізу якщо є
        if self.current_click:
            azimuth, range_val = self.calculate_azimuth_range(
                self.current_click['x'], self.current_click['y']
            )
            self.current_click['azimuth'] = azimuth
            self.current_click['range'] = range_val
        
        self.display_image()
        self.update_results_display()
        self.update_report_data()
        
        # ПОКАЗУЄМО ЗУМ НА НОВІЙ ПОЗИЦІЇ КРАЮ МАСШТАБУ
        if hasattr(self, 'image_label'):
            self.image_label.zoom_widget.show_zoom()
            self.image_label.zoom_widget.update_cursor_position(new_x, new_y)
    
        self.add_result(f"Край масштабу: ({new_x}, {new_y}) | Відстань: {self.custom_scale_distance:.1f}px")

    def toggle_radar_description(self, checked):
        """Перемикання режиму опису РЛС"""
        self.radar_description_enabled = checked
        
        # Активуємо/деактивуємо поля введення
        self.radar_date_edit.setEnabled(checked)
        self.radar_callsign_input.setEnabled(checked)
        self.radar_name_input.setEnabled(checked)
        self.radar_number_input.setEnabled(checked)
        
        if checked:
            self.add_result("✓ Режим опису РЛС активовано")
            self.add_result("  Введені дані будуть застосовані до всіх зображень")
        else:
            self.add_result("✗ Режим опису РЛС деактивовано")
            self.add_result("  Дані РЛС не застосовуватимуться")

    def update_radar_date(self, date):
        """Оновлення дати РЛС"""
        self.radar_date = date
        if self.radar_description_enabled:
            self.add_result(f"📅 Дата РЛС оновлена: {date.toString('dd.MM.yyyy')}")

    def update_radar_callsign(self, text):
        """Оновлення позивного РЛС"""
        self.radar_callsign = text
        if self.radar_description_enabled and text:
            self.add_result(f"📡 Позивний РЛС: {text}")

    def update_radar_name(self, text):
        """Оновлення назви РЛС"""
        self.radar_name = text
        if self.radar_description_enabled and text:
            self.add_result(f"📋 Назва РЛС: {text}")

    def update_radar_number(self, text):
        """Оновлення номера РЛС"""
        self.radar_number = text
        if self.radar_description_enabled and text:
            self.add_result(f"🔢 Номер РЛС: {text}")

    def get_radar_description_data(self):
        """Отримання даних опису РЛС"""
        if not self.radar_description_enabled:
            return None
        
        return {
            'enabled': True,
            'date': self.radar_date,
            'callsign': self.radar_callsign,
            'name': self.radar_name,
            'number': self.radar_number
        }
    
    def tr(self, key):
        return self.translations.get(self.current_language, key)
    
    def set_language(self, language):
        self.current_language = language
        self.update_interface_text()
        
        for lang, action in self.language_actions.items():
            action.setChecked(lang == language)
    
    def update_interface_text(self):
        self.setWindowTitle(self.tr("window_title"))
        self.create_menu_bar()
        
        if hasattr(self, 'control_title'):
            self.control_title.setText(self.tr("controls"))
        if hasattr(self, 'report_title'):
            self.report_title.setText(self.tr("report_data"))
        if hasattr(self, 'browser_label'):
            self.browser_label.setText(self.tr("photo_browser"))
        
        if hasattr(self, 'open_image_btn'):
            self.open_image_btn.setText(self.tr("open_image"))
        if hasattr(self, 'open_folder_btn'):
            self.open_folder_btn.setText(self.tr("open_folder"))
        if hasattr(self, 'save_image_btn'):
            self.save_image_btn.setText(self.tr("save_current_image"))
        if hasattr(self, 'export_new_btn'):
            self.export_new_btn.setText(self.tr("create_new_album"))
        if hasattr(self, 'export_add_btn'):
            self.export_add_btn.setText(self.tr("add_to_existing_album"))
        if hasattr(self, 'scale_edge_btn'):
            self.scale_edge_btn.setText(self.tr("set_scale_edge"))
        if hasattr(self, 'set_center_btn'):
            self.set_center_btn.setText(self.tr("set_center"))
        
        if hasattr(self, 'file_ops_label'):
            self.file_ops_label.setText(self.tr("file_operations"))
        if hasattr(self, 'azimuth_grid_label'):
            self.azimuth_grid_label.setText(self.tr("azimuth_grid"))
        if hasattr(self, 'move_center_label'):
            self.move_center_label.setText(self.tr("move_center"))
        if hasattr(self, 'results_label'):
            self.results_label.setText(self.tr("results"))
        
        # Batch processing переклади
        if hasattr(self, 'batch_label'):
            self.batch_label.setText(self.tr("batch_processing"))
        if hasattr(self, 'save_current_btn'):
            self.save_current_btn.setText(self.tr("save_current_image_data"))
        
        if not self.current_image_path:
            self.image_label.setText(self.tr("open_instruction"))
        
        if self.processor:
            self.update_results_display()

    def init_ui(self):
        self.setWindowTitle(self.tr("window_title"))
        
        default_width = 1400
        default_height = 900
        min_width = 1000  
        min_height = 700
        self.setMinimumSize(min_width, min_height)
        self.setWindowFlags(Qt.Window)
        self.resize(default_width, default_height)
        self.showMaximized()
        
        self.create_menu_bar()
        
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        main_layout = QHBoxLayout()
        central_widget.setLayout(main_layout)
        
        main_splitter = QSplitter(Qt.Horizontal)
        main_layout.addWidget(main_splitter)
        
        self.create_control_panel(main_splitter)
        self.create_vertical_browser_panel(main_splitter)
        self.create_image_panel(main_splitter)
        self.create_report_panel(main_splitter)
        
        main_splitter.setStretchFactor(0, 0)
        main_splitter.setStretchFactor(1, 0)
        main_splitter.setStretchFactor(2, 1)
        main_splitter.setStretchFactor(3, 0)
        
        main_splitter.setSizes([220, 0, 960, 220])  # 220+960+220 = 1400px
        self.main_splitter = main_splitter

    def resizeEvent(self, event):
        """Оновлений метод зміни розміру з правильними розмірами браузера"""
        super().resizeEvent(event)
        
        if hasattr(self, 'processor') and self.processor:
            QTimer.singleShot(50, self.update_image_display_after_resize)
        
        if hasattr(self, 'main_splitter') and hasattr(self, 'browser_widget'):
            current_sizes = self.main_splitter.sizes()
            total_width = sum(current_sizes)
            
            if total_width > 800:
                if self.browser_widget.isVisible():
                    # ПРАВИЛЬНІ розміри: 220 + 280 + ? + 220 = total_width
                    new_image_width = max(450, total_width - 220 - 280 - 220)  # Мінімум 450px
                    self.main_splitter.setSizes([220, 280, new_image_width, 220])
                else:
                    # 220 + 0 + ? + 220 = total_width  
                    new_image_width = max(450, total_width - 220 - 220)  # Мінімум 450px
                    self.main_splitter.setSizes([220, 0, new_image_width, 220])
    
    def update_image_display_after_resize(self):
        if hasattr(self, 'processor') and self.processor:
            self.display_image()
    
    def create_menu_bar(self):
        menubar = self.menuBar()
        menubar.clear()
        
        settings_menu = menubar.addMenu(self.tr("settings"))
        language_menu = settings_menu.addMenu(self.tr("language"))
        
        english_action = language_menu.addAction("English")
        english_action.triggered.connect(lambda: self.set_language('ENGLISH'))
        
        ukrainian_action = language_menu.addAction("Українська")
        ukrainian_action.triggered.connect(lambda: self.set_language('UKRAINIAN'))
        
        english_action.setCheckable(True)
        ukrainian_action.setCheckable(True)
        
        english_action.setChecked(self.current_language == 'ENGLISH')
        ukrainian_action.setChecked(self.current_language == 'UKRAINIAN')
        
        self.language_actions = {
            'ENGLISH': english_action,
            'UKRAINIAN': ukrainian_action
        }

        settings_menu.addSeparator()
        
        help_action = settings_menu.addAction("Документація")
        help_action.triggered.connect(self.show_documentation)
        
        about_action = settings_menu.addAction("Про програму")
        about_action.triggered.connect(self.show_about)

    def show_documentation(self):
        '''Відкриття HTML документації'''
        try:
            success = self.doc_manager.show_documentation()
            if not success:
                QMessageBox.warning(self, "Помилка", 
                                "Не вдалося відкрити документацію.\\n"
                                "Перевірте налаштування браузера.")
        except Exception as e:
            QMessageBox.critical(self, "Помилка", f"Помилка відкриття документації: {str(e)}")

    def show_about(self):
        '''Показати діалог "Про програму"'''
        dialog = AboutDialog(self)
        dialog.exec_()

    def create_control_panel(self, parent):
        """Ліва панель з основними кнопками"""
        control_widget = QWidget()
        control_widget.setFixedWidth(250)
        control_widget.setStyleSheet("""
            QWidget {
                background-color: #f5f5f5; 
                border-right: 1px solid #ccc;
            }
            QLabel {
                background: none;
                border: none;
                color: #343a40;
            }
            QPushButton {
                background-color: #f8f9fa;
                border: 1px solid #dee2e6;
                border-radius: 6px;
                padding: 10px 14px;
                font: 500 12pt "Segoe UI", Arial, sans-serif;
                color: #495057;
            }
            QPushButton:hover {
                background-color: #e9ecef;
                border: 1px solid #adb5bd;
                color: #343a40;
            }
            QPushButton:pressed {
                background-color: #dee2e6;
                border: 1px solid #6c757d;
            }
            QPushButton#main_action {
                background-color: #495057;
                color: white;
                font-weight: 600;
                border: 1px solid #212529;
            }
            QPushButton#main_action:hover {
                background-color: #343a40;
            }
            QPushButton#main_action:pressed {
                background-color: #212529;
            }
            QComboBox, QDateEdit {
                border: 1px solid #dee2e6;
                border-radius: 4px;
                padding: 6px 10px;
                background-color: white;
                font: 12pt "Segoe UI", Arial, sans-serif;
                color: #495057;
                min-height: 22px;
            }
            QComboBox:hover, QDateEdit:hover {
                border: 1px solid #adb5bd;
                background-color: #f8f9fa;
            }
            QComboBox:focus, QDateEdit:focus {
                border: 1px solid #6c757d;
                background-color: white;
            }
        """)
        
        layout = QVBoxLayout()
        layout.setContentsMargins(15, 15, 15, 15)
        layout.setSpacing(12)
        control_widget.setLayout(layout)
        
        # Title з новим шрифтом
        self.control_title = QLabel(self.tr("controls"))
        self.control_title.setFont(QFont("Segoe UI", 16, QFont.Bold))
        self.control_title.setAlignment(Qt.AlignCenter)
        self.control_title.setStyleSheet("font-size: 16pt; font-weight: bold; margin-bottom: 12px; color: #343a40;")
        layout.addWidget(self.control_title)
        
        # File operations section
        self.file_ops_label = QLabel(self.tr("file_operations"))
        self.file_ops_label.setFont(QFont("Segoe UI", 12, QFont.Bold))
        self.file_ops_label.setStyleSheet("color: #6c757d; margin-top: 12px; margin-bottom: 8px; font-weight: bold;")
        layout.addWidget(self.file_ops_label)
        
        # File operation buttons (стандартні)
        self.open_image_btn = QPushButton(self.tr("open_image"))
        self.open_image_btn.clicked.connect(self.open_image)
        layout.addWidget(self.open_image_btn)
        
        self.open_folder_btn = QPushButton(self.tr("open_folder"))
        self.open_folder_btn.clicked.connect(self.open_folder)
        layout.addWidget(self.open_folder_btn)
        
        self.save_image_btn = QPushButton(self.tr("save_current_image"))
        self.save_image_btn.clicked.connect(self.save_current_image)
        layout.addWidget(self.save_image_btn)
        
        # Separator
        separator1 = QFrame()
        separator1.setFrameShape(QFrame.HLine)
        separator1.setFrameShadow(QFrame.Sunken)
        separator1.setStyleSheet("color: #dee2e6; margin: 12px 0px;")
        layout.addWidget(separator1)

        # ===== РОЗДІЛ: ЗАПОВНЕННЯ ДОКУМЕНТУ =====
        self.title_page_label = QLabel("Заповнення документу")
        self.title_page_label.setFont(QFont("Segoe UI", 12, QFont.Bold))
        self.title_page_label.setStyleSheet("color: #6c757d; margin-top: 8px; margin-bottom: 8px; font-weight: bold;")
        layout.addWidget(self.title_page_label)

        # 1.ВИБІР ШАБЛОНУ
        template_container = QWidget()
        template_layout = QHBoxLayout()
        template_layout.setContentsMargins(0, 0, 0, 0)
        template_layout.setSpacing(8)
        template_container.setLayout(template_layout)

        template_label = QLabel("Шаблон:")
        template_label.setFont(QFont("Segoe UI", 12))
        template_label.setStyleSheet("color: #495057;")
        template_layout.addWidget(template_label)

        self.template_combo = QComboBox()
        self.template_combo.setStyleSheet("""
            QComboBox {
                border: 1px solid #dee2e6;
                border-radius: 4px;
                padding: 6px 10px;
                background-color: white;
                font: 12pt "Segoe UI", Arial, sans-serif;
                color: #495057;
                min-height: 20px;
            }
            QComboBox:hover {
                border: 1px solid #adb5bd;
            }
            QComboBox::drop-down {
                border: none;
                width: 20px;
            }
            QComboBox::down-arrow {
                image: none;
                border-left: 5px solid transparent;
                border-right: 5px solid transparent;
                border-top: 5px solid #6c757d;
            }
        """)
        self.template_combo.currentTextChanged.connect(self.on_template_changed)
        template_layout.addWidget(self.template_combo)

        layout.addWidget(template_container)

        # 2. ДАТА
        date_container = QWidget()
        date_layout = QHBoxLayout()
        date_layout.setContentsMargins(0, 0, 0, 0)
        date_layout.setSpacing(8)
        date_container.setLayout(date_layout)

        # Віджет вибору дати документу
        self.document_date_edit = QDateEdit()
        self.document_date_edit.setDate(self.document_date)
        self.document_date_edit.setCalendarPopup(True)
        self.document_date_edit.setDisplayFormat("dd.MM.yyyy")
        self.document_date_edit.setFixedHeight(32)
        self.document_date_edit.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)

        self.document_date_edit.setStyleSheet("""
            QDateEdit {
                border: 1px solid #dee2e6;
                border-radius: 4px;
                padding: 6px 10px;
                background-color: white;
                font: 12pt "Segoe UI", Arial, sans-serif;
                color: #495057;
                min-height: 22px;
            }
            QDateEdit:hover {
                border: 1px solid #adb5bd;
                background-color: #f8f9fa;
            }
            QDateEdit:focus {
                border: 1px solid #6c757d;
            }
            QDateEdit:disabled {
                background-color: #f5f5f5;
                color: #6c757d;
                border: 1px solid #e9ecef;
            }
            QDateEdit::drop-down {
                border: none;
                background-color: transparent;
                width: 18px;
                margin-right: 4px;
            }
            QDateEdit::drop-down:hover {
                background-color: #f8f9fa;
                border-radius: 3px;
            }
            QDateEdit::down-arrow {
                image: none;
                border-left: 4px solid transparent;
                border-right: 4px solid transparent;
                border-top: 5px solid #6c757d;
                margin-top: 1px;
            }
            QDateEdit::down-arrow:hover {
                border-top-color: #495057;
            }
            QCalendarWidget {
                background-color: white;
                color: #495057;
                border: 1px solid #dee2e6;
                border-radius: 6px;
                font-family: "Segoe UI", Arial, sans-serif;
            }
            QCalendarWidget QWidget#qt_calendar_navigationbar {
                background-color: #f8f9fa;
                border-bottom: 1px solid #e9ecef;
                border-top-left-radius: 6px;
                border-top-right-radius: 6px;
                padding: 4px;
            }
            QCalendarWidget QToolButton {
                color: #6c757d;
                background-color: transparent;
                border: 1px solid transparent;
                border-radius: 4px;
                margin: 1px;
                padding: 4px 8px;
                font-weight: normal;
                min-width: 24px;
                min-height: 24px;
            }
            QCalendarWidget QToolButton:hover {
                background-color: #e9ecef;
                color: #495057;
                border: 1px solid #adb5bd;
            }
            QCalendarWidget QToolButton#qt_calendar_prevmonth {
                qproperty-text: "‹";
                font-size: 16pt;
                font-weight: bold;
            }
            QCalendarWidget QToolButton#qt_calendar_nextmonth {
                qproperty-text: "›";
                font-size: 16pt;
                font-weight: bold;
            }
            QCalendarWidget QSpinBox {
                color: #495057;
                background-color: white;
                border: 1px solid #dee2e6;
                border-radius: 3px;
                font-size: 12pt;
                font-weight: normal;
                padding: 2px 6px;
                min-width: 60px;
                margin: 1px;
            }
            QCalendarWidget QSpinBox:hover {
                border: 1px solid #adb5bd;
                background-color: #f8f9fa;
            }
            QCalendarWidget QHeaderView::section {
                color: #6c757d;
                background-color: #f8f9fa;
                border: none;
                border-bottom: 1px solid #e9ecef;
                font-weight: 500;
                padding: 6px 2px;
                font-size: 11pt;
            }
            QCalendarWidget QAbstractItemView {
                background-color: white;
                color: #495057;
                font-size: 12pt;
                border: none;
                selection-background-color: #495057;
                selection-color: white;
            }
            QCalendarWidget QAbstractItemView:item:selected {
                background-color: #495057;
                color: white;
                border-radius: 3px;
                font-weight: normal;
            }
            QCalendarWidget QAbstractItemView:item:focus {
                background-color: #e9ecef;
                color: #495057;
                border: 1px solid #adb5bd;
                border-radius: 3px;
            }""")
        self.document_date_edit.dateChanged.connect(self.update_document_date)
        date_layout.addWidget(self.document_date_edit)

        self.today_btn = QPushButton("Сьогодні")
        self.today_btn.clicked.connect(self.set_document_date_today)
        self.today_btn.setFixedHeight(32)  # Та ж висота що й дата
        self.today_btn.setFixedWidth(85)   # ТРОХИ ШИРШЕ ДЛЯ КОМФОРТУ
        self.today_btn.setStyleSheet("""
            QPushButton {
                background-color: #6c757d;
                color: white;
                border: 1px solid #495057;
                border-radius: 4px;
                padding: 6px 8px;
                font: 500 10pt "Segoe UI", Arial, sans-serif;  /* ЗМЕНШЕНИЙ ШРИФТ */
                font-weight: 500;
            }
            QPushButton:hover {
                background-color: #495057;
                border: 1px solid #343a40;
            }
            QPushButton:pressed {
                background-color: #343a40;
            }
        """)
        date_layout.addWidget(self.today_btn)
        layout.addWidget(date_container)

        # 3. КНОПКИ УПРАВЛІННЯ ШАБЛОНАМИ
        self.create_template_btn = QPushButton("Створити новий шаблон")
        self.create_template_btn.clicked.connect(self.create_new_template)
        layout.addWidget(self.create_template_btn)

        self.edit_template_btn = QPushButton("Редагувати шаблон")
        self.edit_template_btn.clicked.connect(self.edit_current_template)
        layout.addWidget(self.edit_template_btn)
        
        # Separator після заповнення документу
        separator2 = QFrame()
        separator2.setFrameShape(QFrame.HLine)
        separator2.setFrameShadow(QFrame.Sunken)
        separator2.setStyleSheet("color: #dee2e6; margin: 12px 0px;")
        layout.addWidget(separator2)
        
        # ===== ПАКЕТНА ОБРОБКА =====
        self.batch_label = QLabel(self.tr("batch_processing"))
        self.batch_label.setFont(QFont("Segoe UI", 12, QFont.Bold))
        self.batch_label.setStyleSheet("color: #6c757d; margin-top: 8px; margin-bottom: 8px; font-weight: bold;")
        layout.addWidget(self.batch_label)

        self.save_current_btn = QPushButton(self.tr("save_current_image_data"))
        self.save_current_btn.clicked.connect(self.save_current_image_data)
        layout.addWidget(self.save_current_btn)

        # Кнопки скасування змін
        self.cancel_current_btn = QPushButton("Скасувати зміни")
        self.cancel_current_btn.clicked.connect(self.cancel_current_changes)
        self.cancel_current_btn.setStyleSheet("""
            QPushButton {
                background-color: #fd7e14;
                color: white;
                border: 1px solid #e76a00;
                font-weight: 600;
            }
            QPushButton:hover {
                background-color: #e76a00;
                border: 1px solid #dc6502;
            }
        """)
        layout.addWidget(self.cancel_current_btn)

        self.clear_all_btn = QPushButton("Очистити все")
        self.clear_all_btn.clicked.connect(self.clear_all_changes)
        self.clear_all_btn.setStyleSheet("""
            QPushButton {
                background-color: #dc3545;
                color: white;
                border: 1px solid #c82333;
            }
            QPushButton:hover {
                background-color: #c82333;
            }
        """)
        layout.addWidget(self.clear_all_btn)

        # ГОЛОВНА кнопка для створення альбому
        self.create_new_structure_btn = QPushButton(self.tr("create_new_album"))
        self.create_new_structure_btn.clicked.connect(self.create_batch_album_new_structure)
        self.create_new_structure_btn.setObjectName("main_action")  # СПЕЦІАЛЬНИЙ ID для стилю
        layout.addWidget(self.create_new_structure_btn)
        
        # Results area
        self.results_label = QLabel(self.tr("results"))
        self.results_label.setFont(QFont("Segoe UI", 12, QFont.Bold))
        self.results_label.setStyleSheet("color: #6c757d; margin-top: 12px; margin-bottom: 8px; font-weight: bold;")
        layout.addWidget(self.results_label)
        
        self.results_text = QTextEdit()
        self.results_text.setMaximumHeight(120)
        self.results_text.setStyleSheet("""
            QTextEdit {
                background-color: white;
                border: 1px solid #dee2e6;
                font-family: 'Consolas', 'Courier New', monospace;
                font-size: 11pt;
                color: #495057;
                border-radius: 4px;
                padding: 8px;
            }
        """)
        layout.addWidget(self.results_text)

        layout.addStretch()
        parent.addWidget(control_widget)
        
        # Ініціалізація шаблонів
        self.init_title_page_templates()

    def update_document_date(self, date):
        """Оновлення дати документу"""
        self.document_date = date
        self.add_result(f"📅 Дата документу: {date.toString('dd.MM.yyyy')}")
        print(f"Document date updated: {date.toString('dd.MM.yyyy')}")

    def set_document_date_today(self):
        """Встановлення поточної дати як дати документу"""
        today = QDate.currentDate()
        self.document_date_edit.setDate(today)
        self.add_result(f"📅 Встановлено сьогоднішню дату: {today.toString('dd.MM.yyyy')}")

    def get_document_date_for_title_page(self):
        """Отримання дати для титульної сторінки з пріоритетом"""
        # Пріоритет: Дата документу > Дата РЛС > Поточна дата
        
        if hasattr(self, 'document_date') and self.document_date.isValid():
            # Використовуємо дату документу (найвищий пріоритет)
            return self.document_date.toPyDate()
        elif (hasattr(self, 'radar_description_enabled') and 
            self.radar_description_enabled and 
            hasattr(self, 'radar_date') and 
            self.radar_date.isValid()):
            # Використовуємо дату РЛС якщо опис РЛС включений
            return self.radar_date.toPyDate()
        else:
            # Використовуємо поточну дату як fallback
            from datetime import datetime
            return datetime.now().date()


    def init_title_page_templates(self):
        """Ініціалізація системи шаблонів титульних сторінок"""
        # Папка для збереження шаблонів
        self.templates_dir = os.path.join(os.path.expanduser("~"), "PhotoControl_Templates")
        if not os.path.exists(self.templates_dir):
            os.makedirs(self.templates_dir)
        
        # Поточний шаблон
        self.current_template = None
        
        # Завантажуємо існуючі шаблони
        self.load_templates()
        
        # Якщо немає шаблонів - створюємо базовий
        if self.template_combo.count() == 0:
            self.create_default_template()

    def load_templates(self):
        """Завантаження списку шаблонів"""
        self.template_combo.clear()
        
        try:
            template_files = [f for f in os.listdir(self.templates_dir) 
                            if f.endswith('.json')]
            
            for template_file in sorted(template_files):
                template_name = template_file.replace('.json', '')
                self.template_combo.addItem(template_name)
            
            print(f"Завантажено {len(template_files)} шаблонів")
            
        except Exception as e:
            print(f"Помилка завантаження шаблонів: {e}")

    def create_default_template(self):
        """Створення базового шаблону з ПРАВИЛЬНИМИ полями та підписом"""
        default_template = {
            'name': 'Базовий шаблон',
            'unit_name': 'А0000',
            'date': '01.01.2024',
            'commander': {
                'rank': 'полковник',
                'name': 'П.П. ПЕТРЕНКО'
            },
            'chief_of_staff': {
                'rank': 'підполковник', 
                'name': 'С.С. СИДОРЕНКО'
            },
            # ===== НОВИЙ БЛОК: ПІДПИС =====
            'signature_info': {
                'rank': 'головний сержант',
                'name': 'Анна Петренко'
            },
            'margins': {
                'top': 2.25,     # ВИПРАВЛЕНО: 2.25см замість 2.5см
                'left': 2.5,     # ВИПРАВЛЕНО: 2.5см замість 2.25см
                'bottom': 0,     # Залишається 0см
                'right': 0.75    # Залишається 0.75см
            },
            'version': '1.1'
        }
        self.save_template('Базовий шаблон', default_template)

    def on_template_changed(self, template_name):
        """Обробка зміни поточного шаблону"""
        if template_name:
            self.current_template = self.load_template(template_name)
            self.add_result(f"Обрано шаблон: {template_name}")

    def create_new_template(self):
        """Створення нового шаблону"""
        dialog = TitlePageTemplateDialog(self, mode='create')
        if dialog.exec_() == QDialog.Accepted:
            template_data = dialog.get_template_data()
            template_name = template_data['name']
            
            self.save_template(template_name, template_data)
            self.load_templates()
            
            # Встановлюємо новий шаблон як поточний
            index = self.template_combo.findText(template_name)
            if index >= 0:
                self.template_combo.setCurrentIndex(index)

    def edit_current_template(self):
        """Редагування поточного шаблону"""
        current_name = self.template_combo.currentText()
        if not current_name:
            QMessageBox.warning(self, "Попередження", "Немає обраного шаблону для редагування")
            return
        
        current_template = self.load_template(current_name)
        if not current_template:
            QMessageBox.warning(self, "Помилка", "Не вдалося завантажити шаблон")
            return
        
        dialog = TitlePageTemplateDialog(self, mode='edit', template_data=current_template)
        if dialog.exec_() == QDialog.Accepted:
            template_data = dialog.get_template_data()
            new_name = template_data['name']
            
            # Якщо назва змінилась - видаляємо старий файл
            if new_name != current_name:
                old_path = os.path.join(self.templates_dir, f"{current_name}.json")
                if os.path.exists(old_path):
                    os.remove(old_path)
            
            self.save_template(new_name, template_data)
            self.load_templates()
            
            # Встановлюємо оновлений шаблон як поточний
            index = self.template_combo.findText(new_name)
            if index >= 0:
                self.template_combo.setCurrentIndex(index)

    def save_template(self, name, template_data):
        """Збереження шаблону у файл"""
        try:
            file_path = os.path.join(self.templates_dir, f"{name}.json")
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(template_data, f, indent=2, ensure_ascii=False)
            
            print(f"Шаблон збережено: {file_path}")
            
        except Exception as e:
            print(f"Помилка збереження шаблону: {e}")

    def load_template(self, name):
        """Завантаження шаблону з файлу"""
        try:
            file_path = os.path.join(self.templates_dir, f"{name}.json")
            with open(file_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        
        except Exception as e:
            print(f"Помилка завантаження шаблону: {e}")
            return None

    def get_title_page_data_from_template(self):
        """Отримання даних титульної сторінки з оновленою логікою дати та підписом"""
        if not self.current_template:
            # Базові дані якщо немає шаблону
            return {
                'date': self.get_document_date_for_title_page(),
                'unit_info': 'А0000',
                'commander_info': {
                    'rank': 'полковник',
                    'name': 'П.П. ПЕТРЕНКО'
                },
                'chief_of_staff_info': {
                    'rank': 'підполковник',
                    'name': 'С.С. СИДОРЕНКО'
                },
                'signature_info': {  # НОВИЙ БЛОК
                    'rank': 'головний сержант',
                    'name': 'Анна Петренко'
                },
                'margins': {
                    'top': 2.25,
                    'left': 2.5,
                    'bottom': 0,
                    'right': 0.75
                }
            }
        
        # Дані з шаблону
        return {
            'date': self.get_document_date_for_title_page(),
            'unit_info': self.current_template['unit_name'],
            'commander_info': self.current_template['commander'],
            'chief_of_staff_info': self.current_template['chief_of_staff'],
            'signature_info': self.current_template.get('signature_info', {  # НОВИЙ БЛОК З FALLBACK
                'rank': 'головний сержант',
                'name': 'Анна Петренко'
            }),
            'margins': self.current_template['margins']
        }

    def create_vertical_browser_panel(self, parent):
        browser_widget = QWidget()
        # ЗБІЛЬШУЄМО ширину з 180px до 280px
        browser_widget.setFixedWidth(280)
        browser_widget.setStyleSheet("""
            QWidget {
                background-color: #f8f8f8; 
                border-right: 1px solid #ccc;
                border-left: 1px solid #ccc;
            }
            QLabel {
                background: none;
                border: none;
                color: #333;
            }
        """)
        
        layout = QVBoxLayout()
        layout.setContentsMargins(0, 15, 0, 15)
        layout.setSpacing(10)
        browser_widget.setLayout(layout)
        
        self.browser_label = QLabel(self.tr("photo_browser"))
        self.browser_label.setFont(QFont("Arial", 10, QFont.Bold))
        self.browser_label.setStyleSheet("color: #666; margin-bottom: 5px; padding: 0 10px;")
        self.browser_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(self.browser_label)
        
        self.thumbnail_scroll = QScrollArea()
        self.thumbnail_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.thumbnail_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self.thumbnail_scroll.setWidgetResizable(False)
        self.thumbnail_scroll.setStyleSheet("border: none; background: transparent;")
        
        # ЗБІЛЬШУЄМО ширину віджета мініатюр з 160px до 260px
        self.thumbnail_widget = VerticalThumbnailWidget(thumbnail_width=260)
        self.thumbnail_widget.image_selected.connect(self.load_image_from_browser)
        self.thumbnail_scroll.setWidget(self.thumbnail_widget)
        
        layout.addWidget(self.thumbnail_scroll)
        
        browser_widget.hide()
        self.browser_widget = browser_widget
        
        parent.addWidget(browser_widget)

    def create_image_panel(self, parent):
        """Центральна панель для відображення зображень"""
        image_widget = QWidget()
        layout = QVBoxLayout()
        image_widget.setLayout(layout)
        
        self.image_label = ClickableLabel()
        self.image_label.clicked.connect(self.on_image_click)
        self.image_label.dragged.connect(self.on_image_drag)
        self.image_label.mouse_moved.connect(self.on_mouse_hover)
        self.image_label.setText(self.tr("open_instruction"))
        
        scroll_area = QScrollArea()
        scroll_area.setWidget(self.image_label)
        scroll_area.setWidgetResizable(True)
        scroll_area.setMinimumHeight(400)
        layout.addWidget(scroll_area)
        
        parent.addWidget(image_widget)

    def open_image(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, self.tr("select_image"), "",
            f"{self.tr('image_files')} (*.jpg *.jpeg *.png *.bmp *.gif *.tiff);;{self.tr('all_files')} (*.*)"
        )
        
        if file_path:
            self.load_image(file_path)
            self.report_widget.show()
    
    def load_image_from_browser(self, file_path):
        """Завантаження зображення з браузера з виділенням обраного"""
        self.load_image(file_path)
        self.add_result(self.tr("loaded_from_browser").format(name=os.path.basename(file_path)))
        
        # Виділяємо поточне зображення в браузері
        if hasattr(self, 'thumbnail_widget'):
            self.thumbnail_widget.set_selected_image(file_path)

    def open_folder(self):
        """Відкриття папки з зображеннями з правильними розмірами браузера"""
        folder_path = QFileDialog.getExistingDirectory(self, self.tr("select_folder"))
        
        if folder_path:
            self.current_folder = folder_path
            self.load_folder_thumbnails()
            
            self.browser_widget.show()
            self.main_splitter.setSizes([220, 280, 620, 220])
            
            image_extensions = ('.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff', '.tif')
            image_count = sum(1 for f in os.listdir(folder_path) 
                            if f.lower().endswith(image_extensions))
            
            self.add_result(f"{self.tr('loaded_folder')}: {os.path.basename(folder_path)}")
            self.add_result(self.tr("found_images").format(count=image_count))
            
            self.report_widget.show()

    def load_folder_thumbnails(self):
        """Виправлена функція завантаження мініатюр з правильними розмірами"""
        print("🟡 === load_folder_thumbnails STARTED ===")
        
        if not self.current_folder:
            print("❌ current_folder is None!")
            return
        
        print(f"🔍 Loading thumbnails from: {self.current_folder}")
        
        # Перевірка thumbnail_widget
        if not hasattr(self, 'thumbnail_widget'):
            print("❌ thumbnail_widget doesn't exist!")
            return
        
        print(f"✅ thumbnail_widget exists: {type(self.thumbnail_widget)}")
        
        # ВАЖЛИВО: Очищуємо попередні мініатюри
        self.thumbnail_widget.clear_thumbnails()
        
        image_extensions = ('.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff', '.tif')
        image_files = []
        
        try:
            print(f"📁 Scanning folder: {self.current_folder}")
            for filename in os.listdir(self.current_folder):
                if filename.lower().endswith(image_extensions):
                    full_path = os.path.join(self.current_folder, filename)
                    image_files.append(full_path)
                    print(f"✅ Found image: {filename}")
        except Exception as e:
            print(f"❌ Error reading folder: {e}")
            return

        print(f"📊 Total images found: {len(image_files)}")
        
        # Сортуємо файли
        image_files.sort()
        
        if len(image_files) == 0:
            print("📭 No images - adding 'no images' label")
            no_images_label = QLabel(self.tr("no_images_found"))
            no_images_label.setAlignment(Qt.AlignCenter)
            no_images_label.setStyleSheet("color: gray; font-size: 14px; padding: 20px;")
            no_images_label.setWordWrap(True)
            self.thumbnail_widget.layout.addWidget(no_images_label)
            return
        
        print(f"🔄 Creating thumbnails for {len(image_files)} images...")
        
        # ВИПРАВЛЕННЯ: Створюємо мініатюри тільки ОДИН раз
        for i, image_path in enumerate(image_files):
            try:
                print(f"🖼️ Creating thumbnail {i+1}/{len(image_files)}: {os.path.basename(image_path)}")
                self.thumbnail_widget.add_thumbnail(image_path)
                print(f"✅ Thumbnail {i+1} created successfully")
            except Exception as e:
                print(f"❌ Error creating thumbnail {i+1}: {e}")
                import traceback
                traceback.print_exc()
        
        # ВИПРАВЛЕНІ розміри віджета для більших мініатюр
        widget_height = len(image_files) * 190 + 20  # Збільшена висота для мініатюр 240x180
        self.thumbnail_widget.setMinimumHeight(widget_height)
        self.thumbnail_widget.resize(260, widget_height)  # Ширина 260px
        
        print(f"🟢 === load_folder_thumbnails COMPLETED ===")
        print(f"📋 Final result: {len(image_files)} unique thumbnails created")


    # НОВІ МЕТОДИ ДЛЯ СТВОРЕННЯ АЛЬБОМІВ З ТАБЛИЦЯМИ

    def create_batch_album_new_structure(self):
        """Створення альбому з використанням шаблону титульної сторінки та підписом на сторінці опису"""
        if not DOCX_AVAILABLE:
            QMessageBox.warning(self, "Warning", "python-docx library not installed")
            return
        
        if not self.processed_images:
            QMessageBox.warning(self, "Warning", "No processed images to export")
            return
        
        # Зберігаємо поточне зображення якщо є
        if self.processor and self.current_click:
            self.save_current_image_data()
        
        # Діалог вибору файлу
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Create Complete Album with Template and Signature", "",
            "Word Documents (*.docx);;All files (*.*)"
        )
        
        if file_path:
            try:
                # ===== ОТРИМУЄМО ДАНІ З ПОТОЧНОГО ШАБЛОНУ (З ПІДПИСОМ) =====
                title_data = self.get_title_page_data_from_template()
                
                print(f"📋 Using template data with signature:")
                print(f"   Unit: {title_data['unit_info']}")
                print(f"   Commander: {title_data['commander_info']['rank']} {title_data['commander_info']['name']}")
                print(f"   Chief of staff: {title_data['chief_of_staff_info']['rank']} {title_data['chief_of_staff_info']['name']}")
                print(f"   Signature: {title_data['signature_info']['rank']} {title_data['signature_info']['name']}")
                print(f"   Margins: {title_data['margins']}")
                
                # Створюємо повний альбом з шаблоном та підписом
                success = create_complete_album(self.processed_images, title_data, file_path)
                
                if success:
                    current_template_name = self.template_combo.currentText()
                    QMessageBox.information(self, "Success", 
                                        f"Complete album created with template and signature!\n"
                                        f"• Template: {current_template_name}\n"
                                        f"• Unit: {title_data['unit_info']}\n"
                                        f"• Commander: {title_data['commander_info']['name']}\n"
                                        f"• Chief of Staff: {title_data['chief_of_staff_info']['name']}\n"
                                        f"• Signature: {title_data['signature_info']['name']}\n"
                                        f"• Images: {len(self.processed_images)}\n"
                                        f"• Custom margins applied\n\n"
                                        f"Saved: {os.path.basename(file_path)}")
                    
                    self.add_result(f"✓ Album created with template: {current_template_name}")
                    self.add_result(f"✓ Contains {len(self.processed_images)} processed images")
                    self.add_result(f"✓ Signature: {title_data['signature_info']['rank']} {title_data['signature_info']['name']}")
                else:
                    QMessageBox.critical(self, "Error", "Failed to create album with template and signature")
                    
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Could not create album: {str(e)}")
                print(f"Error creating album with template and signature: {e}")
                import traceback
                traceback.print_exc()


    def get_title_page_data_from_gui(self):
        """Отримання даних для титульної сторінки з GUI"""
        from datetime import datetime
        
        title_data = {
            'target_no': self.current_target_number,
            'date': self.radar_date if self.radar_description_enabled else datetime.now(),
            'unit_info': DefaultTemplateData.get_default_unit_info(),
            'commander_info': DefaultTemplateData.get_default_commander_info()
        }
        
        return title_data
        
    def create_report_panel(self, parent):
        """Права панель з азимутальними контролами та описом РЛС (ПОВНІСТЮ ОНОВЛЕНА)"""
        report_widget = QWidget()
        report_widget.setFixedWidth(220)
        report_widget.setStyleSheet("""
            QWidget {
                background-color: #f9f9f9; 
                border-left: 1px solid #dee2e6;
            }
            QLabel {
                background: none;
                border: none;
                color: #343a40;
                font: 12pt "Segoe UI", Arial, sans-serif;
            }
            QPushButton {
                background-color: #f8f9fa;
                border: 2px solid #dee2e6;  /* Товща рамка замість тіні */
                border-radius: 6px;
                padding: 10px 14px;
                font: 500 12pt "Segoe UI", Arial, sans-serif;
                color: #495057;
            }
            QPushButton:hover {
                background-color: #e9ecef;
                border: 2px solid #adb5bd;  /* Акцентна рамка замість тіні */
                color: #343a40;
            }
            QPushButton:pressed {
                background-color: #dee2e6;
                border: 2px solid #6c757d;
                border-style: inset;  /* Вдавлений ефект замість тіні */
            }
            QPushButton:checked {
                background-color: #495057;
                color: white;
                border: 2px solid #343a40;
            }
            QPushButton:checked:hover {
                background-color: #343a40;
                border: 2px solid #212529;
            }
            QLineEdit, QComboBox {
                border: 1px solid #dee2e6;
                border-radius: 4px;
                padding: 6px 10px;
                background-color: white;
                font: 12pt "Segoe UI", Arial, sans-serif;
                color: #495057;
                min-height: 20px;
            }
            QLineEdit:hover, QComboBox:hover {
                border: 1px solid #adb5bd;
                background-color: #f8f9fa;
            }
            QLineEdit:focus, QComboBox:focus {
                border: 1px solid #6c757d;
                background-color: white;
            }
            QLineEdit:disabled, QComboBox:disabled {
                background-color: #f5f5f5;
                color: #6c757d;
                border: 1px solid #e9ecef;
            }
            QComboBox::drop-down {
                border: none;
                width: 20px;
                border-left: 1px solid #dee2e6;
                border-top-right-radius: 4px;
                border-bottom-right-radius: 4px;
                background-color: #f8f9fa;
            }
            QComboBox::drop-down:hover {
                background-color: #e9ecef;
            }
            QComboBox::down-arrow {
                image: none;
                border-left: 5px solid transparent;
                border-right: 5px solid transparent;
                border-top: 5px solid #6c757d;
            }
            QCheckBox {
                color: #495057;
                font: 500 12pt "Segoe UI", Arial, sans-serif;
                padding: 6px;
                font-weight: 500;
            }
            QCheckBox::indicator {
                width: 18px;
                height: 18px;
            }
            QCheckBox::indicator:unchecked {
                border: 2px solid #dee2e6;
                background-color: white;
                border-radius: 3px;
            }
            QCheckBox::indicator:checked {
                border: 2px solid #495057;
                background-color: #495057;
                border-radius: 3px;
            }
            QCheckBox::indicator:unchecked:hover {
                border: 2px solid #adb5bd;
                background-color: #f8f9fa;
            }
            QDateEdit {
                border: 1px solid #dee2e6;
                border-radius: 4px;
                padding: 6px 10px;
                background-color: white;
                font: 12pt "Segoe UI", Arial, sans-serif;
                color: #495057;
                min-height: 22px;
            }
            QDateEdit:hover {
                border: 1px solid #adb5bd;
                background-color: #f8f9fa;
            }
            QDateEdit:focus {
                border: 1px solid #6c757d;
                background-color: white;
            }
            QDateEdit:disabled {
                background-color: #f5f5f5;
                color: #6c757d;
                border: 1px solid #e9ecef;
            }
            QDateEdit::drop-down {
                border-left: 1px solid #dee2e6;
                background-color: #f8f9fa;
                border-top-right-radius: 4px;
                border-bottom-right-radius: 4px;
                width: 20px;
            }
            QDateEdit::drop-down:hover {
                background-color: #e9ecef;
            }
            QDateEdit::drop-down:disabled {
                background-color: #f5f5f5;
            }
            QDateEdit::down-arrow {
            image: none;
            border-left: 4px solid transparent;
            border-right: 4px solid transparent;
            border-top: 4px solid #6c757d;
            margin-top: 2px;
            }
            QDateEdit::down-arrow:hover {
                border-top-color: #495057;
            }
            QDateEdit {
                border: 1px solid #dee2e6;
                border-radius: 4px;
                padding: 6px 10px;
                background-color: white;
                font: 12pt "Segoe UI", Arial, sans-serif;
                color: #495057;
                min-height: 22px;
            }
            QDateEdit:hover {
                border: 1px solid #adb5bd;
                background-color: #f8f9fa;
            }
            QDateEdit:focus {
                border: 1px solid #6c757d;
            }
            QDateEdit:disabled {
                background-color: #f5f5f5;
                color: #6c757d;
                border: 1px solid #e9ecef;
            }
            QDateEdit::drop-down {
                border: none;
                background-color: transparent;
                width: 18px;
                margin-right: 4px;
            }
            QDateEdit::drop-down:hover {
                background-color: #f8f9fa;
                border-radius: 3px;
            }
            QDateEdit::down-arrow {
                image: none;
                border-left: 4px solid transparent;
                border-right: 4px solid transparent;
                border-top: 5px solid #6c757d;
                margin-top: 1px;
            }
            QDateEdit::down-arrow:hover {
                border-top-color: #495057;
            }
            QCalendarWidget {
                background-color: white;
                color: #495057;
                border: 1px solid #dee2e6;
                border-radius: 6px;
                font-family: "Segoe UI", Arial, sans-serif;
            }
            QCalendarWidget QWidget#qt_calendar_navigationbar {
                background-color: #f8f9fa;
                border-bottom: 1px solid #e9ecef;
                border-top-left-radius: 6px;
                border-top-right-radius: 6px;
                padding: 4px;
            }
            QCalendarWidget QToolButton {
                color: #6c757d;
                background-color: transparent;
                border: 1px solid transparent;
                border-radius: 4px;
                margin: 1px;
                padding: 4px 8px;
                font-weight: normal;
                min-width: 24px;
                min-height: 24px;
            }
            QCalendarWidget QToolButton:hover {
                background-color: #e9ecef;
                color: #495057;
                border: 1px solid #adb5bd;
            }
            QCalendarWidget QToolButton#qt_calendar_prevmonth {
                qproperty-text: "‹";
                font-size: 16pt;
                font-weight: bold;
            }
            QCalendarWidget QToolButton#qt_calendar_nextmonth {
                qproperty-text: "›";
                font-size: 16pt;
                font-weight: bold;
            }
            QCalendarWidget QSpinBox {
                color: #495057;
                background-color: white;
                border: 1px solid #dee2e6;
                border-radius: 3px;
                font-size: 12pt;
                font-weight: normal;
                padding: 2px 6px;
                min-width: 60px;
                margin: 1px;
            }
            QCalendarWidget QSpinBox:hover {
                border: 1px solid #adb5bd;
                background-color: #f8f9fa;
            }
            QCalendarWidget QHeaderView::section {
                color: #6c757d;
                background-color: #f8f9fa;
                border: none;
                border-bottom: 1px solid #e9ecef;
                font-weight: 500;
                padding: 6px 2px;
                font-size: 11pt;
            }
            QCalendarWidget QAbstractItemView {
                background-color: white;
                color: #495057;
                font-size: 12pt;
                border: none;
                selection-background-color: #495057;
                selection-color: white;
            }
            QCalendarWidget QAbstractItemView:item:selected {
                background-color: #495057;
                color: white;
                border-radius: 3px;
                font-weight: normal;
            }
            QCalendarWidget QAbstractItemView:item:focus {
                background-color: #e9ecef;
                color: #495057;
                border: 1px solid #adb5bd;
                border-radius: 3px;
            }
        """)
        
        layout = QVBoxLayout()
        layout.setContentsMargins(15, 15, 15, 15)
        layout.setSpacing(12)
        report_widget.setLayout(layout)
        
        # Title з новим шрифтом
        self.report_title = QLabel(self.tr("report_data"))
        self.report_title.setFont(QFont("Segoe UI", 14, QFont.Bold))
        self.report_title.setAlignment(Qt.AlignCenter)
        self.report_title.setStyleSheet("font-size: 14pt; font-weight: bold; margin-bottom: 12px; color: #343a40;")
        layout.addWidget(self.report_title)
        
        # Target data input section (СПОЧАТКУ ДАНІ ПРО ЦІЛЬ)
        manual_group = QFrame()
        manual_group.setStyleSheet("""
            QFrame {
                background-color: white; 
                border: 1px solid #dee2e6; 
                border-radius: 6px; 
                padding: 12px;
            }
        """)
        manual_layout = QVBoxLayout()
        manual_layout.setSpacing(10)
        manual_group.setLayout(manual_layout)
        
        # Номер цілі
        self.target_number_input = QLineEdit()
        self.target_number_input.setPlaceholderText("Номер цілі")
        self.target_number_input.setText(self.current_target_number)
        self.target_number_input.textChanged.connect(self.update_target_number)
        manual_layout.addWidget(self.target_number_input)
        
        # Автоматичні поля (азимут, дальність)
        self.auto_azimuth_label = QLabel("β - --°")
        self.auto_azimuth_label.setStyleSheet("font-weight: 500; color: #495057;")
        manual_layout.addWidget(self.auto_azimuth_label)
        
        self.auto_distance_label = QLabel("D - -- км")
        self.auto_distance_label.setStyleSheet("font-weight: 500; color: #495057;")
        manual_layout.addWidget(self.auto_distance_label)
        
        # Висота (в одному рядку)
        height_container = QWidget()
        height_layout = QHBoxLayout()
        height_layout.setContentsMargins(0, 0, 0, 0)
        height_layout.setSpacing(6)
        height_container.setLayout(height_layout)
        
        height_label = QLabel("H –")
        height_label.setStyleSheet("color: #495057; font-weight: 500;")
        height_layout.addWidget(height_label)
        
        self.height_input = QLineEdit(self.current_height)
        self.height_input.setMaximumWidth(80)
        self.height_input.textChanged.connect(self.update_height)
        height_layout.addWidget(self.height_input)
        
        units_label = QLabel(self.tr("km_unit"))
        units_label.setStyleSheet("color: #6c757d;")
        height_layout.addWidget(units_label)
        
        height_layout.addStretch()
        manual_layout.addWidget(height_container)
        
        # Комбобокси
        self.obstacles_combo = QComboBox()
        self.obstacles_combo.addItems([self.tr("no_obstacles"), self.tr("with_obstacles")])
        self.obstacles_combo.currentTextChanged.connect(self.update_obstacles)
        manual_layout.addWidget(self.obstacles_combo)
        
        self.detection_combo = QComboBox()
        self.detection_combo.addItems([self.tr("detection"), self.tr("tracking"), self.tr("loss")])
        self.detection_combo.currentTextChanged.connect(self.update_detection)
        manual_layout.addWidget(self.detection_combo)
        
        # Масштаб
        self.auto_scale_label = QLabel("M = --")
        self.auto_scale_label.setStyleSheet("font-weight: 500; color: #495057;")
        manual_layout.addWidget(self.auto_scale_label)
        
        layout.addWidget(manual_group)
        
        # Separator після даних про ціль
        separator1 = QFrame()
        separator1.setFrameShape(QFrame.HLine)
        separator1.setFrameShadow(QFrame.Sunken)
        separator1.setStyleSheet("color: #dee2e6; margin: 8px 0px;")
        layout.addWidget(separator1)
        
        # Azimuth Grid section
        self.azimuth_grid_label = QLabel(self.tr("azimuth_grid"))
        self.azimuth_grid_label.setFont(QFont("Segoe UI", 12, QFont.Bold))
        self.azimuth_grid_label.setStyleSheet("color: #6c757d; font-weight: bold; margin-bottom: 8px;")
        layout.addWidget(self.azimuth_grid_label)
        
        # Scale setting (в одному рядку)
        scale_container = QWidget()
        scale_layout = QHBoxLayout()
        scale_layout.setContentsMargins(0, 0, 0, 0)
        scale_layout.setSpacing(8)
        scale_container.setLayout(scale_layout)
        
        scale_label = QLabel(self.tr("scale_setting"))
        scale_label.setStyleSheet("color: #495057; font-weight: 500;")
        scale_layout.addWidget(scale_label)
        
        self.scale_combo = QComboBox()
        self.scale_combo.addItems(["25", "35", "50", "75", "80", "90", "100", "150", "200", "250", "300", "350"])
        self.scale_combo.setCurrentText("300")
        self.scale_combo.currentTextChanged.connect(self.update_scale)
        scale_layout.addWidget(self.scale_combo)
        
        scale_layout.addStretch()
        layout.addWidget(scale_container)
        
        # Кнопки управління сіткою
        self.scale_edge_btn = QPushButton(self.tr("set_scale_edge"))
        self.scale_edge_btn.setCheckable(True)
        self.scale_edge_btn.clicked.connect(self.toggle_scale_edge_mode)
        layout.addWidget(self.scale_edge_btn)
        
        self.set_center_btn = QPushButton(self.tr("set_center"))
        self.set_center_btn.setCheckable(True)
        self.set_center_btn.clicked.connect(self.toggle_center_setting_mode)
        layout.addWidget(self.set_center_btn)
        
        # ===== СЕКЦІЯ: ОПИС РЛС =====
        radar_separator = QFrame()
        radar_separator.setFrameShape(QFrame.HLine)
        radar_separator.setFrameShadow(QFrame.Sunken)
        radar_separator.setStyleSheet("color: #dee2e6; margin: 12px 0px;")
        layout.addWidget(radar_separator)
        
        # Checkbox для активації опису РЛС
        self.radar_description_checkbox = QCheckBox("Додати опис РЛС")
        self.radar_description_checkbox.setFont(QFont("Segoe UI", 12, QFont.Bold))
        self.radar_description_checkbox.setStyleSheet("font-weight: bold; color: #495057; padding: 6px;")
        self.radar_description_checkbox.toggled.connect(self.toggle_radar_description)
        layout.addWidget(self.radar_description_checkbox)
        
        # Група полів для опису РЛС
        radar_group = QFrame()
        radar_group.setStyleSheet("""
            QFrame {
                background-color: white; 
                border: 1px solid #dee2e6; 
                border-radius: 6px; 
                padding: 12px;
            }
        """)
        radar_layout = QVBoxLayout()
        radar_layout.setSpacing(10)
        radar_group.setLayout(radar_layout)
        
        # Вибір дати (БЕЗ СЛОВА "ДАТА:")
        date_container = QWidget()
        date_layout = QHBoxLayout()
        date_layout.setContentsMargins(0, 0, 0, 0)
        date_layout.setSpacing(0)  # Без додаткового простору
        date_container.setLayout(date_layout)

        # ТІЛЬКИ ВІДЖЕТ ДАТИ, БЕЗ ЛЕЙБЛА
        self.radar_date_edit = QDateEdit()
        self.radar_date_edit.setDate(self.radar_date)
        self.radar_date_edit.setCalendarPopup(True)
        self.radar_date_edit.setFixedHeight(32)
        self.radar_date_edit.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        self.radar_date_edit.dateChanged.connect(self.update_radar_date)
        self.radar_date_edit.setEnabled(False)  # Спочатку неактивне
        date_layout.addWidget(self.radar_date_edit)

        radar_layout.addWidget(date_container)

        # Поле для позивного (БЕЗ ЗМІН)
        self.radar_callsign_input = QLineEdit()
        self.radar_callsign_input.setPlaceholderText("Позивний")
        self.radar_callsign_input.setFixedHeight(32)
        self.radar_callsign_input.textChanged.connect(self.update_radar_callsign)
        self.radar_callsign_input.setEnabled(False)
        radar_layout.addWidget(self.radar_callsign_input)

        # Поле для назви РЛС (БЕЗ ЗМІН)
        self.radar_name_input = QLineEdit()
        self.radar_name_input.setPlaceholderText("Назва РЛС")
        self.radar_name_input.setFixedHeight(32)
        self.radar_name_input.textChanged.connect(self.update_radar_name)
        self.radar_name_input.setEnabled(False)
        radar_layout.addWidget(self.radar_name_input)

        # Поле для номера РЛС (БЕЗ ЗМІН)
        self.radar_number_input = QLineEdit()
        self.radar_number_input.setPlaceholderText("Номер РЛС")
        self.radar_number_input.setFixedHeight(32)
        self.radar_number_input.textChanged.connect(self.update_radar_number)
        self.radar_number_input.setEnabled(False)
        radar_layout.addWidget(self.radar_number_input)
        
        layout.addWidget(radar_group)
        
        # Розтягування внизу
        layout.addStretch()
        
        # Приховуємо панель спочатку
        report_widget.hide()
        self.report_widget = report_widget
        parent.addWidget(report_widget)

    def move_center_and_save(self, dx, dy):
        """Переміщення центру з автоматичним збереженням налаштувань"""
        if not self.processor:
            QMessageBox.warning(self, self.tr("warning"), self.tr("no_image_first"))
            return
        
        self.processor.move_center(dx, dy)
        self.save_current_grid_settings()
        
        if self.current_click:
            azimuth, range_val = self.calculate_azimuth_range(
                self.current_click['x'], self.current_click['y']
            )
            self.current_click['azimuth'] = azimuth
            self.current_click['range'] = range_val
        
        self.display_image()
        self.update_results_display()
        self.update_report_data()
        
        self.add_result(self.tr("center_moved").format(x=self.processor.center_x, 
                                                    y=self.processor.center_y))
        self.add_result(self.tr("grid_settings_saved"))
        
        if hasattr(self, 'image_label') and self.image_label.zoom_widget.isVisible():
            self.image_label.zoom_widget.update_cursor_position(self.processor.center_x, self.processor.center_y)

    def add_result(self, text):
        self.results_text.append(text)
    
    
    def load_image(self, file_path):
        """Завантаження зображення зі збереженням налаштувань сітки"""
        try:
            self.current_image_path = file_path
            scale_value = int(self.scale_combo.currentText())
            
            self.processor = AzimuthImageProcessor(file_path, scale=scale_value)
            
            if hasattr(self.processor, 'image') and self.processor.image:
                original_mode = self.processor.image.mode
                
                if original_mode != 'RGB':
                    if original_mode == 'RGBA':
                        rgb_image = Image.new('RGB', self.processor.image.size, (255, 255, 255))
                        rgb_image.paste(self.processor.image, mask=self.processor.image.split()[-1])
                        self.processor.image = rgb_image
                    else:
                        self.processor.image = self.processor.image.convert('RGB')
            
            # Застосовуємо збережені налаштування сітки
            self.apply_saved_grid_settings()
            
            # Очищуємо тільки точку аналізу, але НЕ налаштування сітки
            self.current_click = None
            
            self.display_image()
            self.update_results_display()
            self.update_report_data()
            
            self.add_result(f"{self.tr('loaded')}: {os.path.basename(file_path)}")
            if self.has_saved_grid_settings():
                self.add_result(self.tr("grid_settings_applied"))
            
        except Exception as e:
            QMessageBox.critical(self, self.tr("error"), 
                               self.tr("could_not_load").format(error=str(e)))
    
    def apply_saved_grid_settings(self):
        """Застосувати збережені налаштування сітки до нового зображення"""
        if not self.processor:
            return
        
        # Отримуємо центр нового зображення
        image_center_x = self.processor.image.width // 2
        image_center_y = self.processor.image.height // 2
        
        # Застосовуємо зміщення центру
        if self.saved_grid_settings['center_offset_x'] != 0 or self.saved_grid_settings['center_offset_y'] != 0:
            self.processor.move_center(
                self.saved_grid_settings['center_offset_x'],
                self.saved_grid_settings['center_offset_y']
            )
        
        # Відновлюємо scale edge point відносно центру зображення
        if self.saved_grid_settings.get('scale_edge_relative') and self.saved_grid_settings['custom_scale_distance']:
            edge_relative = self.saved_grid_settings['scale_edge_relative']
            
            # Нові абсолютні координати відносно центру зображення
            new_x = image_center_x + edge_relative['x']
            new_y = image_center_y + edge_relative['y']
            
            # Перевіряємо що координати в межах зображення
            if 0 <= new_x < self.processor.image.width and 0 <= new_y < self.processor.image.height:
                self.scale_edge_point = {'x': new_x, 'y': new_y}
                self.custom_scale_distance = self.saved_grid_settings['custom_scale_distance']
    
        # Відновлюємо масштаб
        if self.saved_grid_settings['scale_value']:
            self.scale_combo.setCurrentText(self.saved_grid_settings['scale_value'])

    def save_current_grid_settings(self):
        """Зберегти поточні налаштування сітки"""
        if not self.processor:
            return
        
        # Зберігаємо зміщення центру відносно центру зображення
        image_center_x = self.processor.image.width // 2
        image_center_y = self.processor.image.height // 2
        
        # Зберігаємо scale edge point відносно центру зображення
        scale_edge_relative = None
        if self.scale_edge_point:
            scale_edge_relative = {
                'x': self.scale_edge_point['x'] - image_center_x,
                'y': self.scale_edge_point['y'] - image_center_y
            }
        
        self.saved_grid_settings = {
            'center_offset_x': self.processor.center_x - image_center_x,
            'center_offset_y': self.processor.center_y - image_center_y,
            'scale_edge_relative': scale_edge_relative,
            'custom_scale_distance': self.custom_scale_distance,
            'scale_value': self.scale_combo.currentText()
        }

    def has_saved_grid_settings(self):
        """Перевірити чи є збережені налаштування сітки"""
        return (self.saved_grid_settings['center_offset_x'] != 0 or 
                self.saved_grid_settings['center_offset_y'] != 0 or
                self.saved_grid_settings.get('scale_edge_relative') is not None or
                self.saved_grid_settings['scale_value'] != "300")
    
    def display_image(self):
        if not self.processor:
            return
        
        pil_image = self.processor.image.copy()
        draw = ImageDraw.Draw(pil_image)
        
        center_x, center_y = self.processor.center_x, self.processor.center_y
        cross_size = 15
        
        # Малюємо червоний хрестик в центрі
        draw.line([center_x - cross_size, center_y, center_x + cross_size, center_y], 
                fill='red', width=2)
        draw.line([center_x, center_y - cross_size, center_x, center_y + cross_size], 
                fill='red', width=2)
        draw.ellipse([center_x - 3, center_y - 3, center_x + 3, center_y + 3], 
                    fill='red', outline='white')
        
        if self.current_click:
            click_x, click_y = self.current_click['x'], self.current_click['y']
            
            # Малюємо синю точку аналізу
            draw.ellipse([click_x - 4, click_y - 4, click_x + 4, click_y + 4], 
                        fill='blue', outline='white', width=1)
            
            # ОНОВЛЕНА ЛІНІЯ: Розраховуємо кінцеву позицію як в документі
            image_width = pil_image.width
            image_height = pil_image.height
            
            # Розрахунок позиції кінця лінії (на рівні підкреслення номера цілі)
            # Використовуємо ту ж логіку що й в create_processed_image_from_data
            right_cell_start_x = int(image_width * 0.85)  # Права комірка починається з 85% ширини
            underline_y = int(image_height * 0.1)         # Позиція підкреслення: 10% висоти зверху
            
            # Кінцева точка лінії: самий правий край на рівні підкреслення
            end_x = image_width - 1  # Самий край зображення
            end_y = underline_y      # На висоті підкреслення номера цілі
            
            # Малюємо оновлену лінію від точки аналізу до розрахованої позиції
            draw.line([click_x, click_y, end_x, end_y], fill='blue', width=3)
        
        # Малюємо зелену точку та лінію для scale edge (якщо є)
        if self.scale_edge_point:
            edge_x, edge_y = self.scale_edge_point['x'], self.scale_edge_point['y']
            
            draw.ellipse([edge_x - 5, edge_y - 5, edge_x + 5, edge_y + 5], 
                        fill='green', outline='white', width=2)
            
            draw.line([center_x, center_y, edge_x, edge_y], fill='green', width=2)
            
            # Перпендикулярна лінія на кінці
            dx = edge_x - center_x
            dy = edge_y - center_y
            length = math.sqrt(dx*dx + dy*dy)
            if length > 0:
                nx, ny = -dy/length, dx/length
                perp_size = 8
                draw.line([
                    edge_x + nx*perp_size, edge_y + ny*perp_size,
                    edge_x - nx*perp_size, edge_y - ny*perp_size
                ], fill='green', width=2)
        
        # Зберігаємо та відображаємо зображення
        with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as temp_file:
            temp_path = temp_file.name
            pil_image.save(temp_path, 'JPEG', quality=95)
        
        pixmap = QPixmap(temp_path)
        
        try:
            os.remove(temp_path)
        except:
            pass
        
        widget_width = self.image_label.width()
        widget_height = self.image_label.height()
        
        scaled_pixmap = pixmap.scaled(widget_width, widget_height, 
                                    Qt.KeepAspectRatio, Qt.SmoothTransformation)
        
        original_width = pixmap.width()
        original_height = pixmap.height()
        scaled_width = scaled_pixmap.width()
        scaled_height = scaled_pixmap.height()
        
        self.scale_factor_x = scaled_width / original_width
        self.scale_factor_y = scaled_height / original_height
        
        self.offset_x = (widget_width - scaled_width) // 2
        self.offset_y = (widget_height - scaled_height) // 2
        
        self.image_label.update_image_geometry(
            original_width, original_height,
            self.scale_factor_x, self.scale_factor_y,
            self.offset_x, self.offset_y
        )
        
        self.image_label.set_zoom_source_image(pixmap)
        self.image_label.setPixmap(scaled_pixmap)
    
    def on_image_click(self, x, y):
        if not self.processor:
            return

        if self.scale_edge_mode:
            self.set_scale_edge_point(x, y)
            return

        if self.center_setting_mode:
            self.set_center_point(x, y)
            return

        if self.current_click:
            existing_widget_x = self.current_click['x'] * self.scale_factor_x + self.offset_x
            existing_widget_y = self.current_click['y'] * self.scale_factor_y + self.offset_y
        
            clicked_widget_x = x * self.scale_factor_x + self.offset_x
            clicked_widget_y = y * self.scale_factor_y + self.offset_y
        
            distance = ((clicked_widget_x - existing_widget_x)**2 + (clicked_widget_y - existing_widget_y)**2)**0.5
        
            if distance <= 15:
                return

        self.place_analysis_point(x, y)
    
    def on_image_drag(self, x, y):
        """Перетягування точки аналізу при утримуванні кнопки миші"""
        if not self.processor:
            return
        
        # НЕ обробляємо перетягування в спеціальних режимах
        if self.scale_edge_mode or self.center_setting_mode:
            return
        
        # Перетягуємо точку аналізу
        self.place_analysis_point(x, y)
    
    def on_mouse_hover(self, x, y):
        if not self.processor or not self.current_click:
            QToolTip.hideText()
            return
        
        hover_widget_x = x * self.scale_factor_x + self.offset_x
        hover_widget_y = y * self.scale_factor_y + self.offset_y
        
        existing_widget_x = self.current_click['x'] * self.scale_factor_x + self.offset_x
        existing_widget_y = self.current_click['y'] * self.scale_factor_y + self.offset_y
        
        distance = ((hover_widget_x - existing_widget_x)**2 + (hover_widget_y - existing_widget_y)**2)**0.5
        
        if distance <= 15:
            azimuth = self.current_click['azimuth']
            range_val = self.current_click['range']
            tooltip_text = f"{self.tr('azimuth')}: {azimuth:.0f}°\n{self.tr('range')}: {range_val:.0f} км"
            
            point = self.image_label.mapToGlobal(self.image_label.rect().topLeft())
            tooltip_x = point.x() + hover_widget_x + 15
            tooltip_y = point.y() + hover_widget_y - 10
            
            QToolTip.showText(QPoint(int(tooltip_x), int(tooltip_y)), tooltip_text)
        else:
            QToolTip.hideText()
    
    def place_analysis_point(self, x, y):
        if not self.processor:
            return
        
        x = max(0, min(x, self.processor.image.width - 1))
        y = max(0, min(y, self.processor.image.height - 1))
        
        try:
            azimuth, range_val = self.calculate_azimuth_range(x, y)
            
            self.current_click = {
                'x': x, 'y': y,
                'azimuth': azimuth, 'range': range_val
            }
            
            self.display_image()
            self.update_results_display()
            self.update_report_data()
            
        except Exception as e:
            QMessageBox.critical(self, self.tr("error"), f"Could not process point: {str(e)}")
    
    def calculate_azimuth_range(self, x, y):
        dx = x - self.processor.center_x
        dy = self.processor.center_y - y
        
        range_pixels = math.sqrt(dx**2 + dy**2)
        
        if self.custom_scale_distance:
            scale_value = int(self.scale_combo.currentText())
            range_actual = (range_pixels / self.custom_scale_distance) * scale_value
        else:
            bottom_edge_distance = self.processor.image.height - self.processor.center_y
            scale_value = int(self.scale_combo.currentText())
            range_actual = (range_pixels / bottom_edge_distance) * scale_value
        
        azimuth_radians = math.atan2(dx, dy)
        azimuth_degrees = math.degrees(azimuth_radians)
        
        if azimuth_degrees < 0:
            azimuth_degrees += 360
            
        return azimuth_degrees, range_actual
    
    def toggle_center_setting_mode(self):
        self.center_setting_mode = self.set_center_btn.isChecked()
        
        if self.center_setting_mode and self.scale_edge_mode:
            self.scale_edge_mode = False
            self.scale_edge_btn.setChecked(False)
            self.scale_edge_btn.setStyleSheet("")
        
        self.image_label.set_center_setting_mode(self.center_setting_mode)
        
        if self.center_setting_mode:
            self.set_center_btn.setStyleSheet("background-color: #495057; color: white; border: 2px solid #343a40;")
            
            #  ДАЄМО ФОКУС ЗОБРАЖЕННЮ ДЛЯ ЗБЕРЕЖЕННЯ ФУНКЦІОНАЛЬНОСТІ МИШІ
            if hasattr(self, 'image_label'):
                self.image_label.setFocus()
                # Встановлюємо политику фокусу
                self.image_label.setFocusPolicy(Qt.StrongFocus)
            
            #  ПОКАЗУЄМО ЗУМ НА ПОТОЧНІЙ ПОЗИЦІЇ ЦЕНТРУ
            if hasattr(self, 'image_label') and self.processor:
                self.image_label.zoom_widget.show_zoom()
                self.image_label.zoom_widget.update_cursor_position(
                    self.processor.center_x, self.processor.center_y
                )
            self.add_result("🎯 Режим центру: ←→↑↓ для переміщення, Esc для виходу")
            self.add_result("   Shift+стрілка = швидше, Ctrl+стрілка = точніше")
        else:
            self.set_center_btn.setStyleSheet("")
            #  ВІДНОВЛЮЄМО СТАНДАРТНУ ПОЛІТИКУ ФОКУСУ
            if hasattr(self, 'image_label'):
                self.image_label.setFocusPolicy(Qt.ClickFocus)
            #  ХОВАЄМО ЗУМ
            if hasattr(self, 'image_label'):
                self.image_label.zoom_widget.hide_zoom()

    def toggle_scale_edge_mode(self):
        self.scale_edge_mode = self.scale_edge_btn.isChecked()
        
        if self.scale_edge_mode and self.center_setting_mode:
            self.center_setting_mode = False
            self.set_center_btn.setChecked(False)
            self.set_center_btn.setStyleSheet("")
        
        self.image_label.set_scale_edge_mode(self.scale_edge_mode)
        
        if self.scale_edge_mode:
            self.scale_edge_btn.setStyleSheet("background-color: #495057; color: white; border: 2px solid #343a40;")
            
            #  ДАЄМО ФОКУС ЗОБРАЖЕННЮ ДЛЯ ЗБЕРЕЖЕННЯ ФУНКЦІОНАЛЬНОСТІ МИШІ
            if hasattr(self, 'image_label'):
                self.image_label.setFocus()
                # Встановлюємо политику фокусу
                self.image_label.setFocusPolicy(Qt.StrongFocus)
            
            #  ПОКАЗУЄМО ЗУМ НА ПОЗИЦІЇ КРАЮ МАСШТАБУ
            if hasattr(self, 'image_label') and self.processor:
                if self.scale_edge_point:
                    self.image_label.zoom_widget.show_zoom()
                    self.image_label.zoom_widget.update_cursor_position(
                        self.scale_edge_point['x'], self.scale_edge_point['y']
                    )
                else:
                    # Якщо немає точки, показуємо біля центру
                    self.image_label.zoom_widget.show_zoom()
                    self.image_label.zoom_widget.update_cursor_position(
                        self.processor.center_x + 50, self.processor.center_y + 50
                    )
            self.add_result("📏 Режим масштабу: ←→↑↓ для переміщення, Esc для виходу")
            self.add_result("   Shift+стрілка = швидше, Ctrl+стрілка = точніше")
        else:
            self.scale_edge_btn.setStyleSheet("")
            #  ВІДНОВЛЮЄМО СТАНДАРТНУ ПОЛІТИКУ ФОКУСУ
            if hasattr(self, 'image_label'):
                self.image_label.setFocusPolicy(Qt.ClickFocus)
            #  ХОВАЄМО ЗУМ
            if hasattr(self, 'image_label'):
                self.image_label.zoom_widget.hide_zoom()
    
    def set_center_point(self, x, y):
        """Встановлення центру з збереженням налаштувань та зумом"""
        if not self.processor:
            return
        
        x = max(0, min(x, self.processor.image.width - 1))
        y = max(0, min(y, self.processor.image.height - 1))
        
        current_center_x = self.processor.center_x
        current_center_y = self.processor.center_y
        
        dx = x - current_center_x
        dy = y - current_center_y
        
        self.processor.move_center(dx, dy)
        self.save_current_grid_settings()
        
        #  ОНОВЛЮЄМО ЗУМ НА НОВІЙ ПОЗИЦІЇ
        if hasattr(self, 'image_label'):
            self.image_label.zoom_widget.update_cursor_position(
                self.processor.center_x, self.processor.center_y
            )
        
        self.display_image()
        
        if self.current_click:
            azimuth, range_val = self.calculate_azimuth_range(
                self.current_click['x'], self.current_click['y']
            )
            self.current_click['azimuth'] = azimuth
            self.current_click['range'] = range_val
        
        self.update_results_display()
        self.update_report_data()
        self.add_result(f"Центр встановлено: ({self.processor.center_x}, {self.processor.center_y})")

    
    def set_scale_edge_point(self, x, y):
        """Встановлення scale edge з збереженням налаштувань та зумом"""
        x = max(0, min(x, self.processor.image.width - 1))
        y = max(0, min(y, self.processor.image.height - 1))
        
        dx = x - self.processor.center_x
        dy = y - self.processor.center_y
        distance = math.sqrt(dx*dx + dy*dy)
        
        self.scale_edge_point = {'x': x, 'y': y}
        self.custom_scale_distance = distance
        self.save_current_grid_settings()
        
        #  ОНОВЛЮЄМО ЗУМ НА НОВІЙ ПОЗИЦІЇ
        if hasattr(self, 'image_label'):
            self.image_label.zoom_widget.update_cursor_position(x, y)
        
        self.display_image()
        
        if self.current_click:
            azimuth, range_val = self.calculate_azimuth_range(
                self.current_click['x'], self.current_click['y']
            )
            self.current_click['azimuth'] = azimuth
            self.current_click['range'] = range_val
        
        self.update_results_display()
        self.update_report_data()
        self.add_result(f"Край масштабу: ({x}, {y}) | Відстань: {distance:.1f}px")

    def update_target_number(self, text):
        self.current_target_number = text
    
    def update_height(self, text):
        self.current_height = text
    
    def update_obstacles(self, text):
        self.current_obstacles = text
    
    def update_detection(self, text):
        self.current_detection = text
    
    def update_report_data(self):
        """Оновлення даних в правій панелі (азимут, дальність, масштаб)"""
        if not self.processor:
            self.auto_azimuth_label.setText("β - --°")
            self.auto_distance_label.setText("D - -- км")
            self.auto_scale_label.setText("M = --")
            return
            
        if self.current_click:
            azimuth = self.current_click['azimuth']
            distance = self.current_click['range']
            scale = int(self.scale_combo.currentText())
            
            self.auto_azimuth_label.setText(f"β - {azimuth:.0f}ᴼ")
            self.auto_distance_label.setText(f"D - {distance:.0f} км")  # ← ЗМІНЕНО .1f на .0f
            self.auto_scale_label.setText(f"M = {scale}")
        else:
            self.auto_azimuth_label.setText("β - --ᴼ")
            self.auto_distance_label.setText("D - -- км")
            
            if hasattr(self, 'scale_combo'):
                scale = int(self.scale_combo.currentText())
                self.auto_scale_label.setText(f"M = {scale}")
            else:
                self.auto_scale_label.setText("M = --")

    def update_results_display(self):
        self.results_text.clear()
        
        if self.processor:
            self.add_result(self.tr("image_info").format(name=os.path.basename(self.current_image_path)))
            self.add_result(self.tr("size").format(width=self.processor.image.width, 
                                                  height=self.processor.image.height))
            self.add_result(self.tr("scale_info").format(scale=int(self.scale_combo.currentText())))
            self.add_result(self.tr("center_info").format(x=self.processor.center_x, 
                                                         y=self.processor.center_y))
            
            if self.custom_scale_distance:
                self.add_result(f"Custom scale edge: {self.custom_scale_distance:.1f} px = {self.scale_combo.currentText()} units")
            else:
                bottom_distance = self.processor.image.height - self.processor.center_y
                self.add_result(self.tr("bottom_edge").format(scale=int(self.scale_combo.currentText())))
                self.add_result(self.tr("pixels_south").format(pixels=bottom_distance))
            self.add_result("")
        
        if self.current_click:
            self.add_result(self.tr("analysis_point"))
            self.add_result(f"{self.tr('position')}: ({self.current_click['x']}, {self.current_click['y']})")
            self.add_result(f"{self.tr('azimuth')}: {self.current_click['azimuth']:.0f}ᴼ")
            self.add_result(f"{self.tr('range')}: {self.current_click['range']:.0f} км")  # ← ЗМІНЕНО .0f
            self.add_result("")
            self.add_result(self.tr("click_to_place"))
            self.add_result(self.tr("drag_to_move"))
            self.add_result(self.tr("line_connects"))
        else:
            self.add_result(self.tr("click_on_image"))

    def update_scale(self):
        """Оновлення масштабу з збереженням налаштувань"""
        if self.processor:
            new_scale = int(self.scale_combo.currentText())
            
            # ЗБЕРЕГТИ налаштування сітки
            self.save_current_grid_settings()
            
            if self.current_click:
                azimuth, range_val = self.calculate_azimuth_range(
                    self.current_click['x'], self.current_click['y']
                )
                self.current_click['azimuth'] = azimuth
                self.current_click['range'] = range_val
            
            self.update_results_display()
            self.update_report_data()
            self.add_result(self.tr("scale_updated").format(scale=new_scale))
            self.add_result("Grid settings saved for next images")
    
    def move_center(self, dx, dy):
        if not self.processor:
            QMessageBox.warning(self, self.tr("warning"), self.tr("no_image_first"))
            return
        
        self.processor.move_center(dx, dy)
        
        if self.current_click:
            azimuth, range_val = self.calculate_azimuth_range(
                self.current_click['x'], self.current_click['y']
            )
            self.current_click['azimuth'] = azimuth
            self.current_click['range'] = range_val
        
        self.display_image()
        self.update_results_display()
        self.update_report_data()
        self.add_result(self.tr("center_moved").format(x=self.processor.center_x, 
                                                    y=self.processor.center_y))
    
        if self.image_label.zoom_widget.isVisible():
            self.image_label.zoom_widget.update_cursor_position(self.processor.center_x, self.processor.center_y)
    
    def clear_results(self):
        self.current_click = None
        if self.processor:
            self.display_image()
        self.update_results_display()
        self.update_report_data()

    def save_current_image(self):
        if not self.processor:
            QMessageBox.warning(self, self.tr("warning"), self.tr("no_image_first"))
            return
        
        if not self.current_click:
            QMessageBox.warning(self, self.tr("warning"), self.tr("no_analysis_point"))
            return
        
        file_path, _ = QFileDialog.getSaveFileName(
            self, self.tr("save_processed_image"), "",
            f"{self.tr('jpeg_files')} (*.jpg);;{self.tr('png_files')} (*.png);;{self.tr('all_files')} (*.*)"
        )
        
        if file_path:
            try:
                final_image = self.processor.image.copy()
                
                if file_path.lower().endswith(('.jpg', '.jpeg')):
                    if final_image.mode != 'RGB':
                        if final_image.mode == 'RGBA':
                            rgb_image = Image.new('RGB', final_image.size, (255, 255, 255))
                            rgb_image.paste(final_image, mask=final_image.split()[-1])
                            final_image = rgb_image
                        else:
                            final_image = final_image.convert('RGB')
                
                draw = ImageDraw.Draw(final_image)
                
                draw.line([
                    (self.current_click['x'], self.current_click['y']), 
                    (final_image.width - 1, 0)
                ], fill='blue', width=3)
                
                if file_path.lower().endswith('.png'):
                    final_image.save(file_path, 'PNG')
                else:
                    final_image.save(file_path, 'JPEG', quality=95)
                
                self.add_result(f"{self.tr('saved')}: {os.path.basename(file_path)}")
                self.add_result(self.tr("saved_clean"))
                QMessageBox.information(self, self.tr("success"), self.tr("image_saved_successfully"))
                
            except Exception as e:
                QMessageBox.critical(self, self.tr("error"), 
                                   self.tr("could_not_save").format(error=str(e)))

    def save_current_image_data(self):
        if not self.processor or not self.current_click:
            QMessageBox.warning(self, "Warning", "No image or analysis point to save")
            return False
        
        image_data = {
            'image_path': self.current_image_path,
            'image_name': os.path.basename(self.current_image_path),
            'analysis_point': {
                'x': self.current_click['x'],
                'y': self.current_click['y'],
                'azimuth': self.current_click['azimuth'],
                'range': self.current_click['range']
            },
            'target_data': {
                'number': self.current_target_number,
                'height': self.current_height,
                'obstacles': self.current_obstacles,
                'detection': self.current_detection,
                'scale': int(self.scale_combo.currentText())
            },
            'processor_settings': {
                'center_x': self.processor.center_x,
                'center_y': self.processor.center_y,
                'scale_edge_point': self.scale_edge_point,
                'custom_scale_distance': self.custom_scale_distance
            }
        }

        # ===== ДОДАЄМО ДАНІ ОПИСУ РЛС =====
        radar_data = self.get_radar_description_data()
        if radar_data and radar_data['enabled']:
            image_data['radar_description'] = radar_data
            self.add_result("📡 Дані РЛС додано до зображення")
            print(f"✓ Radar data included: {radar_data}")
        else:
            print("ℹ No radar data (disabled or empty)")
        
        existing_index = -1
        for i, saved_data in enumerate(self.processed_images):
            if saved_data['image_path'] == self.current_image_path:
                existing_index = i
                break
        
        if existing_index >= 0:
            self.processed_images[existing_index] = image_data
            self.add_result(f"Updated data for: {image_data['image_name']}")
        else:
            self.processed_images.append(image_data)
            self.add_result(f"Saved data for: {image_data['image_name']}")

        # Позначаємо зображення як оброблене
        self.processed_image_paths.add(self.current_image_path)
        
        # Оновлюємо мініатюри
        self.update_thumbnail_processed_status()
        
        self.add_result(f"Total processed images: {len(self.processed_images)}")
        return True

    def update_thumbnail_processed_status(self):
        """Оновлення візуального статусу оброблених мініатюр"""
        if hasattr(self, 'thumbnail_widget'):
            self.thumbnail_widget.mark_image_as_processed(self.current_image_path)

    def cancel_current_changes(self):
        """Скасування змін для поточного зображення"""
        if not self.current_image_path:
            QMessageBox.warning(self, "Попередження", "Немає завантаженого зображення")
            return
        
        # Видаляємо з оброблених зображень якщо є
        self.processed_images = [img for img in self.processed_images 
                               if img['image_path'] != self.current_image_path]
        
        # Видаляємо з множини оброблених шляхів
        self.processed_image_paths.discard(self.current_image_path)
        
        # Оновлюємо візуальний стан мініатюри
        if hasattr(self, 'thumbnail_widget'):
            self.thumbnail_widget.mark_image_as_unprocessed(self.current_image_path)
        
        # Очищуємо поточну точку аналізу
        self.current_click = None
        
        # Скидаємо дані форми до базових значень
        self.reset_form_data()
        
        # Оновлюємо відображення
        self.display_image()
        self.update_results_display()
        self.update_report_data()
        
        self.add_result(f"✗ Скасовано зміни для: {os.path.basename(self.current_image_path)}")
        self.add_result(f"Залишилось оброблених зображень: {len(self.processed_images)}")

    def clear_all_changes(self):
        """Очищення всіх оброблених зображень"""
        if not self.processed_images:
            QMessageBox.information(self, "Інформація", "Немає оброблених зображень для очищення")
            return
        
        # Підтвердження дії
        reply = QMessageBox.question(self, "Підтвердження", 
                                   f"Ви впевнені, що хочете очистити всі {len(self.processed_images)} оброблених зображень?",
                                   QMessageBox.Yes | QMessageBox.No,
                                   QMessageBox.No)
        
        if reply == QMessageBox.Yes:
            # Очищуємо всі дані
            self.processed_images.clear()
            self.processed_image_paths.clear()
            
            # Оновлюємо всі мініатюри
            if hasattr(self, 'thumbnail_widget'):
                self.thumbnail_widget.clear_all_processed_status()
            
            # Очищуємо поточну точку аналізу
            self.current_click = None
            
            # Скидаємо дані форми
            self.reset_form_data()
            
            # Оновлюємо відображення
            self.display_image()
            self.update_results_display()
            self.update_report_data()
            
            self.add_result("🗑️ Очищено всі оброблені зображення")
            self.add_result("Всі зображення повернуто до необробленого стану")

    def reset_form_data(self):
        """Скидання даних форми до базових значень"""
        # Скидаємо номер цілі до базового
        self.current_target_number = "0001"
        self.target_number_input.setText(self.current_target_number)
        
        # Скидаємо висоту
        self.current_height = "0.0"
        self.height_input.setText(self.current_height)
        
        # Скидаємо випадні списки до перших значень
        self.obstacles_combo.setCurrentIndex(0)
        self.detection_combo.setCurrentIndex(0)
        
        # Оновлюємо внутрішні змінні
        self.current_obstacles = self.obstacles_combo.currentText()
        self.current_detection = self.detection_combo.currentText()

# ===== ГОЛОВНА ФУНКЦІЯ ТА ЗАПУСК ПРОГРАМИ =====

def main():
    app = QApplication(sys.argv)
    window = AzimuthGUI()
    window.show()
    sys.exit(app.exec_())

if __name__ == "__main__":
    main()