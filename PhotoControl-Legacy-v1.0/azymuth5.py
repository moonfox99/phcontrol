#!/usr/bin/env python3
"""
Azimuth Image Processor - PyQt5 Version with Ukrainian Translation
Professional interface for processing azimuth grid images
"""

import sys
import os
import math
import tempfile
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QHBoxLayout, 
                             QVBoxLayout, QGridLayout, QPushButton, QLabel, 
                             QComboBox, QTextEdit, QScrollArea, QFrame,
                             QFileDialog, QMessageBox, QSplitter, QToolTip, QLineEdit,
                             QCheckBox, QDateEdit, QSizePolicy)
from PyQt5.QtCore import Qt, QSize, pyqtSignal, QTimer, QPoint, QDate  # Додати QDate
from PyQt5.QtGui import QPixmap, QFont
from PIL import Image, ImageDraw

from translations import Translations
from widgets import ClickableLabel, VerticalThumbnailWidget
from image_processor import AzimuthImageProcessor

try:
    from docx import Document
    from docx.shared import Inches, Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_BREAK
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.section import WD_SECTION_START
    from docx.oxml.shared import OxmlElement, qn
    from docx.oxml.ns import qn
    from docx.oxml import parse_xml
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ===== НОВІ КОНСТАНТИ ДЛЯ ТАБЛИЦЬ 9x9 =====

ALBUM_LAYOUT = {
    # Розміри сторінки A4 в міліметрах
    'PAGE_WIDTH': 210,
    'PAGE_HEIGHT': 297,
    
    # Стандартні поля для титульної та описової сторінок
    'STANDARD_LEFT_MARGIN': 25,
    'STANDARD_RIGHT_MARGIN': 25, 
    'STANDARD_TOP_MARGIN': 10,
    'STANDARD_BOTTOM_MARGIN': 10,

    # НОВІ ПОЛЯ ДЛЯ СТОРІНОК З ТАБЛИЦЯМИ (20мм зверху, 5мм справа/знизу, 2.5мм зліва)
    'TABLE_PAGES_LEFT_MARGIN': 2.5,    # 2.5мм зліва
    'TABLE_PAGES_RIGHT_MARGIN': 5,     # 5мм справа
    'TABLE_PAGES_TOP_MARGIN': 20,      # 20мм зверху
    'TABLE_PAGES_BOTTOM_MARGIN': 5,    # 5мм знизу
    
    # НОВІ ПАРАМЕТРИ ДЛЯ ОКРЕМИХ ТАБЛИЦЬ
    'TABLE_ROWS': 1,  # Тепер тільки 1 рядок на таблицю
    'TABLE_COLS': 3,
    
    # Висоти елементів (в мм)
    'PARAGRAPH_HEIGHT': 5,    # Параграф-розділювач 10мм
    'TABLE_HEIGHT': 130,       # Висота таблиці 130мм
    
    # Ширини колонок (в мм) - залишаються без змін
    'COL_1_WIDTH': 25,      # Перша колонка
    'COL_2_WIDTH': 150,     # Друга колонка (для зображень)
    'COL_3_WIDTH': 30,      # Третя колонка
    
    # Параметри зображень - без змін
    'IMAGE_ASPECT_RATIO': 15/13,
    'IMAGE_WIDTH_CM': 14,  # Ширина зображення в см
}


# Константи для титульної сторінки (БЕЗ ЗМІН)
TITLE_PAGE = {
    'MAIN_FONT_SIZE': 28,
    'FOOTER_FONT_SIZE': 16,
    'TOP_SPACING': 3,
    'FOOTER_SPACING': 10,
    'FOOTER_LEFT_INDENT': 3.75,
}

# Константи для сторінки опису (БЕЗ ЗМІН)
DESCRIPTION_PAGE = {
    'HEADING_FONT_SIZE': 22,
    'TABLE_HEIGHT': 1.7,
    'COL_WIDTHS': [0.84, 3.43, 2.69, 2.54, 2.77, 2.27],
    'HEADERS': [
        "№ пп",
        "Назва документа", 
        "Обліковий номер",
        "Кількість аркушів (знімків)",
        "Гриф секретності",
        "Примітка"
    ]
}

def mm_to_cm(mm):
    """Перетворення міліметрів в сантиметри для docx"""
    return mm / 10.0

def format_ukrainian_date(date_obj):
    """Форматування дати по-українськи для титульної сторінки"""
    try:
        if isinstance(date_obj, str):
            from datetime import datetime
            date_obj = datetime.strptime(date_obj, "%Y-%m-%d")
        
        months = {
            1: 'січня', 2: 'лютого', 3: 'березня', 4: 'квітня',
            5: 'травня', 6: 'червня', 7: 'липня', 8: 'серпня',
            9: 'вересня', 10: 'жовтня', 11: 'листопада', 12: 'грудня'
        }
        
        day = date_obj.day
        month = months[date_obj.month]
        year = date_obj.year
        
        return f'"{day}" {month} {year} року'
    except:
        return "дата не вказана"

def get_default_unit_info():
    """Базова інформація про частину"""
    return "А0000"

def get_default_commander_info():
    """Базова інформація про командира"""
    return {
        'rank': 'полковник',
        'name': 'П.П. ПЕТРЕНКО'
    }

def get_default_responsible_info():
    """Базова інформація про відповідальних осіб"""
    return {
        'responsible': {'rank': 'майор', 'name': 'І.І. ІВАНЕНКО'},
        'assistant': {'rank': 'старший лейтенант', 'name': 'С.С. СИДОРЕНКО'},
        'secretary': {'rank': 'старший лейтенант', 'name': 'С.С. СИДОРЕНКО'}
    }

def set_a4_page_format(section):
    """Встановлення точного формату A4 для секції документу"""
    try:
        from docx.shared import Cm
        from docx.enum.section import WD_ORIENTATION
        
        # Точні розміри A4 згідно з ISO 216
        section.page_width = Cm(21.0)   # 210мм = 21.0см
        section.page_height = Cm(29.7)  # 297мм = 29.7см
        section.orientation = WD_ORIENTATION.PORTRAIT
        
        print(f"✓ A4 format set: {section.page_width.cm:.1f} x {section.page_height.cm:.1f} cm")
        
    except Exception as e:
        print(f"✗ Error setting A4 format: {e}")

def create_complete_album(processed_images, title_data, file_path):
    """Створення повного альбому з новою структурою"""
    try:
        doc = Document()
        
        print("=== Creating Complete Album with New Structure ===")
        
        # 1. ТИТУЛЬНА СТОРІНКА (стандартні поля)
        print("1. Creating title page with standard margins...")
        section = doc.sections[0]
        set_a4_page_format(section)
        section.left_margin = Cm(mm_to_cm(ALBUM_LAYOUT['STANDARD_LEFT_MARGIN']))
        section.right_margin = Cm(mm_to_cm(ALBUM_LAYOUT['STANDARD_RIGHT_MARGIN']))
        section.top_margin = Cm(mm_to_cm(ALBUM_LAYOUT['STANDARD_TOP_MARGIN']))
        section.bottom_margin = Cm(mm_to_cm(ALBUM_LAYOUT['STANDARD_BOTTOM_MARGIN']))
        create_title_page(doc, title_data)
        
        # 2. СТОРІНКА ОПИСУ (стандартні поля)
        print("2. Creating description page...")
        create_description_page(doc)
        
        # 3. СТОРІНКИ З НОВОЮ СТРУКТУРОЮ
        print("3. Creating table pages with new structure...")
        create_new_structure_pages(doc, processed_images)
        
        doc.save(file_path)
        print(f"✓ Complete album with new structure saved: {file_path}")
        
        return True
        
    except Exception as e:
        print(f"✗ Error creating complete album: {e}")
        return False
    
def create_new_structure_pages(doc, processed_images):
    """Створення сторінок з новою структурою та таблицями БЕЗ лівого поля"""
    try:
        if not processed_images:
            print("No processed images to add")
            return
        
        # Додаємо нову секцію для таблиць
        table_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
        
        # Встановлюємо A4 формат
        set_a4_page_format(table_section)
        
        # ОНОВЛЕНІ ПОЛЯ: 0мм зліва (таблиці впритул до краю!), 20мм зверху, 5мм справа/знизу
        table_section.left_margin = Cm(mm_to_cm(ALBUM_LAYOUT['TABLE_PAGES_LEFT_MARGIN']))  # 0мм!
        table_section.right_margin = Cm(mm_to_cm(ALBUM_LAYOUT['TABLE_PAGES_RIGHT_MARGIN']))
        table_section.top_margin = Cm(mm_to_cm(ALBUM_LAYOUT['TABLE_PAGES_TOP_MARGIN']))
        table_section.bottom_margin = Cm(mm_to_cm(ALBUM_LAYOUT['TABLE_PAGES_BOTTOM_MARGIN']))
        
        print(f"✓ NEW margins set: left=0mm (no margin!), top=20mm, right=5mm, bottom=5mm")
        
        # Обробляємо зображення парами (по 2 на сторінку)
        for i in range(0, len(processed_images), 2):
            first_image = processed_images[i]
            second_image = processed_images[i + 1] if i + 1 < len(processed_images) else None
            
            print(f"\n=== Creating page for images {i+1}-{i+2 if second_image else i+1} ===")
            
            # НОВА СТРУКТУРА СТОРІНКИ:
            
            # 1. Параграф-розділювач 5мм
            create_spacer_paragraph(doc, ALBUM_LAYOUT['PARAGRAPH_HEIGHT'])
            print("✓ Added top spacer paragraph (5mm)")
            
            # 2. Перша таблиця 130мм (БЕЗ лівого відступу)
            create_single_image_table(doc, first_image)
            print("✓ Added first table (130mm, aligned to LEFT EDGE)")
            
            # 3. Параграф-розділювач 5мм
            create_spacer_paragraph(doc, ALBUM_LAYOUT['PARAGRAPH_HEIGHT'])
            print("✓ Added middle spacer paragraph (5mm)")
            
            # 4. Друга таблиця 130мм (якщо є друге зображення)
            if second_image:
                create_single_image_table(doc, second_image)
                print("✓ Added second table (130mm, aligned to LEFT EDGE)")
            else:
                # Якщо немає другого зображення, додаємо порожню таблицю
                create_empty_table_placeholder(doc)
                print("✓ Added empty table placeholder")
            
            # Розрив сторінки тільки якщо НЕ остання пара зображень
            #if i + 2 < len(processed_images):
            #    doc.add_page_break()
            #    print("📄 Added page break")
        
        total_pages = (len(processed_images) + 1) // 2
        print(f"\n✅ Final result: {len(processed_images)} images in {total_pages} pages with LEFT-ALIGNED tables")
        
    except Exception as e:
        print(f"✗ Error creating new structure pages: {e}")

def create_spacer_paragraph(doc, height_mm):
    """Створення параграфа-розділювача з точною висотою"""
    try:
        para = doc.add_paragraph()
        para.text = ""  # Порожній параграф
        
        # Встановлюємо точну висоту параграфа
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = Cm(mm_to_cm(height_mm))
        para.paragraph_format.line_spacing_rule = 0  # Точна висота
        
        print(f"✓ Created spacer paragraph: {height_mm}mm")
        
    except Exception as e:
        print(f"✗ Error creating spacer paragraph: {e}")

def create_single_image_table(doc, image_data):
    """Створення таблиці 1x3 для одного зображення"""
    try:
        print(f"🔨 Creating single image table...")
        
        # Створюємо таблицю 1x3 (1 рядок, 3 колонки)
        table = doc.add_table(rows=1, cols=ALBUM_LAYOUT['TABLE_COLS'])
        table.style = None
        
        # Встановлюємо точну ширину таблиці
        set_table_width(table)
        
        # Налаштовуємо колонки (ширини)
        setup_column_widths(table)
        
        # Налаштовуємо єдиний рядок з висотою 130мм
        setup_image_row(table, 0, image_data)
        print(f"✅ Single image table created successfully")
        
        return table
        
    except Exception as e:
        print(f"❌ Error creating single image table: {e}")
        return None

def create_empty_table_placeholder(doc):
    """Створення порожньої таблиці-заповнювача для випадків з непарною кількістю зображень"""
    try:
        table = doc.add_table(rows=1, cols=ALBUM_LAYOUT['TABLE_COLS'])
        table.style = None
        
        set_table_width(table)
        setup_column_widths(table)
        
        # Налаштовуємо порожній рядок
        row = table.rows[0]
        row.height = Cm(mm_to_cm(ALBUM_LAYOUT['TABLE_HEIGHT']))  # 130мм
        
        for col_idx in range(ALBUM_LAYOUT['TABLE_COLS']):
            cell = row.cells[col_idx]
            cell.text = ""
            set_cell_background(cell, "FFFFFF")
            # Прозорі границі для порожньої таблиці
            set_cell_borders(cell, top=False, bottom=False, left=False, right=False)
        
        print(f"✓ Created empty table placeholder")
        
    except Exception as e:
        print(f"✗ Error creating empty table placeholder: {e}")

def set_table_width(table):
    """Встановлення фіксованої ширини таблиці та ЛІВОГО вирівнювання"""
    try:
        tbl = table._tbl
        tblPr = tbl.tblPr
        
        # Фіксований layout
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)
        
        # ТОЧНА ширина таблиці: 30+145+30 = 205мм
        total_width_mm = ALBUM_LAYOUT['COL_1_WIDTH'] + ALBUM_LAYOUT['COL_2_WIDTH'] + ALBUM_LAYOUT['COL_3_WIDTH']
        width_dxa = int(total_width_mm * 56.7)  # 205мм в DXA
        
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:type'), 'dxa')
        tblW.set(qn('w:w'), str(width_dxa))
        tblPr.append(tblW)
        
        # ВАЖЛИВО: Вирівнювання по ЛІВОМУ краю (таблиця притиснута до лівого краю аркуша)
        tblJc = OxmlElement('w:jc')
        tblJc.set(qn('w:val'), 'left')  # Ліве вирівнювання
        tblPr.append(tblJc)
        
        # ДОДАТКОВО: Забезпечуємо що таблиця починається з самого лівого краю
        tblInd = OxmlElement('w:tblInd')
        tblInd.set(qn('w:w'), '0')  # Нульовий відступ зліва
        tblInd.set(qn('w:type'), 'dxa')
        tblPr.append(tblInd)
        
        print(f"✓ Table: {total_width_mm}mm width, LEFT-aligned to page edge (no left margin)")
        
    except Exception as e:
        print(f"✗ Error setting table width and alignment: {e}")
def setup_column_widths(table):
    """Налаштування ширин колонок для таблиці 1x3"""
    try:
        col_1_width = ALBUM_LAYOUT['COL_1_WIDTH']      # 30мм - Індикатор ЗРЛ
        col_2_width = ALBUM_LAYOUT['COL_2_WIDTH']      # 145мм - зображення  
        col_3_width = ALBUM_LAYOUT['COL_3_WIDTH']      # 30мм - дані
        
        # Встановлюємо ширини колонок
        for col_idx in range(ALBUM_LAYOUT['TABLE_COLS']):
            row = table.rows[0]  # Тільки один рядок
            cell = row.cells[col_idx]
            
            if col_idx == 0:
                set_cell_width_mm(cell, col_1_width)    # 30мм
            elif col_idx == 1:
                set_cell_width_mm(cell, col_2_width)    # 145мм
            elif col_idx == 2:
                set_cell_width_mm(cell, col_3_width)    # 30мм
        
        print(f"✓ Column widths set: 30mm, 145mm, 30mm")
        
    except Exception as e:
        print(f"✗ Error setting column widths: {e}")

def setup_image_row(table, row_idx, image_data):
    """Налаштування рядка з зображенням (висота 130мм)"""
    try:
        row = table.rows[row_idx]
        
        # Встановлюємо висоту рядка 130мм
        row.height = Cm(mm_to_cm(ALBUM_LAYOUT['TABLE_HEIGHT']))  # 130мм
        
        # Фіксована висота рядка
        trPr = row._tr.trPr
        if trPr is None:
            trPr = OxmlElement('w:trPr')
            row._tr.append(trPr)
        
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(int(130 * 56.7)))  # 130мм в DXA
        trHeight.set(qn('w:hRule'), 'exact')  # ТОЧНА висота
        trPr.append(trHeight)

        # Налаштовуємо комірки
        for col_idx in range(ALBUM_LAYOUT['TABLE_COLS']):
            cell = row.cells[col_idx]
            
            if col_idx == 0:
                setup_first_cell(cell, True)  # Показуємо границі
            elif col_idx == 1:
                setup_image_cell(cell, image_data, True)  # Показуємо границі
            elif col_idx == 2:
                setup_data_cell(cell, image_data, True)  # Показуємо границі
        
        print(f"✓ Image row {row_idx + 1} configured (130mm height)")
        
    except Exception as e:
        print(f"✗ Error setting up image row {row_idx}: {e}")

# ===== ДОПОМІЖНІ ФУНКЦІЇ (використовуються з попередніх версій) =====

def setup_first_cell(cell, show_borders):
    """Налаштування першої комірки (ширина 30мм, текст 'Індикатор ЗРЛ')"""
    try:
        # Вертикальне центрування
        set_cell_vertical_center(cell)
        
        # Текст
        cell.text = ""

        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = para.add_run("Індикатор ЗРЛ")
        run.font.name = 'Arial'
        run.font.size = Pt(12)

        # 🎯 ДОДАЙТЕ ВНУТРІШНІ ПОЛЯ КОМІРКИ
        tc = cell._tc
        tcPr = tc.tcPr
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            tc.append(tcPr)
        
        tcMar = OxmlElement('w:tcMar')
        # Налаштування відступів для лівої комірки
        margins = {
            'top': '0',        # 0 зверху
            'left': '140',     # 140 DXA = ~2.5мм зліва
            'bottom': '0',     # 0 знизу  
            'right': '0'       # 0 справа
        }
        
        for margin_name, value in margins.items():
            margin = OxmlElement(f'w:{margin_name}')
            margin.set(qn('w:w'), value)
            margin.set(qn('w:type'), 'dxa')
            tcMar.append(margin)
        
        tcPr.append(tcMar)
        
        # Границі: лише права (якщо показуємо границі)
        if show_borders:
            set_cell_borders(cell, top=False, bottom=False, left=False, right=True)
        else:
            set_cell_borders(cell, top=False, bottom=False, left=False, right=False)
        
    except Exception as e:
        print(f"✗ Error configuring first cell: {e}")

def setup_image_cell(cell, image_data, show_borders):
    """Налаштування комірки з зображенням ВПРИТУЛ до меж, але не за них"""
    try:
        # Вертикальне центрування
        set_cell_vertical_center(cell)
        
        # Очищуємо комірку
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # НУЛЬОВІ ВІДСТУПИ У ПАРАГРАФА
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0) 
        para.paragraph_format.left_indent = Pt(0)
        para.paragraph_format.right_indent = Pt(0)
        para.paragraph_format.line_spacing = 1.0
        
        if image_data:
            # Створюємо оброблене зображення
            processed_image = create_processed_image_from_data(image_data)
            
            if processed_image:
                # Зберігаємо у тимчасовий файл
                with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as temp_file:
                    temp_path = temp_file.name
                    processed_image.save(temp_path, 'JPEG', quality=95)
                
                # 🎯 РОЗМІР ВПРИТУЛ ДО МЕЖ: товщина границі ≈ 0.5мм з кожної сторони
                # Відступаємо по 0.5мм від кожного краю, щоб НЕ перекривати границі
                
                border_thickness_mm = 1.0  # Товщина границі + невеликий буфер
                
                effective_width = ALBUM_LAYOUT['COL_2_WIDTH'] - border_thickness_mm    # 149мм
                effective_height = ALBUM_LAYOUT['TABLE_HEIGHT'] - border_thickness_mm  # 129мм
                
                # Додаємо зображення з розрахованими розмірами
                run = para.add_run()
                run.add_picture(temp_path, 
                               width=Cm(effective_width / 10.0),   # 14.9см
                               height=Cm(effective_height / 10.0)) # 12.9см
                
                # Видаляємо тимчасовий файл
                try:
                    os.remove(temp_path)
                except:
                    pass
                
                print(f"✓ Image sized to fit: {effective_width}mm x {effective_height}mm (within borders)")
        
        # 🎯 НУЛЬОВІ ПОЛЯ КОМІРКИ (зображення контролюється розміром)
        tc = cell._tc
        tcPr = tc.tcPr
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            tc.append(tcPr)
        
        tcMar = OxmlElement('w:tcMar')
        for margin_name in ['top', 'left', 'bottom', 'right']:
            margin = OxmlElement(f'w:{margin_name}')
            margin.set(qn('w:w'), '0')  # Нульові поля
            margin.set(qn('w:type'), 'dxa')
            tcMar.append(margin)
        tcPr.append(tcMar)
        
        # Границі
        if show_borders:
            set_cell_borders(cell, top=True, bottom=True, left=True, right=True)
        else:
            set_cell_borders(cell, top=False, bottom=False, left=False, right=False)
        
        # Білий фон
        set_cell_background(cell, "FFFFFF")
        
        print(f"✓ Image fits within cell borders with {border_thickness_mm}mm clearance")
        
    except Exception as e:
        print(f"✗ Error configuring image cell: {e}")
        
def setup_data_cell(cell, image_data, show_borders):
    """Налаштування комірки з даними з лівим відступом"""
    try:
        # Очищуємо комірку
        cell.text = ""

        # ВСТАНОВЛЮЄМО ВНУТРІШНІ ПОЛЯ КОМІРКИ ЗАМІСТЬ ВІДСТУПІВ ПАРАГРАФІВ
        tc = cell._tc
        tcPr = tc.tcPr
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            tc.append(tcPr)
        
        tcMar = OxmlElement('w:tcMar')
        # Налаштування відступів для правої комірки
        margins = {
            'left': '140',     # 140 DXA = ~2.5мм (замість 8pt відступу параграфа)
            'bottom': '0',     # 0мм знизу
        }
        
        for margin_name, value in margins.items():
            margin = OxmlElement(f'w:{margin_name}')
            margin.set(qn('w:w'), value)
            margin.set(qn('w:type'), 'dxa')
            tcMar.append(margin)
        
        tcPr.append(tcMar)
        
        if image_data:
            # Перший параграф з відступами
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(30)  # 30pt зверху
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.15
            
            # НОВИЙ ЛІВИЙ ВІДСТУП
            para.paragraph_format.left_indent = Pt(0)  # 0pt = ~0мм лівий відступ
            para.paragraph_format.right_indent = Pt(0)  # Без правого відступу
            
            # Отримуємо дані
            target_data = image_data['target_data']
            analysis_point = image_data['analysis_point']
            target_no = target_data['number']
            
            # 🎯 ФОРМУЄМО СПИСОК З ВИСОТОЮ
            data_lines = [
                (f"{target_no}", 12, True, True),
                (f"β – {analysis_point['azimuth']:.0f}ᴼ", 12, True, False),
                (f"D – {analysis_point['range']:.0f} км", 12, True, False),
            ]
            
            # 🎯 ДОДАЄМО ВИСОТУ ТІЛЬКИ ЯКЩО НЕ ДОРІВНЮЄ "0.0"
            height_value = target_data.get('height', '0.0')
            try:
                height_float = float(height_value)
                if height_float != 0.0:
                    data_lines.append((f"Н – {height_value} м", 12, True, False))
            except (ValueError, TypeError):
                # Якщо висота не число, але не "0" або "0.0" - додаємо
                if height_value not in ['0', '0.0', '', None]:
                    data_lines.append((f"Н – {height_value} м", 12, True, False))
            
            # Продовжуємо з рештою даних
            data_lines.extend([
                ("без перешкод", 9, True, False),
                (f"{target_data['detection']}", 9, True, False),
                (f"М – {target_data['scale']}", 9, True, False)
            ])
            
            # Додаємо рядки
            for i, (text, font_size, italic, underline) in enumerate(data_lines):
                if i > 0:
                    para = cell.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    para.paragraph_format.line_spacing = 1.15
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    
                    # Лівий відступ для всіх параграфів
                    para.paragraph_format.left_indent = Pt(0)
                    para.paragraph_format.right_indent = Pt(0)
                
                run = para.add_run(text)
                run.font.name = 'Arial'
                run.font.size = Pt(font_size)
                run.italic = italic
                run.underline = underline
        
        # Границі: всі сторони крім лівої (якщо показуємо границі)
        if show_borders:
            set_cell_borders(cell, top=True, bottom=True, left=False, right=True)
        else:
            set_cell_borders(cell, top=False, bottom=False, left=False, right=False)
        
    except Exception as e:
        print(f"✗ Error configuring data cell: {e}")


def set_cell_width_mm(cell, width_mm):
    """Встановлення ФІКСОВАНОЇ ширини комірки БЕЗ внутрішніх відступів"""
    try:
        width_dxa = int(width_mm * 56.7)
        
        tc = cell._tc
        tcPr = tc.tcPr
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            tc.append(tcPr)
        
        # ФІКСОВАНА ширина
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:type'), 'dxa')
        tcW.set(qn('w:w'), str(width_dxa))
        tcPr.append(tcW)
        
        # ЗАБОРОНА автоматичного підгону
        tcFitText = OxmlElement('w:tcFitText')
        tcFitText.set(qn('w:val'), '0')
        tcPr.append(tcFitText)
        
        # ВАЖЛИВО: НЕ додаємо tcMar тут для комірок з зображеннями
        # Відступи налаштовуються окремо в setup_9x9_image_cell
        
        print(f"✓ Fixed cell width: {width_mm}mm (NO internal margins)")
        
    except Exception as e:
        print(f"✗ Error setting fixed cell width: {e}")

def set_cell_vertical_center(cell):
    """Вертикальне центрування комірки"""
    try:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    except:
        try:
            tc = cell._tc
            tcPr = tc.tcPr
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.append(tcPr)
            
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            tcPr.append(vAlign)
        except Exception as e:
            print(f"Warning: Could not set vertical alignment: {e}")

def set_cell_borders(cell, top=False, bottom=False, left=False, right=False):
    """Налаштування рамок комірки"""
    try:
        tc = cell._tc
        tcPr = tc.tcPr
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            tc.append(tcPr)
        
        tcBorders = OxmlElement('w:tcBorders')
        
        border_settings = [
            ('top', top), ('bottom', bottom), ('left', left), ('right', right)
        ]
        
        for position, show_border in border_settings:
            border_element = OxmlElement(f'w:{position}')
            if show_border:
                border_element.set(qn('w:val'), 'single')
                border_element.set(qn('w:sz'), '4')
                border_element.set(qn('w:color'), '000000')
            else:
                border_element.set(qn('w:val'), 'none')
            
            tcBorders.append(border_element)
        
        tcPr.append(tcBorders)
        
    except Exception as e:
        print(f"Warning: Could not set cell borders: {e}")

def set_cell_background(cell, color_hex):
    """Встановлення кольору фону комірки"""
    try:
        tc = cell._tc
        tcPr = tc.tcPr
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            tc.append(tcPr)
        
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color_hex)
        tcPr.append(shd)
        
    except Exception as e:
        print(f"Warning: Could not set cell background: {e}")

def create_processed_image_from_data(image_data):
    """Створення обробленого зображення з даними"""
    try:
        with Image.open(image_data['image_path']) as original_image:
            if original_image.mode != 'RGB':
                if original_image.mode == 'RGBA':
                    rgb_image = Image.new('RGB', original_image.size, (255, 255, 255))
                    rgb_image.paste(original_image, mask=original_image.split()[-1])
                    final_image = rgb_image
                else:
                    final_image = original_image.convert('RGB')
            else:
                final_image = original_image.copy()
        
        draw = ImageDraw.Draw(final_image)
        analysis_point = image_data['analysis_point']
        
        # Розрахунок позиції кінця лінії (вирівняної з підкресленням номера цілі)
        image_width = final_image.width
        image_height = final_image.height
        
        # Права комірка починається приблизно з 85% ширини зображення
        right_cell_start_x = int(image_width * 0.85)
        
        # Позиція підкреслення номера цілі: 30pt зверху + висота тексту
        underline_y = int(image_height * 0.12)  # 12% висоти зображення

        # Кінцева точка лінії: посередині правої комірки, на рівні підкреслення
        end_x = image_width - 1  # Самий край
        end_y = underline_y      # На висоті підкреслення
        
        # Малюємо лінію від точки аналізу до розрахованої позиції
        draw.line([
            (analysis_point['x'], analysis_point['y']), 
            (end_x, end_y)
        ], fill='black', width=3)
        
        print(f"✓ Line drawn to calculated underline position: ({end_x}, {end_y})")
        return final_image
        
    except Exception as e:
        print(f"Error creating processed image: {e}")
        return None

# ===== ФУНКЦІЇ ДЛЯ ТИТУЛЬНОЇ СТОРІНКИ (БЕЗ ЗМІН) =====

def create_title_page(doc, title_data):
    """Створення титульної сторінки"""
    target_no = title_data['target_no']
    date = title_data['date']
    unit_info = title_data.get('unit_info', get_default_unit_info())
    commander_info = title_data.get('commander_info', get_default_commander_info())
    
    try:
        # Основний заголовок з центруванням
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_before = Cm(TITLE_PAGE['TOP_SPACING'])
        
        # Основний заголовок "АЛЬБОМ" 
        main_title_run = title_para.add_run("АЛЬБОМ")
        main_title_run.font.size = Pt(TITLE_PAGE['MAIN_FONT_SIZE'])
        main_title_run.font.bold = True
        main_title_run.font.name = 'Arial'
        
        # Підзаголовки з переносами рядків
        subtitle_lines = [
            "фотознімків засобів радіолокації",
            f"військової частини {unit_info}",
            format_ukrainian_date(date),
            f"по цілі №{target_no}"
        ]
        
        for line in subtitle_lines:
            title_para.add_run(f"\n{line}")
            # Налаштовуємо шрифт для кожного рядка
            for run in title_para.runs[-1:]:
                run.font.size = Pt(TITLE_PAGE['MAIN_FONT_SIZE'])
                run.font.bold = True
                run.font.name = 'Arial'
        
        # Нижня частина титульної сторінки з великим відступом
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        footer_para.paragraph_format.space_before = Cm(TITLE_PAGE['FOOTER_SPACING'])
        footer_para.paragraph_format.left_indent = Cm(TITLE_PAGE['FOOTER_LEFT_INDENT'])
        
        # Налаштування табуляції
        try:
            tab_stops = footer_para.paragraph_format.tab_stops
            tab_stops.clear_all()
            tab_stops.add_tab_stop(Cm(13), WD_TAB_ALIGNMENT.RIGHT)
        except Exception as tab_error:
            print(f"Warning: Could not set tab stops: {tab_error}")
        
        # Текст про командира
        commander_text_run = footer_para.add_run(f"Командир військової частини {unit_info}")
        commander_text_run.font.size = Pt(TITLE_PAGE['FOOTER_FONT_SIZE'])
        commander_text_run.font.bold = True
        commander_text_run.font.name = 'Arial'
        
        # Перенос рядка та звання командира
        rank_run = footer_para.add_run(f"\n{commander_info['rank']}")
        rank_run.font.size = Pt(TITLE_PAGE['FOOTER_FONT_SIZE'])
        rank_run.font.bold = True
        rank_run.font.name = 'Arial'
        
        # Табуляція та ім'я командира
        name_run = footer_para.add_run(f"\t{commander_info['name']}")
        name_run.font.size = Pt(TITLE_PAGE['FOOTER_FONT_SIZE'])
        name_run.font.bold = True
        name_run.font.name = 'Arial'
        
        # Розрив сторінки для переходу до наступної сторінки
        doc.add_page_break()
        
        print("Title page created successfully")
        
    except Exception as e:
        print(f"Error creating title page: {e}")
        # Додаємо базову титульну сторінку у випадку помилки
        fallback_title = doc.add_paragraph(f"АЛЬБОМ\nфотознімків засобів радіолокації\nпо цілі №{target_no}")
        fallback_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in fallback_title.runs:
            run.font.size = Pt(24)
            run.font.bold = True
            run.font.name = 'Arial'
        doc.add_page_break()

def create_description_page(doc):
    """Створення сторінки опису згідно з точними вимогами"""
    try:
        # 🎯 НАЛАШТУВАННЯ ПОЛІВ СТОРІНКИ
        # Ліве поле: 2см, Верхнє поле: 2см, Праве поле: 1см
        section = doc.sections[-1]  # Поточна секція
        section.left_margin = Cm(2.0)    # 2см ліве поле
        section.top_margin = Cm(2.0)     # 2см верхнє поле  
        section.right_margin = Cm(1.0)   # 1см праве поле
        section.bottom_margin = Cm(1.0)  # 1см нижнє поле (стандартне)
        
        print(f"✓ Description page margins: left=2cm, top=2cm, right=1cm")
        
        # 🎯 ЗАГОЛОВОК БЕЗ ВІДСТУПУ ЗВЕРХУ
        heading = doc.add_paragraph("Опис альбому фотознімків")
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # НУЛЬОВИЙ відступ зверху
        heading.paragraph_format.space_before = Pt(0)  # НЕ МАЄ відступу зверху
        heading.paragraph_format.space_after = Pt(0)   # НЕ МАЄ відступу знизу
        
        # Налаштування шрифту заголовку
        heading_run = heading.runs[0]
        heading_run.font.size = Pt(DESCRIPTION_PAGE['HEADING_FONT_SIZE'])  # 22pt
        heading_run.font.bold = True
        heading_run.font.name = 'Arial'
        
        print(f"✓ Title added with zero top spacing")
        
        # 🎯 ТАБЛИЦЯ БЕЗ РОЗРИВУ З ЗАГОЛОВКОМ + ВІДСТУП 1СМ ЗЛІВА
        table = doc.add_table(rows=6, cols=6)  # 1 заголовковий + 5 рядків даних
        table.style = None  # Без стилю за замовчуванням
        table.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # ВІДСТУП ТАБЛИЦІ 1СМ ВІД ЛІВОГО ПОЛЯ
        tbl = table._tbl
        tblPr = tbl.tblPr
        
        # Встановлюємо відступ таблиці зліва
        tblInd = OxmlElement('w:tblInd')
        tblInd.set(qn('w:w'), str(int(1.0 * 567)))  # 1см = 567 DXA
        tblInd.set(qn('w:type'), 'dxa')
        tblPr.append(tblInd)
        
        print(f"✓ Table indented 1cm from left margin")
        
        # 🎯 НАЛАШТУВАННЯ ЗАГОЛОВКОВОГО РЯДКА (висота 1.72см)
        header_row = table.rows[0]
        header_row.height = Cm(1.72)  # Точна висота 1.72см
        
        # Встановлюємо точну висоту рядка
        trPr = header_row._tr.trPr
        if trPr is None:
            trPr = OxmlElement('w:trPr')
            header_row._tr.append(trPr)
        
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(int(1.72 * 567)))  # 1.72см в DXA
        trHeight.set(qn('w:hRule'), 'exact')  # Точна висота
        trPr.append(trHeight)
        
        # 🎯 ОНОВЛЕНІ ЗАГОЛОВКИ З ПРАВИЛЬНОЮ ШИРИНОЮ
        updated_headers = [
            "№ з/п",           # Перший стовпчик
            "Назва документа", 
            "Обліковий номер",
            "Кількість аркушів (знімків)",
            "Гриф секретності",
            "Примітка"
        ]
        
        # Ширини стовпчиків: перший 1.25см, решта по 3см
        column_widths = [1.25, 3.0, 3.0, 3.0, 3.0, 3.0]  # в см
        
        # Заповнення заголовків
        for i, (cell, header_text, col_width) in enumerate(
            zip(header_row.cells, updated_headers, column_widths)
        ):
            # Встановлення ширини колонки
            cell.width = Cm(col_width)
            
            # Вертикальне центрування
            try:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            except AttributeError:
                pass
            
            # Налаштування тексту
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.text = header_text
            
            # 🎯 ШРИФТ ЗАГОЛОВКІВ НЕ ЖИРНИЙ
            for run in para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(12)
                run.font.bold = False  # НЕ жирний!
            
            # Додавання рамок до заголовкових комірок
            set_cell_borders(cell, top=True, bottom=True, left=True, right=True)
        
        print(f"✓ Header row configured: 1.72cm height, non-bold font")
        
        # 🎯 НАЛАШТУВАННЯ РЯДКІВ ДАНИХ (висота 0.5см + нумерація)
        for row_index in range(1, 6):  # Рядки 1-5
            row = table.rows[row_index]
            
            # Встановлення висоти рядка 0.5см
            row.height = Cm(0.5)
            
            # Встановлюємо точну висоту рядка
            trPr = row._tr.trPr
            if trPr is None:
                trPr = OxmlElement('w:trPr')
                row._tr.append(trPr)
            
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), str(int(0.5 * 567)))  # 0.5см в DXA
            trHeight.set(qn('w:hRule'), 'exact')  # Точна висота
            trPr.append(trHeight)
            
            # Налаштування комірок
            for col_index, cell in enumerate(row.cells):
                # Встановлення ширини (така ж як в заголовку)
                cell.width = Cm(column_widths[col_index])
                
                # Вертикальне центрування
                try:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                except AttributeError:
                    pass
                
                # 🎯 НУМЕРАЦІЯ В ПЕРШОМУ СТОВПЧИКУ
                if col_index == 0:
                    # Додаємо номер: 1., 2., 3., 4., 5.
                    cell.text = f"{row_index}."
                    para = cell.paragraphs[0]
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Налаштування шрифту номерів
                    for run in para.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(12)
                        run.font.bold = False
                else:
                    # Порожні комірки для інших стовпчиків
                    cell.text = ""
                
                # Додавання рамок
                set_cell_borders(cell, top=True, bottom=True, left=True, right=True)
        
        print(f"✓ Data rows configured: 0.5cm height, numbered 1-5")
        
        # Розрив сторінки для переходу до наступної сторінки
        doc.add_page_break()
        
        print("✓ Description page created successfully with exact specifications")
        
    except Exception as e:
        print(f"✗ Error creating description page: {e}")
        
        # Fallback - проста сторінка опису
        fallback_heading = doc.add_paragraph("Опис альбому фотознімків")
        fallback_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fallback_heading.runs[0].font.size = Pt(22)
        fallback_heading.runs[0].font.bold = True
        
        # Простий текст замість таблиці
        doc.add_paragraph("Таблиця опису документів")
        doc.add_page_break()

# ===== ОСНОВНИЙ КЛАС GUI =====

class AzimuthGUI(QMainWindow):
    def __init__(self):
        super().__init__()
        self.processor = None
        self.current_image_path = None
        self.current_click = None
        self.scale_factor_x = 1.0
        self.scale_factor_y = 1.0
        self.offset_x = 0
        self.offset_y = 0
        self.current_folder = None
        
        self.current_language = 'UKRAINIAN'  # Default language
        self.translations = Translations()
        
        self.scale_edge_mode = False
        self.scale_edge_point = None
        self.custom_scale_distance = None
        self.center_setting_mode = False
        
        self.current_target_number = "0001"
        self.current_height = "0.0"
        self.current_obstacles = "без перешкод"
        self.current_detection = "Виявлення"

        self.radar_description_enabled = False
        self.radar_date = QDate.currentDate()
        self.radar_callsign = ""
        self.radar_name = ""
        self.radar_number = ""
        
        self.processed_images = []
        self.current_folder_images = []
        self.current_image_index = -1

        self.processed_image_paths = set()  # Множина шляхів оброблених зображень
        
        # Змінні для збереження налаштувань сітки між зображеннями
        self.saved_grid_settings = {
            'center_offset_x': 0,
            'center_offset_y': 0,
            'scale_edge_relative': None,
            'custom_scale_distance': None,
            'scale_value': "300"
        }
        
        self.init_ui()

        # ===== МЕТОДИ ДЛЯ ОБРОБКИ ОПИСУ РЛС =====

    def keyPressEvent(self, event):
        """Обробка натискань клавіш для переміщення центру/краю масштабу"""
        # ПЕРЕВІРЯЄМО ЧИ АКТИВНИЙ СПЕЦІАЛЬНИЙ РЕЖИМ
        if not (self.center_setting_mode or self.scale_edge_mode):
            super().keyPressEvent(event)
            return
        
        if not self.processor:
            super().keyPressEvent(event)
            return
        
        # ⭐ ЗАБИРАЄМО ФОКУС З ПОЛІВ ВВЕДЕННЯ, АЛЕ ПЕРЕДАЄМО ЗОБРАЖЕННЮ
        focused_widget = self.focusWidget()
        if focused_widget and isinstance(focused_widget, (QLineEdit, QComboBox, QDateEdit)):
            # Передаємо фокус зображенню для збереження функціональності миші
            if hasattr(self, 'image_label'):
                self.image_label.setFocus()
            else:
                self.setFocus()
            print(f"🎯 Focus moved from {type(focused_widget).__name__} to image")
        
        # Визначаємо крок переміщення
        step = 1  # Базовий крок
        if event.modifiers() & Qt.ShiftModifier:
            step = 5  # Великий крок з Shift
        elif event.modifiers() & Qt.ControlModifier:
            step = 0.5  # Малий крок з Ctrl
        
        # Обробляємо стрілки
        dx, dy = 0, 0
        if event.key() == Qt.Key_Left:
            dx = -step
        elif event.key() == Qt.Key_Right:
            dx = step
        elif event.key() == Qt.Key_Up:
            dy = -step
        elif event.key() == Qt.Key_Down:
            dy = step
        elif event.key() == Qt.Key_Escape:
            # Escape для виходу з режиму
            if self.center_setting_mode:
                self.toggle_center_setting_mode()
            elif self.scale_edge_mode:
                self.toggle_scale_edge_mode()
            return
        else:
            super().keyPressEvent(event)
            return
        
        # Виконуємо переміщення
        if dx != 0 or dy != 0:
            if self.center_setting_mode:
                self.move_center_with_keyboard(dx, dy)
            elif self.scale_edge_mode:
                self.move_scale_edge_with_keyboard(dx, dy)
        
        event.accept()

    def move_center_with_keyboard(self, dx, dy):
        """Переміщення центру з клавіатури з зумом"""
        if not self.processor:
            return
        
        self.processor.move_center(dx, dy)
        self.save_current_grid_settings()
        
        # Оновлюємо точку аналізу якщо є
        if self.current_click:
            azimuth, range_val = self.calculate_azimuth_range(
                self.current_click['x'], self.current_click['y']
            )
            self.current_click['azimuth'] = azimuth
            self.current_click['range'] = range_val
        
        self.display_image()
        self.update_results_display()
        self.update_report_data()
        
        # ⭐ ПОКАЗУЄМО ЗУМ НА НОВІЙ ПОЗИЦІЇ ЦЕНТРУ
        if hasattr(self, 'image_label'):
            self.image_label.zoom_widget.show_zoom()
            self.image_label.zoom_widget.update_cursor_position(
                self.processor.center_x, self.processor.center_y
            )
    
        self.add_result(f"Центр: ({self.processor.center_x}, {self.processor.center_y}) | ←→↑↓")

    def move_scale_edge_with_keyboard(self, dx, dy):
        """Переміщення краю масштабу з клавіатури з зумом"""
        if not self.processor:
            return
        
        # Ініціалізуємо scale edge point якщо не існує
        if not self.scale_edge_point:
            self.scale_edge_point = {
                'x': self.processor.center_x + 50,
                'y': self.processor.center_y + 50
            }
        
        # Переміщуємо точку
        new_x = max(0, min(self.scale_edge_point['x'] + dx, self.processor.image.width - 1))
        new_y = max(0, min(self.scale_edge_point['y'] + dy, self.processor.image.height - 1))
        
        self.scale_edge_point['x'] = new_x
        self.scale_edge_point['y'] = new_y
        
        # Перераховуємо відстань
        dx_scale = new_x - self.processor.center_x
        dy_scale = new_y - self.processor.center_y
        self.custom_scale_distance = math.sqrt(dx_scale*dx_scale + dy_scale*dy_scale)
        
        self.save_current_grid_settings()
        
        # Оновлюємо точку аналізу якщо є
        if self.current_click:
            azimuth, range_val = self.calculate_azimuth_range(
                self.current_click['x'], self.current_click['y']
            )
            self.current_click['azimuth'] = azimuth
            self.current_click['range'] = range_val
        
        self.display_image()
        self.update_results_display()
        self.update_report_data()
        
        # ⭐ ПОКАЗУЄМО ЗУМ НА НОВІЙ ПОЗИЦІЇ КРАЮ МАСШТАБУ
        if hasattr(self, 'image_label'):
            self.image_label.zoom_widget.show_zoom()
            self.image_label.zoom_widget.update_cursor_position(new_x, new_y)
    
        self.add_result(f"Край масштабу: ({new_x}, {new_y}) | Відстань: {self.custom_scale_distance:.1f}px")

    def toggle_radar_description(self, checked):
        """Перемикання режиму опису РЛС"""
        self.radar_description_enabled = checked
        
        # Активуємо/деактивуємо поля введення
        self.radar_date_edit.setEnabled(checked)
        self.radar_callsign_input.setEnabled(checked)
        self.radar_name_input.setEnabled(checked)
        self.radar_number_input.setEnabled(checked)
        
        if checked:
            self.add_result("✓ Режим опису РЛС активовано")
            self.add_result("  Введені дані будуть застосовані до всіх зображень")
        else:
            self.add_result("✗ Режим опису РЛС деактивовано")
            self.add_result("  Дані РЛС не застосовуватимуться")

    def update_radar_date(self, date):
        """Оновлення дати РЛС"""
        self.radar_date = date
        if self.radar_description_enabled:
            self.add_result(f"📅 Дата РЛС оновлена: {date.toString('dd.MM.yyyy')}")

    def update_radar_callsign(self, text):
        """Оновлення позивного РЛС"""
        self.radar_callsign = text
        if self.radar_description_enabled and text:
            self.add_result(f"📡 Позивний РЛС: {text}")

    def update_radar_name(self, text):
        """Оновлення назви РЛС"""
        self.radar_name = text
        if self.radar_description_enabled and text:
            self.add_result(f"📋 Назва РЛС: {text}")

    def update_radar_number(self, text):
        """Оновлення номера РЛС"""
        self.radar_number = text
        if self.radar_description_enabled and text:
            self.add_result(f"🔢 Номер РЛС: {text}")

    def get_radar_description_data(self):
        """Отримання даних опису РЛС"""
        if not self.radar_description_enabled:
            return None
        
        return {
            'enabled': True,
            'date': self.radar_date,
            'callsign': self.radar_callsign,
            'name': self.radar_name,
            'number': self.radar_number
        }

    # =====================================
    
    def tr(self, key):
        return self.translations.get(self.current_language, key)
    
    def set_language(self, language):
        self.current_language = language
        self.update_interface_text()
        
        for lang, action in self.language_actions.items():
            action.setChecked(lang == language)
    
    def update_interface_text(self):
        self.setWindowTitle(self.tr("window_title"))
        self.create_menu_bar()
        
        if hasattr(self, 'control_title'):
            self.control_title.setText(self.tr("controls"))
        if hasattr(self, 'report_title'):
            self.report_title.setText(self.tr("report_data"))
        if hasattr(self, 'browser_label'):
            self.browser_label.setText(self.tr("photo_browser"))
        
        if hasattr(self, 'open_image_btn'):
            self.open_image_btn.setText(self.tr("open_image"))
        if hasattr(self, 'open_folder_btn'):
            self.open_folder_btn.setText(self.tr("open_folder"))
        if hasattr(self, 'save_image_btn'):
            self.save_image_btn.setText(self.tr("save_current_image"))
        if hasattr(self, 'export_new_btn'):
            self.export_new_btn.setText(self.tr("create_new_album"))
        if hasattr(self, 'export_add_btn'):
            self.export_add_btn.setText(self.tr("add_to_existing_album"))
        if hasattr(self, 'scale_edge_btn'):
            self.scale_edge_btn.setText(self.tr("set_scale_edge"))
        if hasattr(self, 'set_center_btn'):
            self.set_center_btn.setText(self.tr("set_center"))
        
        if hasattr(self, 'file_ops_label'):
            self.file_ops_label.setText(self.tr("file_operations"))
        if hasattr(self, 'azimuth_grid_label'):
            self.azimuth_grid_label.setText(self.tr("azimuth_grid"))
        if hasattr(self, 'move_center_label'):
            self.move_center_label.setText(self.tr("move_center"))
        if hasattr(self, 'results_label'):
            self.results_label.setText(self.tr("results"))
        
        # Batch processing переклади
        if hasattr(self, 'batch_label'):
            self.batch_label.setText(self.tr("batch_processing"))
        if hasattr(self, 'save_current_btn'):
            self.save_current_btn.setText(self.tr("save_current_image_data"))
        if hasattr(self, 'prev_btn'):
            self.prev_btn.setText(f"← {self.tr('previous')}")
        if hasattr(self, 'next_btn'):
            self.next_btn.setText(f"{self.tr('next')} →")
        
        if not self.current_image_path:
            self.image_label.setText(self.tr("open_instruction"))
        
        if self.processor:
            self.update_results_display()

    def init_ui(self):
        self.setWindowTitle(self.tr("window_title"))
        
        default_width = 1400
        default_height = 900
        min_width = 1000  
        min_height = 700
        self.setMinimumSize(min_width, min_height)
        self.setWindowFlags(Qt.Window)
        self.resize(default_width, default_height)
        self.showMaximized()
        
        self.create_menu_bar()
        
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        main_layout = QHBoxLayout()
        central_widget.setLayout(main_layout)
        
        main_splitter = QSplitter(Qt.Horizontal)
        main_layout.addWidget(main_splitter)
        
        self.create_control_panel(main_splitter)
        self.create_vertical_browser_panel(main_splitter)
        self.create_image_panel(main_splitter)
        self.create_report_panel(main_splitter)
        
        main_splitter.setStretchFactor(0, 0)
        main_splitter.setStretchFactor(1, 0)
        main_splitter.setStretchFactor(2, 1)
        main_splitter.setStretchFactor(3, 0)
        
        main_splitter.setSizes([220, 0, 960, 220])  # 220+960+220 = 1400px
        self.main_splitter = main_splitter

    def resizeEvent(self, event):
        """Оновлений метод зміни розміру з правильними розмірами браузера"""
        super().resizeEvent(event)
        
        if hasattr(self, 'processor') and self.processor:
            QTimer.singleShot(50, self.update_image_display_after_resize)
        
        if hasattr(self, 'main_splitter') and hasattr(self, 'browser_widget'):
            current_sizes = self.main_splitter.sizes()
            total_width = sum(current_sizes)
            
            if total_width > 800:
                if self.browser_widget.isVisible():
                    # ПРАВИЛЬНІ розміри: 220 + 280 + ? + 220 = total_width
                    new_image_width = max(450, total_width - 220 - 280 - 220)  # Мінімум 450px
                    self.main_splitter.setSizes([220, 280, new_image_width, 220])
                else:
                    # 220 + 0 + ? + 220 = total_width  
                    new_image_width = max(450, total_width - 220 - 220)  # Мінімум 450px
                    self.main_splitter.setSizes([220, 0, new_image_width, 220])
    
    def update_image_display_after_resize(self):
        if hasattr(self, 'processor') and self.processor:
            self.display_image()
    
    def create_menu_bar(self):
        menubar = self.menuBar()
        menubar.clear()
        
        settings_menu = menubar.addMenu(self.tr("settings"))
        language_menu = settings_menu.addMenu(self.tr("language"))
        
        english_action = language_menu.addAction("English")
        english_action.triggered.connect(lambda: self.set_language('ENGLISH'))
        
        ukrainian_action = language_menu.addAction("Українська")
        ukrainian_action.triggered.connect(lambda: self.set_language('UKRAINIAN'))
        
        english_action.setCheckable(True)
        ukrainian_action.setCheckable(True)
        
        english_action.setChecked(self.current_language == 'ENGLISH')
        ukrainian_action.setChecked(self.current_language == 'UKRAINIAN')
        
        self.language_actions = {
            'ENGLISH': english_action,
            'UKRAINIAN': ukrainian_action
        }

    def create_control_panel(self, parent):
        """Ліва панель з основними кнопками"""
        control_widget = QWidget()
        control_widget.setFixedWidth(250)
        control_widget.setStyleSheet("""
            QWidget {
                background-color: #f5f5f5; 
                border-right: 1px solid #ccc;
            }
            QLabel {
                background: none;
                border: none;
                color: #333;
            }
            QPushButton {
                background-color: #e1e1e1;
                border: 1px solid #c1c1c1;
                border-radius: 6px;
                padding: 8px 12px;
                font-size: 11px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #d1d1d1;
                border: 1px solid #b1b1b1;
            }
            QPushButton:pressed {
                background-color: #c1c1c1;
            }
            QComboBox {
                border: 1px solid #c1c1c1;
                border-radius: 4px;
                padding: 4px 8px;
                background-color: white;
                font-size: 11px;
            }
        """)
        
        layout = QVBoxLayout()
        layout.setContentsMargins(15, 15, 15, 15)
        layout.setSpacing(10)
        control_widget.setLayout(layout)
        
        # Title
        self.control_title = QLabel(self.tr("controls"))
        self.control_title.setFont(QFont("Arial", 14, QFont.Bold))
        self.control_title.setAlignment(Qt.AlignCenter)
        self.control_title.setStyleSheet("font-size: 14px; font-weight: bold; margin-bottom: 10px;")
        layout.addWidget(self.control_title)
        
        # File operations section
        self.file_ops_label = QLabel(self.tr("file_operations"))
        self.file_ops_label.setFont(QFont("Arial", 10, QFont.Bold))
        self.file_ops_label.setStyleSheet("color: #666; margin-top: 10px; margin-bottom: 5px;")
        layout.addWidget(self.file_ops_label)
        
        # File operation buttons
        self.open_image_btn = QPushButton(self.tr("open_image"))
        self.open_image_btn.clicked.connect(self.open_image)
        self.open_image_btn.setStyleSheet("""
            QPushButton {
                background-color: #e1e1e1;
                border: 2px solid #c1c1c1;
                border-radius: 6px;
                padding: 8px 12px;
                font-size: 11px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #d1d1d1;
                border: 2px solid #b1b1b1;
            }
            QPushButton:pressed {
                background-color: #c1c1c1;
                border: 2px solid #a1a1a1;
            }
        """)
        layout.addWidget(self.open_image_btn)
        
        self.open_folder_btn = QPushButton(self.tr("open_folder"))
        self.open_folder_btn.clicked.connect(self.open_folder)
        self.open_folder_btn.setStyleSheet("""
            QPushButton {
                background-color: #e1e1e1;
                border: 2px solid #c1c1c1;
                border-radius: 6px;
                padding: 8px 12px;
                font-size: 11px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #d1d1d1;
                border: 2px solid #b1b1b1;
            }
            QPushButton:pressed {
                background-color: #c1c1c1;
                border: 2px solid #a1a1a1;
            }
        """)
        layout.addWidget(self.open_folder_btn)
        
        self.save_image_btn = QPushButton(self.tr("save_current_image"))
        self.save_image_btn.clicked.connect(self.save_current_image)
        self.save_image_btn.setStyleSheet("""
            QPushButton {
                background-color: #e1e1e1;
                border: 2px solid #c1c1c1;
                border-radius: 6px;
                padding: 8px 12px;
                font-size: 11px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #d1d1d1;
                border: 2px solid #b1b1b1;
            }
            QPushButton:pressed {
                background-color: #c1c1c1;
                border: 2px solid #a1a1a1;
            }
        """)
        layout.addWidget(self.save_image_btn)
        
        # Separator
        separator1 = QFrame()
        separator1.setFrameShape(QFrame.HLine)
        separator1.setFrameShadow(QFrame.Sunken)
        separator1.setStyleSheet("color: #ccc; margin: 10px 0px;")
        layout.addWidget(separator1)
        
        # Batch processing section
        self.batch_label = QLabel(self.tr("batch_processing"))
        self.batch_label.setFont(QFont("Arial", 10, QFont.Bold))
        self.batch_label.setStyleSheet("color: #666; margin-top: 5px; margin-bottom: 5px;")
        layout.addWidget(self.batch_label)

        self.save_current_btn = QPushButton(self.tr("save_current_image_data"))
        self.save_current_btn.clicked.connect(self.save_current_image_data)
        self.save_current_btn.setStyleSheet("""
            QPushButton {
                background-color: #e1e1e1;
                border: 2px solid #c1c1c1;
                border-radius: 6px;
                padding: 8px 12px;
                font-size: 11px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #d1d1d1;
                border: 2px solid #b1b1b1;
            }
            QPushButton:pressed {
                background-color: #c1c1c1;
                border: 2px solid #a1a1a1;
            }
        """)
        layout.addWidget(self.save_current_btn)

        nav_layout = QHBoxLayout()
        self.prev_btn = QPushButton(f"← {self.tr('previous')}")
        self.prev_btn.clicked.connect(self.previous_image_in_batch)
        self.prev_btn.setStyleSheet("""
            QPushButton {
                background-color: #e1e1e1;
                border: 2px solid #c1c1c1;
                border-radius: 6px;
                padding: 8px 12px;
                font-size: 11px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #d1d1d1;
                border: 2px solid #b1b1b1;
            }
            QPushButton:pressed {
                background-color: #c1c1c1;
                border: 2px solid #a1a1a1;
            }
        """)
        nav_layout.addWidget(self.prev_btn)

        self.next_btn = QPushButton(f"{self.tr('next')} →")
        self.next_btn.clicked.connect(self.next_image_in_batch)
        self.next_btn.setStyleSheet("""
            QPushButton {
                background-color: #e1e1e1;
                border: 2px solid #c1c1c1;
                border-radius: 6px;
                padding: 8px 12px;
                font-size: 11px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #d1d1d1;
                border: 2px solid #b1b1b1;
            }
            QPushButton:pressed {
                background-color: #c1c1c1;
                border: 2px solid #a1a1a1;
            }
        """)
        nav_layout.addWidget(self.next_btn)

        nav_widget = QWidget()
        nav_widget.setLayout(nav_layout)
        layout.addWidget(nav_widget)

        # ГОЛОВНА кнопка для створення альбому
        self.create_new_structure_btn = QPushButton(self.tr("create_new_album"))
        self.create_new_structure_btn.clicked.connect(self.create_batch_album_new_structure)
        self.create_new_structure_btn.setStyleSheet("""
            QPushButton {
                font-weight: bold; 
                background-color: #4CAF50; 
                color: white;
                border: 2px solid #45a049;
                border-radius: 6px;
                padding: 8px 12px;
            }
            QPushButton:hover {
                background-color: #45a049;
                border: 2px solid #3d8b40;
            }
            QPushButton:pressed {
                background-color: #3d8b40;
                border: 2px solid #2e6b30;
            }
        """)
        layout.addWidget(self.create_new_structure_btn)
        
        # Results area
        self.results_label = QLabel(self.tr("results"))
        self.results_label.setFont(QFont("Arial", 10, QFont.Bold))
        self.results_label.setStyleSheet("color: #666; margin-top: 10px; margin-bottom: 5px;")
        layout.addWidget(self.results_label)
        
        self.results_text = QTextEdit()
        self.results_text.setMaximumHeight(120)
        self.results_text.setStyleSheet("""
            QTextEdit {
                background-color: white;
                border: 1px solid #ccc;
                font-family: 'Courier New', monospace;
                font-size: 10px;
                border-radius: 4px;
            }
        """)
        layout.addWidget(self.results_text)

        layout.addStretch()
        parent.addWidget(control_widget)

    def create_vertical_browser_panel(self, parent):
        browser_widget = QWidget()
        # ЗБІЛЬШУЄМО ширину з 180px до 280px
        browser_widget.setFixedWidth(280)
        browser_widget.setStyleSheet("""
            QWidget {
                background-color: #f8f8f8; 
                border-right: 1px solid #ccc;
                border-left: 1px solid #ccc;
            }
            QLabel {
                background: none;
                border: none;
                color: #333;
            }
        """)
        
        layout = QVBoxLayout()
        layout.setContentsMargins(0, 15, 0, 15)
        layout.setSpacing(10)
        browser_widget.setLayout(layout)
        
        self.browser_label = QLabel(self.tr("photo_browser"))
        self.browser_label.setFont(QFont("Arial", 10, QFont.Bold))
        self.browser_label.setStyleSheet("color: #666; margin-bottom: 5px; padding: 0 10px;")
        self.browser_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(self.browser_label)
        
        self.thumbnail_scroll = QScrollArea()
        self.thumbnail_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.thumbnail_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self.thumbnail_scroll.setWidgetResizable(False)
        self.thumbnail_scroll.setStyleSheet("border: none; background: transparent;")
        
        # ЗБІЛЬШУЄМО ширину віджета мініатюр з 160px до 260px
        self.thumbnail_widget = VerticalThumbnailWidget(thumbnail_width=260)
        self.thumbnail_widget.image_selected.connect(self.load_image_from_browser)
        self.thumbnail_scroll.setWidget(self.thumbnail_widget)
        
        layout.addWidget(self.thumbnail_scroll)
        
        browser_widget.hide()
        self.browser_widget = browser_widget
        
        parent.addWidget(browser_widget)

    def create_image_panel(self, parent):
        """Центральна панель для відображення зображень"""
        image_widget = QWidget()
        layout = QVBoxLayout()
        image_widget.setLayout(layout)
        
        self.image_label = ClickableLabel()
        self.image_label.clicked.connect(self.on_image_click)
        self.image_label.dragged.connect(self.on_image_drag)
        self.image_label.mouse_moved.connect(self.on_mouse_hover)
        self.image_label.setText(self.tr("open_instruction"))
        
        scroll_area = QScrollArea()
        scroll_area.setWidget(self.image_label)
        scroll_area.setWidgetResizable(True)
        scroll_area.setMinimumHeight(400)
        layout.addWidget(scroll_area)
        
        parent.addWidget(image_widget)

    def open_image(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, self.tr("select_image"), "",
            f"{self.tr('image_files')} (*.jpg *.jpeg *.png *.bmp *.gif *.tiff);;{self.tr('all_files')} (*.*)"
        )
        
        if file_path:
            self.load_image(file_path)
            self.report_widget.show()
    
    def load_image_from_browser(self, file_path):
        """Завантаження зображення з браузера"""
        self.load_image(file_path)
        self.add_result(self.tr("loaded_from_browser").format(name=os.path.basename(file_path)))

    def open_folder(self):
        """Відкриття папки з зображеннями з правильними розмірами браузера"""
        folder_path = QFileDialog.getExistingDirectory(self, self.tr("select_folder"))
        
        if folder_path:
            self.current_folder = folder_path
            self.load_folder_thumbnails()
            
            self.browser_widget.show()
            # ВИПРАВЛЕНІ розміри: 220 + 280 + ? + 220 = total_width
            self.main_splitter.setSizes([220, 280, 620, 220])
            
            image_extensions = ('.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff', '.tif')
            image_count = sum(1 for f in os.listdir(folder_path) 
                            if f.lower().endswith(image_extensions))
            
            self.add_result(f"{self.tr('loaded_folder')}: {os.path.basename(folder_path)}")
            self.add_result(self.tr("found_images").format(count=image_count))
            
            self.report_widget.show()

    def load_folder_thumbnails(self):
        """Виправлена функція завантаження мініатюр з правильними розмірами"""
        print("🟡 === load_folder_thumbnails STARTED ===")
        
        if not self.current_folder:
            print("❌ current_folder is None!")
            return
        
        print(f"🔍 Loading thumbnails from: {self.current_folder}")
        
        # Перевірка thumbnail_widget
        if not hasattr(self, 'thumbnail_widget'):
            print("❌ thumbnail_widget doesn't exist!")
            return
        
        print(f"✅ thumbnail_widget exists: {type(self.thumbnail_widget)}")
        
        # ВАЖЛИВО: Очищуємо попередні мініатюри
        self.thumbnail_widget.clear_thumbnails()
        
        image_extensions = ('.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff', '.tif')
        image_files = []
        
        try:
            print(f"📁 Scanning folder: {self.current_folder}")
            for filename in os.listdir(self.current_folder):
                if filename.lower().endswith(image_extensions):
                    full_path = os.path.join(self.current_folder, filename)
                    image_files.append(full_path)
                    print(f"✅ Found image: {filename}")
        except Exception as e:
            print(f"❌ Error reading folder: {e}")
            return

        print(f"📊 Total images found: {len(image_files)}")
        
        # Сортуємо файли
        image_files.sort()
        self.current_folder_images = image_files
        
        if len(image_files) == 0:
            print("📭 No images - adding 'no images' label")
            no_images_label = QLabel(self.tr("no_images_found"))
            no_images_label.setAlignment(Qt.AlignCenter)
            no_images_label.setStyleSheet("color: gray; font-size: 14px; padding: 20px;")
            no_images_label.setWordWrap(True)
            self.thumbnail_widget.layout.addWidget(no_images_label)
            return
        
        print(f"🔄 Creating thumbnails for {len(image_files)} images...")
        
        # ВИПРАВЛЕННЯ: Створюємо мініатюри тільки ОДИН раз
        for i, image_path in enumerate(image_files):
            try:
                print(f"🖼️ Creating thumbnail {i+1}/{len(image_files)}: {os.path.basename(image_path)}")
                self.thumbnail_widget.add_thumbnail(image_path)
                print(f"✅ Thumbnail {i+1} created successfully")
            except Exception as e:
                print(f"❌ Error creating thumbnail {i+1}: {e}")
                import traceback
                traceback.print_exc()
        
        # ВИПРАВЛЕНІ розміри віджета для більших мініатюр
        widget_height = len(image_files) * 190 + 20  # Збільшена висота для мініатюр 240x180
        self.thumbnail_widget.setMinimumHeight(widget_height)
        self.thumbnail_widget.resize(260, widget_height)  # Ширина 260px
        
        print(f"🟢 === load_folder_thumbnails COMPLETED ===")
        print(f"📋 Final result: {len(image_files)} unique thumbnails created")


    # НОВІ МЕТОДИ ДЛЯ СТВОРЕННЯ АЛЬБОМІВ З ТАБЛИЦЯМИ

    def create_batch_album_new_structure(self):
        """НОВИЙ МЕТОД з використанням нової структури"""
        if not DOCX_AVAILABLE:
            QMessageBox.warning(self, "Warning", "python-docx library not installed")
            return
        
        if not self.processed_images:
            QMessageBox.warning(self, "Warning", "No processed images to export")
            return
        
        # Зберігаємо поточне зображення якщо є
        if self.processor and self.current_click:
            self.save_current_image_data()
        
        # Діалог вибору файлу
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Create Complete Album with New Structure", "",
            "Word Documents (*.docx);;All files (*.*)"
        )
        
        if file_path:
            try:
                # Отримуємо дані для титульної сторінки
                title_data = self.get_title_page_data_from_gui()
                
                # Створюємо повний альбом з новою структурою
                success = create_complete_album(self.processed_images, title_data, file_path)
                
                if success:
                    QMessageBox.information(self, "Success", 
                                        f"Complete album with new structure created!\n"
                                        f"- Title page (standard margins)\n"
                                        f"- Description page (standard margins)\n" 
                                        f"- {len(self.processed_images)} images in new structure\n"
                                        f"- New layout: paragraph + table + paragraph + table + paragraph\n"
                                        f"- Page margins: 20mm top, 5mm sides/bottom\n"
                                        f"Saved: {os.path.basename(file_path)}")
                    
                    self.add_result(f"✓ Complete new structure album created: {os.path.basename(file_path)}")
                    self.add_result(f"✓ Contains {len(self.processed_images)} processed images")
                else:
                    QMessageBox.critical(self, "Error", "Failed to create new structure album")
                    
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Could not create new structure album: {str(e)}")
                print(f"Error creating new structure album: {e}")

    def get_title_page_data_from_gui(self):
        """Отримання даних для титульної сторінки з GUI"""
        from datetime import datetime
        
        title_data = {
            'target_no': self.current_target_number,
            'date': datetime.now(),
            'unit_info': get_default_unit_info(),
            'commander_info': get_default_commander_info()
        }
        
        return title_data
        
    def create_report_panel(self, parent):
        """Права панель з азимутальними контролами та описом РЛС"""
        report_widget = QWidget()
        report_widget.setFixedWidth(220)
        report_widget.setStyleSheet("background-color: #f9f9f9; border-left: 1px solid #ccc;")
        
        layout = QVBoxLayout()
        report_widget.setLayout(layout)
        
        # Title
        self.report_title = QLabel(self.tr("report_data"))
        self.report_title.setFont(QFont("Arial", 12, QFont.Bold))
        self.report_title.setAlignment(Qt.AlignCenter)
        layout.addWidget(self.report_title)
        
        # Target data input section (СПОЧАТКУ ДАНІ ПРО ЦІЛЬ)
        manual_group = QFrame()
        manual_group.setStyleSheet("background-color: #fff; border: 1px solid #ddd; border-radius: 5px; padding: 10px;")
        manual_layout = QVBoxLayout()
        manual_group.setLayout(manual_layout)
        
        self.target_number_input = QLineEdit()
        self.target_number_input.setPlaceholderText("Номер цілі")
        self.target_number_input.setText(self.current_target_number)
        self.target_number_input.textChanged.connect(self.update_target_number)
        manual_layout.addWidget(self.target_number_input)
        
        self.auto_azimuth_label = QLabel("β - --°")
        manual_layout.addWidget(self.auto_azimuth_label)
        
        self.auto_distance_label = QLabel("D - -- км")
        manual_layout.addWidget(self.auto_distance_label)
        
        height_layout = QHBoxLayout()
        height_layout.addWidget(QLabel("H –"))
        self.height_input = QLineEdit(self.current_height)
        self.height_input.setMaximumWidth(80)
        self.height_input.textChanged.connect(self.update_height)
        height_layout.addWidget(self.height_input)
        height_layout.addWidget(QLabel(self.tr("km_unit")))
        height_layout.addStretch()
        manual_layout.addLayout(height_layout)
        
        self.obstacles_combo = QComboBox()
        self.obstacles_combo.addItems([self.tr("no_obstacles"), self.tr("with_obstacles")])
        self.obstacles_combo.currentTextChanged.connect(self.update_obstacles)
        manual_layout.addWidget(self.obstacles_combo)
        
        self.detection_combo = QComboBox()
        self.detection_combo.addItems([self.tr("detection"), self.tr("tracking"), self.tr("loss")])
        self.detection_combo.currentTextChanged.connect(self.update_detection)
        manual_layout.addWidget(self.detection_combo)
        
        self.auto_scale_label = QLabel("M = --")
        manual_layout.addWidget(self.auto_scale_label)
        
        layout.addWidget(manual_group)
        
        # Separator після даних про ціль
        grid_separator = QFrame()
        grid_separator.setFrameShape(QFrame.HLine)
        grid_separator.setFrameShadow(QFrame.Sunken)
        grid_separator.setStyleSheet("color: #ccc; margin: 8px 0px;")
        layout.addWidget(grid_separator)
        
        # Azimuth Grid section (ПІСЛЯ ДАНИХ ПРО ЦІЛЬ)
        self.azimuth_grid_label = QLabel(self.tr("azimuth_grid"))
        self.azimuth_grid_label.setFont(QFont("Arial", 10, QFont.Bold))
        self.azimuth_grid_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(self.azimuth_grid_label)
        
        # Scale setting
        scale_layout = QHBoxLayout()
        scale_label = QLabel(self.tr("scale_setting"))
        scale_label.setStyleSheet("color: #333; margin-right: 5px;")
        scale_layout.addWidget(scale_label)
        self.scale_combo = QComboBox()
        self.scale_combo.addItems(["50", "100", "150", "200", "300"])
        self.scale_combo.setCurrentText("300")
        self.scale_combo.currentTextChanged.connect(self.update_scale)
        self.scale_combo.setStyleSheet("""
            QComboBox {
                border: 2px solid #c1c1c1;
                border-radius: 4px;
                padding: 4px 8px;
                background-color: white;
                font-size: 11px;
            }
            QComboBox:hover {
                border: 2px solid #b1b1b1;
            }
            QComboBox:focus {
                border: 2px solid #007acc;
            }
        """)
        scale_layout.addWidget(self.scale_combo)
        scale_layout.addStretch()
        layout.addLayout(scale_layout)
        
        # Scale edge tool
        self.scale_edge_btn = QPushButton(self.tr("set_scale_edge"))
        self.scale_edge_btn.setCheckable(True)
        self.scale_edge_btn.clicked.connect(self.toggle_scale_edge_mode)
        self.scale_edge_btn.setStyleSheet("""
            QPushButton {
                background-color: #e1e1e1;
                border: 2px solid #c1c1c1;
                border-radius: 6px;
                padding: 8px 12px;
                font-size: 11px;
                font-weight: bold;
                margin-top: 5px;
            }
            QPushButton:hover {
                background-color: #d1d1d1;
                border: 2px solid #b1b1b1;
            }
            QPushButton:pressed {
                background-color: #c1c1c1;
                border: 2px solid #a1a1a1;
            }
            QPushButton:checked {
                background-color: #4CAF50;
                color: white;
                border: 2px solid #45a049;
            }
        """)
        layout.addWidget(self.scale_edge_btn)
        
        # Center setting tool
        self.set_center_btn = QPushButton(self.tr("set_center"))
        self.set_center_btn.setCheckable(True)
        self.set_center_btn.clicked.connect(self.toggle_center_setting_mode)
        self.set_center_btn.setStyleSheet("""
            QPushButton {
                background-color: #e1e1e1;
                border: 2px solid #c1c1c1;
                border-radius: 6px;
                padding: 8px 12px;
                font-size: 11px;
                font-weight: bold;
                margin-top: 5px;
            }
            QPushButton:hover {
                background-color: #d1d1d1;
                border: 2px solid #b1b1b1;
            }
            QPushButton:pressed {
                background-color: #c1c1c1;
                border: 2px solid #a1a1a1;
            }
            QPushButton:checked {
                background-color: #FF9800;
                color: white;
                border: 2px solid #f57c00;
            }
        """)
        layout.addWidget(self.set_center_btn)
        
        # ===== СЕКЦІЯ: ОПИС РЛС =====
        radar_separator = QFrame()
        radar_separator.setFrameShape(QFrame.HLine)
        radar_separator.setFrameShadow(QFrame.Sunken)
        radar_separator.setStyleSheet("color: #ccc; margin: 8px 0px;")
        layout.addWidget(radar_separator)
        
        # Checkbox для активації опису РЛС
        self.radar_description_checkbox = QCheckBox("Додати опис РЛС")
        self.radar_description_checkbox.setFont(QFont("Arial", 10, QFont.Bold))
        self.radar_description_checkbox.setStyleSheet("""
            QCheckBox {
                color: #333;
                padding: 5px;
            }
            QCheckBox::indicator {
                width: 16px;
                height: 16px;
            }
            QCheckBox::indicator:unchecked {
                border: 2px solid #ccc;
                background-color: white;
                border-radius: 3px;
            }
            QCheckBox::indicator:checked {
                border: 2px solid #4CAF50;
                background-color: #4CAF50;
                border-radius: 3px;
            }
        """)
        self.radar_description_checkbox.toggled.connect(self.toggle_radar_description)
        layout.addWidget(self.radar_description_checkbox)
        
        # Група полів для опису РЛС
        radar_group = QFrame()
        radar_group.setStyleSheet("background-color: #fff; border: 1px solid #ddd; border-radius: 5px; padding: 8px;")
        radar_layout = QVBoxLayout()
        radar_group.setLayout(radar_layout)
        
        # Вибір дати
        date_layout = QHBoxLayout()
        date_layout.addWidget(QLabel("Дата:"))
        self.radar_date_edit = QDateEdit()
        self.radar_date_edit.setDate(self.radar_date)
        self.radar_date_edit.setCalendarPopup(True)
        self.radar_date_edit.setFixedHeight(30)  # ЗБІЛЬШЕНА ВИСОТА з стандартної до 30px
        self.radar_date_edit.setStyleSheet("""
            QDateEdit {
                border: 1px solid #c1c1c1;
                border-radius: 4px;
                padding: 4px 8px;
                background-color: white;
                font-size: 10px;
            }
        """)
        self.radar_date_edit.dateChanged.connect(self.update_radar_date)
        self.radar_date_edit.setEnabled(False)  # Спочатку неактивне
        date_layout.addWidget(self.radar_date_edit)
        radar_layout.addLayout(date_layout)
        
        # Поле для позивного
        self.radar_callsign_input = QLineEdit()
        self.radar_callsign_input.setPlaceholderText("Позивний")
        self.radar_callsign_input.setFixedHeight(30)  # ЗБІЛЬШЕНА ВИСОТА з стандартної до 30px
        self.radar_callsign_input.textChanged.connect(self.update_radar_callsign)
        self.radar_callsign_input.setEnabled(False)  # Спочатку неактивне
        self.radar_callsign_input.setStyleSheet("""
            QLineEdit {
                border: 1px solid #c1c1c1;
                border-radius: 4px;
                padding: 4px 8px;
                background-color: white;
                font-size: 10px;
            }
        """)
        radar_layout.addWidget(self.radar_callsign_input)
        
        # Поле для назви РЛС
        self.radar_name_input = QLineEdit()
        self.radar_name_input.setPlaceholderText("Назва РЛС")
        self.radar_name_input.setFixedHeight(30)  # ЗБІЛЬШЕНА ВИСОТА з стандартної до 30px
        self.radar_name_input.textChanged.connect(self.update_radar_name)
        self.radar_name_input.setEnabled(False)  # Спочатку неактивне
        self.radar_name_input.setStyleSheet("""
            QLineEdit {
                border: 1px solid #c1c1c1;
                border-radius: 4px;
                padding: 4px 8px;
                background-color: white;
                font-size: 10px;
            }
        """)
        radar_layout.addWidget(self.radar_name_input)
        
        # Поле для номера РЛС
        self.radar_number_input = QLineEdit()
        self.radar_number_input.setPlaceholderText("Номер РЛС")
        self.radar_number_input.setFixedHeight(30)  # ЗБІЛЬШЕНА ВИСОТА з стандартної до 30px
        self.radar_number_input.textChanged.connect(self.update_radar_number)
        self.radar_number_input.setEnabled(False)  # Спочатку неактивне
        self.radar_number_input.setStyleSheet("""
            QLineEdit {
                border: 1px solid #c1c1c1;
                border-radius: 4px;
                padding: 4px 8px;
                background-color: white;
                font-size: 10px;
            }
        """)
        radar_layout.addWidget(self.radar_number_input)
        
        layout.addWidget(radar_group)
        
        # Додаємо простір знизу, але не розтягуємо елементи
        layout.addStretch()
        
        report_widget.hide()
        self.report_widget = report_widget
        parent.addWidget(report_widget)

    def move_center_and_save(self, dx, dy):
        """Переміщення центру з автоматичним збереженням налаштувань"""
        if not self.processor:
            QMessageBox.warning(self, self.tr("warning"), self.tr("no_image_first"))
            return
        
        self.processor.move_center(dx, dy)
        self.save_current_grid_settings()
        
        if self.current_click:
            azimuth, range_val = self.calculate_azimuth_range(
                self.current_click['x'], self.current_click['y']
            )
            self.current_click['azimuth'] = azimuth
            self.current_click['range'] = range_val
        
        self.display_image()
        self.update_results_display()
        self.update_report_data()
        
        self.add_result(self.tr("center_moved").format(x=self.processor.center_x, 
                                                    y=self.processor.center_y))
        self.add_result(self.tr("grid_settings_saved"))
        
        if hasattr(self, 'image_label') and self.image_label.zoom_widget.isVisible():
            self.image_label.zoom_widget.update_cursor_position(self.processor.center_x, self.processor.center_y)

    def add_result(self, text):
        self.results_text.append(text)
    
    
    def load_image(self, file_path):
        """Завантаження зображення зі збереженням налаштувань сітки"""
        try:
            self.current_image_path = file_path
            scale_value = int(self.scale_combo.currentText())
            
            self.processor = AzimuthImageProcessor(file_path, scale=scale_value)
            
            if hasattr(self.processor, 'image') and self.processor.image:
                original_mode = self.processor.image.mode
                
                if original_mode != 'RGB':
                    if original_mode == 'RGBA':
                        rgb_image = Image.new('RGB', self.processor.image.size, (255, 255, 255))
                        rgb_image.paste(self.processor.image, mask=self.processor.image.split()[-1])
                        self.processor.image = rgb_image
                    else:
                        self.processor.image = self.processor.image.convert('RGB')
            
            # Застосовуємо збережені налаштування сітки
            self.apply_saved_grid_settings()
            
            # Очищуємо тільки точку аналізу, але НЕ налаштування сітки
            self.current_click = None
            
            self.display_image()
            self.update_results_display()
            self.update_report_data()
            
            self.add_result(f"{self.tr('loaded')}: {os.path.basename(file_path)}")
            if self.has_saved_grid_settings():
                self.add_result(self.tr("grid_settings_applied"))
            
        except Exception as e:
            QMessageBox.critical(self, self.tr("error"), 
                               self.tr("could_not_load").format(error=str(e)))
    
    def apply_saved_grid_settings(self):
        """Застосувати збережені налаштування сітки до нового зображення"""
        if not self.processor:
            return
        
        # Отримуємо центр нового зображення
        image_center_x = self.processor.image.width // 2
        image_center_y = self.processor.image.height // 2
        
        # Застосовуємо зміщення центру
        if self.saved_grid_settings['center_offset_x'] != 0 or self.saved_grid_settings['center_offset_y'] != 0:
            self.processor.move_center(
                self.saved_grid_settings['center_offset_x'],
                self.saved_grid_settings['center_offset_y']
            )
        
        # Відновлюємо scale edge point відносно центру зображення
        if self.saved_grid_settings.get('scale_edge_relative') and self.saved_grid_settings['custom_scale_distance']:
            edge_relative = self.saved_grid_settings['scale_edge_relative']
            
            # Нові абсолютні координати відносно центру зображення
            new_x = image_center_x + edge_relative['x']
            new_y = image_center_y + edge_relative['y']
            
            # Перевіряємо що координати в межах зображення
            if 0 <= new_x < self.processor.image.width and 0 <= new_y < self.processor.image.height:
                self.scale_edge_point = {'x': new_x, 'y': new_y}
                self.custom_scale_distance = self.saved_grid_settings['custom_scale_distance']
    
        # Відновлюємо масштаб
        if self.saved_grid_settings['scale_value']:
            self.scale_combo.setCurrentText(self.saved_grid_settings['scale_value'])

    def save_current_grid_settings(self):
        """Зберегти поточні налаштування сітки"""
        if not self.processor:
            return
        
        # Зберігаємо зміщення центру відносно центру зображення
        image_center_x = self.processor.image.width // 2
        image_center_y = self.processor.image.height // 2
        
        # Зберігаємо scale edge point відносно центру зображення
        scale_edge_relative = None
        if self.scale_edge_point:
            scale_edge_relative = {
                'x': self.scale_edge_point['x'] - image_center_x,
                'y': self.scale_edge_point['y'] - image_center_y
            }
        
        self.saved_grid_settings = {
            'center_offset_x': self.processor.center_x - image_center_x,
            'center_offset_y': self.processor.center_y - image_center_y,
            'scale_edge_relative': scale_edge_relative,
            'custom_scale_distance': self.custom_scale_distance,
            'scale_value': self.scale_combo.currentText()
        }

    def has_saved_grid_settings(self):
        """Перевірити чи є збережені налаштування сітки"""
        return (self.saved_grid_settings['center_offset_x'] != 0 or 
                self.saved_grid_settings['center_offset_y'] != 0 or
                self.saved_grid_settings.get('scale_edge_relative') is not None or
                self.saved_grid_settings['scale_value'] != "300")
    
    def display_image(self):
        if not self.processor:
            return
        
        pil_image = self.processor.image.copy()
        draw = ImageDraw.Draw(pil_image)
        
        center_x, center_y = self.processor.center_x, self.processor.center_y
        cross_size = 15
        
        # Малюємо червоний хрестик в центрі
        draw.line([center_x - cross_size, center_y, center_x + cross_size, center_y], 
                fill='red', width=2)
        draw.line([center_x, center_y - cross_size, center_x, center_y + cross_size], 
                fill='red', width=2)
        draw.ellipse([center_x - 3, center_y - 3, center_x + 3, center_y + 3], 
                    fill='red', outline='white')
        
        if self.current_click:
            click_x, click_y = self.current_click['x'], self.current_click['y']
            
            # Малюємо синю точку аналізу
            draw.ellipse([click_x - 4, click_y - 4, click_x + 4, click_y + 4], 
                        fill='blue', outline='white', width=1)
            
            # ОНОВЛЕНА ЛІНІЯ: Розраховуємо кінцеву позицію як в документі
            image_width = pil_image.width
            image_height = pil_image.height
            
            # Розрахунок позиції кінця лінії (на рівні підкреслення номера цілі)
            # Використовуємо ту ж логіку що й в create_processed_image_from_data
            right_cell_start_x = int(image_width * 0.85)  # Права комірка починається з 85% ширини
            underline_y = int(image_height * 0.1)         # Позиція підкреслення: 10% висоти зверху
            
            # Кінцева точка лінії: самий правий край на рівні підкреслення
            end_x = image_width - 1  # Самий край зображення
            end_y = underline_y      # На висоті підкреслення номера цілі
            
            # Малюємо оновлену лінію від точки аналізу до розрахованої позиції
            draw.line([click_x, click_y, end_x, end_y], fill='blue', width=3)
        
        # Малюємо зелену точку та лінію для scale edge (якщо є)
        if self.scale_edge_point:
            edge_x, edge_y = self.scale_edge_point['x'], self.scale_edge_point['y']
            
            draw.ellipse([edge_x - 5, edge_y - 5, edge_x + 5, edge_y + 5], 
                        fill='green', outline='white', width=2)
            
            draw.line([center_x, center_y, edge_x, edge_y], fill='green', width=2)
            
            # Перпендикулярна лінія на кінці
            dx = edge_x - center_x
            dy = edge_y - center_y
            length = math.sqrt(dx*dx + dy*dy)
            if length > 0:
                nx, ny = -dy/length, dx/length
                perp_size = 8
                draw.line([
                    edge_x + nx*perp_size, edge_y + ny*perp_size,
                    edge_x - nx*perp_size, edge_y - ny*perp_size
                ], fill='green', width=2)
        
        # Зберігаємо та відображаємо зображення
        with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as temp_file:
            temp_path = temp_file.name
            pil_image.save(temp_path, 'JPEG', quality=95)
        
        pixmap = QPixmap(temp_path)
        
        try:
            os.remove(temp_path)
        except:
            pass
        
        widget_width = self.image_label.width()
        widget_height = self.image_label.height()
        
        scaled_pixmap = pixmap.scaled(widget_width, widget_height, 
                                    Qt.KeepAspectRatio, Qt.SmoothTransformation)
        
        original_width = pixmap.width()
        original_height = pixmap.height()
        scaled_width = scaled_pixmap.width()
        scaled_height = scaled_pixmap.height()
        
        self.scale_factor_x = scaled_width / original_width
        self.scale_factor_y = scaled_height / original_height
        
        self.offset_x = (widget_width - scaled_width) // 2
        self.offset_y = (widget_height - scaled_height) // 2
        
        self.image_label.update_image_geometry(
            original_width, original_height,
            self.scale_factor_x, self.scale_factor_y,
            self.offset_x, self.offset_y
        )
        
        self.image_label.set_zoom_source_image(pixmap)
        self.image_label.setPixmap(scaled_pixmap)
    
    def on_image_click(self, x, y):
        if not self.processor:
            return

        if self.scale_edge_mode:
            self.set_scale_edge_point(x, y)
            return

        if self.center_setting_mode:
            self.set_center_point(x, y)
            return

        if self.current_click:
            existing_widget_x = self.current_click['x'] * self.scale_factor_x + self.offset_x
            existing_widget_y = self.current_click['y'] * self.scale_factor_y + self.offset_y
        
            clicked_widget_x = x * self.scale_factor_x + self.offset_x
            clicked_widget_y = y * self.scale_factor_y + self.offset_y
        
            distance = ((clicked_widget_x - existing_widget_x)**2 + (clicked_widget_y - existing_widget_y)**2)**0.5
        
            if distance <= 15:
                return

        self.place_analysis_point(x, y)
    
    def on_image_drag(self, x, y):
        """Перетягування точки аналізу при утримуванні кнопки миші"""
        if not self.processor:
            return
        
        # НЕ обробляємо перетягування в спеціальних режимах
        if self.scale_edge_mode or self.center_setting_mode:
            return
        
        # Перетягуємо точку аналізу
        self.place_analysis_point(x, y)
    
    def on_mouse_hover(self, x, y):
        if not self.processor or not self.current_click:
            QToolTip.hideText()
            return
        
        hover_widget_x = x * self.scale_factor_x + self.offset_x
        hover_widget_y = y * self.scale_factor_y + self.offset_y
        
        existing_widget_x = self.current_click['x'] * self.scale_factor_x + self.offset_x
        existing_widget_y = self.current_click['y'] * self.scale_factor_y + self.offset_y
        
        distance = ((hover_widget_x - existing_widget_x)**2 + (hover_widget_y - existing_widget_y)**2)**0.5
        
        if distance <= 15:
            azimuth = self.current_click['azimuth']
            range_val = self.current_click['range']
            tooltip_text = f"{self.tr('azimuth')}: {azimuth:.0f}°\n{self.tr('range')}: {range_val:.0f} км"
            
            point = self.image_label.mapToGlobal(self.image_label.rect().topLeft())
            tooltip_x = point.x() + hover_widget_x + 15
            tooltip_y = point.y() + hover_widget_y - 10
            
            QToolTip.showText(QPoint(int(tooltip_x), int(tooltip_y)), tooltip_text)
        else:
            QToolTip.hideText()
    
    def place_analysis_point(self, x, y):
        if not self.processor:
            return
        
        x = max(0, min(x, self.processor.image.width - 1))
        y = max(0, min(y, self.processor.image.height - 1))
        
        try:
            azimuth, range_val = self.calculate_azimuth_range(x, y)
            
            self.current_click = {
                'x': x, 'y': y,
                'azimuth': azimuth, 'range': range_val
            }
            
            self.display_image()
            self.update_results_display()
            self.update_report_data()
            
        except Exception as e:
            QMessageBox.critical(self, self.tr("error"), f"Could not process point: {str(e)}")
    
    def calculate_azimuth_range(self, x, y):
        dx = x - self.processor.center_x
        dy = self.processor.center_y - y
        
        range_pixels = math.sqrt(dx**2 + dy**2)
        
        if self.custom_scale_distance:
            scale_value = int(self.scale_combo.currentText())
            range_actual = (range_pixels / self.custom_scale_distance) * scale_value
        else:
            bottom_edge_distance = self.processor.image.height - self.processor.center_y
            scale_value = int(self.scale_combo.currentText())
            range_actual = (range_pixels / bottom_edge_distance) * scale_value
        
        azimuth_radians = math.atan2(dx, dy)
        azimuth_degrees = math.degrees(azimuth_radians)
        
        if azimuth_degrees < 0:
            azimuth_degrees += 360
            
        return azimuth_degrees, range_actual
    
    def toggle_center_setting_mode(self):
        self.center_setting_mode = self.set_center_btn.isChecked()
        
        if self.center_setting_mode and self.scale_edge_mode:
            self.scale_edge_mode = False
            self.scale_edge_btn.setChecked(False)
            self.scale_edge_btn.setStyleSheet("")
        
        self.image_label.set_center_setting_mode(self.center_setting_mode)
        
        if self.center_setting_mode:
            self.set_center_btn.setStyleSheet("background-color: #FF9800; color: white;")
            
            # ⭐ ДАЄМО ФОКУС ЗОБРАЖЕННЮ ДЛЯ ЗБЕРЕЖЕННЯ ФУНКЦІОНАЛЬНОСТІ МИШІ
            if hasattr(self, 'image_label'):
                self.image_label.setFocus()
                # Встановлюємо политику фокусу
                self.image_label.setFocusPolicy(Qt.StrongFocus)
            
            # ⭐ ПОКАЗУЄМО ЗУМ НА ПОТОЧНІЙ ПОЗИЦІЇ ЦЕНТРУ
            if hasattr(self, 'image_label') and self.processor:
                self.image_label.zoom_widget.show_zoom()
                self.image_label.zoom_widget.update_cursor_position(
                    self.processor.center_x, self.processor.center_y
                )
            self.add_result("🎯 Режим центру: ←→↑↓ для переміщення, Esc для виходу")
            self.add_result("   Shift+стрілка = швидше, Ctrl+стрілка = точніше")
        else:
            self.set_center_btn.setStyleSheet("")
            # ⭐ ВІДНОВЛЮЄМО СТАНДАРТНУ ПОЛІТИКУ ФОКУСУ
            if hasattr(self, 'image_label'):
                self.image_label.setFocusPolicy(Qt.ClickFocus)
            # ⭐ ХОВАЄМО ЗУМ
            if hasattr(self, 'image_label'):
                self.image_label.zoom_widget.hide_zoom()

    def toggle_scale_edge_mode(self):
        self.scale_edge_mode = self.scale_edge_btn.isChecked()
        
        if self.scale_edge_mode and self.center_setting_mode:
            self.center_setting_mode = False
            self.set_center_btn.setChecked(False)
            self.set_center_btn.setStyleSheet("")
        
        self.image_label.set_scale_edge_mode(self.scale_edge_mode)
        
        if self.scale_edge_mode:
            self.scale_edge_btn.setStyleSheet("background-color: #4CAF50; color: white;")
            
            # ⭐ ДАЄМО ФОКУС ЗОБРАЖЕННЮ ДЛЯ ЗБЕРЕЖЕННЯ ФУНКЦІОНАЛЬНОСТІ МИШІ
            if hasattr(self, 'image_label'):
                self.image_label.setFocus()
                # Встановлюємо политику фокусу
                self.image_label.setFocusPolicy(Qt.StrongFocus)
            
            # ⭐ ПОКАЗУЄМО ЗУМ НА ПОЗИЦІЇ КРАЮ МАСШТАБУ
            if hasattr(self, 'image_label') and self.processor:
                if self.scale_edge_point:
                    self.image_label.zoom_widget.show_zoom()
                    self.image_label.zoom_widget.update_cursor_position(
                        self.scale_edge_point['x'], self.scale_edge_point['y']
                    )
                else:
                    # Якщо немає точки, показуємо біля центру
                    self.image_label.zoom_widget.show_zoom()
                    self.image_label.zoom_widget.update_cursor_position(
                        self.processor.center_x + 50, self.processor.center_y + 50
                    )
            self.add_result("📏 Режим масштабу: ←→↑↓ для переміщення, Esc для виходу")
            self.add_result("   Shift+стрілка = швидше, Ctrl+стрілка = точніше")
        else:
            self.scale_edge_btn.setStyleSheet("")
            # ⭐ ВІДНОВЛЮЄМО СТАНДАРТНУ ПОЛІТИКУ ФОКУСУ
            if hasattr(self, 'image_label'):
                self.image_label.setFocusPolicy(Qt.ClickFocus)
            # ⭐ ХОВАЄМО ЗУМ
            if hasattr(self, 'image_label'):
                self.image_label.zoom_widget.hide_zoom()
    
    def set_center_point(self, x, y):
        """Встановлення центру з збереженням налаштувань та зумом"""
        if not self.processor:
            return
        
        x = max(0, min(x, self.processor.image.width - 1))
        y = max(0, min(y, self.processor.image.height - 1))
        
        current_center_x = self.processor.center_x
        current_center_y = self.processor.center_y
        
        dx = x - current_center_x
        dy = y - current_center_y
        
        self.processor.move_center(dx, dy)
        self.save_current_grid_settings()
        
        # ⭐ ОНОВЛЮЄМО ЗУМ НА НОВІЙ ПОЗИЦІЇ
        if hasattr(self, 'image_label'):
            self.image_label.zoom_widget.update_cursor_position(
                self.processor.center_x, self.processor.center_y
            )
        
        self.display_image()
        
        if self.current_click:
            azimuth, range_val = self.calculate_azimuth_range(
                self.current_click['x'], self.current_click['y']
            )
            self.current_click['azimuth'] = azimuth
            self.current_click['range'] = range_val
        
        self.update_results_display()
        self.update_report_data()
        self.add_result(f"Центр встановлено: ({self.processor.center_x}, {self.processor.center_y})")

    
    def set_scale_edge_point(self, x, y):
        """Встановлення scale edge з збереженням налаштувань та зумом"""
        x = max(0, min(x, self.processor.image.width - 1))
        y = max(0, min(y, self.processor.image.height - 1))
        
        dx = x - self.processor.center_x
        dy = y - self.processor.center_y
        distance = math.sqrt(dx*dx + dy*dy)
        
        self.scale_edge_point = {'x': x, 'y': y}
        self.custom_scale_distance = distance
        self.save_current_grid_settings()
        
        # ⭐ ОНОВЛЮЄМО ЗУМ НА НОВІЙ ПОЗИЦІЇ
        if hasattr(self, 'image_label'):
            self.image_label.zoom_widget.update_cursor_position(x, y)
        
        self.display_image()
        
        if self.current_click:
            azimuth, range_val = self.calculate_azimuth_range(
                self.current_click['x'], self.current_click['y']
            )
            self.current_click['azimuth'] = azimuth
            self.current_click['range'] = range_val
        
        self.update_results_display()
        self.update_report_data()
        self.add_result(f"Край масштабу: ({x}, {y}) | Відстань: {distance:.1f}px")

    def update_target_number(self, text):
        self.current_target_number = text
    
    def update_height(self, text):
        self.current_height = text
    
    def update_obstacles(self, text):
        self.current_obstacles = text
    
    def update_detection(self, text):
        self.current_detection = text
    
    def update_report_data(self):
        """Оновлення даних в правій панелі (азимут, дальність, масштаб)"""
        if not self.processor:
            self.auto_azimuth_label.setText("β - --°")
            self.auto_distance_label.setText("D - -- км")
            self.auto_scale_label.setText("M = --")
            return
            
        if self.current_click:
            azimuth = self.current_click['azimuth']
            distance = self.current_click['range']
            scale = int(self.scale_combo.currentText())
            
            self.auto_azimuth_label.setText(f"β - {azimuth:.0f}ᴼ")
            self.auto_distance_label.setText(f"D - {distance:.0f} км")  # ← ЗМІНЕНО .1f на .0f
            self.auto_scale_label.setText(f"M = {scale}")
        else:
            self.auto_azimuth_label.setText("β - --ᴼ")
            self.auto_distance_label.setText("D - -- км")
            
            if hasattr(self, 'scale_combo'):
                scale = int(self.scale_combo.currentText())
                self.auto_scale_label.setText(f"M = {scale}")
            else:
                self.auto_scale_label.setText("M = --")

    def update_results_display(self):
        self.results_text.clear()
        
        if self.processor:
            self.add_result(self.tr("image_info").format(name=os.path.basename(self.current_image_path)))
            self.add_result(self.tr("size").format(width=self.processor.image.width, 
                                                  height=self.processor.image.height))
            self.add_result(self.tr("scale_info").format(scale=int(self.scale_combo.currentText())))
            self.add_result(self.tr("center_info").format(x=self.processor.center_x, 
                                                         y=self.processor.center_y))
            
            if self.custom_scale_distance:
                self.add_result(f"Custom scale edge: {self.custom_scale_distance:.1f} px = {self.scale_combo.currentText()} units")
            else:
                bottom_distance = self.processor.image.height - self.processor.center_y
                self.add_result(self.tr("bottom_edge").format(scale=int(self.scale_combo.currentText())))
                self.add_result(self.tr("pixels_south").format(pixels=bottom_distance))
            self.add_result("")
        
        if self.current_click:
            self.add_result(self.tr("analysis_point"))
            self.add_result(f"{self.tr('position')}: ({self.current_click['x']}, {self.current_click['y']})")
            self.add_result(f"{self.tr('azimuth')}: {self.current_click['azimuth']:.0f}ᴼ")
            self.add_result(f"{self.tr('range')}: {self.current_click['range']:.0f} км")  # ← ЗМІНЕНО .0f
            self.add_result("")
            self.add_result(self.tr("click_to_place"))
            self.add_result(self.tr("drag_to_move"))
            self.add_result(self.tr("line_connects"))
        else:
            self.add_result(self.tr("click_on_image"))

    def update_scale(self):
        """Оновлення масштабу з збереженням налаштувань"""
        if self.processor:
            new_scale = int(self.scale_combo.currentText())
            
            # ЗБЕРЕГТИ налаштування сітки
            self.save_current_grid_settings()
            
            if self.current_click:
                azimuth, range_val = self.calculate_azimuth_range(
                    self.current_click['x'], self.current_click['y']
                )
                self.current_click['azimuth'] = azimuth
                self.current_click['range'] = range_val
            
            self.update_results_display()
            self.update_report_data()
            self.add_result(self.tr("scale_updated").format(scale=new_scale))
            self.add_result("Grid settings saved for next images")
    
    def move_center(self, dx, dy):
        if not self.processor:
            QMessageBox.warning(self, self.tr("warning"), self.tr("no_image_first"))
            return
        
        self.processor.move_center(dx, dy)
        
        if self.current_click:
            azimuth, range_val = self.calculate_azimuth_range(
                self.current_click['x'], self.current_click['y']
            )
            self.current_click['azimuth'] = azimuth
            self.current_click['range'] = range_val
        
        self.display_image()
        self.update_results_display()
        self.update_report_data()
        self.add_result(self.tr("center_moved").format(x=self.processor.center_x, 
                                                    y=self.processor.center_y))
    
        if self.image_label.zoom_widget.isVisible():
            self.image_label.zoom_widget.update_cursor_position(self.processor.center_x, self.processor.center_y)
    
    def clear_results(self):
        self.current_click = None
        if self.processor:
            self.display_image()
        self.update_results_display()
        self.update_report_data()

    def save_current_image(self):
        if not self.processor:
            QMessageBox.warning(self, self.tr("warning"), self.tr("no_image_first"))
            return
        
        if not self.current_click:
            QMessageBox.warning(self, self.tr("warning"), self.tr("no_analysis_point"))
            return
        
        file_path, _ = QFileDialog.getSaveFileName(
            self, self.tr("save_processed_image"), "",
            f"{self.tr('jpeg_files')} (*.jpg);;{self.tr('png_files')} (*.png);;{self.tr('all_files')} (*.*)"
        )
        
        if file_path:
            try:
                final_image = self.processor.image.copy()
                
                if file_path.lower().endswith(('.jpg', '.jpeg')):
                    if final_image.mode != 'RGB':
                        if final_image.mode == 'RGBA':
                            rgb_image = Image.new('RGB', final_image.size, (255, 255, 255))
                            rgb_image.paste(final_image, mask=final_image.split()[-1])
                            final_image = rgb_image
                        else:
                            final_image = final_image.convert('RGB')
                
                draw = ImageDraw.Draw(final_image)
                
                draw.line([
                    (self.current_click['x'], self.current_click['y']), 
                    (final_image.width - 1, 0)
                ], fill='blue', width=3)
                
                if file_path.lower().endswith('.png'):
                    final_image.save(file_path, 'PNG')
                else:
                    final_image.save(file_path, 'JPEG', quality=95)
                
                self.add_result(f"{self.tr('saved')}: {os.path.basename(file_path)}")
                self.add_result(self.tr("saved_clean"))
                QMessageBox.information(self, self.tr("success"), self.tr("image_saved_successfully"))
                
            except Exception as e:
                QMessageBox.critical(self, self.tr("error"), 
                                   self.tr("could_not_save").format(error=str(e)))

    def save_current_image_data(self):
        if not self.processor or not self.current_click:
            QMessageBox.warning(self, "Warning", "No image or analysis point to save")
            return False
        
        image_data = {
            'image_path': self.current_image_path,
            'image_name': os.path.basename(self.current_image_path),
            'analysis_point': {
                'x': self.current_click['x'],
                'y': self.current_click['y'],
                'azimuth': self.current_click['azimuth'],
                'range': self.current_click['range']
            },
            'target_data': {
                'number': self.current_target_number,
                'height': self.current_height,
                'obstacles': self.current_obstacles,
                'detection': self.current_detection,
                'scale': int(self.scale_combo.currentText())
            },
            'processor_settings': {
                'center_x': self.processor.center_x,
                'center_y': self.processor.center_y,
                'scale_edge_point': self.scale_edge_point,
                'custom_scale_distance': self.custom_scale_distance
            }
        }

        # ===== ДОДАЄМО ДАНІ ОПИСУ РЛС =====
        radar_data = self.get_radar_description_data()
        if radar_data:
            image_data['radar_description'] = radar_data
            self.add_result("📡 Дані РЛС додано до зображення")
        
        existing_index = -1
        for i, saved_data in enumerate(self.processed_images):
            if saved_data['image_path'] == self.current_image_path:
                existing_index = i
                break
        
        if existing_index >= 0:
            self.processed_images[existing_index] = image_data
            self.add_result(f"Updated data for: {image_data['image_name']}")
        else:
            self.processed_images.append(image_data)
            self.add_result(f"Saved data for: {image_data['image_name']}")

        # Позначаємо зображення як оброблене
        self.processed_image_paths.add(self.current_image_path)
        
        # Оновлюємо мініатюри
        self.update_thumbnail_processed_status()
        
        self.add_result(f"Total processed images: {len(self.processed_images)}")
        return True

    def next_image_in_batch(self):
        if not self.current_folder_images:
            QMessageBox.warning(self, "Warning", "No folder loaded for batch processing")
            return
        
        if self.processor and self.current_click:
            self.save_current_image_data()
        
        self.current_image_index = (self.current_image_index + 1) % len(self.current_folder_images)
        next_image = self.current_folder_images[self.current_image_index]
        
        self.load_image(next_image)
        self.add_result(f"Loaded image {self.current_image_index + 1}/{len(self.current_folder_images)}: {os.path.basename(next_image)}")

    def previous_image_in_batch(self):
        if not self.current_folder_images:
            QMessageBox.warning(self, "Warning", "No folder loaded for batch processing")
            return
        
        if self.processor and self.current_click:
            self.save_current_image_data()
        
        self.current_image_index = (self.current_image_index - 1) % len(self.current_folder_images)
        prev_image = self.current_folder_images[self.current_image_index]
        
        self.load_image(prev_image)
        self.add_result(f"Loaded image {self.current_image_index + 1}/{len(self.current_folder_images)}: {os.path.basename(prev_image)}")

    def update_thumbnail_processed_status(self):
        """Оновлення візуального статусу оброблених мініатюр"""
        if hasattr(self, 'thumbnail_widget'):
            self.thumbnail_widget.mark_image_as_processed(self.current_image_path)

# ===== ГОЛОВНА ФУНКЦІЯ ТА ЗАПУСК ПРОГРАМИ =====

def main():
    app = QApplication(sys.argv)
    window = AzimuthGUI()
    window.show()
    sys.exit(app.exec_())

if __name__ == "__main__":
    main()