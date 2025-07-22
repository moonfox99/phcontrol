#!/usr/bin/env python3
"""
Azimuth Image Processor - PyQt5 Version with Ukrainian Translation
Professional interface for processing azimuth grid images
"""

import sys
import os
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QHBoxLayout, 
                             QVBoxLayout, QGridLayout, QPushButton, QLabel, 
                             QComboBox, QTextEdit, QScrollArea, QFrame,
                             QFileDialog, QMessageBox, QSizePolicy, QSplitter, QToolTip, QLineEdit)
from PyQt5.QtCore import Qt, QSize, pyqtSignal, QTimer, QPoint
from PyQt5.QtGui import QPixmap, QPainter, QPen, QBrush, QFont, QCursor
from PIL import Image, ImageDraw
import math
from image_processor import AzimuthImageProcessor

# Try to import python-docx for Word document generation
try:
    from docx import Document
    from docx.shared import Inches, Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. Install with: pip install python-docx")

class Translations:
    """Language translations for the application"""
    
    def __init__(self):
        self.languages = {
            'ENGLISH': {
                # Window and main interface
                "window_title": "Azimuth Image Processor - PyQt5",
                "controls": "Image Processing",
                "photo_browser": "Photo Browser:",
                "results": "Results:",
                "settings": "Settings",
                "language": "Language:",
                
                # File operations
                "file_operations": "Image Selection:",
                "open_image": "Open Image",
                "open_folder": "Open Folder",
                "save_current_image": "Save Current Image",
                "export_to_word": "Export to Word Album",
                "create_new_album": "Create New Album",
                "add_to_existing_album": "Add to Existing Album",
                
                # Azimuth grid
                "azimuth_grid": "Azimuth Grid:",
                "scale_setting": "Scale:",
                
                # Center adjustment
                "center_adjustment": "Grid Center:",
                "move_center": "Move Center:",
                
                # Controls
                "clear_analysis_point": "Clear Analysis Point",
                "set_scale_edge": "Set Scale Edge",
                "scale_edge_mode": "Scale Edge Mode",
                "scale_edge_set": "Scale edge set at distance: {distance:.1f} px",
                "click_to_set_edge": "Click on image to set scale edge",
                "scale_edge_active": "Scale edge mode active - click to set edge",
                
                # Report data panel
                "report_data": "Target Data",
                "target_number": "Target №:",
                "height": "H -",
                "obstacles": "Obstacles:",
                "detection_type": "Detection:",
                "no_obstacles": "без перешкод",
                "with_obstacles": "з перешкодами", 
                "detection": "Виявлення",
                "tracking": "Супров-ня",
                "loss": "Втрата",
                "place_point_first": "Place analysis point to enable data",
                "km_unit": "км",
                
                # Messages and results
                "analysis_point": "=== ANALYSIS POINT ===",
                "position": "Position",
                "azimuth": "Azimuth",
                "range": "Range",
                "units": "units",
                "click_to_place": "• Click to place new point",
                "drag_to_move": "• Drag blue dot to move", 
                "line_connects": "• Line connects to top-right corner",
                "click_on_image": "Click on image to place analysis point",
                
                # File dialogs
                "select_image": "Select Image",
                "select_folder": "Select Photo Folder",
                "save_processed_image": "Save Processed Image",
                "image_files": "Image files",
                "jpeg_files": "JPEG files",
                "png_files": "PNG files",
                "all_files": "All files",
                
                # Status messages
                "loaded": "Loaded",
                "loaded_folder": "Loaded folder",
                "found_images": "Found {count} images",
                "saved": "Saved",
                "saved_clean": "Saved with clean line only",
                "image_info": "Image: {name}",
                "size": "Size: {width}x{height}",
                "scale_info": "Scale: {scale} (range to bottom edge)",
                "center_info": "Center: ({x}, {y})",
                "bottom_edge": "Bottom edge: {scale} units",
                "pixels_south": "(at {pixels} pixels south of center)",
                "loaded_from_browser": "Loaded from browser: {name}",
                "center_moved": "Center moved to: ({x}, {y})",
                "scale_updated": "Scale updated to: {scale}",
                
                # Warnings and errors
                "warning": "Warning",
                "error": "Error",
                "success": "Success",
                "no_image_first": "Please open an image first",
                "no_analysis_point": "No analysis point placed",
                "could_not_load": "Could not load image: {error}",
                "could_not_save": "Could not save image: {error}",
                "could_not_open_folder": "Could not open folder: {error}",
                "image_saved_successfully": "Image saved successfully!",
                "word_album_created": "Word album created successfully!",
                "added_to_word_album": "Added to Word album successfully!",
                "docx_not_available": "python-docx library not installed.\nInstall with: pip install python-docx",
                "load_image_and_point": "Please load an image and set analysis point first",
                
                # Instructions
                "open_instruction": "Open an image or folder to start",
                "no_images_found": "No images found in this folder"
            },
            
            'UKRAINIAN': {
                # Window and main interface
                "window_title": "Обробник зображень азимуту - PyQt5",
                "controls": "Обробка зображення",
                "photo_browser": "Перегляд фото:",
                "results": "Результати:",
                "settings": "Налаштування",
                "language": "Мова:",
                
                # File operations  
                "file_operations": "Вибір зображень:",
                "open_image": "Відкрити зображення",
                "open_folder": "Відкрити папку",
                "save_current_image": "Зберегти поточне зображення",
                "export_to_word": "Експорт в Word альбом",
                "create_new_album": "Створити новий альбом",
                "add_to_existing_album": "Додати до існуючого альбому",
                
                # Azimuth grid
                "azimuth_grid": "Азимутальна сітка:",
                "scale_setting": "Масштаб:",
                
                # Center adjustment
                "center_adjustment": "Центр сітки:",
                "move_center": "Перемістити центр:",
                
                # Controls
                "clear_analysis_point": "Очистити точку аналізу",
                "set_scale_edge": "Встановити межу сітки",
                "scale_edge_mode": "Режим вибору межі сітки",
                "scale_edge_set": "Край масштабу встановлено на відстані: {distance:.1f} пікс",
                "click_to_set_edge": "Клацніть на зображенні, щоб встановити край масштабу",
                "scale_edge_active": "Режим краю масштабу активний - клацніть для встановлення",
                
                # Report data panel
                "report_data": "Дані про ціль",
                "target_number": "Номер цілі:",
                "height": "H -",
                "obstacles": "Перешкоди:",
                "detection_type": "Тип:",
                "no_obstacles": "без перешкод",
                "with_obstacles": "з перешкодами",
                "detection": "Виявлення",
                "tracking": "Супров-ня",
                "loss": "Втрата",
                "place_point_first": "Розмістіть точку аналізу для активації даних",
                "km_unit": "км",
                
                # Messages and results
                "analysis_point": "=== ТОЧКА АНАЛІЗУ ===",
                "position": "Позиція",
                "azimuth": "Азимут",
                "range": "Дальність",
                "units": "одиниць",
                "click_to_place": "• Клацніть, щоб розмістити нову точку",
                "drag_to_move": "• Перетягніть синю точку для переміщення",
                "line_connects": "• Лінія з'єднується з правим верхнім кутом",
                "click_on_image": "Клацніть на зображенні, щоб розмістити точку аналізу",
                
                # File dialogs
                "select_image": "Виберіть зображення",
                "select_folder": "Виберіть папку з фото",
                "save_processed_image": "Зберегти оброблене зображення",
                "image_files": "Файли зображень",
                "jpeg_files": "JPEG файли",
                "png_files": "PNG файли", 
                "all_files": "Всі файли",
                
                # Status messages
                "loaded": "Завантажено",
                "loaded_folder": "Завантажено папку",
                "found_images": "Знайдено {count} зображень",
                "saved": "Збережено",
                "saved_clean": "Збережено з чистою лінією",
                "image_info": "Зображення: {name}",
                "size": "Розмір: {width}x{height}",
                "scale_info": "Масштаб: {scale} (дальність до нижнього краю)",
                "center_info": "Центр: ({x}, {y})",
                "bottom_edge": "Нижній край: {scale} одиниць",
                "pixels_south": "(на {pixels} пікселів південніше центру)",
                "loaded_from_browser": "Завантажено з браузера: {name}",
                "center_moved": "Центр переміщено до: ({x}, {y})",
                "scale_updated": "Масштаб оновлено до: {scale}",
                
                # Warnings and errors
                "warning": "Попередження",
                "error": "Помилка",
                "success": "Успіх",
                "no_image_first": "Спочатку відкрийте зображення",
                "no_analysis_point": "Точка аналізу не розміщена",
                "could_not_load": "Не вдалося завантажити зображення: {error}",
                "could_not_save": "Не вдалося зберегти зображення: {error}",
                "could_not_open_folder": "Не вдалося відкрити папку: {error}",
                "image_saved_successfully": "Зображення успішно збережено!",
                "word_album_created": "Word альбом успішно створено!",
                "added_to_word_album": "Успішно додано до Word альбому!",
                "docx_not_available": "Бібліотека python-docx не встановлена.\nВстановіть командою: pip install python-docx",
                "load_image_and_point": "Спочатку завантажте зображення та встановіть точку аналізу",
                
                # Instructions
                "open_instruction": "Відкрийте зображення або папку для початку",
                "no_images_found": "У цій папці не знайдено зображень"
            }
        }

class ClickableLabel(QLabel):
    """Custom QLabel that maintains strict 15:13 aspect ratio"""
    clicked = pyqtSignal(int, int)
    dragged = pyqtSignal(int, int)
    mouse_moved = pyqtSignal(int, int)
    
    def __init__(self):
        super().__init__()
        # Set minimum size with 15:13 aspect ratio
        self.setMinimumSize(450, 390)  # 450 * 13/15 = 390
        self.setStyleSheet("border: 1px solid gray; background-color: white;")
        self.setAlignment(Qt.AlignCenter)
        self.setText("Open an image or folder to start")
        self.setScaledContents(False)  # We'll handle scaling manually
        self.dragging = False
        self.setMouseTracking(True)
        self.scale_edge_mode = False
        
        # Critical: Store the target aspect ratio
        self.target_aspect_ratio = 15.0 / 13.0
        
        # Initialize display area variables
        self.display_width = 450
        self.display_height = 390
        self.display_x = 0
        self.display_y = 0
    
    def set_scale_edge_mode(self, enabled):
        """Enable/disable scale edge setting mode"""
        self.scale_edge_mode = enabled
        if enabled:
            self.setCursor(Qt.CrossCursor)
        else:
            self.setCursor(Qt.ArrowCursor)
    
    def sizeHint(self):
        """Provide size hint that maintains 15:13 ratio"""
        # Always suggest a size that maintains the aspect ratio
        width = max(450, self.width())
        height = int(width * 13 / 15)
        return QSize(width, height)
    
    def heightForWidth(self, width):
        """Maintain aspect ratio: height should be width * 13/15"""
        return int(width * 13 / 15)
    
    def hasHeightForWidth(self):
        """Indicate that this widget maintains aspect ratio"""
        return True
    
    def resizeEvent(self, event):
        """Maintain 15:13 aspect ratio within available space"""
        super().resizeEvent(event)
        
        # Get available space
        available_width = event.size().width()
        available_height = event.size().height()
        
        # Calculate the largest 15:13 rectangle that fits
        ratio_width = available_height * self.target_aspect_ratio
        ratio_height = available_width / self.target_aspect_ratio
        
        if ratio_width <= available_width:
            # Height is the limiting dimension
            self.display_width = int(ratio_width)
            self.display_height = available_height
        else:
            # Width is the limiting dimension
            self.display_width = available_width
            self.display_height = int(ratio_height)
        
        # Center the display area
        self.display_x = (available_width - self.display_width) // 2
        self.display_y = (available_height - self.display_height) // 2
        
        print(f"ClickableLabel: Available({available_width}x{available_height}) -> "
              f"Display({self.display_width}x{self.display_height}) at ({self.display_x}, {self.display_y})")
    
    def paintEvent(self, event):
        """Custom paint - removed debug border for cleaner look"""
        super().paintEvent(event)
        # Debug border removed for production use
        # Users can still resize and see the title bar clearly
    
    def mousePressEvent(self, event):
        if event.button() == Qt.LeftButton:
            # Only process clicks within the display area
            x, y = event.x(), event.y()
            if (self.display_x <= x <= self.display_x + self.display_width and
                self.display_y <= y <= self.display_y + self.display_height):
                
                # Convert to display area coordinates
                adjusted_x = x - self.display_x
                adjusted_y = y - self.display_y
                
                self.clicked.emit(adjusted_x, adjusted_y)
                if not self.scale_edge_mode:
                    self.dragging = True
    
    def mouseMoveEvent(self, event):
        x, y = event.x(), event.y()
        
        # Convert to display area coordinates (even if outside for hover)
        adjusted_x = x - self.display_x
        adjusted_y = y - self.display_y
        
        self.mouse_moved.emit(adjusted_x, adjusted_y)
        
        # Only emit drag if within display area
        if (self.dragging and event.buttons() & Qt.LeftButton and not self.scale_edge_mode and
            self.display_x <= x <= self.display_x + self.display_width and
            self.display_y <= y <= self.display_y + self.display_height):
            self.dragged.emit(adjusted_x, adjusted_y)
    
    def mouseReleaseEvent(self, event):
        self.dragging = False

class ThumbnailWidget(QWidget):
    """Widget to display image thumbnails in a horizontal scroll area"""
    image_selected = pyqtSignal(str)
    
    def __init__(self):
        super().__init__()
        self.setMinimumHeight(160)  # Changed from setFixedHeight
        self.setMaximumHeight(180)  # Added maximum height
        
        # Use horizontal layout with proper sizing
        self.layout = QHBoxLayout()
        self.layout.setContentsMargins(10, 10, 10, 10)
        self.layout.setSpacing(10)
        self.layout.setAlignment(Qt.AlignLeft)
        self.setLayout(self.layout)
        
        # Important: Set size policy to expand horizontally
        self.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        
        # Style
        self.setStyleSheet("background-color: #f0f0f0; border: 1px solid #ccc;")
    
    def clear_thumbnails(self):
        """Remove all thumbnail widgets"""
        print(f"Clearing {self.layout.count()} existing thumbnails")
        while self.layout.count():
            child = self.layout.takeAt(0)
            if child.widget():
                child.widget().deleteLater()
        print("Thumbnails cleared")
    
    def add_thumbnail(self, image_path):
        """Add a thumbnail for the given image"""
        try:
            # Create thumbnail frame
            thumb_frame = QFrame()
            thumb_frame.setFixedSize(130, 150)
            thumb_frame.setStyleSheet("""
                QFrame {
                    border: 2px solid #ddd;
                    border-radius: 5px;
                    background-color: white;
                    margin: 5px;
                }
                QFrame:hover {
                    border-color: #4CAF50;
                    background-color: #f9f9f9;
                }
            """)
            
            # Layout for frame
            frame_layout = QVBoxLayout()
            frame_layout.setContentsMargins(5, 5, 5, 5)
            thumb_frame.setLayout(frame_layout)
            
            # Create thumbnail image - FIXED FOR ALL IMAGE MODES
            with Image.open(image_path) as pil_image:
                # Handle different image modes properly
                if pil_image.mode == 'RGBA':
                    # PNG with transparency - convert to RGB with white background
                    rgb_image = Image.new('RGB', pil_image.size, (255, 255, 255))
                    rgb_image.paste(pil_image, mask=pil_image.split()[-1])  # Use alpha channel as mask
                    pil_image = rgb_image
                elif pil_image.mode == 'L':
                    # Grayscale - convert to RGB
                    pil_image = pil_image.convert('RGB')
                elif pil_image.mode == 'P':
                    # Palette mode - convert to RGB
                    pil_image = pil_image.convert('RGB')
                elif pil_image.mode not in ['RGB', 'RGBA']:
                    # Any other mode - convert to RGB
                    pil_image = pil_image.convert('RGB')
                
                # Now we have RGB, create thumbnail
                pil_image.thumbnail((120, 120), Image.Resampling.LANCZOS)
                
                # Save to temporary file as PNG (supports all modes)
                import tempfile
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
                    temp_path = temp_file.name
                    pil_image.save(temp_path, 'PNG')
                
                # Load QPixmap from temp file
                qimg = QPixmap(temp_path)
                
                # Clean up temp file
                try:
                    os.remove(temp_path)
                except:
                    pass
            
            # Create clickable label
            thumb_label = QLabel()
            if not qimg.isNull():
                thumb_label.setPixmap(qimg.scaled(120, 120, Qt.KeepAspectRatio, Qt.SmoothTransformation))
                print(f"✓ Successfully created thumbnail for {os.path.basename(image_path)}")  # Debug
            else:
                thumb_label.setText("Error\nLoading")
                thumb_label.setStyleSheet("color: red; text-align: center;")
                print(f"✗ Failed to create QPixmap for {os.path.basename(image_path)}")  # Debug
            
            thumb_label.setAlignment(Qt.AlignCenter)
            thumb_label.mousePressEvent = lambda event, path=image_path: self.image_selected.emit(path)
            thumb_label.setCursor(Qt.PointingHandCursor)
            
            # Filename label
            filename = os.path.basename(image_path)
            if len(filename) > 15:
                filename = filename[:12] + "..."
            name_label = QLabel(filename)
            name_label.setAlignment(Qt.AlignCenter)
            name_label.setFont(QFont("Arial", 8))
            
            # Add to frame
            frame_layout.addWidget(thumb_label)
            frame_layout.addWidget(name_label)
            
            # Add frame to main layout
            self.layout.addWidget(thumb_frame)
            print(f"✓ Added thumbnail widget to layout for {os.path.basename(image_path)}")
            
            # Force immediate update
            thumb_frame.show()
            self.update()
            
        except Exception as e:
            print(f"✗ Error creating thumbnail for {os.path.basename(image_path)}: {e}")
            
            # Create simple text-based thumbnail as fallback
            simple_frame = QFrame()
            simple_frame.setFixedSize(130, 150)
            simple_frame.setStyleSheet("border: 2px solid red; background-color: #ffe6e6;")
            
            simple_layout = QVBoxLayout()
            simple_frame.setLayout(simple_layout)
            
            # Just show filename
            filename = os.path.basename(image_path)
            if len(filename) > 15:
                filename = filename[:12] + "..."
            
            file_label = QLabel(filename)
            file_label.setAlignment(Qt.AlignCenter)
            file_label.setWordWrap(True)
            file_label.setFont(QFont("Arial", 10))
            simple_layout.addWidget(file_label)
            
            # Make it clickable
            simple_frame.mousePressEvent = lambda event, path=image_path: self.image_selected.emit(path)
            simple_frame.setCursor(Qt.PointingHandCursor)
            
            self.layout.addWidget(simple_frame)
            print(f"✓ Added fallback thumbnail for {os.path.basename(image_path)}")

class AzimuthGUI(QMainWindow):
    def __init__(self):
        super().__init__()
        self.processor = None
        self.current_image_path = None
        self.current_click = None
        self.scale_factor_x = 1.0  # X-axis scale factor for coordinate conversion
        self.scale_factor_y = 1.0  # Y-axis scale factor for coordinate conversion  
        self.offset_x = 0          # X offset for centering
        self.offset_y = 0          # Y offset for centering
        self.current_folder = None
        
        # Language settings
        self.current_language = 'ENGLISH'  # Default language
        self.translations = Translations()
        
        # Scale edge settings
        self.scale_edge_mode = False
        self.scale_edge_point = None  # Will store the scale edge coordinates
        self.custom_scale_distance = None  # Custom scale distance in pixels
        
        # Report data settings
        self.current_target_number = "0001"  # Persistent target number
        self.current_height = "0.0"  # Current height value
        self.current_obstacles = "без перешкод"  # Current obstacles status
        self.current_detection = "Виявлення"  # Current detection type
        
        # Word document settings
        self.current_word_document = None  # Current Word document for album
        self.word_document_path = None  # Path to current Word document
        
        self.init_ui()
    
    def tr(self, key):
        """Get translation for the current language"""
        lang_dict = self.translations.languages[self.current_language]
        return lang_dict.get(key, key)  # Return key if translation not found
    
    def set_language(self, language):
        """Change the interface language"""
        self.current_language = language
        self.update_interface_text()
        
        # Update menu checkboxes
        for lang, action in self.language_actions.items():
            action.setChecked(lang == language)
    
    def update_interface_text(self):
        """Update all interface text with current language"""
        # Update window title
        self.setWindowTitle(self.tr("window_title"))
        
        # Update menu bar
        self.create_menu_bar()  # Recreate menu with new language
        
        # Update title of control panel
        if hasattr(self, 'control_title'):
            self.control_title.setText(self.tr("controls"))
        
        # Update title of report panel
        if hasattr(self, 'report_title'):
            self.report_title.setText(self.tr("report_data"))
        
        # Update all buttons and labels
        if hasattr(self, 'open_image_btn'):
            self.open_image_btn.setText(self.tr("open_image"))
        if hasattr(self, 'open_folder_btn'):
            self.open_folder_btn.setText(self.tr("open_folder"))
        if hasattr(self, 'save_image_btn'):
            self.save_image_btn.setText(self.tr("save_current_image"))
        if hasattr(self, 'export_new_btn'):
            self.export_new_btn.setText(self.tr("create_new_album"))
        if hasattr(self, 'export_add_btn'):
            self.export_add_btn.setText(self.tr("add_to_existing_album"))
        if hasattr(self, 'clear_btn'):
            self.clear_btn.setText(self.tr("clear_analysis_point"))
        if hasattr(self, 'scale_edge_btn'):
            self.scale_edge_btn.setText(self.tr("set_scale_edge"))
        
        # Update section labels
        if hasattr(self, 'file_ops_label'):
            self.file_ops_label.setText(self.tr("file_operations"))
        if hasattr(self, 'azimuth_grid_label'):
            self.azimuth_grid_label.setText(self.tr("azimuth_grid"))
        if hasattr(self, 'move_center_label'):
            self.move_center_label.setText(self.tr("move_center"))
        if hasattr(self, 'scale_edge_label'):
            self.scale_edge_label.setText(f"{self.tr('scale_edge_mode')}:")
        if hasattr(self, 'results_label'):
            self.results_label.setText(self.tr("results"))
        if hasattr(self, 'browser_label'):
            self.browser_label.setText(self.tr("photo_browser"))
        
        # Update labels that are created dynamically
        if not self.current_image_path:
            self.image_label.setText(self.tr("open_instruction"))
        
        # Update results if there's current data
        if self.processor:
            self.update_results_display()

    def init_ui(self):
        """Initialize the user interface"""
        self.setWindowTitle(self.tr("window_title"))
        
        # Set window size with 15:13 aspect ratio
        window_width = 1500
        window_height = 1300
        
        # Get screen geometry excluding taskbar
        screen = QApplication.primaryScreen()
        screen_geometry = screen.availableGeometry()  # This excludes taskbar/dock areas
        
        # Calculate safe position - ensure window fits completely on screen with title bar access
        max_x = screen_geometry.width() - window_width
        max_y = screen_geometry.height() - window_height
        
        # Center horizontally, position in upper area but leave room for title bar
        x = screen_geometry.x() + max(0, max_x // 2)
        y = screen_geometry.y() + max(30, min(100, max_y // 4))  # At least 30px from top, max 100px
        
        # Ensure window doesn't go off screen
        x = max(screen_geometry.x(), min(x, screen_geometry.x() + max_x))
        y = max(screen_geometry.y(), min(y, screen_geometry.y() + max_y))
        
        print(f"Screen available: {screen_geometry.width()}x{screen_geometry.height()} at ({screen_geometry.x()}, {screen_geometry.y()})")
        print(f"Window position: ({x}, {y}) size: {window_width}x{window_height}")
        
        # Ensure normal window behavior (title bar, close button, drag capability)
        self.setWindowFlags(Qt.Window)  # Standard window with all controls
        self.setGeometry(x, y, window_width, window_height)
        
        # Set minimum size to maintain proportions (reasonable minimum)
        min_width = 900  
        min_height = int(min_width * 13 / 15)  # Maintain 15:13 ratio
        self.setMinimumSize(min_width, min_height)
        
        # Create menu bar
        self.create_menu_bar()
        
        # Central widget
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        # Main layout
        main_layout = QHBoxLayout()
        central_widget.setLayout(main_layout)
        
        # Create splitter for resizable panels
        main_splitter = QSplitter(Qt.Horizontal)
        main_layout.addWidget(main_splitter)
        
        # Left panel - Controls (fixed width)
        self.create_control_panel(main_splitter)
        
        # Center panel - Image and thumbnails (expandable but constrained)
        self.create_image_panel(main_splitter)
        
        # Right panel - Report data (fixed width)
        self.create_report_panel(main_splitter)
        
        # Set splitter proportions and constraints
        main_splitter.setStretchFactor(0, 0)  # Controls panel - never stretch
        main_splitter.setStretchFactor(1, 1)  # Image panel - can stretch
        main_splitter.setStretchFactor(2, 0)  # Report panel - never stretch
        
        # Set initial sizes: Controls(300) | Image(900) | Report(300)
        main_splitter.setSizes([300, 900, 300])
        
        # Store reference to splitter for proportion management
        self.main_splitter = main_splitter
    
    def resizeEvent(self, event):
        """Handle window resize to maintain image area proportions"""
        super().resizeEvent(event)
        
        # Calculate how much space is available for the image area
        total_width = event.size().width()
        total_height = event.size().height()
        
        # Account for fixed panels (left: 300px, right: 300px) and margins
        fixed_width = 600 + 40  # 300 + 300 + margins
        available_width = max(300, total_width - fixed_width)  # Minimum 300px for image
        
        # Calculate height based on 15:13 ratio, accounting for thumbnails and margins
        available_height = max(200, total_height - 200)  # Account for thumbnails and margins
        
        # The image widget should maintain 15:13 ratio within available space
        target_ratio = 15.0 / 13.0
        
        # Calculate ideal dimensions
        width_for_height = available_height * target_ratio
        height_for_width = available_width / target_ratio
        
        if width_for_height <= available_width:
            # Height is limiting factor
            ideal_image_width = int(width_for_height)
            ideal_image_height = available_height
        else:
            # Width is limiting factor  
            ideal_image_width = available_width
            ideal_image_height = int(height_for_width)
        
        print(f"Window resize: {total_width}x{total_height} -> Image area: {ideal_image_width}x{ideal_image_height}")
        
        # Update splitter sizes to maintain proportions
        # Keep side panels fixed, adjust center panel
        current_sizes = self.main_splitter.sizes()
        new_center_size = ideal_image_width + 40  # Add some padding
        self.main_splitter.setSizes([300, new_center_size, 300])
        
        # If image is loaded, update its display after a short delay
        if hasattr(self, 'processor') and self.processor:
            QTimer.singleShot(100, self.update_image_display_after_resize)
    
    def update_image_display_after_resize(self):
        """Update image display after window resize to maintain proportions"""
        if hasattr(self, 'processor') and self.processor:
            self.display_image()
    
    def create_menu_bar(self):
        """Create menu bar with Settings"""
        menubar = self.menuBar()
        menubar.clear()  # Clear existing menu
        
        # Settings menu
        settings_menu = menubar.addMenu(self.tr("settings"))
        
        # Language submenu
        language_menu = settings_menu.addMenu(self.tr("language"))
        
        # Language actions
        english_action = language_menu.addAction("English")
        english_action.triggered.connect(lambda: self.set_language('ENGLISH'))
        
        ukrainian_action = language_menu.addAction("Українська")
        ukrainian_action.triggered.connect(lambda: self.set_language('UKRAINIAN'))
        
        # Make actions checkable
        english_action.setCheckable(True)
        ukrainian_action.setCheckable(True)
        
        # Set current language as checked
        english_action.setChecked(self.current_language == 'ENGLISH')
        ukrainian_action.setChecked(self.current_language == 'UKRAINIAN')
        
        # Store actions for later updates
        self.language_actions = {
            'ENGLISH': english_action,
            'UKRAINIAN': ukrainian_action
        }
    
    def create_control_panel(self, parent):
        """Create the left control panel"""
        control_widget = QWidget()
        control_widget.setFixedWidth(300)
        control_widget.setStyleSheet("""
            QWidget {
                background-color: #f5f5f5; 
                border-right: 1px solid #ccc;
            }
            QLabel {
                background: none;
                border: none;
                color: #333;
            }
            QPushButton {
                background-color: #e1e1e1;
                border: 1px solid #c1c1c1;
                border-radius: 4px;
                padding: 8px 12px;
                font-size: 11px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #d1d1d1;
                border: 1px solid #b1b1b1;
            }
            QPushButton:pressed {
                background-color: #c1c1c1;
            }
            QComboBox {
                border: 1px solid #c1c1c1;
                border-radius: 4px;
                padding: 4px 8px;
                background-color: white;
                font-size: 11px;
            }
        """)
        
        layout = QVBoxLayout()
        layout.setContentsMargins(15, 15, 15, 15)
        layout.setSpacing(10)
        control_widget.setLayout(layout)
        
        # Title
        self.control_title = QLabel(self.tr("controls"))
        self.control_title.setFont(QFont("Arial", 14, QFont.Bold))
        self.control_title.setAlignment(Qt.AlignCenter)
        self.control_title.setStyleSheet("font-size: 14px; font-weight: bold; margin-bottom: 10px;")
        layout.addWidget(self.control_title)
        
        # Image Selection section
        self.file_ops_label = QLabel(self.tr("file_operations"))
        self.file_ops_label.setFont(QFont("Arial", 10, QFont.Bold))
        self.file_ops_label.setStyleSheet("color: #666; margin-top: 10px; margin-bottom: 5px;")
        layout.addWidget(self.file_ops_label)
        
        # Buttons
        self.open_image_btn = QPushButton(self.tr("open_image"))
        self.open_image_btn.clicked.connect(self.open_image)
        layout.addWidget(self.open_image_btn)
        
        self.open_folder_btn = QPushButton(self.tr("open_folder"))
        self.open_folder_btn.clicked.connect(self.open_folder)
        layout.addWidget(self.open_folder_btn)
        
        self.save_image_btn = QPushButton(self.tr("save_current_image"))
        self.save_image_btn.clicked.connect(self.save_current_image)
        layout.addWidget(self.save_image_btn)
        
        # Export to Word buttons
        self.export_new_btn = QPushButton(self.tr("create_new_album"))
        self.export_new_btn.clicked.connect(self.create_new_word_album)
        layout.addWidget(self.export_new_btn)
        
        self.export_add_btn = QPushButton(self.tr("add_to_existing_album"))
        self.export_add_btn.clicked.connect(self.add_to_existing_album)
        layout.addWidget(self.export_add_btn)
        
        # Separator line
        separator1 = QFrame()
        separator1.setFrameShape(QFrame.HLine)
        separator1.setFrameShadow(QFrame.Sunken)
        separator1.setStyleSheet("color: #ccc; margin: 10px 0px;")
        layout.addWidget(separator1)
        
        # Azimuth Grid section
        self.azimuth_grid_label = QLabel(self.tr("azimuth_grid"))
        self.azimuth_grid_label.setFont(QFont("Arial", 10, QFont.Bold))
        self.azimuth_grid_label.setStyleSheet("color: #666; margin-top: 5px; margin-bottom: 5px;")
        layout.addWidget(self.azimuth_grid_label)
        
        # Scale setting
        scale_layout = QHBoxLayout()
        scale_label = QLabel(self.tr("scale_setting"))
        scale_label.setStyleSheet("color: #333; margin-right: 5px;")
        scale_layout.addWidget(scale_label)
        self.scale_combo = QComboBox()
        self.scale_combo.addItems(["50", "100", "150", "200", "300"])
        self.scale_combo.setCurrentText("300")
        self.scale_combo.currentTextChanged.connect(self.update_scale)
        scale_layout.addWidget(self.scale_combo)
        scale_layout.addStretch()
        
        scale_widget = QWidget()
        scale_widget.setLayout(scale_layout)
        layout.addWidget(scale_widget)
        
        # Scale edge tool
        self.scale_edge_label = QLabel(f"{self.tr('scale_edge_mode')}:")
        self.scale_edge_label.setStyleSheet("color: #333; margin-top: 5px;")
        layout.addWidget(self.scale_edge_label)
        self.scale_edge_btn = QPushButton(self.tr("set_scale_edge"))
        self.scale_edge_btn.setCheckable(True)
        self.scale_edge_btn.clicked.connect(self.toggle_scale_edge_mode)
        layout.addWidget(self.scale_edge_btn)
        
        # Center movement
        self.move_center_label = QLabel(self.tr("move_center"))
        self.move_center_label.setStyleSheet("color: #333; margin-top: 10px;")
        layout.addWidget(self.move_center_label)
        
        move_layout = QGridLayout()
        move_layout.setSpacing(2)
        
        btn_up = QPushButton("↑")
        btn_up.setMaximumWidth(40)
        btn_up.clicked.connect(lambda: self.move_center(0, -10))
        move_layout.addWidget(btn_up, 0, 1)
        
        btn_left = QPushButton("←")
        btn_left.setMaximumWidth(40)
        btn_left.clicked.connect(lambda: self.move_center(-10, 0))
        move_layout.addWidget(btn_left, 1, 0)
        
        btn_center = QPushButton("●")
        btn_center.setMaximumWidth(40)
        btn_center.clicked.connect(self.show_center_preview)
        move_layout.addWidget(btn_center, 1, 1)
        
        btn_right = QPushButton("→")
        btn_right.setMaximumWidth(40)
        btn_right.clicked.connect(lambda: self.move_center(10, 0))
        move_layout.addWidget(btn_right, 1, 2)
        
        btn_down = QPushButton("↓")
        btn_down.setMaximumWidth(40)
        btn_down.clicked.connect(lambda: self.move_center(0, 10))
        move_layout.addWidget(btn_down, 2, 1)
        
        move_widget = QWidget()
        move_widget.setLayout(move_layout)
        layout.addWidget(move_widget)
        
        # Separator line
        separator2 = QFrame()
        separator2.setFrameShape(QFrame.HLine)
        separator2.setFrameShadow(QFrame.Sunken)
        separator2.setStyleSheet("color: #ccc; margin: 10px 0px;")
        layout.addWidget(separator2)
        
        # Clear button
        self.clear_btn = QPushButton(self.tr("clear_analysis_point"))
        self.clear_btn.clicked.connect(self.clear_results)
        layout.addWidget(self.clear_btn)
        
        # Results area
        self.results_label = QLabel(self.tr("results"))
        self.results_label.setFont(QFont("Arial", 10, QFont.Bold))
        self.results_label.setStyleSheet("color: #666; margin-top: 10px; margin-bottom: 5px;")
        layout.addWidget(self.results_label)
        
        self.results_text = QTextEdit()
        self.results_text.setMaximumHeight(200)
        self.results_text.setStyleSheet("""
            QTextEdit {
                background-color: white;
                border: 1px solid #ccc;
                font-family: 'Courier New', monospace;
                font-size: 10px;
                border-radius: 4px;
            }
        """)
        layout.addWidget(self.results_text)
        
        # Stretch at bottom
        layout.addStretch()
        
        parent.addWidget(control_widget)
    
    def create_image_panel(self, parent):
        """Create the center panel for image display and thumbnails"""
        image_widget = QWidget()
        layout = QVBoxLayout()
        image_widget.setLayout(layout)
        
        # Image display area
        self.image_label = ClickableLabel()
        self.image_label.clicked.connect(self.on_image_click)
        self.image_label.dragged.connect(self.on_image_drag)
        self.image_label.mouse_moved.connect(self.on_mouse_hover)
        self.image_label.setText(self.tr("open_instruction"))
        
        # Scroll area for image
        scroll_area = QScrollArea()
        scroll_area.setWidget(self.image_label)
        scroll_area.setWidgetResizable(True)
        scroll_area.setMinimumHeight(400)
        layout.addWidget(scroll_area)
        
        # Thumbnail browser
        self.browser_label = QLabel(self.tr("photo_browser"))
        self.browser_label.setFont(QFont("Arial", 10, QFont.Bold))
        self.browser_label.setStyleSheet("color: #666; margin-top: 10px; margin-bottom: 5px;")
        layout.addWidget(self.browser_label)
        
        # Create scroll area that properly handles the content
        self.thumbnail_scroll = QScrollArea()
        self.thumbnail_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self.thumbnail_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.thumbnail_scroll.setMinimumHeight(180)  # Increased from fixed height
        self.thumbnail_scroll.setMaximumHeight(200)  # Added maximum height
        
        # Important: Set widget resizable to false and use a proper container
        self.thumbnail_scroll.setWidgetResizable(False)
        
        self.thumbnail_widget = ThumbnailWidget()
        self.thumbnail_widget.image_selected.connect(self.load_image_from_browser)
        self.thumbnail_scroll.setWidget(self.thumbnail_widget)
        
        layout.addWidget(self.thumbnail_scroll)
        
        # Initially hide thumbnail browser and its label
        self.thumbnail_scroll.hide()
        self.browser_label.hide()
        
        parent.addWidget(image_widget)
    
    def create_report_panel(self, parent):
        """Create the right panel for report data entry"""
        report_widget = QWidget()
        report_widget.setFixedWidth(300)
        report_widget.setStyleSheet("background-color: #f9f9f9; border-left: 1px solid #ccc;")
        
        layout = QVBoxLayout()
        report_widget.setLayout(layout)
        
        # Title
        self.report_title = QLabel(self.tr("report_data"))
        self.report_title.setFont(QFont("Arial", 14, QFont.Bold))
        self.report_title.setAlignment(Qt.AlignCenter)
        layout.addWidget(self.report_title)
        
        # Current analysis status
        self.analysis_status_label = QLabel(self.tr("place_point_first"))
        self.analysis_status_label.setStyleSheet("color: #666; font-style: italic; padding: 10px;")
        self.analysis_status_label.setWordWrap(True)
        layout.addWidget(self.analysis_status_label)
        
        # Manual input section
        manual_group = QFrame()
        manual_group.setStyleSheet("background-color: #fff; border: 1px solid #ddd; border-radius: 5px; padding: 10px;")
        manual_layout = QVBoxLayout()
        manual_group.setLayout(manual_layout)
        
        # Target number with placeholder
        self.target_number_input = QLineEdit()
        self.target_number_input.setPlaceholderText("Номер цілі")
        self.target_number_input.setText(self.current_target_number)
        self.target_number_input.textChanged.connect(self.update_target_number)
        manual_layout.addWidget(self.target_number_input)
        
        # Auto β label
        self.auto_azimuth_label = QLabel("β - --°")
        manual_layout.addWidget(self.auto_azimuth_label)
        
        # Auto D label  
        self.auto_distance_label = QLabel("D - -- км")
        manual_layout.addWidget(self.auto_distance_label)
        
        # Height input with H prefix
        height_layout = QHBoxLayout()
        height_layout.addWidget(QLabel("H –"))
        self.height_input = QLineEdit(self.current_height)
        self.height_input.setMaximumWidth(80)
        self.height_input.textChanged.connect(self.update_height)
        height_layout.addWidget(self.height_input)
        height_layout.addWidget(QLabel(self.tr("km_unit")))
        height_layout.addStretch()
        manual_layout.addLayout(height_layout)
        
        # Obstacles dropdown (no label)
        self.obstacles_combo = QComboBox()
        self.obstacles_combo.addItems([self.tr("no_obstacles"), self.tr("with_obstacles")])
        self.obstacles_combo.currentTextChanged.connect(self.update_obstacles)
        manual_layout.addWidget(self.obstacles_combo)
        
        # Detection type dropdown (no label)
        self.detection_combo = QComboBox()
        self.detection_combo.addItems([self.tr("detection"), self.tr("tracking"), self.tr("loss")])
        self.detection_combo.currentTextChanged.connect(self.update_detection)
        manual_layout.addWidget(self.detection_combo)
        
        # Auto M label
        self.auto_scale_label = QLabel("M = --")
        manual_layout.addWidget(self.auto_scale_label)
        
        layout.addWidget(manual_group)
        
        # Stretch to push everything to top
        layout.addStretch()
        
        # Initially hide the panel
        report_widget.hide()
        self.report_widget = report_widget
        
        parent.addWidget(report_widget)
    
    def add_result(self, text):
        """Add text to results area"""
        self.results_text.append(text)
    
    def open_image(self):
        """Open a single image file"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, self.tr("select_image"), "",
            f"{self.tr('image_files')} (*.jpg *.jpeg *.png *.bmp *.gif *.tiff);;{self.tr('all_files')} (*.*)"
        )
        
        if file_path:
            self.load_image(file_path)
            # Show report panel when image is loaded
            self.report_widget.show()
    
    def open_folder(self):
        """Open a folder and display thumbnails"""
        folder_path = QFileDialog.getExistingDirectory(self, self.tr("select_folder"))
        
        if folder_path:
            self.current_folder = folder_path
            self.load_folder_thumbnails()
            
            # Show thumbnail browser and its label
            self.thumbnail_scroll.show()
            self.browser_label.show()
            print("Made thumbnail scroll visible")
            
            # Count images
            image_extensions = ('.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff', '.tif')
            image_count = sum(1 for f in os.listdir(folder_path) 
                             if f.lower().endswith(image_extensions))
            
            self.add_result(f"{self.tr('loaded_folder')}: {os.path.basename(folder_path)}")
            self.add_result(self.tr("found_images").format(count=image_count))
            
            # Show report panel when folder is loaded
            self.report_widget.show()
            
            # Debug: Check if thumbnail browser is actually visible
            print(f"Thumbnail scroll visible: {self.thumbnail_scroll.isVisible()}")
            print(f"Thumbnail widget size: {self.thumbnail_widget.size()}")
            print(f"Thumbnail scroll size: {self.thumbnail_scroll.size()}")

    def load_folder_thumbnails(self):
        """Load thumbnails for all images in the current folder"""
        if not self.current_folder:
            return
        
        print(f"Loading thumbnails from folder: {self.current_folder}")
        
        # Clear existing thumbnails
        self.thumbnail_widget.clear_thumbnails()
        
        # Find image files
        image_extensions = ('.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff', '.tif')
        image_files = []
        
        try:
            for filename in os.listdir(self.current_folder):
                if filename.lower().endswith(image_extensions):
                    full_path = os.path.join(self.current_folder, filename)
                    image_files.append(full_path)
                    print(f"Found image: {filename}")
        except Exception as e:
            print(f"Error reading folder: {e}")
            return
        
        # Sort files
        image_files.sort()
        print(f"Total images found: {len(image_files)}")
        
        if len(image_files) == 0:
            print("No image files found in folder!")
            # Add a message widget
            no_images_label = QLabel(self.tr("no_images_found"))
            no_images_label.setAlignment(Qt.AlignCenter)
            no_images_label.setStyleSheet("color: gray; font-size: 14px; padding: 20px;")
            self.thumbnail_widget.layout.addWidget(no_images_label)
            return
        
        # Create thumbnails
        success_count = 0
        for i, image_path in enumerate(image_files):
            print(f"Creating thumbnail {i+1}/{len(image_files)}: {os.path.basename(image_path)}")
            try:
                self.thumbnail_widget.add_thumbnail(image_path)
                success_count += 1
            except Exception as e:
                print(f"Failed to create thumbnail for {os.path.basename(image_path)}: {e}")
        
        print(f"Successfully created {success_count}/{len(image_files)} thumbnails")
        
        # IMPORTANT: Update the widget size after adding all thumbnails
        widget_width = len(image_files) * 140 + 20  # 130 + 10 spacing per thumbnail + margins
        self.thumbnail_widget.setMinimumWidth(widget_width)
        self.thumbnail_widget.resize(widget_width, 160)
        
        print(f"Set thumbnail widget width to: {widget_width}")
        
        # Check if thumbnails are actually added
        widget_count = self.thumbnail_widget.layout.count()
        print(f"Number of widgets in thumbnail layout: {widget_count}")
        print(f"Final thumbnail widget size: {self.thumbnail_widget.size()}")

    def load_image_from_browser(self, file_path):
        """Load image selected from thumbnail browser"""
        self.load_image(file_path)
        self.add_result(self.tr("loaded_from_browser").format(name=os.path.basename(file_path)))
    
    def load_image(self, file_path):
        """Load and display an image with proper format handling"""
        try:
            self.current_image_path = file_path
            scale_value = int(self.scale_combo.currentText())
            
            # Initialize processor
            self.processor = AzimuthImageProcessor(file_path, scale=scale_value)
            
            # CRITICAL: Force convert processor image to RGB immediately after loading
            if hasattr(self.processor, 'image') and self.processor.image:
                original_mode = self.processor.image.mode
                print(f"Original image mode from processor: {original_mode}")  # Debug
                
                if original_mode != 'RGB':
                    print(f"Converting processor image from {original_mode} to RGB")  # Debug
                    if original_mode == 'RGBA':
                        # Handle transparency
                        rgb_image = Image.new('RGB', self.processor.image.size, (255, 255, 255))
                        rgb_image.paste(self.processor.image, mask=self.processor.image.split()[-1])
                        self.processor.image = rgb_image
                    else:
                        # Convert any other mode to RGB
                        self.processor.image = self.processor.image.convert('RGB')
                    print(f"Processor image converted to: {self.processor.image.mode}")  # Debug
            
            self.display_image()
            self.current_click = None
            self.update_results_display()
            self.update_report_data()  # Update report panel
            
            self.add_result(f"{self.tr('loaded')}: {os.path.basename(file_path)}")
            
        except Exception as e:
            print(f"Error loading image: {e}")  # Debug
            QMessageBox.critical(self, self.tr("error"), 
                               self.tr("could_not_load").format(error=str(e)))
    
    def display_image(self):
        """Display the current image maintaining strict 15:13 proportions"""
        if not self.processor:
            return
        
        # Get PIL image and convert to QPixmap
        pil_image = self.processor.image.copy()
        
        # Draw center crosshairs and current click
        draw = ImageDraw.Draw(pil_image)
        
        # Center crosshairs
        center_x, center_y = self.processor.center_x, self.processor.center_y
        cross_size = 15
        
        draw.line([center_x - cross_size, center_y, center_x + cross_size, center_y], 
                  fill='red', width=2)
        draw.line([center_x, center_y - cross_size, center_x, center_y + cross_size], 
                  fill='red', width=2)
        draw.ellipse([center_x - 3, center_y - 3, center_x + 3, center_y + 3], 
                     fill='red', outline='white')
        
        # Current analysis point
        if self.current_click:
            click_x, click_y = self.current_click['x'], self.current_click['y']
            
            # Draw smaller analysis point (assistance only - not saved)
            draw.ellipse([click_x - 4, click_y - 4, click_x + 4, click_y + 4], 
                         fill='blue', outline='white', width=1)
            
            # Draw line to top-right corner (no triangle)
            draw.line([click_x, click_y, pil_image.width - 1, 0], fill='blue', width=3)
        
        # Draw scale edge point if set
        if self.scale_edge_point:
            edge_x, edge_y = self.scale_edge_point['x'], self.scale_edge_point['y']
            
            # Draw scale edge point in green
            draw.ellipse([edge_x - 5, edge_y - 5, edge_x + 5, edge_y + 5], 
                         fill='green', outline='white', width=2)
            
            # Draw line from center to scale edge
            draw.line([center_x, center_y, edge_x, edge_y], fill='green', width=2)
            
            # Draw small lines perpendicular to the scale line for better visibility
            dx = edge_x - center_x
            dy = edge_y - center_y
            length = math.sqrt(dx*dx + dy*dy)
            if length > 0:
                # Normalize and create perpendicular
                nx, ny = -dy/length, dx/length
                perp_size = 8
                draw.line([
                    edge_x + nx*perp_size, edge_y + ny*perp_size,
                    edge_x - nx*perp_size, edge_y - ny*perp_size
                ], fill='green', width=2)
        
        # Convert PIL to QPixmap using temporary file method
        import tempfile
        with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as temp_file:
            temp_path = temp_file.name
            pil_image.save(temp_path, 'JPEG', quality=95)
        
        pixmap = QPixmap(temp_path)
        
        # Clean up temp file
        try:
            os.remove(temp_path)
        except:
            pass
        
        # Get the display area dimensions from ClickableLabel
        display_width = getattr(self.image_label, 'display_width', self.image_label.width())
        display_height = getattr(self.image_label, 'display_height', self.image_label.height())
        
        # Scale pixmap to fit the display area while maintaining aspect ratio
        scaled_pixmap = pixmap.scaled(display_width, display_height, 
                                     Qt.KeepAspectRatio, Qt.SmoothTransformation)
        
        # Calculate scale factors for coordinate conversion
        original_width = pixmap.width()
        original_height = pixmap.height()
        scaled_width = scaled_pixmap.width()
        scaled_height = scaled_pixmap.height()
        
        # Store scale factors for coordinate conversion
        self.scale_factor_x = scaled_width / original_width
        self.scale_factor_y = scaled_height / original_height
        
        # Calculate offsets within the display area (image may be smaller than display area)
        self.image_offset_x = (display_width - scaled_width) // 2
        self.image_offset_y = (display_height - scaled_height) // 2
        
        # Add display area offset from ClickableLabel
        display_x = getattr(self.image_label, 'display_x', 0)
        display_y = getattr(self.image_label, 'display_y', 0)
        
        # Total offsets (display area offset + image centering offset)
        self.offset_x = display_x + self.image_offset_x
        self.offset_y = display_y + self.image_offset_y
        
        print(f"Display: Original({original_width}x{original_height}) -> "
              f"Scaled({scaled_width}x{scaled_height}) -> "
              f"DisplayArea({display_width}x{display_height}) -> "
              f"TotalOffset({self.offset_x},{self.offset_y})")  # Debug
        
        self.image_label.setPixmap(scaled_pixmap)
    
    def on_image_click(self, x, y):
        """Handle click on image with proper coordinate conversion"""
        if not self.processor:
            return
        
        print(f"Click at widget coordinates: ({x}, {y})")  # Debug
        
        # Convert widget coordinates to image coordinates
        # Account for centering offset and scaling
        image_x = (x - self.offset_x) / self.scale_factor_x
        image_y = (y - self.offset_y) / self.scale_factor_y
        
        # Ensure coordinates are within image bounds
        actual_x = max(0, min(int(image_x), self.processor.image.width - 1))
        actual_y = max(0, min(int(image_y), self.processor.image.height - 1))
        
        print(f"Converted to image coordinates: ({actual_x}, {actual_y})")  # Debug
        
        # Handle scale edge mode
        if self.scale_edge_mode:
            self.set_scale_edge_point(actual_x, actual_y)
            return
        
        # Check if clicking near existing point (for better UX)
        if self.current_click:
            # Convert existing point to widget coordinates for distance calculation
            existing_widget_x = self.current_click['x'] * self.scale_factor_x + self.offset_x
            existing_widget_y = self.current_click['y'] * self.scale_factor_y + self.offset_y
            distance = ((x - existing_widget_x)**2 + (y - existing_widget_y)**2)**0.5
            
            if distance <= 15:  # Within 15 pixels - start dragging
                print("Starting drag mode")  # Debug
                return
        
        self.place_analysis_point(actual_x, actual_y)
    
    def on_image_drag(self, x, y):
        """Handle drag on image with proper coordinate conversion"""
        if not self.processor:
            return
        
        print(f"Drag to widget coordinates: ({x}, {y})")  # Debug
        
        # Convert widget coordinates to image coordinates
        image_x = (x - self.offset_x) / self.scale_factor_x
        image_y = (y - self.offset_y) / self.scale_factor_y
        
        # Ensure coordinates are within bounds
        actual_x = max(0, min(int(image_x), self.processor.image.width - 1))
        actual_y = max(0, min(int(image_y), self.processor.image.height - 1))
        
        print(f"Converted to image coordinates: ({actual_x}, {actual_y})")  # Debug
        
        self.place_analysis_point(actual_x, actual_y)
    
    def on_mouse_hover(self, x, y):
        """Handle mouse hover over image - show tooltip near analysis point"""
        if not self.processor or not self.current_click:
            QToolTip.hideText()
            return
        
        # Convert existing analysis point to widget coordinates
        existing_widget_x = self.current_click['x'] * self.scale_factor_x + self.offset_x
        existing_widget_y = self.current_click['y'] * self.scale_factor_y + self.offset_y
        distance = ((x - existing_widget_x)**2 + (y - existing_widget_y)**2)**0.5
        
        if distance <= 15:  # Within 15 pixels of analysis point
            # Show tooltip with current analysis data
            azimuth = self.current_click['azimuth']
            range_val = self.current_click['range']
            tooltip_text = f"{self.tr('azimuth')}: {azimuth:.1f}°\n{self.tr('range')}: {range_val:.1f} {self.tr('units')}"
            
            # Calculate global position for tooltip
            point = self.image_label.mapToGlobal(self.image_label.rect().topLeft())
            tooltip_x = point.x() + x + 15
            tooltip_y = point.y() + y - 10
            
            QToolTip.showText(QPoint(tooltip_x, tooltip_y), tooltip_text)
        else:
            # Hide tooltip when not near analysis point
            QToolTip.hideText()
    
    def place_analysis_point(self, x, y):
        """Place analysis point at specified coordinates"""
        if not self.processor:
            return
        
        # Ensure coordinates are within bounds
        x = max(0, min(x, self.processor.image.width - 1))
        y = max(0, min(y, self.processor.image.height - 1))
        
        try:
            # Calculate azimuth and range using custom scale if available
            azimuth, range_val = self.calculate_azimuth_range(x, y)
            
            # Update current click
            self.current_click = {
                'x': x, 'y': y,
                'azimuth': azimuth, 'range': range_val
            }
            
            # Redraw the display
            self.display_image()
            
            # Update results display
            self.update_results_display()
            self.update_report_data()  # Update report panel
            
        except Exception as e:
            QMessageBox.critical(self, self.tr("error"), f"Could not process point: {str(e)}")
    
    def update_report_data(self):
        """Update the report data panel"""
        if not self.processor:
            self.analysis_status_label.setText(self.tr("place_point_first"))
            self.auto_azimuth_label.setText("β - --°")
            self.auto_distance_label.setText("D - -- км")
            self.auto_scale_label.setText("M = --")
            return
        
        if self.current_click:
            # Update auto-filled data
            azimuth = self.current_click['azimuth']
            distance = self.current_click['range']
            scale = int(self.scale_combo.currentText())
            
            self.analysis_status_label.setText("Analysis point set ✓")
            self.auto_azimuth_label.setText(f"β - {azimuth:.0f}°")
            self.auto_distance_label.setText(f"D - {distance:.1f} км")
            self.auto_scale_label.setText(f"M = {scale}")
        else:
            self.analysis_status_label.setText(self.tr("place_point_first"))
            self.auto_azimuth_label.setText("β - --°")
            self.auto_distance_label.setText("D - -- км")
            
            # Scale is still available even without analysis point
            if hasattr(self, 'scale_combo'):
                scale = int(self.scale_combo.currentText())
                self.auto_scale_label.setText(f"M = {scale}")
            else:
                self.auto_scale_label.setText("M = --")
    
    def update_target_number(self, text):
        """Update persistent target number"""
        self.current_target_number = text
    
    def update_height(self, text):
        """Update height value"""
        self.current_height = text
    
    def update_obstacles(self, text):
        """Update obstacles status"""
        self.current_obstacles = text
    
    def update_detection(self, text):
        """Update detection type"""
        self.current_detection = text
    
    def calculate_azimuth_range(self, x, y):
        """Calculate azimuth and range using custom scale if available"""
        # Calculate relative coordinates from center
        dx = x - self.processor.center_x
        dy = self.processor.center_y - y  # Invert Y axis
        
        # Calculate range (distance from center)
        range_pixels = math.sqrt(dx**2 + dy**2)
        
        # Use custom scale distance if available, otherwise use default bottom edge method
        if self.custom_scale_distance:
            # Use the custom scale distance set by the user
            scale_value = int(self.scale_combo.currentText())
            range_actual = (range_pixels / self.custom_scale_distance) * scale_value
        else:
            # Fall back to original bottom edge calculation
            bottom_edge_distance = self.processor.image.height - self.processor.center_y
            scale_value = int(self.scale_combo.currentText())
            range_actual = (range_pixels / bottom_edge_distance) * scale_value
        
        # Calculate azimuth (angle from north, clockwise)
        azimuth_radians = math.atan2(dx, dy)
        azimuth_degrees = math.degrees(azimuth_radians)
        
        # Normalize azimuth to 0-360 degrees
        if azimuth_degrees < 0:
            azimuth_degrees += 360
            
        return azimuth_degrees, range_actual
    
    def toggle_scale_edge_mode(self):
        """Toggle scale edge setting mode"""
        self.scale_edge_mode = self.scale_edge_btn.isChecked()
        
        # Update cursor and visual feedback
        self.image_label.set_scale_edge_mode(self.scale_edge_mode)
        
        if self.scale_edge_mode:
            self.scale_edge_btn.setStyleSheet("background-color: #4CAF50; color: white;")
            self.add_result(self.tr("scale_edge_active"))
        else:
            self.scale_edge_btn.setStyleSheet("")
    
    def set_scale_edge_point(self, x, y):
        """Set the scale edge point"""
        # Ensure coordinates are within bounds
        x = max(0, min(x, self.processor.image.width - 1))
        y = max(0, min(y, self.processor.image.height - 1))
        
        # Calculate distance from center
        dx = x - self.processor.center_x
        dy = y - self.processor.center_y
        distance = math.sqrt(dx*dx + dy*dy)
        
        # Store the scale edge point and distance
        self.scale_edge_point = {'x': x, 'y': y}
        self.custom_scale_distance = distance
        
        # Auto-disable scale edge mode after setting
        self.scale_edge_mode = False
        self.scale_edge_btn.setChecked(False)
        self.scale_edge_btn.setStyleSheet("")
        self.image_label.set_scale_edge_mode(False)
        
        # Update display
        self.display_image()
        
        # Recalculate current analysis point if exists
        if self.current_click:
            azimuth, range_val = self.calculate_azimuth_range(
                self.current_click['x'], self.current_click['y']
            )
            self.current_click['azimuth'] = azimuth
            self.current_click['range'] = range_val
        
        # Update results
        self.update_results_display()
        self.update_report_data()  # Update report panel
        self.add_result(self.tr("scale_edge_set").format(distance=distance))
    
    def update_results_display(self):
        """Update the results area"""
        self.results_text.clear()
        
        if self.processor:
            self.add_result(self.tr("image_info").format(name=os.path.basename(self.current_image_path)))
            self.add_result(self.tr("size").format(width=self.processor.image.width, 
                                                  height=self.processor.image.height))
            self.add_result(self.tr("scale_info").format(scale=int(self.scale_combo.currentText())))
            self.add_result(self.tr("center_info").format(x=self.processor.center_x, 
                                                         y=self.processor.center_y))
            
            # Show scale reference information
            if self.custom_scale_distance:
                self.add_result(f"Custom scale edge: {self.custom_scale_distance:.1f} px = {self.scale_combo.currentText()} units")
            else:
                bottom_distance = self.processor.image.height - self.processor.center_y
                self.add_result(self.tr("bottom_edge").format(scale=int(self.scale_combo.currentText())))
                self.add_result(self.tr("pixels_south").format(pixels=bottom_distance))
            self.add_result("")
        
        if self.current_click:
            self.add_result(self.tr("analysis_point"))
            self.add_result(f"{self.tr('position')}: ({self.current_click['x']}, {self.current_click['y']})")
            self.add_result(f"{self.tr('azimuth')}: {self.current_click['azimuth']:.1f}°")
            self.add_result(f"{self.tr('range')}: {self.current_click['range']:.1f} {self.tr('units')}")
            self.add_result("")
            self.add_result(self.tr("click_to_place"))
            self.add_result(self.tr("drag_to_move"))
            self.add_result(self.tr("line_connects"))
        else:
            self.add_result(self.tr("click_on_image"))
    
    def update_scale(self):
        """Update scale value"""
        if self.processor:
            new_scale = int(self.scale_combo.currentText())
            
            if self.current_click:
                azimuth, range_val = self.calculate_azimuth_range(
                    self.current_click['x'], self.current_click['y']
                )
                self.current_click['azimuth'] = azimuth
                self.current_click['range'] = range_val
            
            self.update_results_display()
            self.update_report_data()  # Update report panel
            self.add_result(self.tr("scale_updated").format(scale=new_scale))
    
    def move_center(self, dx, dy):
        """Move center point"""
        if not self.processor:
            QMessageBox.warning(self, self.tr("warning"), self.tr("no_image_first"))
            return
        
        self.processor.move_center(dx, dy)
        
        if self.current_click:
            azimuth, range_val = self.calculate_azimuth_range(
                self.current_click['x'], self.current_click['y']
            )
            self.current_click['azimuth'] = azimuth
            self.current_click['range'] = range_val
        
        self.display_image()
        self.update_results_display()
        self.update_report_data()  # Update report panel
        self.add_result(self.tr("center_moved").format(x=self.processor.center_x, 
                                                      y=self.processor.center_y))
    
    def show_center_preview(self):
        """Show center position"""
        if self.processor:
            self.display_image()
            self.update_results_display()
    
    def clear_results(self):
        """Clear analysis point"""
        self.current_click = None
        # Also clear scale edge point if user wants to reset everything
        # self.scale_edge_point = None  # Uncomment if you want to clear scale edge too
        # self.custom_scale_distance = None
        if self.processor:
            self.display_image()
        self.update_results_display()
        self.update_report_data()  # Update report panel
    
    def save_current_image(self):
        """Save processed image"""
        if not self.processor:
            QMessageBox.warning(self, self.tr("warning"), self.tr("no_image_first"))
            return
        
        if not self.current_click:
            QMessageBox.warning(self, self.tr("warning"), self.tr("no_analysis_point"))
            return
        
        file_path, _ = QFileDialog.getSaveFileName(
            self, self.tr("save_processed_image"), "",
            f"{self.tr('jpeg_files')} (*.jpg);;{self.tr('png_files')} (*.png);;{self.tr('all_files')} (*.*)"
        )
        
        if file_path:
            try:
                # Create final image with ONLY the line
                final_image = self.processor.image.copy()
                
                print(f"Original image mode: {final_image.mode}")  # Debug
                
                # Force convert to RGB for JPEG, keep original for PNG
                if file_path.lower().endswith(('.jpg', '.jpeg')):
                    if final_image.mode != 'RGB':
                        print(f"Converting {final_image.mode} to RGB for JPEG")  # Debug
                        if final_image.mode == 'RGBA':
                            # Handle transparency
                            rgb_image = Image.new('RGB', final_image.size, (255, 255, 255))
                            rgb_image.paste(final_image, mask=final_image.split()[-1])
                            final_image = rgb_image
                        else:
                            # Convert any other mode to RGB
                            final_image = final_image.convert('RGB')
                
                draw = ImageDraw.Draw(final_image)
                
                # Draw ONLY the line to top-right corner
                draw.line([
                    (self.current_click['x'], self.current_click['y']), 
                    (final_image.width - 1, 0)
                ], fill='blue', width=3)
                
                # Save with appropriate format
                if file_path.lower().endswith('.png'):
                    final_image.save(file_path, 'PNG')
                    print(f"Saved as PNG with mode: {final_image.mode}")  # Debug
                else:
                    final_image.save(file_path, 'JPEG', quality=95)
                    print(f"Saved as JPEG with mode: {final_image.mode}")  # Debug
                
                self.add_result(f"{self.tr('saved')}: {os.path.basename(file_path)}")
                self.add_result(self.tr("saved_clean"))
                QMessageBox.information(self, self.tr("success"), self.tr("image_saved_successfully"))
                
            except Exception as e:
                print(f"Error saving image: {e}")  # Debug
                QMessageBox.critical(self, self.tr("error"), 
                                   self.tr("could_not_save").format(error=str(e)))
    
    def create_new_word_album(self):
        """Create a new Word album document"""
        if not DOCX_AVAILABLE:
            QMessageBox.warning(self, self.tr("warning"), 
                               self.tr("docx_not_available"))
            return
        
        if not self.processor or not self.current_click:
            QMessageBox.warning(self, self.tr("warning"), 
                               self.tr("load_image_and_point"))
            return
        
        # Ask user for save location
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Create New Word Album", "",
            "Word Documents (*.docx);;All files (*.*)"
        )
        
        if file_path:
            try:
                # Create new document
                self.current_word_document = Document()
                self.word_document_path = file_path
                
                # Set default font for entire document to Arial
                self.set_document_default_font()
                
                # Add current processed image to document
                self.add_image_to_word_document()
                
                # Save document
                self.current_word_document.save(file_path)
                
                QMessageBox.information(self, self.tr("success"), 
                                       f"New album created: {os.path.basename(file_path)}")
                self.add_result(f"Created new album: {os.path.basename(file_path)}")
                
            except Exception as e:
                QMessageBox.critical(self, self.tr("error"), 
                                   f"Could not create Word album: {str(e)}")
    
    def set_document_default_font(self):
        """Set Arial as the default font and landscape orientation for the entire document"""
        try:
            # Access document styles
            styles = self.current_word_document.styles
            
            # Set default paragraph font to Arial
            default_style = styles['Normal']
            default_style.font.name = 'Arial'
            default_style.font.size = Pt(12)  # Default size
            
            # Set document to landscape orientation
            section = self.current_word_document.sections[0]
            section.orientation = 1  # 1 = Landscape, 0 = Portrait
            
            # Swap width and height for landscape
            new_width, new_height = section.page_height, section.page_width
            section.page_width = new_width
            section.page_height = new_height
            
            print(f"Document set to landscape orientation: {new_width} x {new_height}")  # Debug
            
            # Also set the document-wide default
            doc = self.current_word_document
            doc_element = doc._element
            
            # Create document defaults element
            docDefaults = OxmlElement('w:docDefaults')
            
            # Character defaults
            rPrDefault = OxmlElement('w:rPrDefault')
            rPr = OxmlElement('w:rPr')
            
            # Set Arial as default font
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Arial')
            rFonts.set(qn('w:hAnsi'), 'Arial')
            rFonts.set(qn('w:eastAsia'), 'Arial')
            rFonts.set(qn('w:cs'), 'Arial')
            rPr.append(rFonts)
            
            rPrDefault.append(rPr)
            docDefaults.append(rPrDefault)
            
            # Add to document
            doc_element.insert(0, docDefaults)
            
        except Exception as e:
            print(f"Error setting document defaults: {e}")
            # If setting defaults fails, we'll still set fonts individually
    
    def add_to_existing_album(self):
        """Add current image to existing Word album"""
        if not DOCX_AVAILABLE:
            QMessageBox.warning(self, self.tr("warning"), 
                               self.tr("docx_not_available"))
            return
        
        if not self.processor or not self.current_click:
            QMessageBox.warning(self, self.tr("warning"), 
                               self.tr("load_image_and_point"))
            return
        
        # Ask user to select existing document
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Select Existing Word Album", "",
            "Word Documents (*.docx);;All files (*.*)"
        )
        
        if file_path:
            try:
                # Open existing document
                self.current_word_document = Document(file_path)
                self.word_document_path = file_path
                
                # Ensure Arial is set as default (in case existing document has different defaults)
                self.set_document_default_font()
                
                # Add current processed image to document
                self.add_image_to_word_document()
                
                # Save document
                self.current_word_document.save(file_path)
                
                QMessageBox.information(self, self.tr("success"), 
                                       f"Added to album: {os.path.basename(file_path)}")
                self.add_result(f"Added to album: {os.path.basename(file_path)}")
                
            except Exception as e:
                QMessageBox.critical(self, self.tr("error"), 
                                   f"Could not add to Word album: {str(e)}")
    
    def add_image_to_word_document(self):
        """Add current processed image and data to Word document"""
        if not self.current_word_document or not self.processor or not self.current_click:
            return
        
        # Create table with 1 row and 2 columns
        table = self.current_word_document.add_table(rows=1, cols=2)
        
        # Remove default table style and borders to match template
        table.style = None
        
        # Set table properties to match template exactly
        tbl = table._tbl
        tblPr = tbl.tblPr
        
        # Set table borders (thick borders like in template)
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '12')  # Thick border like in template
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        tblPr.append(tblBorders)
        
        # Set table width to 100%
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:type'), 'pct')
        tblW.set(qn('w:w'), '5000')  # 100% width
        tblPr.append(tblW)
        
        # Get cells
        left_cell = table.cell(0, 0)
        right_cell = table.cell(0, 1)
        
        # Set cell widths - left cell much larger for image, right cell narrow for data
        left_cell.width = Cm(12)  # Larger for image
        right_cell.width = Cm(3)   # Narrow for data (exactly like template)
        
        # Set cell properties to match template
        for cell in [left_cell, right_cell]:
            # Set cell margins
            cellPr = cell._tc.tcPr
            if cellPr is None:
                cellPr = OxmlElement('w:tcPr')
                cell._tc.append(cellPr)
            
            # Set cell margins to match template
            tcMar = OxmlElement('w:tcMar')
            for margin in ['top', 'left', 'bottom', 'right']:
                mar = OxmlElement(f'w:{margin}')
                mar.set(qn('w:w'), '108')  # Small margins
                mar.set(qn('w:type'), 'dxa')
                tcMar.append(mar)
            cellPr.append(tcMar)
            
            # Set vertical alignment to center
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            cellPr.append(vAlign)
        
        # Add processed image to left cell
        self.add_processed_image_to_cell(left_cell)
        
        # Add data to right cell  
        self.add_data_to_cell(right_cell)
        
        # Add spacing after table (not page break unless many tables)
        if len(self.current_word_document.tables) > 1:
            para = self.current_word_document.add_paragraph()
            para.text = ""
    
    def add_image_to_word_document(self):
        """Add current processed image and data to Word document"""
        if not self.current_word_document or not self.processor or not self.current_click:
            return
        
        # Create table with 1 row and 2 columns
        table = self.current_word_document.add_table(rows=1, cols=2)
        
        # Remove default table style and borders to match template
        table.style = None
        
        # Set table properties to match template exactly
        tbl = table._tbl
        tblPr = tbl.tblPr
        
        # Set table borders (thick borders like in template)
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '12')  # Thick border like in template
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        tblPr.append(tblBorders)
        
        # Set table width to 100%
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:type'), 'pct')
        tblW.set(qn('w:w'), '5000')  # 100% width
        tblPr.append(tblW)
        
        # Get cells
        left_cell = table.cell(0, 0)
        right_cell = table.cell(0, 1)
        
        # Set cell widths - left cell much larger for image, right cell narrow for data
        left_cell.width = Cm(12)  # Larger for image
        right_cell.width = Cm(3)   # Narrow for data (exactly like template)
        
        # Set cell properties to match template
        for cell in [left_cell, right_cell]:
            # Set cell margins
            cellPr = cell._tc.tcPr
            if cellPr is None:
                cellPr = OxmlElement('w:tcPr')
                cell._tc.append(cellPr)
            
            # Set cell margins to match template
            tcMar = OxmlElement('w:tcMar')
            for margin in ['top', 'left', 'bottom', 'right']:
                mar = OxmlElement(f'w:{margin}')
                mar.set(qn('w:w'), '108')  # Small margins
                mar.set(qn('w:type'), 'dxa')
                tcMar.append(mar)
            cellPr.append(tcMar)
            
            # Set vertical alignment to center
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            cellPr.append(vAlign)
        
        # Add processed image and "Індикатор ЗРЛ" to left cell
        self.add_processed_image_and_label_to_cell(left_cell)
        
        # Add data to right cell  
        self.add_data_to_cell(right_cell)
        
        # Add spacing after table (not page break unless many tables)
        if len(self.current_word_document.tables) > 1:
            para = self.current_word_document.add_paragraph()
            para.text = ""
    
    def add_image_to_word_document(self):
        """Add current processed image and data to Word document"""
        if not self.current_word_document or not self.processor or not self.current_click:
            return
        
        # Create table with 1 row and 3 columns (left invisible column for "Індикатор ЗРЛ", image, data)
        table = self.current_word_document.add_table(rows=1, cols=3)
        
        # Remove default table style
        table.style = None
        
        # Set table properties
        tbl = table._tbl
        tblPr = tbl.tblPr
        
        # Set table borders (thin borders)
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # Thin border (was 12)
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        tblPr.append(tblBorders)
        
        # Set table width to 100%
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:type'), 'pct')
        tblW.set(qn('w:w'), '5000')
        tblPr.append(tblW)
        
        # Get cells
        label_cell = table.cell(0, 0)    # Left cell for "Індикатор ЗРЛ"
        image_cell = table.cell(0, 1)    # Middle cell for image
        data_cell = table.cell(0, 2)     # Right cell for data
        
        # Set cell widths for landscape orientation (more width available)
        label_cell.width = Cm(3)   # Narrow for label
        image_cell.width = Cm(20)  # Much larger for image in landscape
        data_cell.width = Cm(4)    # Slightly wider for data
        
        # Configure label cell (make borders invisible except right border)
        self.configure_label_cell(label_cell)
        
        # Set standard cell properties for image and data cells
        for cell in [image_cell, data_cell]:
            cellPr = cell._tc.tcPr
            if cellPr is None:
                cellPr = OxmlElement('w:tcPr')
                cell._tc.append(cellPr)
            
            # Set small cell margins
            tcMar = OxmlElement('w:tcMar')
            for margin in ['top', 'left', 'bottom', 'right']:
                mar = OxmlElement(f'w:{margin}')
                mar.set(qn('w:w'), '108')
                mar.set(qn('w:type'), 'dxa')
                tcMar.append(mar)
            cellPr.append(tcMar)
            
            # Center vertical alignment
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            cellPr.append(vAlign)
        
        # Add content to cells
        self.add_label_to_cell(label_cell)
        self.add_processed_image_to_cell(image_cell)
        self.add_data_to_cell(data_cell)
        
        # Add spacing after table
        if len(self.current_word_document.tables) > 1:
            para = self.current_word_document.add_paragraph()
            para.text = ""
    
    def configure_label_cell(self, cell):
        """Configure label cell with invisible borders except right border"""
        cellPr = cell._tc.tcPr
        if cellPr is None:
            cellPr = OxmlElement('w:tcPr')
            cell._tc.append(cellPr)
        
        # Set cell borders - transparent except right border
        tcBorders = OxmlElement('w:tcBorders')
        
        # Invisible borders
        for border_name in ['top', 'left', 'bottom']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')  # Invisible
            tcBorders.append(border)
        
        # Visible right border
        right_border = OxmlElement('w:right')
        right_border.set(qn('w:val'), 'single')
        right_border.set(qn('w:sz'), '4')  # Thin border
        right_border.set(qn('w:space'), '0')
        right_border.set(qn('w:color'), '000000')
        tcBorders.append(right_border)
        
        cellPr.append(tcBorders)
        
        # Set small margins
        tcMar = OxmlElement('w:tcMar')
        for margin in ['top', 'left', 'bottom', 'right']:
            mar = OxmlElement(f'w:{margin}')
            mar.set(qn('w:w'), '108')
            mar.set(qn('w:type'), 'dxa')
            tcMar.append(mar)
        cellPr.append(tcMar)
        
        # Center vertical alignment
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        cellPr.append(vAlign)
    
    def add_label_to_cell(self, cell):
        """Add 'Індикатор ЗРЛ' label to left cell"""
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add the label text
        run = para.add_run("Індикатор ЗРЛ")
        run.font.size = Pt(12)  # Arial is inherited from document default
    
    def add_processed_image_to_cell(self, cell):
        """Add processed image to cell with bulletproof format handling"""
        try:
            # Get image from processor and ensure it's RGB
            if not self.processor or not hasattr(self.processor, 'image'):
                print("No processor or image available")  # Debug
                return
            
            final_image = self.processor.image.copy()
            print(f"Image mode before processing: {final_image.mode}")  # Debug
            
            # BULLETPROOF conversion to RGB
            if final_image.mode != 'RGB':
                print(f"Converting from {final_image.mode} to RGB")  # Debug
                try:
                    if final_image.mode == 'RGBA':
                        # Create white background for transparency
                        rgb_image = Image.new('RGB', final_image.size, (255, 255, 255))
                        rgb_image.paste(final_image, mask=final_image.split()[-1])
                        final_image = rgb_image
                    elif final_image.mode in ('P', 'L', '1'):
                        # Palette, Grayscale, or Monochrome
                        final_image = final_image.convert('RGB')
                    else:
                        # Any other exotic mode
                        final_image = final_image.convert('RGB')
                    print(f"Successfully converted to: {final_image.mode}")  # Debug
                except Exception as conv_error:
                    print(f"Conversion error: {conv_error}")  # Debug
                    # Force conversion using PIL's most compatible method
                    final_image = final_image.convert('RGB')
            
            # Verify we have RGB before proceeding
            if final_image.mode != 'RGB':
                raise ValueError(f"Failed to convert image to RGB, still in {final_image.mode} mode")
            
            # Now safely draw on the RGB image
            draw = ImageDraw.Draw(final_image)
            
            # Draw ONLY the line to top-right corner
            draw.line([
                (self.current_click['x'], self.current_click['y']), 
                (final_image.width - 1, 0)
            ], fill='blue', width=3)
            
            # Save to temporary file as JPEG
            import tempfile
            with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as temp_file:
                temp_path = temp_file.name
                print(f"Saving RGB image to: {temp_path}")  # Debug
                final_image.save(temp_path, 'JPEG', quality=95)
                print(f"Successfully saved JPEG")  # Debug
            
            # Add to Word document
            cell.text = ""
            top_spacing_para = cell.paragraphs[0]
            top_spacing_para.text = ""
            
            # Set spacing
            pPr = top_spacing_para._element.get_or_add_pPr()
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:before'), '0')
            spacing.set(qn('w:after'), '720')
            pPr.append(spacing)
            
            # Add image paragraph
            img_paragraph = cell.add_paragraph()
            img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add image to document
            run = img_paragraph.add_run()
            run.add_picture(temp_path, width=Cm(19))  # Larger image for landscape orientation
            print(f"Successfully added image to Word document")  # Debug
            
            # Clean up temp file
            try:
                os.remove(temp_path)
                print(f"Cleaned up temp file")  # Debug
            except:
                pass
                
        except Exception as e:
            print(f"Error in add_processed_image_to_cell: {e}")
            import traceback
            traceback.print_exc()  # Full error traceback for debugging
    
    def add_data_to_cell(self, cell):
        """Add analysis data to Word document cell with tight formatting (no empty lines)"""
        try:
            # Clear existing content
            cell.text = ""
            
            # Get data
            target_num = self.current_target_number
            azimuth = self.current_click['azimuth']
            distance = self.current_click['range']
            height = self.current_height
            obstacles = self.current_obstacles
            detection = self.current_detection
            scale = int(self.scale_combo.currentText())
            
            # Add top spacing (equivalent to one enter at font size 30)
            top_spacing_para = cell.paragraphs[0]
            top_spacing_para.text = ""
            pPr = top_spacing_para._element.get_or_add_pPr()
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:before'), '0')
            spacing.set(qn('w:after'), '720')  # Spacing equivalent to font size 30
            pPr.append(spacing)
            
            # Create data items WITHOUT empty lines between them
            # Format: (text, font_size, italic, underline)
            data_items = [
                (f"{target_num}", 12, True, True),          # Target number - 12pt, italic, underlined
                (f"β - {azimuth:.0f}°", 12, True, False),   # Azimuth - 12pt, italic
                (f"D - {distance:.1f} км", 12, True, False), # Distance - 12pt, italic
                (f"Н - {height} км", 12, True, False),     # Height - 12pt, italic
                (f"{obstacles}", 9, True, False),           # Obstacles - 9pt, italic
                (f"{detection}", 9, True, False),           # Detection type - 9pt, italic
                (f"М={scale}", 9, True, False)              # Scale - 9pt, italic
            ]
            
            # Add each line with tight formatting (no empty lines)
            for text, font_size, italic, underline in data_items:
                # Add new paragraph for each line
                p = cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Left alignment
                
                # Set tight paragraph spacing - no gaps between lines
                pPr = p._element.get_or_add_pPr()
                spacing = OxmlElement('w:spacing')
                spacing.set(qn('w:before'), '0')
                spacing.set(qn('w:after'), '0')  # No space after each line
                spacing.set(qn('w:line'), '200')    # Very tight line spacing
                spacing.set(qn('w:lineRule'), 'auto')
                pPr.append(spacing)
                
                # Add text with formatting
                run = p.add_run(text)
                run.font.size = Pt(font_size)
                run.italic = italic
                run.underline = underline
                
                # Set character spacing (slightly condensed)
                rPr = run._element.rPr
                if rPr is None:
                    rPr = OxmlElement('w:rPr')
                    run._element.insert(0, rPr)
                spacing = OxmlElement('w:spacing')
                spacing.set(qn('w:val'), '-5')  # Slightly condensed
                rPr.append(spacing)
            
        except Exception as e:
            print(f"Error adding data to cell: {e}")

def main():
    app = QApplication(sys.argv)
    window = AzimuthGUI()
    window.show()
    sys.exit(app.exec_())

if __name__ == "__main__":
    main()