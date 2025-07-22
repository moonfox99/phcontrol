#!/usr/bin/env python3
"""
Azimuth Image Processor - PyQt5 Version with Ukrainian Translation
Professional interface for processing azimuth grid images
"""

import sys
import os
import math
import tempfile
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QHBoxLayout, 
                             QVBoxLayout, QGridLayout, QPushButton, QLabel, 
                             QComboBox, QTextEdit, QScrollArea, QFrame,
                             QFileDialog, QMessageBox, QSplitter, QToolTip, QLineEdit)
from PyQt5.QtCore import Qt, QSize, pyqtSignal, QTimer, QPoint
from PyQt5.QtGui import QPixmap, QFont
from PIL import Image, ImageDraw

from translations import Translations
from widgets import ClickableLabel, VerticalThumbnailWidget
from image_processor import AzimuthImageProcessor

try:
    from docx import Document
    from docx.shared import Inches, Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_BREAK
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.section import WD_SECTION_START
    from docx.oxml.shared import OxmlElement, qn
    from docx.oxml.ns import qn
    from docx.oxml import parse_xml
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ===== ОНОВЛЕНІ КОНСТАНТИ =====

ALBUM_LAYOUT = {
    # Розміри сторінки A4 в міліметрах
    'PAGE_WIDTH': 210,
    'PAGE_HEIGHT': 297,
    
    # Стандартні поля для титульної та описової сторінок
    'STANDARD_LEFT_MARGIN': 25,
    'STANDARD_RIGHT_MARGIN': 25, 
    'STANDARD_TOP_MARGIN': 10,
    'STANDARD_BOTTOM_MARGIN': 10,
    
    # Поля для сторінок з таблицями (НУЛЬОВІ!)
    'TABLE_PAGES_LEFT_MARGIN': 0,
    'TABLE_PAGES_RIGHT_MARGIN': 0,
    'TABLE_PAGES_TOP_MARGIN': 0,
    'TABLE_PAGES_BOTTOM_MARGIN': 0,
    
    # Відстані для таблиць (від краю сторінки при нульових полях)
    'TOP_MARGIN_TO_FIRST_TABLE': 25,      # 25мм від верху сторінки до першої таблиці
    'RIGHT_MARGIN_FROM_TABLE': 5,         # 5мм від таблиці до правого краю сторінки
    'BOTTOM_MARGIN_FROM_LAST_TABLE': 5,   # 5мм від останньої таблиці до низу сторінки
    'SPACING_BETWEEN_TABLES': 5,          # 5мм між таблицями
    
    # Розміри комірок таблиці
    'LEFT_CELL_WIDTH': 30,        # Ширина комірки "Індикатор ЗРЛ"
    'RIGHT_CELL_WIDTH': 30,       # Ширина правої комірки з даними
    'RIGHT_CELL_TOP_PADDING': 15, # Відстань від верху до тексту в правій комірці
    
    # Параметри розташування
    'TABLES_PER_PAGE': 2,         # Кількість таблиць на сторінку
    'IMAGE_ASPECT_RATIO': 15/13,  # Пропорції зображення (ширина/висота)
}

# Константи для титульної сторінки (БЕЗ ЗМІН)
TITLE_PAGE = {
    'MAIN_FONT_SIZE': 28,
    'FOOTER_FONT_SIZE': 16,
    'TOP_SPACING': 3,
    'FOOTER_SPACING': 10,
    'FOOTER_LEFT_INDENT': 3.75,
}

# Константи для сторінки опису (БЕЗ ЗМІН)
DESCRIPTION_PAGE = {
    'HEADING_FONT_SIZE': 22,
    'TABLE_HEIGHT': 1.7,
    'COL_WIDTHS': [0.84, 3.43, 2.69, 2.54, 2.77, 2.27],
    'HEADERS': [
        "№ пп",
        "Назва документа", 
        "Обліковий номер",
        "Кількість аркушів (знімків)",
        "Гриф секретності",
        "Примітка"
    ]
}

def mm_to_cm(mm):
    """Перетворення міліметрів в сантиметри для docx"""
    return mm / 10.0

def format_ukrainian_date(date_obj):
    """Форматування дати по-українськи для титульної сторінки"""
    try:
        if isinstance(date_obj, str):
            from datetime import datetime
            date_obj = datetime.strptime(date_obj, "%Y-%m-%d")
        
        months = {
            1: 'січня', 2: 'лютого', 3: 'березня', 4: 'квітня',
            5: 'травня', 6: 'червня', 7: 'липня', 8: 'серпня',
            9: 'вересня', 10: 'жовтня', 11: 'листопада', 12: 'грудня'
        }
        
        day = date_obj.day
        month = months[date_obj.month]
        year = date_obj.year
        
        return f'"{day}" {month} {year} року'
    except:
        return "дата не вказана"

def get_default_unit_info():
    """Базова інформація про частину"""
    return "А0000"

def get_default_commander_info():
    """Базова інформація про командира"""
    return {
        'rank': 'полковник',
        'name': 'П.П. ПЕТРЕНКО'
    }

def get_default_responsible_info():
    """Базова інформація про відповідальних осіб"""
    return {
        'responsible': {'rank': 'майор', 'name': 'І.І. ІВАНЕНКО'},
        'assistant': {'rank': 'старший лейтенант', 'name': 'С.С. СИДОРЕНКО'},
        'secretary': {'rank': 'старший лейтенант', 'name': 'С.С. СИДОРЕНКО'}
    }

def create_complete_album(processed_images, title_data, file_path):
    """
    Створення повного альбому з правильними розмірами та полями
    
    Args:
        processed_images: Список оброблених зображень
        title_data: Дані для титульної сторінки
        file_path: Шлях для збереження файлу
    """
    try:
        # Створюємо новий документ
        doc = Document()
        
        print("=== Creating Complete Album ===")
        
        # 1. ТИТУЛЬНА СТОРІНКА (стандартні поля)
        print("1. Creating title page with standard margins...")
        set_standard_margins(doc)
        create_title_page(doc, title_data)
        
        # 2. СТОРІНКА ОПИСУ (стандартні поля)
        print("2. Creating description page with standard margins...")
        create_description_page(doc)
        
        # 3. СТОРІНКИ З ТАБЛИЦЯМИ (нульові поля)
        print("3. Creating table pages with zero margins...")
        create_table_pages_with_zero_margins(doc, processed_images)
        
        # Зберігаємо документ
        doc.save(file_path)
        print(f"✓ Complete album saved: {file_path}")
        
        return True
        
    except Exception as e:
        print(f"✗ Error creating complete album: {e}")
        return False

def set_standard_margins(doc):
    """Встановлення стандартних полів для титульної та описової сторінок"""
    try:
        section = doc.sections[0]
        section.left_margin = Cm(mm_to_cm(ALBUM_LAYOUT['STANDARD_LEFT_MARGIN']))
        section.right_margin = Cm(mm_to_cm(ALBUM_LAYOUT['STANDARD_RIGHT_MARGIN']))
        section.top_margin = Cm(mm_to_cm(ALBUM_LAYOUT['STANDARD_TOP_MARGIN']))
        section.bottom_margin = Cm(mm_to_cm(ALBUM_LAYOUT['STANDARD_BOTTOM_MARGIN']))
        
        print(f"✓ Standard margins set: {ALBUM_LAYOUT['STANDARD_LEFT_MARGIN']}mm all around")
        
    except Exception as e:
        print(f"✗ Error setting standard margins: {e}")

def create_table_pages_with_zero_margins(doc, processed_images):
    """Створення сторінок з таблицями з нульовими полями"""
    try:
        if not processed_images:
            print("No processed images to add")
            return
        
        # Додаємо нову секцію з нульовими полями для таблиць
        table_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
        table_section.left_margin = Cm(mm_to_cm(ALBUM_LAYOUT['TABLE_PAGES_LEFT_MARGIN']))
        table_section.right_margin = Cm(mm_to_cm(ALBUM_LAYOUT['TABLE_PAGES_RIGHT_MARGIN']))
        table_section.top_margin = Cm(mm_to_cm(ALBUM_LAYOUT['TABLE_PAGES_TOP_MARGIN']))
        table_section.bottom_margin = Cm(mm_to_cm(ALBUM_LAYOUT['TABLE_PAGES_BOTTOM_MARGIN']))
        
        print(f"✓ Zero margins set for table pages")
        
        # Додаємо таблиці
        add_image_tables_to_section(doc, processed_images)
        
    except Exception as e:
        print(f"✗ Error creating table pages: {e}")

def add_image_tables_to_section(doc, processed_images):
    """Додавання таблиць зображень до секції з нульовими полями"""
    
    tables_count = 0
    
    for i, image_data in enumerate(processed_images):
        # Визначаємо чи це перша таблиця на сторінці
        is_first_on_page = (tables_count % ALBUM_LAYOUT['TABLES_PER_PAGE']) == 0
        
        # Отримуємо номер цілі
        target_no = image_data['target_data']['number']
        
        # Створюємо таблицю з точними розмірами
        table = create_precise_image_table(doc, image_data, target_no, is_first_on_page)
        
        if table:
            tables_count += 1
            print(f"✓ Added table {tables_count} for target {target_no}")
        
        # Розрив сторінки після кожних двох таблиць (крім останньої)
        if (tables_count % ALBUM_LAYOUT['TABLES_PER_PAGE'] == 0 and 
            i < len(processed_images) - 1):
            doc.add_page_break()
            print(f"Page break after {tables_count} tables")
    
    print(f"✓ Added {tables_count} tables on {(tables_count + 1) // 2} pages")

def create_precise_image_table(doc, image_data, target_no, is_first_on_page=True):
    """
    Створення таблиці зображення з точними розмірами згідно вимог
    При нульових полях всі відстані рахуються від краю сторінки
    """
    try:
        # Розрахунок ширин при нульових полях
        total_page_width = ALBUM_LAYOUT['PAGE_WIDTH']  # 210мм
        right_margin = ALBUM_LAYOUT['RIGHT_MARGIN_FROM_TABLE']  # 5мм від таблиці до краю
        left_cell_width = ALBUM_LAYOUT['LEFT_CELL_WIDTH']   # 30мм
        right_cell_width = ALBUM_LAYOUT['RIGHT_CELL_WIDTH'] # 30мм
        
        # Ширина середньої комірки = вся сторінка - право поле - ліва комірка - права комірка
        middle_cell_width = total_page_width - right_margin - left_cell_width - right_cell_width
        total_table_width = left_cell_width + middle_cell_width + right_cell_width
        
        print(f"Table widths: Left={left_cell_width}mm, Middle={middle_cell_width}mm, Right={right_cell_width}mm")
        print(f"Total table width: {total_table_width}mm (page width - {right_margin}mm right margin)")
        
        # Додаємо відступ зверху
        if is_first_on_page:
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Cm(mm_to_cm(ALBUM_LAYOUT['TOP_MARGIN_TO_FIRST_TABLE']))
        else:
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Cm(mm_to_cm(ALBUM_LAYOUT['SPACING_BETWEEN_TABLES']))
        
        # Створюємо таблицю
        table = doc.add_table(rows=1, cols=3)
        table.style = None
        
        # Встановлюємо точну ширину таблиці
        set_absolute_table_width(table, total_table_width)
        
        row = table.rows[0]
        
        # Висота таблиці (13см для збереження пропорцій 15:13)
        table_height_cm = 13
        row.height = Cm(table_height_cm)
        
        # Налаштовуємо комірки
        left_cell = row.cells[0]
        middle_cell = row.cells[1]
        right_cell = row.cells[2]
        
        # Встановлюємо точні ширини комірок
        set_absolute_cell_width(left_cell, left_cell_width)
        set_absolute_cell_width(middle_cell, middle_cell_width)
        set_absolute_cell_width(right_cell, right_cell_width)
        
        # Налаштовуємо вміст комірок
        setup_left_cell(left_cell)
        setup_middle_cell(middle_cell, image_data, middle_cell_width, table_height_cm)
        setup_right_cell(right_cell, image_data, target_no)
        
        print(f"✓ Created precise table for target {target_no}")
        return table
        
    except Exception as e:
        print(f"✗ Error creating precise table: {e}")
        return None
    
# ===== ФУНКЦІЇ НАЛАШТУВАННЯ ТАБЛИЦЬ ТА КОМІРОК =====

def set_absolute_table_width(table, width_mm):
    """Встановлення абсолютної ширини таблиці"""
    try:
        tbl = table._tbl
        tblPr = tbl.tblPr
        
        # Фіксований layout
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)
        
        # Абсолютна ширина в DXA (1мм = 56.7 DXA)
        width_dxa = int(width_mm * 56.7)
        
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:type'), 'dxa')
        tblW.set(qn('w:w'), str(width_dxa))
        tblPr.append(tblW)
        
        # Вирівнювання по лівому краю (при нульових полях = край сторінки)
        tblJc = OxmlElement('w:jc')
        tblJc.set(qn('w:val'), 'left')
        tblPr.append(tblJc)
        
        print(f"✓ Table width set: {width_mm}mm ({width_dxa} DXA)")
        
    except Exception as e:
        print(f"✗ Error setting table width: {e}")

def set_absolute_cell_width(cell, width_mm):
    """Встановлення абсолютної ширини комірки"""
    try:
        width_dxa = int(width_mm * 56.7)
        
        tc = cell._tc
        tcPr = tc.tcPr
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            tc.append(tcPr)
        
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:type'), 'dxa')
        tcW.set(qn('w:w'), str(width_dxa))
        tcPr.append(tcW)
        
        print(f"✓ Cell width: {width_mm}mm")
        
    except Exception as e:
        print(f"✗ Error setting cell width: {e}")

def setup_left_cell(cell):
    """Налаштування лівої комірки"""
    try:
        # Вертикальне центрування
        set_cell_vertical_center(cell)
        
        # Текст
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = para.add_run("Індикатор ЗРЛ")
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        
        # Рамки: тільки права
        set_cell_borders(cell, top=False, bottom=False, left=False, right=True)
        
        print("✓ Left cell configured")
        
    except Exception as e:
        print(f"✗ Error configuring left cell: {e}")

def setup_middle_cell(cell, image_data, cell_width_mm, cell_height_cm):
    """Налаштування середньої комірки з зображенням"""
    try:
        # Вертикальне центрування
        set_cell_vertical_center(cell)
        
        # Очищуємо комірку
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Створюємо оброблене зображення
        processed_image = create_processed_image_from_data(image_data)
        
        if processed_image:
            # Розрахунок розмірів з дотриманням пропорцій 15:13
            max_width_cm = mm_to_cm(cell_width_mm * 0.95)
            max_height_cm = cell_height_cm * 0.95
            
            target_ratio = ALBUM_LAYOUT['IMAGE_ASPECT_RATIO']  # 15/13
            
            if max_width_cm / target_ratio <= max_height_cm:
                image_width_cm = max_width_cm
                image_height_cm = max_width_cm / target_ratio
            else:
                image_height_cm = max_height_cm
                image_width_cm = max_height_cm * target_ratio
            
            # Зберігаємо у тимчасовий файл
            with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as temp_file:
                temp_path = temp_file.name
                processed_image.save(temp_path, 'JPEG', quality=95)
            
            # Додаємо зображення
            run = para.add_run()
            run.add_picture(temp_path, width=Cm(image_width_cm), height=Cm(image_height_cm))
            
            # Видаляємо тимчасовий файл
            try:
                os.remove(temp_path)
            except:
                pass
            
            print(f"✓ Image: {image_width_cm:.1f}x{image_height_cm:.1f}cm (ratio: {image_width_cm/image_height_cm:.2f})")
            
        else:
            # Fallback
            para.text = "Помилка завантаження зображення"
            para.runs[0].font.size = Pt(10)
        
        # Рамки: всі сторони
        set_cell_borders(cell, top=True, bottom=True, left=True, right=True)
        
        # Білий фон
        set_cell_background(cell, "FFFFFF")
        
        print("✓ Middle cell configured")
        
    except Exception as e:
        print(f"✗ Error configuring middle cell: {e}")

def setup_right_cell(cell, image_data, target_no):
    """Налаштування правої комірки з відступом 15мм зверху"""
    try:
        # Очищуємо комірку
        cell.text = ""
        
        # Перший параграф з відступом 15мм зверху
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Cm(mm_to_cm(ALBUM_LAYOUT['RIGHT_CELL_TOP_PADDING']))
        
        # Отримуємо дані
        target_data = image_data['target_data']
        analysis_point = image_data['analysis_point']
        
        # Дані для відображення
        data_lines = [
            (f"{target_no}", 12, True, True),
            (f"β – {analysis_point['azimuth']:.0f}°", 12, True, False),
            (f"D – {analysis_point['range']:.1f} км", 12, True, False),
            ("без перешкод", 9, True, False),
            (f"{target_data['detection']}", 9, True, False),
            (f"М – {target_data['scale']}", 9, True, False)
        ]
        
        # Додаємо рядки
        for i, (text, font_size, italic, underline) in enumerate(data_lines):
            if i > 0:
                para = cell.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.space_before = Cm(0)
                para.paragraph_format.space_after = Cm(0)
            
            run = para.add_run(text)
            run.font.name = 'Arial'
            run.font.size = Pt(font_size)
            run.italic = italic
            run.underline = underline
        
        # Рамки: всі крім лівої
        set_cell_borders(cell, top=True, bottom=True, left=False, right=True)
        
        print("✓ Right cell configured")
        
    except Exception as e:
        print(f"✗ Error configuring right cell: {e}")

def create_processed_image_from_data(image_data):
    """Створення обробленого зображення з лінією аналізу"""
    try:
        with Image.open(image_data['image_path']) as original_image:
            if original_image.mode != 'RGB':
                if original_image.mode == 'RGBA':
                    rgb_image = Image.new('RGB', original_image.size, (255, 255, 255))
                    rgb_image.paste(original_image, mask=original_image.split()[-1])
                    final_image = rgb_image
                else:
                    final_image = original_image.convert('RGB')
            else:
                final_image = original_image.copy()
        
        # Додаємо лінію аналізу
        draw = ImageDraw.Draw(final_image)
        analysis_point = image_data['analysis_point']
        
        draw.line([
            (analysis_point['x'], analysis_point['y']), 
            (final_image.width - 1, 0)
        ], fill='black', width=3)
        
        return final_image
        
    except Exception as e:
        print(f"Error creating processed image: {e}")
        return None
    
# ===== ДОПОМІЖНІ ФУНКЦІЇ =====

def set_cell_vertical_center(cell):
    """Вертикальне центрування комірки"""
    try:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    except:
        try:
            tc = cell._tc
            tcPr = tc.tcPr
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.append(tcPr)
            
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            tcPr.append(vAlign)
        except Exception as e:
            print(f"Warning: Could not set vertical alignment: {e}")

def set_cell_borders(cell, top=False, bottom=False, left=False, right=False):
    """Налаштування рамок комірки"""
    try:
        tc = cell._tc
        tcPr = tc.tcPr
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            tc.append(tcPr)
        
        tcBorders = OxmlElement('w:tcBorders')
        
        border_settings = [
            ('top', top), ('bottom', bottom), ('left', left), ('right', right)
        ]
        
        for position, show_border in border_settings:
            border_element = OxmlElement(f'w:{position}')
            if show_border:
                border_element.set(qn('w:val'), 'single')
                border_element.set(qn('w:sz'), '4')
                border_element.set(qn('w:color'), '000000')
            else:
                border_element.set(qn('w:val'), 'none')
            
            tcBorders.append(border_element)
        
        tcPr.append(tcBorders)
        
    except Exception as e:
        print(f"Warning: Could not set cell borders: {e}")

def set_cell_background(cell, color_hex):
    """Встановлення кольору фону комірки"""
    try:
        tc = cell._tc
        tcPr = tc.tcPr
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')
            tc.append(tcPr)
        
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color_hex)
        tcPr.append(shd)
        
    except Exception as e:
        print(f"Warning: Could not set cell background: {e}")

def set_document_margins_and_font(doc):
    """Налаштування марджинів документу та базового шрифту"""
    try:
        # Налаштування марджинів секції
        section = doc.sections[0]
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(0.5)
        
        # Налаштування базового шрифту документу
        styles = doc.styles
        normal_style = styles['Normal']
        normal_style.font.name = 'Arial'
        normal_style.font.size = Pt(12)
        
        print("Document margins and font configured successfully")
        
    except Exception as e:
        print(f"Warning: Could not set document margins: {e}")

# ===== ТИТУЛЬНА СТОРІНКА (БЕЗ ЗМІН) =====

def create_title_page(doc, title_data):
    """Створення титульної сторінки"""
    target_no = title_data['target_no']
    date = title_data['date']
    unit_info = title_data.get('unit_info', get_default_unit_info())
    commander_info = title_data.get('commander_info', get_default_commander_info())
    
    try:
        # Основний заголовок з центруванням
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_before = Cm(TITLE_PAGE['TOP_SPACING'])
        
        # Основний заголовок "АЛЬБОМ" 
        main_title_run = title_para.add_run("АЛЬБОМ")
        main_title_run.font.size = Pt(TITLE_PAGE['MAIN_FONT_SIZE'])
        main_title_run.font.bold = True
        main_title_run.font.name = 'Arial'
        
        # Підзаголовки з переносами рядків
        subtitle_lines = [
            "фотознімків засобів радіолокації",
            f"військової частини {unit_info}",
            format_ukrainian_date(date),
            f"по цілі №{target_no}"
        ]
        
        for line in subtitle_lines:
            title_para.add_run(f"\n{line}")
            # Налаштовуємо шрифт для кожного рядка
            for run in title_para.runs[-1:]:
                run.font.size = Pt(TITLE_PAGE['MAIN_FONT_SIZE'])
                run.font.bold = True
                run.font.name = 'Arial'
        
        # Нижня частина титульної сторінки з великим відступом
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        footer_para.paragraph_format.space_before = Cm(TITLE_PAGE['FOOTER_SPACING'])
        footer_para.paragraph_format.left_indent = Cm(TITLE_PAGE['FOOTER_LEFT_INDENT'])
        
        # Налаштування табуляції
        try:
            tab_stops = footer_para.paragraph_format.tab_stops
            tab_stops.clear_all()
            tab_stops.add_tab_stop(Cm(13), WD_TAB_ALIGNMENT.RIGHT)
        except Exception as tab_error:
            print(f"Warning: Could not set tab stops: {tab_error}")
        
        # Текст про командира
        commander_text_run = footer_para.add_run(f"Командир військової частини {unit_info}")
        commander_text_run.font.size = Pt(TITLE_PAGE['FOOTER_FONT_SIZE'])
        commander_text_run.font.bold = True
        commander_text_run.font.name = 'Arial'
        
        # Перенос рядка та звання командира
        rank_run = footer_para.add_run(f"\n{commander_info['rank']}")
        rank_run.font.size = Pt(TITLE_PAGE['FOOTER_FONT_SIZE'])
        rank_run.font.bold = True
        rank_run.font.name = 'Arial'
        
        # Табуляція та ім'я командира
        name_run = footer_para.add_run(f"\t{commander_info['name']}")
        name_run.font.size = Pt(TITLE_PAGE['FOOTER_FONT_SIZE'])
        name_run.font.bold = True
        name_run.font.name = 'Arial'
        
        # Розрив сторінки для переходу до наступної сторінки
        doc.add_page_break()
        
        print("Title page created successfully")
        
    except Exception as e:
        print(f"Error creating title page: {e}")
        # Додаємо базову титульну сторінку у випадку помилки
        fallback_title = doc.add_paragraph(f"АЛЬБОМ\nфотознімків засобів радіолокації\nпо цілі №{target_no}")
        fallback_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in fallback_title.runs:
            run.font.size = Pt(24)
            run.font.bold = True
            run.font.name = 'Arial'
        doc.add_page_break()

# ===== СТОРІНКА ОПИСУ (БЕЗ ЗМІН) =====

def create_description_page(doc):
    """Створення сторінки опису"""
    try:
        # Заголовок сторінки з центруванням
        heading = doc.add_paragraph("Опис альбому фотознімків")
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Налаштування шрифту заголовку
        heading_run = heading.runs[0]
        heading_run.font.size = Pt(DESCRIPTION_PAGE['HEADING_FONT_SIZE'])
        heading_run.font.bold = True
        heading_run.font.name = 'Arial'
        
        # Порожній параграф для відступу
        doc.add_paragraph()
        
        # Створення таблиці опису з точними розмірами
        table = doc.add_table(rows=6, cols=6)  # 1 заголовковий + 5 порожніх рядків
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.style = None  # Без стилю за замовчуванням
        
        # Налаштування заголовкового рядка
        header_row = table.rows[0]
        header_row.height = Cm(DESCRIPTION_PAGE['TABLE_HEIGHT'])
        
        # Заповнення заголовків з точними розмірами колонок
        for i, (cell, header_text, col_width) in enumerate(
            zip(header_row.cells, DESCRIPTION_PAGE['HEADERS'], DESCRIPTION_PAGE['COL_WIDTHS'])
        ):
            # Встановлення ширини колонки
            cell.width = Cm(col_width)
            
            # Вертикальне центрування
            try:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            except AttributeError:
                print("Warning: WD_ALIGN_VERTICAL.CENTER not available, using manual alignment")
                pass
            
            # Налаштування тексту
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.text = header_text
            
            # Форматування шрифту заголовків
            for run in para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(12)
                run.font.bold = True
            
            # Додавання рамок до заголовкових комірок
            set_cell_borders(cell, top=True, bottom=True, left=True, right=True)
        
        # Налаштування порожніх рядків (рядки 1-5)
        for row_index in range(1, 6):
            row = table.rows[row_index]
            
            # Встановлення висоти рядка
            row.height = Cm(1.0)  # Стандартна висота для порожніх рядків
            
            # Налаштування комірок порожніх рядків
            for col_index, cell in enumerate(row.cells):
                # Встановлення ширини (така ж як в заголовку)
                cell.width = Cm(DESCRIPTION_PAGE['COL_WIDTHS'][col_index])
                
                # Безпечне вертикальне центрування
                try:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                except AttributeError:
                    try:
                        tc = cell._tc
                        tcPr = tc.tcPr
                        if tcPr is None:
                            tcPr = OxmlElement('w:tcPr')
                            tc.append(tcPr)
                        
                        vAlign = OxmlElement('w:vAlign')
                        vAlign.set(qn('w:val'), 'center')
                        tcPr.append(vAlign)
                    except Exception:
                        pass  # Ігноруємо якщо не вдається
                
                # Порожній текст
                cell.text = ""
                
                # Додавання рамок
                set_cell_borders(cell, top=True, bottom=True, left=True, right=True)
        
        # Розрив сторінки для переходу до наступної сторінки
        doc.add_page_break()
        
        print("Description page created successfully")
        
    except Exception as e:
        print(f"Error creating description page: {e}")
        
        # Fallback - проста сторінка опису
        fallback_heading = doc.add_paragraph("Опис альбому фотознімків")
        fallback_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fallback_heading.runs[0].font.size = Pt(22)
        fallback_heading.runs[0].font.bold = True
        
        # Простий текст замість таблиці
        doc.add_paragraph("Таблиця опису документів")
        doc.add_page_break()

# ===== ОСНОВНИЙ КЛАС GUI =====

class AzimuthGUI(QMainWindow):
    def __init__(self):
        super().__init__()
        self.processor = None
        self.current_image_path = None
        self.current_click = None
        self.scale_factor_x = 1.0
        self.scale_factor_y = 1.0
        self.offset_x = 0
        self.offset_y = 0
        self.current_folder = None
        
        self.current_language = 'UKRAINIAN'  # Default language
        self.translations = Translations()
        
        self.scale_edge_mode = False
        self.scale_edge_point = None
        self.custom_scale_distance = None
        self.center_setting_mode = False
        
        self.current_target_number = "0001"
        self.current_height = "0.0"
        self.current_obstacles = "без перешкод"
        self.current_detection = "Виявлення"
        
        self.current_word_document = None
        self.word_document_path = None
        
        self.processed_images = []
        self.current_folder_images = []
        self.current_image_index = -1
        
        # Змінні для збереження налаштувань сітки між зображеннями
        self.saved_grid_settings = {
            'center_offset_x': 0,
            'center_offset_y': 0,
            'scale_edge_relative': None,
            'custom_scale_distance': None,
            'scale_value': "300"
        }
        
        self.init_ui()
    
    def tr(self, key):
        return self.translations.get(self.current_language, key)
    
    def set_language(self, language):
        self.current_language = language
        self.update_interface_text()
        
        for lang, action in self.language_actions.items():
            action.setChecked(lang == language)
    
    def update_interface_text(self):
        self.setWindowTitle(self.tr("window_title"))
        self.create_menu_bar()
        
        if hasattr(self, 'control_title'):
            self.control_title.setText(self.tr("controls"))
        if hasattr(self, 'report_title'):
            self.report_title.setText(self.tr("report_data"))
        if hasattr(self, 'browser_label'):
            self.browser_label.setText(self.tr("photo_browser"))
        
        if hasattr(self, 'open_image_btn'):
            self.open_image_btn.setText(self.tr("open_image"))
        if hasattr(self, 'open_folder_btn'):
            self.open_folder_btn.setText(self.tr("open_folder"))
        if hasattr(self, 'save_image_btn'):
            self.save_image_btn.setText(self.tr("save_current_image"))
        if hasattr(self, 'export_new_btn'):
            self.export_new_btn.setText(self.tr("create_new_album"))
        if hasattr(self, 'export_add_btn'):
            self.export_add_btn.setText(self.tr("add_to_existing_album"))
        if hasattr(self, 'scale_edge_btn'):
            self.scale_edge_btn.setText(self.tr("set_scale_edge"))
        if hasattr(self, 'set_center_btn'):
            self.set_center_btn.setText(self.tr("set_center"))
        
        if hasattr(self, 'file_ops_label'):
            self.file_ops_label.setText(self.tr("file_operations"))
        if hasattr(self, 'azimuth_grid_label'):
            self.azimuth_grid_label.setText(self.tr("azimuth_grid"))
        if hasattr(self, 'move_center_label'):
            self.move_center_label.setText(self.tr("move_center"))
        if hasattr(self, 'results_label'):
            self.results_label.setText(self.tr("results"))
        
        # Batch processing переклади
        if hasattr(self, 'batch_label'):
            self.batch_label.setText(self.tr("batch_processing"))
        if hasattr(self, 'save_current_btn'):
            self.save_current_btn.setText(self.tr("save_current_image_data"))
        if hasattr(self, 'prev_btn'):
            self.prev_btn.setText(f"← {self.tr('previous')}")
        if hasattr(self, 'next_btn'):
            self.next_btn.setText(f"{self.tr('next')} →")
        
        if not self.current_image_path:
            self.image_label.setText(self.tr("open_instruction"))
        
        if self.processor:
            self.update_results_display()

    def init_ui(self):
        self.setWindowTitle(self.tr("window_title"))
        
        default_width = 1400
        default_height = 900
        min_width = 1000  
        min_height = 700
        self.setMinimumSize(min_width, min_height)
        self.setWindowFlags(Qt.Window)
        self.resize(default_width, default_height)
        self.showMaximized()
        
        self.create_menu_bar()
        
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        main_layout = QHBoxLayout()
        central_widget.setLayout(main_layout)
        
        main_splitter = QSplitter(Qt.Horizontal)
        main_layout.addWidget(main_splitter)
        
        self.create_control_panel(main_splitter)
        self.create_vertical_browser_panel(main_splitter)
        self.create_image_panel(main_splitter)
        self.create_report_panel(main_splitter)
        
        main_splitter.setStretchFactor(0, 0)
        main_splitter.setStretchFactor(1, 0)
        main_splitter.setStretchFactor(2, 1)
        main_splitter.setStretchFactor(3, 0)
        
        main_splitter.setSizes([300, 0, 800, 300])
        self.main_splitter = main_splitter
    
    def resizeEvent(self, event):
        super().resizeEvent(event)
        
        if hasattr(self, 'processor') and self.processor:
            QTimer.singleShot(50, self.update_image_display_after_resize)
        
        if hasattr(self, 'main_splitter') and hasattr(self, 'browser_widget'):
            current_sizes = self.main_splitter.sizes()
            total_width = sum(current_sizes)
            
            if total_width > 800:
                if self.browser_widget.isVisible():
                    new_image_width = max(400, total_width - 300 - 180 - 300)
                    self.main_splitter.setSizes([300, 180, new_image_width, 300])
                else:
                    new_image_width = max(400, total_width - 300 - 300)
                    self.main_splitter.setSizes([300, 0, new_image_width, 300])
    
    def update_image_display_after_resize(self):
        if hasattr(self, 'processor') and self.processor:
            self.display_image()
    
    def create_menu_bar(self):
        menubar = self.menuBar()
        menubar.clear()
        
        settings_menu = menubar.addMenu(self.tr("settings"))
        language_menu = settings_menu.addMenu(self.tr("language"))
        
        english_action = language_menu.addAction("English")
        english_action.triggered.connect(lambda: self.set_language('ENGLISH'))
        
        ukrainian_action = language_menu.addAction("Українська")
        ukrainian_action.triggered.connect(lambda: self.set_language('UKRAINIAN'))
        
        english_action.setCheckable(True)
        ukrainian_action.setCheckable(True)
        
        english_action.setChecked(self.current_language == 'ENGLISH')
        ukrainian_action.setChecked(self.current_language == 'UKRAINIAN')
        
        self.language_actions = {
            'ENGLISH': english_action,
            'UKRAINIAN': ukrainian_action
        }

    def create_control_panel(self, parent):
        """Ліва панель з правильними перекладами"""
        control_widget = QWidget()
        control_widget.setFixedWidth(300)
        control_widget.setStyleSheet("""
            QWidget {
                background-color: #f5f5f5; 
                border-right: 1px solid #ccc;
            }
            QLabel {
                background: none;
                border: none;
                color: #333;
            }
            QPushButton {
                background-color: #e1e1e1;
                border: 1px solid #c1c1c1;
                border-radius: 6px;
                padding: 8px 12px;
                font-size: 11px;
                font-weight: bold;
                box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
            }
            QPushButton:hover {
                background-color: #d1d1d1;
                border: 1px solid #b1b1b1;
                box-shadow: 0 3px 6px rgba(0, 0, 0, 0.15);
            }
            QPushButton:pressed {
                background-color: #c1c1c1;
                box-shadow: 0 1px 2px rgba(0, 0, 0, 0.2);
                transform: translateY(1px);
            }
            QComboBox {
                border: 1px solid #c1c1c1;
                border-radius: 4px;
                padding: 4px 8px;
                background-color: white;
                font-size: 11px;
                box-shadow: 0 1px 3px rgba(0, 0, 0, 0.1);
            }
        """)
        
        layout = QVBoxLayout()
        layout.setContentsMargins(15, 15, 15, 15)
        layout.setSpacing(10)
        control_widget.setLayout(layout)
        
        # Title
        self.control_title = QLabel(self.tr("controls"))
        self.control_title.setFont(QFont("Arial", 14, QFont.Bold))
        self.control_title.setAlignment(Qt.AlignCenter)
        self.control_title.setStyleSheet("font-size: 14px; font-weight: bold; margin-bottom: 10px;")
        layout.addWidget(self.control_title)
        
        # Image Selection section
        self.file_ops_label = QLabel(self.tr("file_operations"))
        self.file_ops_label.setFont(QFont("Arial", 10, QFont.Bold))
        self.file_ops_label.setStyleSheet("color: #666; margin-top: 10px; margin-bottom: 5px;")
        layout.addWidget(self.file_ops_label)
        
        # File operation buttons
        self.open_image_btn = QPushButton(self.tr("open_image"))
        self.open_image_btn.clicked.connect(self.open_image)
        layout.addWidget(self.open_image_btn)
        
        self.open_folder_btn = QPushButton(self.tr("open_folder"))
        self.open_folder_btn.clicked.connect(self.open_folder)
        layout.addWidget(self.open_folder_btn)
        
        self.save_image_btn = QPushButton(self.tr("save_current_image"))
        self.save_image_btn.clicked.connect(self.save_current_image)
        layout.addWidget(self.save_image_btn)
        
        self.export_new_btn = QPushButton(self.tr("create_new_album"))
        self.export_new_btn.clicked.connect(self.create_new_word_album)
        layout.addWidget(self.export_new_btn)
        
        self.export_add_btn = QPushButton(self.tr("add_to_existing_album"))
        self.export_add_btn.clicked.connect(self.add_to_existing_album)
        layout.addWidget(self.export_add_btn)
        
        # Separator
        separator1 = QFrame()
        separator1.setFrameShape(QFrame.HLine)
        separator1.setFrameShadow(QFrame.Sunken)
        separator1.setStyleSheet("color: #ccc; margin: 10px 0px;")
        layout.addWidget(separator1)
        
        # Batch processing section (з перекладами)
        self.batch_label = QLabel(self.tr("batch_processing"))
        self.batch_label.setFont(QFont("Arial", 10, QFont.Bold))
        self.batch_label.setStyleSheet("color: #666; margin-top: 5px; margin-bottom: 5px;")
        layout.addWidget(self.batch_label)

        self.save_current_btn = QPushButton(self.tr("save_current_image_data"))
        self.save_current_btn.clicked.connect(self.save_current_image_data)
        layout.addWidget(self.save_current_btn)

        nav_layout = QHBoxLayout()
        self.prev_btn = QPushButton(f"← {self.tr('previous')}")
        self.prev_btn.clicked.connect(self.previous_image_in_batch)
        nav_layout.addWidget(self.prev_btn)

        self.next_btn = QPushButton(f"{self.tr('next')} →")
        self.next_btn.clicked.connect(self.next_image_in_batch)
        nav_layout.addWidget(self.next_btn)

        nav_widget = QWidget()
        nav_widget.setLayout(nav_layout)
        layout.addWidget(nav_widget)

        self.create_batch_btn = QPushButton("Create Batch Album")
        self.create_batch_btn.clicked.connect(self.create_batch_album)
        self.create_batch_btn.setStyleSheet("""
            QPushButton {
                font-weight: bold; 
                background-color: #4CAF50; 
                color: white;
                border: 1px solid #45a049;
                border-radius: 6px;
                padding: 8px 12px;
                box-shadow: 0 2px 4px rgba(0, 0, 0, 0.15);
            }
            QPushButton:hover {
                background-color: #45a049;
                box-shadow: 0 3px 6px rgba(0, 0, 0, 0.2);
            }
            QPushButton:pressed {
                background-color: #3d8b40;
                box-shadow: 0 1px 2px rgba(0, 0, 0, 0.25);
                transform: translateY(1px);
            }
        """)
        layout.addWidget(self.create_batch_btn)
        
        # Separator
        batch_separator = QFrame()
        batch_separator.setFrameShape(QFrame.HLine)
        batch_separator.setFrameShadow(QFrame.Sunken)
        batch_separator.setStyleSheet("color: #ccc; margin: 10px 0px;")
        layout.addWidget(batch_separator)
        
        # Results area
        self.results_label = QLabel(self.tr("results"))
        self.results_label.setFont(QFont("Arial", 10, QFont.Bold))
        self.results_label.setStyleSheet("color: #666; margin-top: 10px; margin-bottom: 5px;")
        layout.addWidget(self.results_label)
        
        self.results_text = QTextEdit()
        self.results_text.setMaximumHeight(200)
        self.results_text.setStyleSheet("""
            QTextEdit {
                background-color: white;
                border: 1px solid #ccc;
                font-family: 'Courier New', monospace;
                font-size: 10px;
                border-radius: 4px;
                box-shadow: inset 0 1px 3px rgba(0, 0, 0, 0.1);
            }
        """)
        layout.addWidget(self.results_text)
        
        layout.addStretch()
        parent.addWidget(control_widget)
    
    def create_vertical_browser_panel(self, parent):
        browser_widget = QWidget()
        browser_widget.setFixedWidth(180)
        browser_widget.setStyleSheet("""
            QWidget {
                background-color: #f8f8f8; 
                border-right: 1px solid #ccc;
                border-left: 1px solid #ccc;
            }
            QLabel {
                background: none;
                border: none;
                color: #333;
            }
        """)
        
        layout = QVBoxLayout()
        layout.setContentsMargins(0, 15, 0, 15)
        layout.setSpacing(10)
        browser_widget.setLayout(layout)
        
        self.browser_label = QLabel(self.tr("photo_browser"))
        self.browser_label.setFont(QFont("Arial", 10, QFont.Bold))
        self.browser_label.setStyleSheet("color: #666; margin-bottom: 5px; padding: 0 10px;")
        self.browser_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(self.browser_label)
        
        self.thumbnail_scroll = QScrollArea()
        self.thumbnail_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.thumbnail_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self.thumbnail_scroll.setWidgetResizable(False)
        self.thumbnail_scroll.setStyleSheet("border: none; background: transparent;")
        
        self.thumbnail_widget = VerticalThumbnailWidget()
        self.thumbnail_widget.image_selected.connect(self.load_image_from_browser)
        self.thumbnail_scroll.setWidget(self.thumbnail_widget)
        
        layout.addWidget(self.thumbnail_scroll)
        
        browser_widget.hide()
        self.browser_widget = browser_widget
        
        parent.addWidget(browser_widget)
    
    def create_image_panel(self, parent):
        image_widget = QWidget()
        layout = QVBoxLayout()
        image_widget.setLayout(layout)
        
        self.image_label = ClickableLabel()
        self.image_label.clicked.connect(self.on_image_click)
        self.image_label.dragged.connect(self.on_image_drag)
        self.image_label.mouse_moved.connect(self.on_mouse_hover)
        self.image_label.setText(self.tr("open_instruction"))
        
        scroll_area = QScrollArea()
        scroll_area.setWidget(self.image_label)
        scroll_area.setWidgetResizable(True)
        scroll_area.setMinimumHeight(400)
        layout.addWidget(scroll_area)
        
        parent.addWidget(image_widget)
    
    def create_report_panel(self, parent):
        """Права панель з азимутальними контролами"""
        report_widget = QWidget()
        report_widget.setFixedWidth(300)
        report_widget.setStyleSheet("background-color: #f9f9f9; border-left: 1px solid #ccc;")
        
        layout = QVBoxLayout()
        report_widget.setLayout(layout)
        
        # Title
        self.report_title = QLabel(self.tr("report_data"))
        self.report_title.setFont(QFont("Arial", 14, QFont.Bold))
        self.report_title.setAlignment(Qt.AlignCenter)
        layout.addWidget(self.report_title)
        
        # Target data input section
        manual_group = QFrame()
        manual_group.setStyleSheet("background-color: #fff; border: 1px solid #ddd; border-radius: 5px; padding: 10px;")
        manual_layout = QVBoxLayout()
        manual_group.setLayout(manual_layout)
        
        self.target_number_input = QLineEdit()
        self.target_number_input.setPlaceholderText("Номер цілі")
        self.target_number_input.setText(self.current_target_number)
        self.target_number_input.textChanged.connect(self.update_target_number)
        manual_layout.addWidget(self.target_number_input)
        
        self.auto_azimuth_label = QLabel("β - --°")
        manual_layout.addWidget(self.auto_azimuth_label)
        
        self.auto_distance_label = QLabel("D - -- км")
        manual_layout.addWidget(self.auto_distance_label)
        
        height_layout = QHBoxLayout()
        height_layout.addWidget(QLabel("H –"))
        self.height_input = QLineEdit(self.current_height)
        self.height_input.setMaximumWidth(80)
        self.height_input.textChanged.connect(self.update_height)
        height_layout.addWidget(self.height_input)
        height_layout.addWidget(QLabel(self.tr("km_unit")))
        height_layout.addStretch()
        manual_layout.addLayout(height_layout)
        
        self.obstacles_combo = QComboBox()
        self.obstacles_combo.addItems([self.tr("no_obstacles"), self.tr("with_obstacles")])
        self.obstacles_combo.currentTextChanged.connect(self.update_obstacles)
        manual_layout.addWidget(self.obstacles_combo)
        
        self.detection_combo = QComboBox()
        self.detection_combo.addItems([self.tr("detection"), self.tr("tracking"), self.tr("loss")])
        self.detection_combo.currentTextChanged.connect(self.update_detection)
        manual_layout.addWidget(self.detection_combo)
        
        self.auto_scale_label = QLabel("M = --")
        manual_layout.addWidget(self.auto_scale_label)
        
        layout.addWidget(manual_group)
        
        # Separator
        grid_separator = QFrame()
        grid_separator.setFrameShape(QFrame.HLine)
        grid_separator.setFrameShadow(QFrame.Sunken)
        grid_separator.setStyleSheet("color: #ccc; margin: 15px 0px 10px 0px;")
        layout.addWidget(grid_separator)
        
        # Azimuth Grid section
        self.azimuth_grid_label = QLabel(self.tr("azimuth_grid"))
        self.azimuth_grid_label.setFont(QFont("Arial", 12, QFont.Bold))
        self.azimuth_grid_label.setStyleSheet("color: #333; margin-bottom: 10px;")
        self.azimuth_grid_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(self.azimuth_grid_label)
        
        # Scale setting
        scale_layout = QHBoxLayout()
        scale_label = QLabel(self.tr("scale_setting"))
        scale_label.setStyleSheet("color: #333; margin-right: 5px;")
        scale_layout.addWidget(scale_label)
        self.scale_combo = QComboBox()
        self.scale_combo.addItems(["50", "100", "150", "200", "300"])
        self.scale_combo.setCurrentText("300")
        self.scale_combo.currentTextChanged.connect(self.update_scale)
        self.scale_combo.setStyleSheet("""
            QComboBox {
                border: 1px solid #c1c1c1;
                border-radius: 4px;
                padding: 4px 8px;
                background-color: white;
                font-size: 11px;
                box-shadow: 0 1px 3px rgba(0, 0, 0, 0.1);
            }
        """)
        scale_layout.addWidget(self.scale_combo)
        scale_layout.addStretch()
        layout.addLayout(scale_layout)
        
        # Scale edge tool
        self.scale_edge_btn = QPushButton(self.tr("set_scale_edge"))
        self.scale_edge_btn.setCheckable(True)
        self.scale_edge_btn.clicked.connect(self.toggle_scale_edge_mode)
        self.scale_edge_btn.setStyleSheet("""
            QPushButton {
                background-color: #e1e1e1;
                border: 1px solid #c1c1c1;
                border-radius: 6px;
                padding: 8px 12px;
                font-size: 11px;
                font-weight: bold;
                box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
                margin-top: 5px;
            }
            QPushButton:hover {
                background-color: #d1d1d1;
                border: 1px solid #b1b1b1;
                box-shadow: 0 3px 6px rgba(0, 0, 0, 0.15);
            }
            QPushButton:pressed {
                background-color: #c1c1c1;
                box-shadow: 0 1px 2px rgba(0, 0, 0, 0.2);
                transform: translateY(1px);
            }
            QPushButton:checked {
                background-color: #4CAF50;
                color: white;
                border: 1px solid #45a049;
            }
        """)
        layout.addWidget(self.scale_edge_btn)
        
        # Center setting tool
        self.set_center_btn = QPushButton(self.tr("set_center"))
        self.set_center_btn.setCheckable(True)
        self.set_center_btn.clicked.connect(self.toggle_center_setting_mode)
        self.set_center_btn.setStyleSheet("""
            QPushButton {
                background-color: #e1e1e1;
                border: 1px solid #c1c1c1;
                border-radius: 6px;
                padding: 8px 12px;
                font-size: 11px;
                font-weight: bold;
                box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
                margin-top: 5px;
            }
            QPushButton:hover {
                background-color: #d1d1d1;
                border: 1px solid #b1b1b1;
                box-shadow: 0 3px 6px rgba(0, 0, 0, 0.15);
            }
            QPushButton:pressed {
                background-color: #c1c1c1;
                box-shadow: 0 1px 2px rgba(0, 0, 0, 0.2);
                transform: translateY(1px);
            }
            QPushButton:checked {
                background-color: #FF9800;
                color: white;
                border: 1px solid #f57c00;
            }
        """)
        layout.addWidget(self.set_center_btn)
        
        # Center movement controls (БЕЗ центральної кнопки)
        self.move_center_label = QLabel(self.tr("move_center"))
        self.move_center_label.setStyleSheet("color: #333; margin-top: 10px; margin-bottom: 5px;")
        layout.addWidget(self.move_center_label)
        
        move_layout = QGridLayout()
        move_layout.setSpacing(4)

        btn_style = """
            QPushButton {
                background-color: #e1e1e1;
                border: 1px solid #c1c1c1;
                border-radius: 6px;
                padding: 8px;
                font-size: 14px;
                font-weight: bold;
                box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
                min-width: 45px;
                min-height: 35px;
            }
            QPushButton:hover {
                background-color: #d1d1d1;
                border: 1px solid #b1b1b1;
                box-shadow: 0 3px 6px rgba(0, 0, 0, 0.15);
            }
            QPushButton:pressed {
                background-color: #c1c1c1;
                box-shadow: 0 1px 2px rgba(0, 0, 0, 0.2);
                transform: translateY(1px);
            }
        """

        # Кнопки навігації БЕЗ центральної
        btn_up = QPushButton("↑")
        btn_up.setStyleSheet(btn_style)
        btn_up.clicked.connect(lambda: self.move_center_and_save(0, -2))
        move_layout.addWidget(btn_up, 0, 1)

        btn_left = QPushButton("←")
        btn_left.setStyleSheet(btn_style)
        btn_left.clicked.connect(lambda: self.move_center_and_save(-2, 0))
        move_layout.addWidget(btn_left, 1, 0)

        btn_right = QPushButton("→")
        btn_right.setStyleSheet(btn_style)
        btn_right.clicked.connect(lambda: self.move_center_and_save(2, 0))
        move_layout.addWidget(btn_right, 1, 2)

        btn_down = QPushButton("↓")
        btn_down.setStyleSheet(btn_style)
        btn_down.clicked.connect(lambda: self.move_center_and_save(0, 2))
        move_layout.addWidget(btn_down, 2, 1)
        
        layout.addLayout(move_layout)
        layout.addStretch()
        
        report_widget.hide()
        self.report_widget = report_widget
        parent.addWidget(report_widget)

    def move_center_and_save(self, dx, dy):
        """Переміщення центру з автоматичним збереженням налаштувань"""
        if not self.processor:
            QMessageBox.warning(self, self.tr("warning"), self.tr("no_image_first"))
            return
        
        self.processor.move_center(dx, dy)
        self.save_current_grid_settings()
        
        if self.current_click:
            azimuth, range_val = self.calculate_azimuth_range(
                self.current_click['x'], self.current_click['y']
            )
            self.current_click['azimuth'] = azimuth
            self.current_click['range'] = range_val
        
        self.display_image()
        self.update_results_display()
        self.update_report_data()
        
        self.add_result(self.tr("center_moved").format(x=self.processor.center_x, 
                                                    y=self.processor.center_y))
        self.add_result(self.tr("grid_settings_saved"))
        
        if hasattr(self, 'image_label') and self.image_label.zoom_widget.isVisible():
            self.image_label.zoom_widget.update_cursor_position(self.processor.center_x, self.processor.center_y)

    def add_result(self, text):
        self.results_text.append(text)
    
    def open_image(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, self.tr("select_image"), "",
            f"{self.tr('image_files')} (*.jpg *.jpeg *.png *.bmp *.gif *.tiff);;{self.tr('all_files')} (*.*)"
        )
        
        if file_path:
            self.load_image(file_path)
            self.report_widget.show()
    
    def open_folder(self):
        folder_path = QFileDialog.getExistingDirectory(self, self.tr("select_folder"))
        
        if folder_path:
            self.current_folder = folder_path
            self.load_folder_thumbnails()
            
            self.browser_widget.show()
            self.main_splitter.setSizes([300, 180, 720, 300])
            
            image_extensions = ('.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff', '.tif')
            image_count = sum(1 for f in os.listdir(folder_path) 
                             if f.lower().endswith(image_extensions))
            
            self.add_result(f"{self.tr('loaded_folder')}: {os.path.basename(folder_path)}")
            self.add_result(self.tr("found_images").format(count=image_count))
            
            self.report_widget.show()

    def load_folder_thumbnails(self):
        if not self.current_folder:
            return
        
        self.thumbnail_widget.clear_thumbnails()
        
        image_extensions = ('.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff', '.tif')
        image_files = []
        
        try:
            for filename in os.listdir(self.current_folder):
                if filename.lower().endswith(image_extensions):
                    full_path = os.path.join(self.current_folder, filename)
                    image_files.append(full_path)
        except Exception as e:
            return

        image_files.sort()
        self.current_folder_images = image_files  # Виправлено: ініціалізація для batch processing
        
        if len(image_files) == 0:
            no_images_label = QLabel(self.tr("no_images_found"))
            no_images_label.setAlignment(Qt.AlignCenter)
            no_images_label.setStyleSheet("color: gray; font-size: 14px; padding: 20px;")
            no_images_label.setWordWrap(True)
            self.thumbnail_widget.layout.addWidget(no_images_label)
            return
        
        for image_path in image_files:
            try:
                self.thumbnail_widget.add_thumbnail(image_path)
            except Exception as e:
                pass
        
        widget_height = len(image_files) * 130 + 20
        self.thumbnail_widget.setMinimumHeight(widget_height)
        self.thumbnail_widget.resize(160, widget_height)

    def load_image_from_browser(self, file_path):
        self.load_image(file_path)
        self.add_result(self.tr("loaded_from_browser").format(name=os.path.basename(file_path)))
    
    def load_image(self, file_path):
        """Завантаження зображення зі збереженням налаштувань сітки"""
        try:
            self.current_image_path = file_path
            scale_value = int(self.scale_combo.currentText())
            
            self.processor = AzimuthImageProcessor(file_path, scale=scale_value)
            
            if hasattr(self.processor, 'image') and self.processor.image:
                original_mode = self.processor.image.mode
                
                if original_mode != 'RGB':
                    if original_mode == 'RGBA':
                        rgb_image = Image.new('RGB', self.processor.image.size, (255, 255, 255))
                        rgb_image.paste(self.processor.image, mask=self.processor.image.split()[-1])
                        self.processor.image = rgb_image
                    else:
                        self.processor.image = self.processor.image.convert('RGB')
            
            # Застосовуємо збережені налаштування сітки
            self.apply_saved_grid_settings()
            
            # Очищуємо тільки точку аналізу, але НЕ налаштування сітки
            self.current_click = None
            
            self.display_image()
            self.update_results_display()
            self.update_report_data()
            
            self.add_result(f"{self.tr('loaded')}: {os.path.basename(file_path)}")
            if self.has_saved_grid_settings():
                self.add_result(self.tr("grid_settings_applied"))
            
        except Exception as e:
            QMessageBox.critical(self, self.tr("error"), 
                               self.tr("could_not_load").format(error=str(e)))
    
    def apply_saved_grid_settings(self):
        """Застосувати збережені налаштування сітки до нового зображення"""
        if not self.processor:
            return
        
        # Отримуємо центр нового зображення
        image_center_x = self.processor.image.width // 2
        image_center_y = self.processor.image.height // 2
        
        # Застосовуємо зміщення центру
        if self.saved_grid_settings['center_offset_x'] != 0 or self.saved_grid_settings['center_offset_y'] != 0:
            self.processor.move_center(
                self.saved_grid_settings['center_offset_x'],
                self.saved_grid_settings['center_offset_y']
            )
        
        # ВИПРАВЛЕННЯ: Відновлюємо scale edge point відносно центру ЗОБРАЖЕННЯ
        if self.saved_grid_settings.get('scale_edge_relative') and self.saved_grid_settings['custom_scale_distance']:
            edge_relative = self.saved_grid_settings['scale_edge_relative']
            
            # Нові абсолютні координати відносно центру зображення
            new_x = image_center_x + edge_relative['x']
            new_y = image_center_y + edge_relative['y']
            
            # Перевіряємо що координати в межах зображення
            if 0 <= new_x < self.processor.image.width and 0 <= new_y < self.processor.image.height:
                self.scale_edge_point = {'x': new_x, 'y': new_y}
                self.custom_scale_distance = self.saved_grid_settings['custom_scale_distance']
    
        # Відновлюємо масштаб
        if self.saved_grid_settings['scale_value']:
            self.scale_combo.setCurrentText(self.saved_grid_settings['scale_value'])

    def save_current_grid_settings(self):
        """Зберегти поточні налаштування сітки"""
        if not self.processor:
            return
        
        # Зберігаємо зміщення центру відносно центру зображення
        image_center_x = self.processor.image.width // 2
        image_center_y = self.processor.image.height // 2
        
        # ВИПРАВЛЕННЯ: Зберігаємо scale edge point відносно центру ЗОБРАЖЕННЯ, а не сітки
        scale_edge_relative = None
        if self.scale_edge_point:
            scale_edge_relative = {
                'x': self.scale_edge_point['x'] - image_center_x,
                'y': self.scale_edge_point['y'] - image_center_y
            }
        
        self.saved_grid_settings = {
            'center_offset_x': self.processor.center_x - image_center_x,
            'center_offset_y': self.processor.center_y - image_center_y,
            'scale_edge_relative': scale_edge_relative,  # Відносно центру зображення
            'custom_scale_distance': self.custom_scale_distance,
            'scale_value': self.scale_combo.currentText()
            }

    def has_saved_grid_settings(self):
        """Перевірити чи є збережені налаштування сітки"""
        return (self.saved_grid_settings['center_offset_x'] != 0 or 
                self.saved_grid_settings['center_offset_y'] != 0 or
                self.saved_grid_settings.get('scale_edge_relative') is not None or  # ВИПРАВЛЕННЯ
                self.saved_grid_settings['scale_value'] != "300")
    
    def display_image(self):
        if not self.processor:
            return
        
        pil_image = self.processor.image.copy()
        draw = ImageDraw.Draw(pil_image)
        
        center_x, center_y = self.processor.center_x, self.processor.center_y
        cross_size = 15
        
        draw.line([center_x - cross_size, center_y, center_x + cross_size, center_y], 
                  fill='red', width=2)
        draw.line([center_x, center_y - cross_size, center_x, center_y + cross_size], 
                  fill='red', width=2)
        draw.ellipse([center_x - 3, center_y - 3, center_x + 3, center_y + 3], 
                     fill='red', outline='white')
        
        if self.current_click:
            click_x, click_y = self.current_click['x'], self.current_click['y']
            
            draw.ellipse([click_x - 4, click_y - 4, click_x + 4, click_y + 4], 
                         fill='blue', outline='white', width=1)
            
            draw.line([click_x, click_y, pil_image.width - 1, 0], fill='blue', width=3)
        
        if self.scale_edge_point:
            edge_x, edge_y = self.scale_edge_point['x'], self.scale_edge_point['y']
            
            draw.ellipse([edge_x - 5, edge_y - 5, edge_x + 5, edge_y + 5], 
                         fill='green', outline='white', width=2)
            
            draw.line([center_x, center_y, edge_x, edge_y], fill='green', width=2)
            
            dx = edge_x - center_x
            dy = edge_y - center_y
            length = math.sqrt(dx*dx + dy*dy)
            if length > 0:
                nx, ny = -dy/length, dx/length
                perp_size = 8
                draw.line([
                    edge_x + nx*perp_size, edge_y + ny*perp_size,
                    edge_x - nx*perp_size, edge_y - ny*perp_size
                ], fill='green', width=2)
        
        with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as temp_file:
            temp_path = temp_file.name
            pil_image.save(temp_path, 'JPEG', quality=95)
        
        pixmap = QPixmap(temp_path)
        
        try:
            os.remove(temp_path)
        except:
            pass
        
        widget_width = self.image_label.width()
        widget_height = self.image_label.height()
        
        scaled_pixmap = pixmap.scaled(widget_width, widget_height, 
                                     Qt.KeepAspectRatio, Qt.SmoothTransformation)
        
        original_width = pixmap.width()
        original_height = pixmap.height()
        scaled_width = scaled_pixmap.width()
        scaled_height = scaled_pixmap.height()
        
        self.scale_factor_x = scaled_width / original_width
        self.scale_factor_y = scaled_height / original_height
        
        self.offset_x = (widget_width - scaled_width) // 2
        self.offset_y = (widget_height - scaled_height) // 2
        
        self.image_label.update_image_geometry(
            original_width, original_height,
            self.scale_factor_x, self.scale_factor_y,
            self.offset_x, self.offset_y
        )
        
        self.image_label.set_zoom_source_image(pixmap)
        self.image_label.setPixmap(scaled_pixmap)
    
    def on_image_click(self, x, y):
        if not self.processor:
            return

        if self.scale_edge_mode:
            self.set_scale_edge_point(x, y)
            return

        if self.center_setting_mode:
            self.set_center_point(x, y)
            return

        if self.current_click:
            existing_widget_x = self.current_click['x'] * self.scale_factor_x + self.offset_x
            existing_widget_y = self.current_click['y'] * self.scale_factor_y + self.offset_y
        
            clicked_widget_x = x * self.scale_factor_x + self.offset_x
            clicked_widget_y = y * self.scale_factor_y + self.offset_y
        
            distance = ((clicked_widget_x - existing_widget_x)**2 + (clicked_widget_y - existing_widget_y)**2)**0.5
        
            if distance <= 15:
                return

        self.place_analysis_point(x, y)
    
    def on_image_drag(self, x, y):
        if not self.processor:
            return
        
        self.place_analysis_point(x, y)
    
    def on_mouse_hover(self, x, y):
        if not self.processor or not self.current_click:
            QToolTip.hideText()
            return
        
        hover_widget_x = x * self.scale_factor_x + self.offset_x
        hover_widget_y = y * self.scale_factor_y + self.offset_y
        
        existing_widget_x = self.current_click['x'] * self.scale_factor_x + self.offset_x
        existing_widget_y = self.current_click['y'] * self.scale_factor_y + self.offset_y
        
        distance = ((hover_widget_x - existing_widget_x)**2 + (hover_widget_y - existing_widget_y)**2)**0.5
        
        if distance <= 15:
            azimuth = self.current_click['azimuth']
            range_val = self.current_click['range']
            tooltip_text = f"{self.tr('azimuth')}: {azimuth:.0f}°\n{self.tr('range')}: {range_val:.0f} км"
            
            point = self.image_label.mapToGlobal(self.image_label.rect().topLeft())
            tooltip_x = point.x() + hover_widget_x + 15
            tooltip_y = point.y() + hover_widget_y - 10
            
            QToolTip.showText(QPoint(int(tooltip_x), int(tooltip_y)), tooltip_text)
        else:
            QToolTip.hideText()
    
    def place_analysis_point(self, x, y):
        if not self.processor:
            return
        
        x = max(0, min(x, self.processor.image.width - 1))
        y = max(0, min(y, self.processor.image.height - 1))
        
        try:
            azimuth, range_val = self.calculate_azimuth_range(x, y)
            
            self.current_click = {
                'x': x, 'y': y,
                'azimuth': azimuth, 'range': range_val
            }
            
            self.display_image()
            self.update_results_display()
            self.update_report_data()
            
        except Exception as e:
            QMessageBox.critical(self, self.tr("error"), f"Could not process point: {str(e)}")
    
    def calculate_azimuth_range(self, x, y):
        dx = x - self.processor.center_x
        dy = self.processor.center_y - y
        
        range_pixels = math.sqrt(dx**2 + dy**2)
        
        if self.custom_scale_distance:
            scale_value = int(self.scale_combo.currentText())
            range_actual = (range_pixels / self.custom_scale_distance) * scale_value
        else:
            bottom_edge_distance = self.processor.image.height - self.processor.center_y
            scale_value = int(self.scale_combo.currentText())
            range_actual = (range_pixels / bottom_edge_distance) * scale_value
        
        azimuth_radians = math.atan2(dx, dy)
        azimuth_degrees = math.degrees(azimuth_radians)
        
        if azimuth_degrees < 0:
            azimuth_degrees += 360
            
        return azimuth_degrees, range_actual
    
    def toggle_center_setting_mode(self):
        self.center_setting_mode = self.set_center_btn.isChecked()
        
        if self.center_setting_mode and self.scale_edge_mode:
            self.scale_edge_mode = False
            self.scale_edge_btn.setChecked(False)
            self.scale_edge_btn.setStyleSheet("")
        
        self.image_label.set_center_setting_mode(self.center_setting_mode)
        
        if self.center_setting_mode:
            self.set_center_btn.setStyleSheet("background-color: #FF9800; color: white;")
            self.add_result("Center setting mode active - click to set new center")
        else:
            self.set_center_btn.setStyleSheet("")
    
    def set_center_point(self, x, y):
        """Встановлення центру з збереженням налаштувань"""
        if not self.processor:
            return
        
        x = max(0, min(x, self.processor.image.width - 1))
        y = max(0, min(y, self.processor.image.height - 1))
        
        current_center_x = self.processor.center_x
        current_center_y = self.processor.center_y
        
        dx = x - current_center_x
        dy = y - current_center_y
        
        self.processor.move_center(dx, dy)
        
        # Зберегти налаштування сітки
        self.save_current_grid_settings()
        
        self.center_setting_mode = False
        self.set_center_btn.setChecked(False)
        self.set_center_btn.setStyleSheet("")
        self.image_label.set_center_setting_mode(False)
        
        self.display_image()
        
        if self.current_click:
            azimuth, range_val = self.calculate_azimuth_range(
                self.current_click['x'], self.current_click['y']
            )
            self.current_click['azimuth'] = azimuth
            self.current_click['range'] = range_val
        
        self.update_results_display()
        self.update_report_data()
        self.add_result(f"Center moved to: ({self.processor.center_x}, {self.processor.center_y})")
        self.add_result(self.tr("grid_settings_saved"))

    def toggle_scale_edge_mode(self):
        self.scale_edge_mode = self.scale_edge_btn.isChecked()
        
        self.image_label.set_scale_edge_mode(self.scale_edge_mode)
        
        if self.scale_edge_mode:
            self.scale_edge_btn.setStyleSheet("background-color: #4CAF50; color: white;")
            self.add_result(self.tr("scale_edge_active"))
        else:
            self.scale_edge_btn.setStyleSheet("")
    
    def set_scale_edge_point(self, x, y):
        """Встановлення scale edge з збереженням налаштувань"""
        x = max(0, min(x, self.processor.image.width - 1))
        y = max(0, min(y, self.processor.image.height - 1))
        
        dx = x - self.processor.center_x
        dy = y - self.processor.center_y
        distance = math.sqrt(dx*dx + dy*dy)
        
        self.scale_edge_point = {'x': x, 'y': y}
        self.custom_scale_distance = distance
        
        # Зберегти налаштування сітки
        self.save_current_grid_settings()
        
        self.scale_edge_mode = False
        self.scale_edge_btn.setChecked(False)
        self.scale_edge_btn.setStyleSheet("")
        self.image_label.set_scale_edge_mode(False)
        
        self.display_image()
        
        if self.current_click:
            azimuth, range_val = self.calculate_azimuth_range(
                self.current_click['x'], self.current_click['y']
            )
            self.current_click['azimuth'] = azimuth
            self.current_click['range'] = range_val
        
        self.update_results_display()
        self.update_report_data()
        self.add_result(self.tr("scale_edge_set").format(distance=distance))
        self.add_result(self.tr("grid_settings_saved"))

    def update_report_data(self):
        if not self.processor:
            self.auto_azimuth_label.setText("β - --°")
            self.auto_distance_label.setText("D - -- км")
            self.auto_scale_label.setText("M = --")
            return
        if self.current_click:
            azimuth = self.current_click['azimuth']
            distance = self.current_click['range']
            scale = int(self.scale_combo.currentText())
            
            self.auto_azimuth_label.setText(f"β - {azimuth:.0f}°")
            self.auto_distance_label.setText(f"D - {distance:.1f} км")
            self.auto_scale_label.setText(f"M = {scale}")
        else:
            self.auto_azimuth_label.setText("β - --°")
            self.auto_distance_label.setText("D - -- км")
            
            if hasattr(self, 'scale_combo'):
                scale = int(self.scale_combo.currentText())
                self.auto_scale_label.setText(f"M = {scale}")
            else:
                self.auto_scale_label.setText("M = --")

    def update_target_number(self, text):
        self.current_target_number = text
    
    def update_height(self, text):
        self.current_height = text
    
    def update_obstacles(self, text):
        self.current_obstacles = text
    
    def update_detection(self, text):
        self.current_detection = text

    def update_results_display(self):
        self.results_text.clear()
        
        if self.processor:
            self.add_result(self.tr("image_info").format(name=os.path.basename(self.current_image_path)))
            self.add_result(self.tr("size").format(width=self.processor.image.width, 
                                                  height=self.processor.image.height))
            self.add_result(self.tr("scale_info").format(scale=int(self.scale_combo.currentText())))
            self.add_result(self.tr("center_info").format(x=self.processor.center_x, 
                                                         y=self.processor.center_y))
            
            if self.custom_scale_distance:
                self.add_result(f"Custom scale edge: {self.custom_scale_distance:.1f} px = {self.scale_combo.currentText()} units")
            else:
                bottom_distance = self.processor.image.height - self.processor.center_y
                self.add_result(self.tr("bottom_edge").format(scale=int(self.scale_combo.currentText())))
                self.add_result(self.tr("pixels_south").format(pixels=bottom_distance))
            self.add_result("")
        
        if self.current_click:
            self.add_result(self.tr("analysis_point"))
            self.add_result(f"{self.tr('position')}: ({self.current_click['x']}, {self.current_click['y']})")
            self.add_result(f"{self.tr('azimuth')}: {self.current_click['azimuth']:.0f}°")
            self.add_result(f"{self.tr('range')}: {self.current_click['range']:.0f} км")
            self.add_result("")
            self.add_result(self.tr("click_to_place"))
            self.add_result(self.tr("drag_to_move"))
            self.add_result(self.tr("line_connects"))
        else:
            self.add_result(self.tr("click_on_image"))

    def update_scale(self):
        """Оновлення масштабу з збереженням налаштувань"""
        if self.processor:
            new_scale = int(self.scale_combo.currentText())
            
            # ЗБЕРЕГТИ налаштування сітки
            self.save_current_grid_settings()
            
            if self.current_click:
                azimuth, range_val = self.calculate_azimuth_range(
                    self.current_click['x'], self.current_click['y']
                )
                self.current_click['azimuth'] = azimuth
                self.current_click['range'] = range_val
            
            self.update_results_display()
            self.update_report_data()
            self.add_result(self.tr("scale_updated").format(scale=new_scale))
            self.add_result("Grid settings saved for next images")
    
    def move_center(self, dx, dy):
        if not self.processor:
            QMessageBox.warning(self, self.tr("warning"), self.tr("no_image_first"))
            return
        
        self.processor.move_center(dx, dy)
        
        if self.current_click:
            azimuth, range_val = self.calculate_azimuth_range(
                self.current_click['x'], self.current_click['y']
            )
            self.current_click['azimuth'] = azimuth
            self.current_click['range'] = range_val
        
        self.display_image()
        self.update_results_display()
        self.update_report_data()
        self.add_result(self.tr("center_moved").format(x=self.processor.center_x, 
                                                    y=self.processor.center_y))
    
        if self.image_label.zoom_widget.isVisible():
            self.image_label.zoom_widget.update_cursor_position(self.processor.center_x, self.processor.center_y)
    
    def clear_results(self):
        self.current_click = None
        if self.processor:
            self.display_image()
        self.update_results_display()
        self.update_report_data()

    def save_current_image(self):
        if not self.processor:
            QMessageBox.warning(self, self.tr("warning"), self.tr("no_image_first"))
            return
        
        if not self.current_click:
            QMessageBox.warning(self, self.tr("warning"), self.tr("no_analysis_point"))
            return
        
        file_path, _ = QFileDialog.getSaveFileName(
            self, self.tr("save_processed_image"), "",
            f"{self.tr('jpeg_files')} (*.jpg);;{self.tr('png_files')} (*.png);;{self.tr('all_files')} (*.*)"
        )
        
        if file_path:
            try:
                final_image = self.processor.image.copy()
                
                if file_path.lower().endswith(('.jpg', '.jpeg')):
                    if final_image.mode != 'RGB':
                        if final_image.mode == 'RGBA':
                            rgb_image = Image.new('RGB', final_image.size, (255, 255, 255))
                            rgb_image.paste(final_image, mask=final_image.split()[-1])
                            final_image = rgb_image
                        else:
                            final_image = final_image.convert('RGB')
                
                draw = ImageDraw.Draw(final_image)
                
                draw.line([
                    (self.current_click['x'], self.current_click['y']), 
                    (final_image.width - 1, 0)
                ], fill='blue', width=3)
                
                if file_path.lower().endswith('.png'):
                    final_image.save(file_path, 'PNG')
                else:
                    final_image.save(file_path, 'JPEG', quality=95)
                
                self.add_result(f"{self.tr('saved')}: {os.path.basename(file_path)}")
                self.add_result(self.tr("saved_clean"))
                QMessageBox.information(self, self.tr("success"), self.tr("image_saved_successfully"))
                
            except Exception as e:
                QMessageBox.critical(self, self.tr("error"), 
                                   self.tr("could_not_save").format(error=str(e)))

    def save_current_image_data(self):
        if not self.processor or not self.current_click:
            QMessageBox.warning(self, "Warning", "No image or analysis point to save")
            return False
        
        image_data = {
            'image_path': self.current_image_path,
            'image_name': os.path.basename(self.current_image_path),
            'analysis_point': {
                'x': self.current_click['x'],
                'y': self.current_click['y'],
                'azimuth': self.current_click['azimuth'],
                'range': self.current_click['range']
            },
            'target_data': {
                'number': self.current_target_number,
                'height': self.current_height,
                'obstacles': self.current_obstacles,
                'detection': self.current_detection,
                'scale': int(self.scale_combo.currentText())
            },
            'processor_settings': {
                'center_x': self.processor.center_x,
                'center_y': self.processor.center_y,
                'scale_edge_point': self.scale_edge_point,
                'custom_scale_distance': self.custom_scale_distance
            }
        }
        
        existing_index = -1
        for i, saved_data in enumerate(self.processed_images):
            if saved_data['image_path'] == self.current_image_path:
                existing_index = i
                break
        
        if existing_index >= 0:
            self.processed_images[existing_index] = image_data
            self.add_result(f"Updated data for: {image_data['image_name']}")
        else:
            self.processed_images.append(image_data)
            self.add_result(f"Saved data for: {image_data['image_name']}")
        
        self.add_result(f"Total processed images: {len(self.processed_images)}")
        return True

    def next_image_in_batch(self):
        if not self.current_folder_images:
            QMessageBox.warning(self, "Warning", "No folder loaded for batch processing")
            return
        
        if self.processor and self.current_click:
            self.save_current_image_data()
        
        self.current_image_index = (self.current_image_index + 1) % len(self.current_folder_images)
        next_image = self.current_folder_images[self.current_image_index]
        
        self.load_image(next_image)
        self.add_result(f"Loaded image {self.current_image_index + 1}/{len(self.current_folder_images)}: {os.path.basename(next_image)}")

    def previous_image_in_batch(self):
        if not self.current_folder_images:
            QMessageBox.warning(self, "Warning", "No folder loaded for batch processing")
            return
        
        if self.processor and self.current_click:
            self.save_current_image_data()
        
        self.current_image_index = (self.current_image_index - 1) % len(self.current_folder_images)
        prev_image = self.current_folder_images[self.current_image_index]
        
        self.load_image(prev_image)
        self.add_result(f"Loaded image {self.current_image_index + 1}/{len(self.current_folder_images)}: {os.path.basename(prev_image)}")

    def create_batch_album(self):
        """ОНОВЛЕНИЙ МЕТОД з використанням нульових полів для таблиць"""
        if not DOCX_AVAILABLE:
            QMessageBox.warning(self, "Warning", "python-docx library not installed")
            return
        
        if not self.processed_images:
            QMessageBox.warning(self, "Warning", "No processed images to export")
            return
        
        # Зберігаємо поточне зображення якщо є
        if self.processor and self.current_click:
            self.save_current_image_data()
        
        # Діалог вибору файлу
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Create Complete Album", "",
            "Word Documents (*.docx);;All files (*.*)"
        )
        
        if file_path:
            try:
                # Отримуємо дані для титульної сторінки
                title_data = self.get_title_page_data_from_gui()
                
                # Створюємо повний альбом з правильними полями
                success = create_complete_album(self.processed_images, title_data, file_path)
                
                if success:
                    QMessageBox.information(self, "Success", 
                                        f"Complete album created!\n"
                                        f"- Title page (standard margins)\n"
                                        f"- Description page (standard margins)\n" 
                                        f"- {len(self.processed_images)} image tables (zero margins)\n"
                                        f"Saved: {os.path.basename(file_path)}")
                    
                    self.add_result(f"✓ Complete album created: {os.path.basename(file_path)}")
                    self.add_result(f"✓ Contains {len(self.processed_images)} processed images")
                else:
                    QMessageBox.critical(self, "Error", "Failed to create album")
                    
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Could not create album: {str(e)}")
                print(f"Error creating album: {e}")

    def get_title_page_data_from_gui(self):
        """Отримання даних для титульної сторінки з GUI"""
        from datetime import datetime
        
        title_data = {
            'target_no': self.current_target_number,
            'date': datetime.now(),
            'unit_info': get_default_unit_info(),
            'commander_info': get_default_commander_info()
        }
        
        return title_data

    def create_new_word_album(self):
        if not DOCX_AVAILABLE:
            QMessageBox.warning(self, self.tr("warning"), 
                               self.tr("docx_not_available"))
            return
        
        if not self.processor or not self.current_click:
            QMessageBox.warning(self, self.tr("warning"), 
                               self.tr("load_image_and_point"))
            return
        
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Create New Word Album", "",
            "Word Documents (*.docx);;All files (*.*)"
        )
        
        if file_path:
            try:
                self.current_word_document = Document()
                self.word_document_path = file_path
                
                self.set_document_defaults()
                self.add_image_to_word_document()
                
                self.current_word_document.save(file_path)
                
                QMessageBox.information(self, self.tr("success"), 
                                       f"New album created: {os.path.basename(file_path)}")
                self.add_result(f"Created new album: {os.path.basename(file_path)}")
                
            except Exception as e:
                QMessageBox.critical(self, self.tr("error"), 
                                   f"Could not create Word album: {str(e)}")
    
    def add_to_existing_album(self):
        if not DOCX_AVAILABLE:
            QMessageBox.warning(self, self.tr("warning"), 
                               self.tr("docx_not_available"))
            return
        
        if not self.processor or not self.current_click:
            QMessageBox.warning(self, self.tr("warning"), 
                               self.tr("load_image_and_point"))
            return
        
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Select Existing Word Album", "",
            "Word Documents (*.docx);;All files (*.*)"
        )
        
        if file_path:
            try:
                self.current_word_document = Document(file_path)
                self.word_document_path = file_path
                
                self.set_document_defaults()
                self.add_image_to_word_document()
                
                self.current_word_document.save(file_path)
                
                QMessageBox.information(self, self.tr("success"), 
                                       f"Added to album: {os.path.basename(file_path)}")
                self.add_result(f"Added to album: {os.path.basename(file_path)}")
                
            except Exception as e:
                QMessageBox.critical(self, self.tr("error"), 
                                   f"Could not add to Word album: {str(e)}")
                
    def set_document_arial_font(self, doc):
        try:
            styles = doc.styles
            normal_style = styles['Normal']
            normal_style.font.name = 'Arial'
            normal_style.font.size = Pt(12)
            
            for style in styles:
                if hasattr(style, 'font'):
                    style.font.name = 'Arial'
                    
        except Exception as e:
            pass
        
    def create_js_style_table(self, doc, image_data):
        try:
            table = doc.add_table(rows=1, cols=3)
            table.style = None
            
            tbl = table._tbl
            tblPr = tbl.tblPr
            
            tblW = OxmlElement('w:tblW')
            tblW.set(qn('w:type'), 'pct')
            tblW.set(qn('w:w'), '5000')
            tblPr.append(tblW)
            
            tblLayout = OxmlElement('w:tblLayout')
            tblLayout.set(qn('w:type'), 'fixed')
            tblPr.append(tblLayout)
            
            row = table.rows[0]
            label_cell = row.cells[0]
            image_cell = row.cells[1] 
            data_cell = row.cells[2]
            
            self.set_js_cell_widths(label_cell, image_cell, data_cell)
            row.height = Cm(13)
            
            self.fill_js_label_cell(label_cell)
            self.fill_js_image_cell(image_cell, image_data)
            self.fill_js_data_cell(data_cell, image_data)
            
        except Exception as e:
            pass

    def set_js_cell_widths(self, label_cell, image_cell, data_cell):
        try:
            label_cell.width = Cm(2.55)
            data_cell.width = Cm(3)
            image_cell.width = Cm(15)
        except Exception as e:
            pass

    def fill_js_label_cell(self, cell):
        try:
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = para.add_run("Індикатор ЗРЛ")
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            self.set_js_label_cell_borders(cell)
        except Exception as e:
            pass

    def set_js_label_cell_borders(self, cell):
        try:
            tc = cell._tc
            tcPr = tc.tcPr
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.append(tcPr)
            
            tcBorders = OxmlElement('w:tcBorders')
            
            left_border = OxmlElement('w:left')
            left_border.set(qn('w:val'), 'none')
            tcBorders.append(left_border)
            
            right_border = OxmlElement('w:right')
            right_border.set(qn('w:val'), 'single')
            right_border.set(qn('w:sz'), '4')
            right_border.set(qn('w:color'), '000000')
            tcBorders.append(right_border)
            
            top_border = OxmlElement('w:top')
            top_border.set(qn('w:val'), 'none')
            tcBorders.append(top_border)
            
            bottom_border = OxmlElement('w:bottom')
            bottom_border.set(qn('w:val'), 'none')
            tcBorders.append(bottom_border)
            
            tcPr.append(tcBorders)
        except Exception as e:
            pass

    def fill_js_image_cell(self, cell, image_data):
        try:
            processed_image = self.create_processed_image_from_data(image_data)
            
            if not processed_image:
                return
            
            with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as temp_file:
                temp_path = temp_file.name
                processed_image.save(temp_path, 'JPEG', quality=95)
            
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            img_width = Cm(14.5)
            img_height = Cm(14.5 * 13/15)
            
            run = para.add_run()
            run.add_picture(temp_path, width=img_width, height=img_height)
            
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            self.set_cell_shading(cell, "FFFFFF")
            self.set_js_image_cell_borders(cell)
            
            try:
                os.remove(temp_path)
            except:
                pass
        except Exception as e:
            pass

    def create_processed_image_from_data(self, image_data):
        try:
            with Image.open(image_data['image_path']) as original_image:
                if original_image.mode != 'RGB':
                    if original_image.mode == 'RGBA':
                        rgb_image = Image.new('RGB', original_image.size, (255, 255, 255))
                        rgb_image.paste(original_image, mask=original_image.split()[-1])
                        final_image = rgb_image
                    else:
                        final_image = original_image.convert('RGB')
                else:
                    final_image = original_image.copy()
            
            draw = ImageDraw.Draw(final_image)
            analysis_point = image_data['analysis_point']
            
            draw.line([
                (analysis_point['x'], analysis_point['y']), 
                (final_image.width - 1, 0)
            ], fill='black', width=3)
            
            return final_image
        except Exception as e:
            return None

    def set_cell_shading(self, cell, color):
        try:
            tc = cell._tc
            tcPr = tc.tcPr
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.append(tcPr)
            
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), color)
            tcPr.append(shd)
        except Exception as e:
            pass

    def set_js_image_cell_borders(self, cell):
        try:
            tc = cell._tc
            tcPr = tc.tcPr
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.append(tcPr)
            
            tcBorders = OxmlElement('w:tcBorders')
            
            for border_name in ['left', 'right', 'top', 'bottom']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:color'), '000000')
                tcBorders.append(border)
            
            tcPr.append(tcBorders)
        except Exception as e:
            pass

    def fill_js_data_cell(self, cell, image_data):
        try:
            cell.text = ""
            
            target_data = image_data['target_data']
            analysis_point = image_data['analysis_point']
            
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            run1 = para.add_run(" ")
            run1.font.name = 'Arial'
            run1.add_break()
            run1.add_break()
            
            run2 = para.add_run(target_data['number'])
            run2.font.name = 'Arial'
            run2.font.size = Pt(12)
            run2.italic = True
            run2.underline = True
            run2.add_break()
            
            run3 = para.add_run(f"β – {analysis_point['azimuth']:.0f}º")
            run3.font.name = 'Arial'
            run3.font.size = Pt(12)
            run3.italic = True
            run3.add_break()
            
            run4 = para.add_run(f"D – {analysis_point['range']:.1f} км")
            run4.font.name = 'Arial'
            run4.font.size = Pt(12)
            run4.italic = True
            run4.add_break()
            
            run5 = para.add_run("без перешкод")
            run5.font.name = 'Arial'
            run5.font.size = Pt(9)
            run5.italic = True
            run5.add_break()
            
            run6 = para.add_run(target_data['detection'])
            run6.font.name = 'Arial'
            run6.font.size = Pt(9)
            run6.italic = True
            run6.add_break()
            
            run7 = para.add_run(f"М – {target_data['scale']}")
            run7.font.name = 'Arial'
            run7.font.size = Pt(9)
            run7.italic = True
            
            self.set_js_data_cell_borders(cell)
        except Exception as e:
            pass

    def set_js_data_cell_borders(self, cell):
        try:
            tc = cell._tc
            tcPr = tc.tcPr
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.append(tcPr)
            
            tcBorders = OxmlElement('w:tcBorders')
            
            left_border = OxmlElement('w:left')
            left_border.set(qn('w:val'), 'none')
            tcBorders.append(left_border)
            
            for border_name in ['right', 'top', 'bottom']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:color'), '000000')
                tcBorders.append(border)
            
            tcPr.append(tcBorders)
        except Exception as e:
            pass

    def set_document_defaults(self):
        try:
            styles = self.current_word_document.styles
            default_style = styles['Normal']
            default_style.font.name = 'Arial'
            default_style.font.size = Pt(12)
            
            section = self.current_word_document.sections[0]
            section.orientation = 1
            
            new_width, new_height = section.page_height, section.page_width
            section.page_width = new_width
            section.page_height = new_height
        except Exception as e:
            pass
    
    def add_image_to_word_document(self):
        if not self.current_word_document or not self.processor or not self.current_click:
            return
        
        table = self.current_word_document.add_table(rows=1, cols=3)
        table.style = None
        
        tbl = table._tbl
        tblPr = tbl.tblPr
        
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        tblPr.append(tblBorders)

        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:type'), 'pct')
        tblW.set(qn('w:w'), '5000')
        tblPr.append(tblW)
        
        label_cell = table.cell(0, 0)
        image_cell = table.cell(0, 1)
        data_cell = table.cell(0, 2)
        
        label_cell.width = Cm(3)
        image_cell.width = Cm(20)
        data_cell.width = Cm(4)
        
        self.configure_label_cell(label_cell)
        
        for cell in [image_cell, data_cell]:
            cellPr = cell._tc.tcPr
            if cellPr is None:
                cellPr = OxmlElement('w:tcPr')
                cell._tc.append(cellPr)
            
            tcMar = OxmlElement('w:tcMar')
            for margin in ['top', 'left', 'bottom', 'right']:
                mar = OxmlElement(f'w:{margin}')
                mar.set(qn('w:w'), '108')
                mar.set(qn('w:type'), 'dxa')
                tcMar.append(mar)
            cellPr.append(tcMar)
            
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), 'center')
            cellPr.append(vAlign)
        
        self.add_label_to_cell(label_cell)
        self.add_processed_image_to_cell(image_cell)
        self.add_data_to_cell(data_cell)
        
        if len(self.current_word_document.tables) > 1:
            para = self.current_word_document.add_paragraph()
            para.text = ""
    
    def configure_label_cell(self, cell):
        cellPr = cell._tc.tcPr
        if cellPr is None:
            cellPr = OxmlElement('w:tcPr')
            cell._tc.append(cellPr)
        
        tcBorders = OxmlElement('w:tcBorders')
        
        for border_name in ['top', 'left', 'bottom']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            tcBorders.append(border)
        
        right_border = OxmlElement('w:right')
        right_border.set(qn('w:val'), 'single')
        right_border.set(qn('w:sz'), '4')
        right_border.set(qn('w:space'), '0')
        right_border.set(qn('w:color'), '000000')
        tcBorders.append(right_border)
        
        cellPr.append(tcBorders)
        
        tcMar = OxmlElement('w:tcMar')
        for margin in ['top', 'left', 'bottom', 'right']:
            mar = OxmlElement(f'w:{margin}')
            mar.set(qn('w:w'), '108')
            mar.set(qn('w:type'), 'dxa')
            tcMar.append(mar)
        cellPr.append(tcMar)
        
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        cellPr.append(vAlign)
    
    def add_label_to_cell(self, cell):
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = para.add_run("Індикатор ЗРЛ")
        run.font.size = Pt(12)
    
    def add_processed_image_to_cell(self, cell):
        try:
            if not self.processor or not hasattr(self.processor, 'image'):
                return
            
            final_image = self.processor.image.copy()
            
            if final_image.mode != 'RGB':
                try:
                    if final_image.mode == 'RGBA':
                        rgb_image = Image.new('RGB', final_image.size, (255, 255, 255))
                        rgb_image.paste(final_image, mask=final_image.split()[-1])
                        final_image = rgb_image
                    else:
                        final_image = final_image.convert('RGB')
                except Exception:
                    final_image = final_image.convert('RGB')
            
            draw = ImageDraw.Draw(final_image)
            
            draw.line([
                (self.current_click['x'], self.current_click['y']), 
                (final_image.width - 1, 0)
            ], fill='blue', width=3)
            
            with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as temp_file:
                temp_path = temp_file.name
                final_image.save(temp_path, 'JPEG', quality=95)
            
            cell.text = ""
            top_spacing_para = cell.paragraphs[0]
            top_spacing_para.text = ""
            
            pPr = top_spacing_para._element.get_or_add_pPr()
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:before'), '0')
            spacing.set(qn('w:after'), '720')
            pPr.append(spacing)
            
            img_paragraph = cell.add_paragraph()
            img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = img_paragraph.add_run()
            run.add_picture(temp_path, width=Cm(19))
            
            try:
                os.remove(temp_path)
            except:
                pass
        except Exception as e:
            pass
    
    def add_data_to_cell(self, cell):
        try:
            cell.text = ""
            
            target_num = self.current_target_number
            azimuth = self.current_click['azimuth']
            distance = self.current_click['range']
            height = self.current_height
            obstacles = self.current_obstacles
            detection = self.current_detection
            scale = int(self.scale_combo.currentText())
            
            top_spacing_para = cell.paragraphs[0]
            top_spacing_para.text = ""
            pPr = top_spacing_para._element.get_or_add_pPr()
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:before'), '0')
            spacing.set(qn('w:after'), '720')
            pPr.append(spacing)
            
            data_items = [
                (f"{target_num}", 12, True, True),
                (f"β - {azimuth:.0f}°", 12, True, False),
                (f"D - {distance:.1f} км", 12, True, False),
                (f"Н - {height} км", 12, True, False),
                (f"{obstacles}", 9, True, False),
                (f"{detection}", 9, True, False),
                (f"М={scale}", 9, True, False)
            ]
            
            for text, font_size, italic, underline in data_items:
                p = cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                pPr = p._element.get_or_add_pPr()
                spacing = OxmlElement('w:spacing')
                spacing.set(qn('w:before'), '0')
                spacing.set(qn('w:after'), '0')
                spacing.set(qn('w:line'), '200')
                spacing.set(qn('w:lineRule'), 'auto')
                pPr.append(spacing)
                
                run = p.add_run(text)
                run.font.size = Pt(font_size)
                run.italic = italic
                run.underline = underline
                
                rPr = run._element.rPr
                if rPr is None:
                    rPr = OxmlElement('w:rPr')
                    run._element.insert(0, rPr)
                spacing = OxmlElement('w:spacing')
                spacing.set(qn('w:val'), '-5')
                rPr.append(spacing)
        except Exception as e:
            pass

def main():
    app = QApplication(sys.argv)
    window = AzimuthGUI()
    window.show()
    sys.exit(app.exec_())

if __name__ == "__main__":
    main()